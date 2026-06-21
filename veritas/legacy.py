#!/usr/bin/env python3
"""Paper Audit - 学术论文自动审查工具 [耿同学版]
基于3个开源项目思路开发：
- wooly99/geng-academic-fraud-detector 耿同学六式
- NeoSpecies/AcademicIntegrityHunter 本地统计算法
- jingshouyan/academic-integrity-geng 五维审查体系
输入论文文件或目录 → 文本提取 → 本地统计检测 + LLM语义分析 → 输出md/html格式报告
用法: python paper_audit.py <paper_path> [--mineru] [--max-chars 8000] [--output report.md]
"""
import re, json, time, argparse, urllib.request, urllib.parse, zlib, math, collections, os, mimetypes, fnmatch, csv, platform, webbrowser, subprocess, sys, requests, builtins, hashlib, html, base64, io, concurrent.futures, signal, threading, datetime
from pathlib import Path
from typing import Tuple, Dict, List, Any, Callable

from .adapter_types import (
    AdapterResult,
    AuditAdapters,
    ImageDetectorAdapter,
    ImageSemanticAdapter,
    MinerUAdapter,
    ReferenceLookupAdapter,
    TextLLMAdapter,
)
from .artifacts import (
    _artifact_base_from_output,
    apply_audit_artifact_type,
    audit_artifact_paths,
    audit_limited_reasons,
    coverage_blocking_failure,
    explicit_output_path_from_args,
    failed_audit_artifact_paths,
)
from .failed_diagnostics import adapter_failure_to_audit_failure, failed_audit_payload, preflight_failure_to_audit_failure
from .fake_adapters import (
    FakeImageDetectorAdapter,
    FakeImageSemanticAdapter,
    FakeMinerUAdapter,
    FakeReferenceLookupAdapter,
    FakeScenarioMixin,
    FakeTextLLMAdapter,
    fake_audit_adapters,
)
from .file_utils import _json_load, _json_save, _load_merged_json_dicts, _safe_name
from .followups import (
    _followup_language_instruction,
    _followup_tone_instruction,
    _normalize_custom_concerns,
    _normalize_followup_issues,
    _split_author_text,
    _followup_draft_filename,
    _followup_output_dir,
    build_followup_generation_context,
    build_followup_prompt,
    generate_and_save_followup_draft_from_namespace,
    generate_followup_draft_from_namespace,
    load_existing_followups,
    normalize_article_identity,
    normalize_followup_language,
    normalize_followup_tone,
    save_followup_artifacts_from_namespace,
)
from .html_utils import _html_escape, _json_for_script_tag
from .image_cache import _image_file_fingerprint_from_namespace, _image_semantic_cache_key_from_namespace
from .image_reporting import (
    _image_detector_display,
    _image_semantic_display,
    format_image_audit_html as _format_image_audit_html,
    format_image_audit_markdown as _format_image_audit_markdown,
    save_image_review_manifest as _save_image_review_manifest,
)
from .image_results import (
    _detector_timeout_result,
    _extract_json_object,
    _glm_error_result,
    _glm_timeout_result,
    _normalize_detector_result,
    _normalize_glm_image_result,
)
from .image_selection import (
    _flush_image_cache,
    _image_audit_sort_key,
    _image_detector_priority_key,
    _image_semantic_priority_key,
)
from .local_analysis import benford_analysis, extract_all_numbers, local_stat_check, smart_chunk_text
from .config import (
    CapabilityConfig,
    RuntimeConfig,
    apply_runtime_config_to_namespace,
    default_runtime_config_from_namespace,
    load_runtime_config_from_namespace,
)
from .desktop_gui import (
    DESKTOP_GUI_ARTIFACT_LABELS,
    DESKTOP_GUI_CONFIG_CAPABILITIES,
    DESKTOP_GUI_CONFIG_DEPENDENCIES,
    DESKTOP_GUI_FOLLOWUP_LABELS,
    DESKTOP_GUI_LLM_CONFIG_FIELDS,
    DESKTOP_GUI_STAGE_LABELS,
    _desktop_gui_preflight_status_label,
    _desktop_gui_report_type_label,
    _desktop_gui_risk_label,
    _desktop_gui_stage_label,
    _desktop_gui_status_label,
    desktop_gui_artifact_preview,
    desktop_gui_checked_config_snapshot_from_namespace,
    desktop_gui_config_file_path,
    desktop_gui_config_snapshot,
    desktop_gui_followup_context_from_namespace,
    desktop_gui_generate_followup_draft_from_namespace,
    desktop_gui_progress_from_log_line,
    desktop_gui_run_summary,
    desktop_gui_start_run,
    desktop_gui_write_llm_config as _desktop_gui_write_llm_config,
    open_desktop_path,
)
from .evidence_rendering import (
    _clean_mineru_table_block,
    _escaped_html_table_fragment_to_html,
    _evidence_contains_table,
    _is_markdown_table_separator,
    _looks_like_markdown_table,
    _markdown_table_to_html,
    _parse_html_table_rows,
    _plain_table_summary_text,
    _render_data_table_html,
    _render_unmarked_evidence_html,
    _split_markdown_table_row,
    render_evidence_html,
    render_evidence_summary_html,
)
from .runtime_metadata import ensure_runtime_meta, runtime_metadata, runtime_utc_year
from . import risk_rules as _risk_rules
from .text_utils import _brief_text, _normalize_title, _title_tokens, _token_similarity
from .models import (
    AuditFailure,
    AuditReportModel,
    CoverageModel,
    EvidenceFinding,
    ImageAuditModel,
    ReferenceAuditModel,
    RunMetadataModel,
)
from .mineru_text import (
    _extract_mineru_structured_text,
    _flatten_mineru_items,
    _format_mineru_content_list,
    _format_mineru_table_block,
    _mineru_block_text,
    _mineru_block_type,
    _mineru_content_text,
    _normalize_markdown_table,
    _normalize_mineru_table_text,
    _table_rows_to_markdown,
)
from .paper_identity import extract_paper_identity
from .preflight import _chat_completions_endpoint, preflight_mineru_from_namespace, preflight_text_llm_from_namespace
from .preflight_types import PreflightResult, run_preflight_once
from .project_files import (
    SUPPORTED_TEXT_FILE_EXTENSIONS,
    _is_missing_meta_value,
    _main_paper_score,
    find_project_files,
    normalize_run_meta,
)
from .production_adapters import (
    ProductionImageDetectorAdapter,
    ProductionImageSemanticAdapter,
    ProductionMinerUAdapter,
    ProductionReferenceLookupAdapter,
    ProductionTextLLMAdapter,
    default_audit_adapters,
)
from .run_types import RunRequest, RunResult
from .report_schema import LLM_REQUIRED_FINDING_FIELDS, _json_string_unescape, normalize_llm_report_schema, parse_report
from .retry_commands import _shell_quote, default_retry_command, retry_command_from_args
from .risk_rule_helpers import (
    _CHECK_STOPWORDS,
    _append_unique_evidence,
    _brief_check_list,
    _build_merged_conclusion,
    _build_merged_summary,
    _check_label_for_summary,
    _check_member_summary,
    _check_merge_key,
    _check_severity,
    _check_similarity,
    _check_text_blob,
    _check_text_for_scoring,
    _downgrade_extraction_red_flags,
    _downgrade_unverified_future_publication_checks,
    _extract_years_from_check,
    _is_future_publication_check,
    _is_extraction_limited_check,
    _max_risk,
    _merge_check_into,
    _normalize_check_terms,
    _risk_index,
    _same_or_similar_check,
    _should_downgrade_extraction_red_flag,
    _soften_extraction_red_flag_language,
    _soften_nonfinal_red_flag_language,
)
from .versions import ADAPTER_VERSION, PROMPT_VERSION, RISK_RULE_VERSION, SCHEMA_VERSION
from .web_runner_paths import (
    _web_runner_common_search_roots,
    _web_runner_is_basename_only_input,
    resolve_web_runner_input_path,
)
from .web_runner import (
    _web_runner_capability_status,
    _web_runner_history_path,
    _web_runner_input_parts,
    _web_runner_now,
    _web_runner_output_base,
    _web_runner_report_summary_from_payload,
    _web_runner_run_id,
    _web_runner_safe_run,
    _web_runner_timestamp,
    dropped_local_path_from_uri_text,
    pick_local_path,
    web_runner_config_status_from_namespace,
    web_runner_default_output_stem_from_namespace,
)
from .workspace import (
    create_run_workspace,
    record_run_workspace_artifacts,
    record_run_workspace_json,
    run_workspace_path,
)

_risk_rules.runtime_utc_year = lambda: runtime_utc_year()
apply_risk_rules = _risk_rules.apply_risk_rules
merge_chunk_reports = _risk_rules.merge_chunk_reports

# Windows/重定向控制台默认GBK时，emoji/中文符号可能触发UnicodeEncodeError；统一兜底为UTF-8。
try:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


# ══════════════════════════════════════════════════════════════
# 运行日志 / 进度条 / 输出目录控制
# ══════════════════════════════════════════════════════════════
_ORIGINAL_PRINT = builtins.print
_RUN_LOG_FILE = None
_RUN_OUTPUT_DIR = None
_RUN_OUTPUT_STEM = None
_RESUME_EVENTS_ENABLED = True


def get_output_base(input_path: Path):
    """返回所有运行产物的基准目录和名称。

    规则：输入文件→文件所在目录；输入目录→该目录本身。
    """
    input_path = Path(input_path)
    if input_path.is_dir():
        return input_path, input_path.name or "audit_report"
    return input_path.parent, input_path.stem


def setup_run_logging(input_path: Path):
    """把print同时写到控制台和同目录log文件。"""
    global _RUN_LOG_FILE, _RUN_OUTPUT_DIR, _RUN_OUTPUT_STEM
    out_dir, stem = get_output_base(Path(input_path))
    out_dir.mkdir(parents=True, exist_ok=True)
    _RUN_OUTPUT_DIR = out_dir
    _RUN_OUTPUT_STEM = _safe_name(stem)
    _RUN_LOG_FILE = out_dir / f"{_RUN_OUTPUT_STEM}.paper_audit.log"
    _RUN_LOG_FILE.write_text(
        f"Paper Audit Log\nSTART {time.strftime('%F %T')}\nINPUT {Path(input_path)}\nOUTPUT_DIR {out_dir}\n\n",
        encoding="utf-8"
    )

    def tee_print(*args, **kwargs):
        _ORIGINAL_PRINT(*args, **kwargs)
        try:
            sep = kwargs.get("sep", " ")
            end = kwargs.get("end", "\n")
            msg = sep.join(str(a) for a in args) + end
            with _RUN_LOG_FILE.open("a", encoding="utf-8", errors="replace") as f:
                f.write(msg)
        except Exception:
            pass
    builtins.print = tee_print
    print(f"🧾 日志文件: {_RUN_LOG_FILE}")
    return _RUN_LOG_FILE


def get_resume_dir(output_dir: Path, output_stem: str):
    """断点续作缓存目录。"""
    d = Path(output_dir) / f".{_safe_name(output_stem)}.paper_audit_resume"
    d.mkdir(parents=True, exist_ok=True)
    return d


def resume_event(resume_dir: Path, step: str, status: str, detail: str = "", **extra):
    """记录可断点续作的步骤清单，同时写入普通log。"""
    if not _RESUME_EVENTS_ENABLED:
        return
    try:
        event = {"time": time.strftime("%F %T"), "step": step, "status": status, "detail": detail}
        event.update(extra)
        manifest = Path(resume_dir) / "resume_manifest.jsonl"
        manifest.parent.mkdir(parents=True, exist_ok=True)
        with manifest.open("a", encoding="utf-8") as f:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")
        print(f"🧭 断点记录: {step} | {status} | {detail}")
    except Exception as e:
        print(f"⚠️ 写入断点记录失败: {e}")


def _allow_llm_cache_read(no_resume=False, llm_cache_only=False):
    return (not bool(no_resume)) or bool(llm_cache_only)


def _text_fingerprint(text: str, extra: str = ""):
    h = hashlib.sha256()
    h.update((text or "").encode("utf-8", errors="ignore"))
    h.update(str(extra).encode("utf-8", errors="ignore"))
    return h.hexdigest()[:16]


# LLM运行参数：由CLI覆盖。默认保守，避免一次请求无限阻塞。
LLM_TIMEOUT = 45
LLM_RETRIES = 1
EXTRACT_CACHE_VERSION = 7
MIN_IMAGE_BYTES = 5000
IMAGE_SEMANTIC_CACHE_VERSION = 3


def progress_bar(current, total, label="", width=28):
    """打印一行文本进度条；日志中保留每次更新。"""
    try:
        total = max(int(total), 1)
        current = max(0, min(int(current), total))
        filled = int(width * current / total)
        bar = "█" * filled + "░" * (width - filled)
        pct = current * 100 / total
        print(f"📊 [{bar}] {current}/{total} {pct:5.1f}% {label}")
    except Exception:
        print(f"📊 {current}/{total} {label}")


def save_mineru_artifacts(zip_url: str, zip_data: bytes, source_name: str, output_dir=None, batch_id=None):
    """把MinerU下载链接和zip保存到与输入文件/目录一致的位置。"""
    out_dir = Path(output_dir) if output_dir else (_RUN_OUTPUT_DIR or Path.cwd())
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = _safe_name(Path(source_name).stem if source_name else (batch_id or "mineru"))
    suffix = f".{_safe_name(batch_id)}" if batch_id else ""
    link_path = out_dir / f"{stem}{suffix}.mineru_url.txt"
    zip_path = out_dir / f"{stem}{suffix}.mineru.zip"
    link_path.write_text(zip_url + "\n", encoding="utf-8")
    zip_path.write_bytes(zip_data)
    print(f"  🔗 MinerU下载链接已保存: {link_path}")
    print(f"  📦 MinerU zip已保存: {zip_path} ({len(zip_data)/1024/1024:.2f}MB)")
    return zip_path, link_path

# 可选依赖：处理Word/Excel/Supplement文件
try:
    from docx import Document
    DOCX_SUPPORTED = True
except ImportError:
    DOCX_SUPPORTED = False
try:
    from openpyxl import load_workbook
    EXCEL_SUPPORTED = True
except ImportError:
    EXCEL_SUPPORTED = False

# ══════════════════════════════════════════════════════════════
# 配置区
# ══════════════════════════════════════════════════════════════
import importlib

LLM_API_KEY = ""
LLM_API_URL = "https://api.openai.com/v1/chat/completions"
LLM_MODEL = "gpt-3.5-turbo"
MINERU_TOKEN = ""
MINERU_BASE = "https://mineru.net"
GLM_API_KEY = ""
GLM_API_URL = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
GLM_VISION_MODEL = "glm-4.6v-flash"

def default_runtime_config() -> RuntimeConfig:
    return default_runtime_config_from_namespace(globals())


def load_runtime_config(config_module_name: str = "config", env=os.environ, verbose: bool = True) -> RuntimeConfig:
    """Load config.py and environment variables explicitly for a CLI run."""
    return load_runtime_config_from_namespace(globals(), config_module_name=config_module_name, env=env, verbose=verbose)


def apply_runtime_config(runtime_config: RuntimeConfig):
    """Apply explicit runtime config to the legacy module globals used by current code."""
    return apply_runtime_config_to_namespace(runtime_config, globals())


def format_failed_audit_markdown(failure: AuditFailure, input_path: Path, meta: Dict[str, Any] = None) -> str:
    """Render a failed audit diagnostic report as Markdown."""
    payload = failed_audit_payload(failure, input_path, meta)
    failed = payload["failure"]
    fix_hints = failed["fix_hints"] or ["检查关键服务配置、网络连通性和服务商返回状态后重试。"]
    completed_stages = failed["completed_stages"] or ["无"]
    retry_command = failed["retry_command"] or f"python paper_audit.py {json.dumps(str(input_path), ensure_ascii=False)} --json"
    lines = [
        "# 学术论文审查失败诊断",
        "",
        "> 未生成完整审查报告。关键审查能力失败，本次运行只生成失败诊断产物。",
        "",
        "## 失败恢复面板",
        "",
        f"**文件**: `{input_path}`",
        f"**产物类型**: failed",
        f"**完整审查报告已生成**: 否",
        f"**失败时间**: {payload['created_at']}",
        "",
        "## 失败能力",
        "",
        f"- 能力: `{failed['capability']}`",
        f"- 错误类别: `{failed['error_class']}`",
        f"- 错误信息: {failed['message']}",
        "",
        "## 已完成阶段",
        "",
    ]
    lines.extend(f"- {stage}" for stage in completed_stages)
    lines.extend([
        "",
        "## 修复建议",
        "",
    ])
    lines.extend(f"- {hint}" for hint in fix_hints)
    lines.extend([
        "",
        "## 重试命令",
        "",
        "```bash",
        retry_command,
        "```",
    ])
    if failed["details"]:
        lines.extend([
            "",
            "## 技术细节",
            "",
            "```json",
            json.dumps(failed["details"], ensure_ascii=False, indent=2),
            "```",
        ])
    meta = payload.get("meta") or {}
    completed_sections = []
    if meta.get("resource_audit") is not None:
        completed_sections.extend(format_resource_audit_markdown(meta.get("resource_audit")))
    if meta.get("reference_audit") is not None:
        completed_sections.extend(format_reference_audit_markdown(meta.get("reference_audit")))
    if completed_sections:
        lines.extend([
            "",
            "## 已完成校检摘要",
            "",
            "> 以下为失败前已经完成并写入 JSON 的正式校检结果；失败能力修复后应重新运行以生成完整报告。",
            "",
        ])
        lines.extend(completed_sections)
    return "\n".join(lines)


def format_failed_audit_html(failure: AuditFailure, input_path: Path, meta: Dict[str, Any] = None) -> str:
    payload = failed_audit_payload(failure, input_path, meta)
    failed = payload["failure"]
    runtime = payload.get("meta", {}).get("runtime") or {}
    stages = failed.get("completed_stages") or ["无"]
    hints = failed.get("fix_hints") or ["检查关键服务配置、网络连通性和服务商返回状态后重试。"]
    retry_command = failed.get("retry_command") or f"python paper_audit.py {_shell_quote(str(input_path))} --json"
    details = failed.get("details") or {}
    cache_state = details.get("resume_dir") or details.get("cache_dir") or runtime.get("resume_dir") or "见运行日志/工作区"
    stages_html = "".join(f"<li>{_html_escape(stage)}</li>" for stage in stages)
    hints_html = "".join(f"<li>{_html_escape(hint)}</li>" for hint in hints)
    details_html = ""
    if details:
        details_html = f"""
    <section>
      <h2>技术细节</h2>
      <pre>{_html_escape(json.dumps(details, ensure_ascii=False, indent=2))}</pre>
    </section>"""
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>审查失败诊断</title>
<style>
body {{ margin:0; padding:24px; font-family: ui-sans-serif, -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Noto Sans SC', sans-serif; background:#f6f6f6; color:#171717; line-height:1.6; }}
main {{ max-width:1080px; margin:0 auto; }}
.hero, section {{ background:#fff; border:1px solid #d4d4d4; border-radius:8px; padding:18px; margin-bottom:14px; }}
.status {{ display:inline-block; border-radius:6px; background:#b42318; color:#fff; padding:6px 10px; font-weight:800; margin-bottom:10px; }}
h1 {{ font-size:24px; margin:0 0 8px; }}
h2 {{ font-size:17px; margin:0 0 10px; border-bottom:1px solid #d4d4d4; padding-bottom:6px; }}
.grid {{ display:grid; grid-template-columns: repeat(3, minmax(0,1fr)); gap:8px; margin-top:12px; }}
.cell {{ border:1px solid #d4d4d4; border-radius:6px; background:#fafafa; padding:9px; font-size:13px; color:#666; }}
.cell strong {{ display:block; color:#171717; overflow-wrap:anywhere; }}
pre, code {{ white-space:pre-wrap; overflow-wrap:anywhere; word-break:break-word; background:#f1f1f1; border:1px solid #d4d4d4; border-radius:6px; padding:10px; display:block; }}
ul {{ padding-left:22px; }}
@media (max-width:760px) {{ body {{ padding:12px; }} .grid {{ grid-template-columns:1fr; }} }}
</style>
</head>
<body>
<main>
  <div class="hero" id="failed-diagnostics">
    <div class="status">失败诊断 failed</div>
    <h1>失败恢复面板</h1>
    <p>关键审查能力失败，本次运行只生成失败诊断产物；修复后可使用下方命令断点续跑。</p>
    <div class="grid">
      <div class="cell"><span>文件</span><strong>{_html_escape(str(input_path))}</strong></div>
      <div class="cell"><span>失败能力</span><strong>{_html_escape(failed.get('capability'))}</strong></div>
      <div class="cell"><span>错误类别</span><strong>{_html_escape(failed.get('error_class'))}</strong></div>
      <div class="cell"><span>运行时间</span><strong>{_html_escape(runtime.get('local_time') or payload.get('created_at'))}</strong></div>
      <div class="cell"><span>运行时 UTC 年份</span><strong>{_html_escape(runtime.get('utc_year'))}</strong></div>
      <div class="cell"><span>缓存/恢复状态</span><strong>{_html_escape(cache_state)}</strong></div>
      <div class="cell"><span>产物类型</span><strong>failed</strong></div>
    </div>
  </div>
  <section>
    <h2>失败原因</h2>
    <p>{_html_escape(failed.get('message'))}</p>
  </section>
  <section>
    <h2>断点续跑命令</h2>
    <code>{_html_escape(retry_command)}</code>
  </section>
  <section>
    <h2>已完成阶段</h2>
    <ul>{stages_html}</ul>
  </section>
  <section>
    <h2>修复建议</h2>
    <ul>{hints_html}</ul>
  </section>
  {details_html}
</main>
</body>
</html>"""


def save_failed_audit_diagnostics(
    failure: AuditFailure,
    input_path: Path,
    output_dir: Path = None,
    output_stem: str = None,
    meta: Dict[str, Any] = None,
) -> Tuple[Path, Path]:
    """Write failed audit Markdown and JSON artifacts to the formal output location."""
    md_path, html_path, json_path = failed_audit_artifact_paths(input_path, output_dir=output_dir, output_stem=output_stem)
    md_path.parent.mkdir(parents=True, exist_ok=True)
    html_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.parent.mkdir(parents=True, exist_ok=True)
    md_path.write_text(format_failed_audit_markdown(failure, input_path, meta=meta), encoding="utf-8")
    html_path.write_text(format_failed_audit_html(failure, input_path, meta=meta), encoding="utf-8")
    json_path.write_text(json.dumps(failed_audit_payload(failure, input_path, meta=meta), ensure_ascii=False, indent=2), encoding="utf-8")
    return md_path, json_path


def _effective_limit(limit, total: int) -> int:
    return max(0, int(total if limit is None else limit))


def preflight_mineru(timeout=10) -> PreflightResult:
    return preflight_mineru_from_namespace(globals(), timeout=timeout)


def preflight_text_llm(timeout=10) -> PreflightResult:
    return preflight_text_llm_from_namespace(globals(), timeout=timeout)


def run_adapter_e2e_audit(
    input_path: Path,
    adapters: AuditAdapters,
    output_dir: Path = None,
    text: str = "Fake audit text with n=20 and p=0.04.",
    references_text: str = "",
    image_paths: List[str] = None,
) -> Dict[str, Any]:
    """Deterministic adapter-driven audit harness for end-to-end tests."""
    input_path = Path(input_path)
    output_dir = Path(output_dir or (input_path if input_path.is_dir() else input_path.parent))
    output_dir.mkdir(parents=True, exist_ok=True)
    output_stem = input_path.name if input_path.is_dir() else input_path.stem
    workspace = create_run_workspace(input_path, output_dir, output_stem)
    completed_stages = ["init"]
    image_paths = list(image_paths or [])
    retry_command = default_retry_command(input_path)

    def fail(capability: str, result: AdapterResult):
        failure = adapter_failure_to_audit_failure(capability, result, retry_command, completed_stages)
        md_path, json_path = save_failed_audit_diagnostics(failure, input_path)
        record_run_workspace_artifacts(workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
        return {
            "outcome": "failed",
            "capability": capability,
            "md_path": md_path,
            "json_path": json_path,
            "workspace": workspace,
        }

    if input_path.suffix.lower() == ".pdf":
        mineru_preflight = adapters.mineru.preflight()
        if not mineru_preflight.ok:
            return fail("mineru", mineru_preflight)
        mineru_result = adapters.mineru.extract(input_path, output_dir=output_dir)
        if not mineru_result.ok:
            return fail("mineru", mineru_result)
        value = mineru_result.value or {}
        text = value.get("text", text) if isinstance(value, dict) else text
        completed_stages.append("mineru_extract")

    if references_text:
        reference_result = adapters.reference_lookup.audit(references_text, online=True)
        if not reference_result.ok:
            return fail("reference_lookup", reference_result)
        completed_stages.append("reference_lookup")

    if image_paths:
        for image_path in image_paths:
            semantic_result = adapters.image_semantic.analyze(image_path)
            if semantic_result.status == "failure":
                return fail("image_semantic", semantic_result)
            detector_result = adapters.image_detector.detect(image_path)
            if detector_result.status == "failure":
                return fail("image_detector", detector_result)
        completed_stages.append("image_audit")

    llm_preflight = adapters.text_llm.preflight()
    if not llm_preflight.ok:
        return fail("text_llm", llm_preflight)
    llm_result = adapters.text_llm.review(text, chunk_info=(0, 1))
    if not llm_result.ok:
        return fail("text_llm", llm_result)
    completed_stages.append("text_llm_review")

    report = parse_report(llm_result.value)
    if report.get("parse_error"):
        return fail("text_llm", AdapterResult.failure("schema_error", "LLM returned invalid report schema", {"raw": llm_result.value}))

    stat_result = local_stat_check(text)
    image_audit = {"image_count": len(image_paths), "images": []}
    report = apply_risk_rules(report, stat_result=stat_result, image_audit=image_audit)
    meta = apply_audit_artifact_type({
        "artifact_type": "complete",
        "total_chars": len(text),
        "extraction_method": "fake_adapter",
        "reference_audit": {"reference_count": 1 if references_text else 0},
        "resource_audit": {"resource_count": 0, "resources": [], "issues": []},
        "image_audit": image_audit,
        "prompt_version": PROMPT_VERSION,
        "schema_version": SCHEMA_VERSION,
        "adapter_version": ADAPTER_VERSION,
        "risk_rule_version": RISK_RULE_VERSION,
    }, [])
    md_path, html_path, json_path = audit_artifact_paths(input_path, artifact_type="complete")
    md_path.write_text(format_report(report, str(input_path), meta, stat_result), encoding="utf-8")
    html_path.write_text(format_html_report(report, str(input_path), meta, stat_result), encoding="utf-8")
    json_path.write_text(
        json.dumps({"report_type": "complete", "llm_report": report, "stat_result": stat_result, "meta": meta}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    record_run_workspace_artifacts(workspace, "complete", [md_path, html_path, json_path], meta={"completed_stages": completed_stages})
    return {
        "outcome": "complete",
        "md_path": md_path,
        "html_path": html_path,
        "json_path": json_path,
        "workspace": workspace,
    }


# ─── 欺诈模式知识库加载 ───
FRAUD_PATTERNS_PATH = Path(__file__).resolve().parent.parent / "fraud_patterns.json"
FRAUD_PATTERNS = []
PATTERN_HINTS = ""
if FRAUD_PATTERNS_PATH.exists():
    try:
        with open(FRAUD_PATTERNS_PATH, "r", encoding="utf-8") as f:
            pattern_data = json.load(f)
            FRAUD_PATTERNS = pattern_data.get("patterns", [])
        print(f"✅ 加载欺诈模式知识库成功，共{len(FRAUD_PATTERNS)}条检测模式")
        # 构建提示词片段
        PATTERN_HINTS = "\n## 最新欺诈模式知识库（社区贡献+PubPeer案例汇总）\n"
        for idx, p in enumerate(FRAUD_PATTERNS, 1):
            PATTERN_HINTS += f"{idx}. [{p['risk_level']}风险] {p['name']}：{p['detection_hint']}\n"
    except Exception as e:
        print(f"⚠️ 知识库加载失败: {e}，使用默认检测规则")

# ══════════════════════════════════════════════════════════════
# 审查体系配置 - LLM System Prompt
# ══════════════════════════════════════════════════════════════
SYSTEM_PROMPT_TPL = """你是一个严厉的学术论文审查专家（耿同学标准）。
你需要结合以下维度对输入的论文文本进行审查，输出严格的JSON格式：

## 审查维度
1. 数据与结果自洽性 — 数字前后矛盾、统计量不一致、图表数据不匹配
2. 图片与图表异常 — 描述性分析图片可疑特征（旋转复用、背景一致、拼接痕迹）
3. 方法论严谨性 — 样本量不足、缺乏多重比较校正、实验设计缺陷
4. 结构与引用规范性 — 自引率异常、引用质量差、逻辑谬误
5. 作者与期刊可信度 — 产出异常、利益冲突未披露、同行评审缺失

## 检查项（耿同学六式 + 7类红旗）
- 耿同学六式：图片复用/数据造假/图片拼接/统计异常/产出异常/方法矛盾
- 7类红旗：引用质量差/逻辑谬误/方法论缺陷/可疑结论/同行评审缺失/利益冲突未披露/语言质量差

## OCR/表格噪声约束
- 输入中可能包含由MinerU/OCR生成的结构化标记，如[[TABLE_START]]、[[TABLE_END]]、[[FIGURE]]、[[EXTRACTION_NOTE]]。
- 不得仅因Markdown表格错位、列名断裂、OCR漏字、分页续表、表格被分块或图片转写不清晰判定为🚩红旗。
- 表格相关红旗必须基于明确学术证据：同一指标跨正文/表格矛盾、样本量/分组不自洽、p值/置信区间/均值标准差逻辑冲突、图表结论与正文明确冲突。
- 对提取不清晰但无明确学术矛盾的内容，应判为⚠️疑点或✅通过，并在detail中写“需人工核对原PDF/表格”，不要写成造假结论。

请按以下JSON格式输出（确保JSON合法，无多余内容）：
{{
  "summary": "一句话总评",
  "risk_level": "高/中/低/严重证据冲突",
  "detection_score": 0,
  "checks": [
    {{
      "category": "数据与结果/图片与图表/方法论/结构与引用/作者与期刊",
      "item": "检查项名称",
      "verdict": "🚩红旗/⚠️疑点/✅通过",
      "source": "证据来源位置，例如正文段落/表格/图片/参考文献编号；没有则写'未定位'",
      "source_text": "必须填写：论文原文中的直接摘录；若无直接证据写'未找到直接原文证据'",
      "evidence": "必须填写：具体证据，引用原文片段并说明所在章节/表图/段落线索",
      "reason": "必须填写：为什么该证据支持此判定，说明可疑逻辑链",
      "recommendation": "必须填写：建议人工复核或后续处理动作",
      "confidence": 0.0,
      "detail": "详细分析说明：包含影响范围、需人工复核的点、若为通过也说明依据"
    }}
  ],
  "conclusion": "综合结论与行动建议"
}}
{pattern_hints}
"""

# 动态构建系统提示词
SYSTEM_PROMPT = SYSTEM_PROMPT_TPL.format(pattern_hints=PATTERN_HINTS)

# ══════════════════════════════════════════════════════════════
# MinerU API 模块 — PDF转Markdown
# ══════════════════════════════════════════════════════════════

def _http_request(url, method="GET", headers=None, data=None, timeout=60):
    """通用HTTP请求封装（使用requests，绕过Cloudflare UA检测）"""
    _BROWSER_UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36"
    if headers is None:
        headers = {}
    headers.setdefault("User-Agent", _BROWSER_UA)
    if method.upper() == "GET":
        resp = requests.get(url, headers=headers, timeout=timeout)
    elif method.upper() == "POST":
        resp = requests.post(url, headers=headers, data=data, timeout=timeout)
    else:
        resp = requests.request(method, url, headers=headers, data=data, timeout=timeout)
    resp.raise_for_status()
    return resp.content, resp.status_code


def mineru_precision_extract_by_url(pdf_url, model_version="vlm", language="ch",
                                     poll_interval=10, poll_timeout=600, output_dir=None):
    """🎯 Precision API — 通过URL解析PDF（需要Token，≤200MB/200页）

    流程：POST创建任务 → GET轮询结果 → 下载zip中的Markdown
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    print(f"  🎯 [MinerU Precision] 提交URL任务: {pdf_url[:80]}...")

    # 1. 创建提取任务
    create_url = f"{MINERU_BASE}/api/v4/extract/task"
    payload = json.dumps({"url": pdf_url, "model_version": model_version}).encode()
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MINERU_TOKEN}"
    }
    try:
        resp_data, status = _http_request(create_url, "POST", headers, payload, timeout=30)
        result = json.loads(resp_data.decode())
    except Exception as e:
        return None, {"error": f"创建任务失败: {e}"}

    if result.get("code") != 0 and not result.get("data", {}).get("batch_id"):
        return None, {"error": f"创建任务返回异常: {result}"}

    batch_id = result.get("data", {}).get("batch_id")
    if not batch_id:
        return None, {"error": f"未获取到batch_id: {result}"}
    print(f"  ✅ 任务已创建: batch_id={batch_id}")

    # 2. 轮询任务状态
    poll_url = f"{MINERU_BASE}/api/v4/extract/task/{batch_id}"
    start = time.time()
    state_labels = {"processing": "处理中", "queued": "排队中"}

    while time.time() - start < poll_timeout:
        try:
            resp_data, _ = _http_request(poll_url, "GET", headers, timeout=30)
            result = json.loads(resp_data.decode())
        except Exception as e:
            print(f"  ⚠️ 轮询异常: {e}")
            time.sleep(poll_interval)
            continue

        task_list = result.get("data", {}).get("task_list", [])
        if not task_list:
            # 单文件模式
            state = result.get("data", {}).get("state", "unknown")
        else:
            state = task_list[0].get("state", "unknown")

        elapsed = int(time.time() - start)

        if state == "done":
            # 获取zip下载链接
            zip_url = task_list[0].get("zip_url") if task_list else result.get("data", {}).get("zip_url")
            if not zip_url:
                return None, {"error": "任务完成但未获取到下载链接"}

            print(f"  ✅ [{elapsed}s] 解析完成，下载Markdown...")
            markdown = _download_zip_and_extract_md(zip_url, output_dir=output_dir, source_name="url_input", batch_id=batch_id)
            if markdown:
                meta = {"source": "mineru_precision", "batch_id": batch_id,
                        "zip_url": zip_url, "zip_saved_dir": str(output_dir) if output_dir else str(_RUN_OUTPUT_DIR) if _RUN_OUTPUT_DIR else None,
                        "model": model_version, "chars": len(markdown)}
                return markdown, meta
            else:
                return None, {"error": "下载或解压zip失败"}

        elif state == "failed":
            err = task_list[0].get("err_msg", "未知") if task_list else "未知"
            return None, {"error": f"任务失败: {err}"}

        label = state_labels.get(state, state)
        print(f"  ⏳ [{elapsed}s] {label}...")
        time.sleep(poll_interval)

    return None, {"error": f"轮询超时({poll_timeout}s), batch_id={batch_id}"}


def mineru_extract_file(file_path, model_version="vlm", language="ch",
                        poll_interval=10, poll_timeout=600, output_dir=None):
    """🎯 MinerU v4 本地文件解析（需要Token，≤200MB/200页）

    流程：POST /api/v4/file-urls/batch → PUT上传至OSS → 轮询任务结果 → 下载zip提取Markdown
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    file_path = Path(file_path)
    file_size_mb = file_path.stat().st_size / 1024 / 1024
    print(f"  🎯 [MinerU v4] 上传文件: {file_path.name} ({file_size_mb:.1f}MB)")

    if not MINERU_TOKEN:
        return None, {"error": "MINERU_TOKEN未配置，无法使用MinerU API"}

    auth_headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MINERU_TOKEN}"
    }

    # 1. 获取上传URL
    try:
        batch_payload = json.dumps({
            "enable_formula": True,
            "language": language,
            "layout_model": "doclayout_yolo",
            "enable_table": True,
            "files": [{"name": file_path.name, "is_ocr": True}]
        }).encode()
        resp_data, _ = _http_request(
            f"{MINERU_BASE}/api/v4/file-urls/batch", "POST",
            auth_headers, batch_payload, timeout=30
        )
        result = json.loads(resp_data.decode())
    except Exception as e:
        return None, {"error": f"获取上传URL失败: {e}"}

    if result.get("code") != 0:
        return None, {"error": f"file-urls/batch返回异常: {result}"}

    batch_id = result["data"]["batch_id"]
    file_urls = result["data"]["file_urls"]
    if not file_urls:
        return None, {"error": "未获取到上传URL"}

    upload_url = file_urls[0]
    print(f"  ✅ 获取上传URL成功: batch_id={batch_id}")

    # 2. 上传文件至OSS（注意：预签名URL不能带Authorization header，也不带Content-Type以免签名不匹配）
    try:
        with open(file_path, "rb") as f:
            file_data = f.read()
        resp = requests.put(upload_url, data=file_data, timeout=120)
        resp.raise_for_status()
        print(f"  ✅ 文件已上传至OSS ({file_size_mb:.1f}MB)")
    except Exception as e:
        return None, {"error": f"上传文件到OSS失败: {e}"}

    # 3. 轮询批量解析结果（file-urls/batch 上传后，用 batch_id 查询 extract-results/batch）
    # 经验验证：/api/v4/extract/task/{batch_id} 会返回 task not found，batch_id 应查批量结果接口。
    poll_url = f"{MINERU_BASE}/api/v4/extract-results/batch/{batch_id}"
    start = time.time()
    state_labels = {"processing": "处理中", "queued": "排队中", "pending": "等待中"}

    while time.time() - start < poll_timeout:
        try:
            resp_data, _ = _http_request(poll_url, "GET", auth_headers, timeout=30)
            result = json.loads(resp_data.decode())
        except Exception as e:
            print(f"  ⚠️ 轮询异常: {e}")
            time.sleep(poll_interval)
            continue

        elapsed = int(time.time() - start)
        if result.get("code") != 0:
            print(f"  ⚠️ [{elapsed}s] 查询结果异常: code={result.get('code')} msg={result.get('msg')}")
            time.sleep(poll_interval)
            continue

        data = result.get("data", {}) or {}
        extract_results = data.get("extract_result") or data.get("task_list") or []
        task = extract_results[0] if extract_results else data
        state = task.get("state", "unknown")

        if state == "done":
            zip_url = task.get("full_zip_url") or data.get("full_zip_url")
            if not zip_url:
                return None, {"error": "任务完成但未获取到下载链接", "batch_id": batch_id,
                              "result": result}

            print(f"  ✅ [{elapsed}s] 解析完成，下载Markdown...")
            markdown = _download_zip_and_extract_md(zip_url, output_dir=output_dir, source_name=file_path.name, batch_id=batch_id)
            if markdown:
                return markdown, {"source": "mineru_v4", "batch_id": batch_id,
                                  "zip_url": zip_url, "model": model_version,
                                  "zip_saved_dir": str(output_dir) if output_dir else str(_RUN_OUTPUT_DIR) if _RUN_OUTPUT_DIR else None,
                                  "chars": len(markdown)}
            return None, {"error": "下载或解压zip失败", "batch_id": batch_id,
                          "zip_url": zip_url}

        elif state == "failed":
            err = task.get("err_msg") or data.get("err_msg") or "未知错误"
            return None, {"error": f"任务失败: {err}", "batch_id": batch_id,
                          "result": result}

        label = state_labels.get(state, state)
        print(f"  ⏳ [{elapsed}s] {label}...")
        time.sleep(poll_interval)

    return None, {"error": f"轮询超时({poll_timeout}s), batch_id={batch_id}",
                  "poll_url": poll_url}


def _download_zip_and_extract_md(zip_url, output_dir=None, source_name=None, batch_id=None):
    """下载zip、按输入同目录保存，并优先提取MinerU结构化文本。"""
    zip_data = None
    last_err = None
    for attempt in range(3):
        try:
            if attempt:
                print(f"  ↻ MinerU zip下载重试 {attempt}/2...")
            zip_data, _ = _http_request(zip_url, "GET", timeout=180)
            if output_dir or _RUN_OUTPUT_DIR:
                save_mineru_artifacts(zip_url, zip_data, source_name or "mineru", output_dir=output_dir, batch_id=batch_id)
            break
        except Exception as e:
            last_err = e
            time.sleep(3 * (attempt + 1))
    if zip_data is None:
        print(f"  ❌ 下载zip失败: {last_err}")
        return None

    # 用 zipfile 从内存解析
    import zipfile, io
    try:
        with zipfile.ZipFile(io.BytesIO(zip_data)) as zf:
            structured = _extract_mineru_structured_text(zf)
            if structured:
                return structured

            # 找到 .md 文件
            md_files = [n for n in zf.namelist() if n.endswith(".md")]
            if not md_files:
                # 降级：找 .txt 或其他文本
                text_files = [n for n in zf.namelist() if n.endswith((".txt", ".mdown", ".markdown"))]
                md_files = text_files
            if not md_files:
                print(f"  ⚠️ zip中未找到Markdown文件: {zf.namelist()[:10]}")
                # 尝试任何非图片文件
                for n in zf.namelist():
                    if not any(n.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".svg", ".html"]):
                        try:
                            content = zf.read(n).decode("utf-8", errors="ignore")
                            if len(content) > 100:
                                return content
                        except:
                            continue
                return None

            # 读取最大的 .md 文件
            best = None
            best_len = 0
            for md_file in md_files:
                content = zf.read(md_file).decode("utf-8", errors="ignore")
                if len(content) > best_len:
                    best = content
                    best_len = len(content)
            return best
    except zipfile.BadZipFile:
        # 不是zip？尝试直接作为文本
        try:
            return zip_data.decode("utf-8", errors="ignore")
        except:
            return None


def mineru_extract(file_path, language="ch", output_dir=None):
    """MinerU统一入口：使用v4 API
    
    本地文件：上传至OSS → 创建任务 → 轮询结果
    URL：直接提交v4/extract/task
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    file_path = Path(file_path)
    if file_path.exists():
        return mineru_extract_file(file_path, language=language, output_dir=output_dir)
    else:
        # 当作URL处理
        return mineru_precision_extract_by_url(str(file_path), language=language, output_dir=output_dir)


def extract_text_from_file(file_path: Path, max_chars_per_file=None, use_mineru=False, mineru_lang="ch", output_dir=None) -> str:
    """从任意支持的文件类型中提取文本

    max_chars_per_file=None 表示不截断。目录级分析默认应先完整提取每个文件，
    再由后续 smart_chunk_text 做全目录分块/合并审查，避免“目录里只分析到一个文件/每文件只取开头”。
    """
    ext = file_path.suffix.lower()
    header = f"=== 文件: {file_path.name} ==="
    text = f"\n\n{header}"
    limit = max_chars_per_file if max_chars_per_file is not None else 999999999
    
    try:
        if ext == ".pdf":
            if use_mineru:
                print(f"  ⚙️  使用MinerU API提取PDF全文内容...")
                md, meta = mineru_extract(file_path, language=mineru_lang, output_dir=output_dir)
                if md:
                    text += "\n" + md
                else:
                    err = meta.get("error", "未知错误") if isinstance(meta, dict) else "未知错误"
                    print(f"  ⚠️  MinerU提取失败: {err}，降级为本地PDF提取")
                    pdf_text, _, _ = extract_pdf_text(file_path, max_chars=limit)
                    text += "\n" + pdf_text
            else:
                pdf_text, _, _ = extract_pdf_text(file_path, max_chars=limit)
                text += "\n" + pdf_text
        elif ext == ".docx" and DOCX_SUPPORTED:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += "\n" + para.text
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join([cell.text for cell in row.cells])
        elif ext in {".xlsx", ".xlsm"} and EXCEL_SUPPORTED:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                text += f"\n[工作表: {sheet_name}]"
                sheet = wb[sheet_name]
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if max_chars_per_file is not None and i > 1000:
                        text += "\n[数据过多，已截断]"
                        break
                    row_str = " | ".join([str(v) for v in row if v is not None])
                    if row_str.strip():
                        text += "\n" + row_str
        elif ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if max_chars_per_file is not None and i > 1000:
                        text += "\n[数据过多，已截断]"
                        break
                    text += "\n" + " | ".join(row)
        elif ext in {".txt", ".md"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text += "\n" + f.read(limit)
    except Exception as e:
        text += f"\n[文件解析失败: {str(e)}]"
    
    if max_chars_per_file is not None and len(text) > max_chars_per_file + len(header) + 4:
        return text[:max_chars_per_file + len(header) + 4] + "\n[文本过长已截断]"
    return text


def optional_dependency_for_extension(ext: str):
    ext = str(ext or "").lower()
    if ext == ".docx" and not DOCX_SUPPORTED:
        return "python-docx", "python3 -m pip install python-docx"
    if ext in {".xlsx", ".xlsm"} and not EXCEL_SUPPORTED:
        return "openpyxl", "python3 -m pip install openpyxl"
    return None, None


def extracted_body_text(file_content: str, file_name: str = "") -> str:
    text = str(file_content or "").strip()
    if file_name:
        text = re.sub(rf"^=+\s*文件:\s*{re.escape(str(file_name))}\s*=+\s*", "", text).strip()
    text = re.sub(r"^\s*=+\s*文件:.*?=+\s*", "", text, count=1).strip()
    return text


# ══════════════════════════════════════════════════════════════
# PDF原始提取模块（MinerU不可用时的降级方案）
# ══════════════════════════════════════════════════════════════

def extract_pdf_text(filepath, max_chars=8000):
    """从PDF文件中提取文本（纯标准库实现，MinerU的降级方案）"""
    with open(filepath, "rb") as f:
        raw = f.read()
    parts = []
    for s in re.findall(rb"stream\r?\n(.*?)\r?\nendstream", raw, re.DOTALL):
        try:
            dec = zlib.decompress(s)
            for x in re.findall(rb"\((.*?)\)\s*Tj", dec):
                d = x.decode("latin-1", errors="ignore")
                if len(d.strip()) > 1:
                    parts.append(d)
            for bt in re.findall(rb"BT(.*?)ET", dec, re.DOTALL):
                for x in re.findall(rb"\((.*?)\)", bt):
                    d = x.decode("latin-1", errors="ignore")
                    if len(d.strip()) > 1:
                        parts.append(d)
        except:
            pass
    text = re.sub(r"\s+", " ", " ".join(parts)).strip()
    meta = {"size_mb": round(len(raw) / 1024 / 1024, 2), "total_chars": len(text),
            "extraction_method": "raw_pdf_stream"}
    return text[:max_chars], meta, raw


# ══════════════════════════════════════════════════════════════
# LLM调用模块
# ══════════════════════════════════════════════════════════════

def call_llm(text, max_retries=None, chunk_info=None, timeout=None):
    """调用OpenAI兼容API进行语义审查。

    改进点：
    - timeout/max_retries可由CLI控制，便于不稳定网关降级；
    - 每次失败打印尝试编号，日志更容易判断是否卡死；
    - 去除重复payload构建，减少维护歧义。
    """
    if max_retries is None:
        max_retries = int(globals().get("LLM_RETRIES", 1))
    if timeout is None:
        timeout = int(globals().get("LLM_TIMEOUT", 45))

    if chunk_info and chunk_info[1] > 1:
        idx, total = chunk_info
        user_msg = (
            f"审查以下论文文本（第{idx+1}/{total}段，请重点关注本段内容，"
            f"同时注意与其他段落的逻辑连贯性）。"
            "只返回一个合法JSON对象；checks最多2条；每个字符串字段不超过120字；"
            "source_text/source/evidence必须短摘录，不能展开长篇推理：\n\n"
            f"{text}"
        )
    else:
        user_msg = (
            "审查以下论文文本。只返回一个合法JSON对象；checks最多3条；"
            "每个字符串字段不超过120字；source_text/source/evidence必须短摘录：\n\n"
            f"{text}"
        )

    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_msg}
        ],
        "temperature": 0.2,
        "max_tokens": 4000,
    }
    _BROWSER_UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36"
    last_err = None
    for attempt in range(max_retries + 1):
        try:
            if attempt:
                print(f"     ↻ API重试 {attempt}/{max_retries}（timeout={timeout}s）")
            resp = requests.post(
                _chat_completions_endpoint(LLM_API_URL), json=payload,
                headers={"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json", "User-Agent": _BROWSER_UA},
                timeout=timeout
            )
            resp.raise_for_status()
            result = resp.json()
            return result["choices"][0]["message"]["content"]
        except Exception as e:
            last_err = e
            if attempt < max_retries:
                print(f"     ⚠️ API尝试 {attempt+1}/{max_retries+1} 失败: {str(e)[:160]}")
                time.sleep(3 * (attempt + 1))
            else:
                raise RuntimeError(f"API调用失败({attempt+1}次, timeout={timeout}s): {str(last_err)[:180]}...")


def call_llm_messages(messages, temperature=0.2, timeout=None, max_tokens=1800):
    if timeout is None:
        timeout = int(globals().get("LLM_TIMEOUT", 60))
    if not LLM_API_KEY:
        raise RuntimeError("LLM_API_KEY未配置")
    payload = {
        "model": LLM_MODEL,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    _BROWSER_UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36"
    resp = requests.post(
        _chat_completions_endpoint(LLM_API_URL),
        json=payload,
        headers={
            "Authorization": f"Bearer {LLM_API_KEY}",
            "Content-Type": "application/json",
            "User-Agent": _BROWSER_UA,
        },
        timeout=timeout,
    )
    resp.raise_for_status()
    result = resp.json()
    return result["choices"][0]["message"]["content"]


def save_followup_artifacts(kind, context, language, text):
    return save_followup_artifacts_from_namespace(globals(), kind, context, language, text)


def generate_and_save_followup_draft(
    kind,
    context,
    language="zh",
    identity=None,
    selected_issues=None,
    custom_concerns=None,
    tone="conservative",
    disclaimer_confirmed=False,
    timeout=None,
):
    return generate_and_save_followup_draft_from_namespace(
        globals(),
        kind,
        context,
        language=language,
        identity=identity,
        selected_issues=selected_issues,
        custom_concerns=custom_concerns,
        tone=tone,
        disclaimer_confirmed=disclaimer_confirmed,
        timeout=timeout,
    )


def generate_followup_draft(kind, context, language="zh", tone=None, timeout=None):
    return generate_followup_draft_from_namespace(globals(), kind, context, language=language, tone=tone, timeout=timeout)


def report_action_service_url(host="127.0.0.1", port=8765):
    return f"http://{host}:{int(port)}"


def report_action_service_health(host="127.0.0.1", port=8765, timeout=0.5):
    """Return the local report action service health payload, or None when unavailable."""
    try:
        with urllib.request.urlopen(f"{report_action_service_url(host, port)}/health", timeout=timeout) as resp:
            if resp.status != 200:
                return None
            payload = json.loads(resp.read().decode("utf-8", errors="replace") or "{}")
            if payload.get("ok"):
                return payload
    except Exception:
        return None
    return None


def _report_action_entrypoint():
    candidate = Path(__file__).resolve().parents[1] / "paper_audit.py"
    if candidate.exists():
        return candidate
    return Path(sys.argv[0]).resolve()


def ensure_report_action_service(host="127.0.0.1", port=8765, log_path: Path = None, startup_timeout=2.0):
    """Start or reuse the localhost action service used by generated HTML reports."""
    existing = report_action_service_health(host=host, port=port, timeout=0.3)
    if existing:
        return {"ok": True, "status": "already_running", "url": report_action_service_url(host, port), "health": existing}

    command = [
        sys.executable,
        str(_report_action_entrypoint()),
        "--serve-report-actions",
        "--report-actions-port",
        str(int(port)),
    ]
    popen_kwargs = {
        "stdin": subprocess.DEVNULL,
        "start_new_session": True,
    }
    if log_path:
        Path(log_path).parent.mkdir(parents=True, exist_ok=True)
        log_file = open(log_path, "a", encoding="utf-8")
        popen_kwargs["stdout"] = log_file
        popen_kwargs["stderr"] = subprocess.STDOUT
    else:
        log_file = None
        popen_kwargs["stdout"] = subprocess.DEVNULL
        popen_kwargs["stderr"] = subprocess.DEVNULL

    try:
        process = subprocess.Popen(command, **popen_kwargs)
    except Exception as e:
        if log_file:
            log_file.close()
        return {"ok": False, "status": "start_failed", "url": report_action_service_url(host, port), "error": f"{type(e).__name__}: {_brief_text(str(e), 240)}"}
    finally:
        if log_file:
            log_file.close()

    deadline = time.time() + float(startup_timeout)
    while time.time() < deadline:
        health = report_action_service_health(host=host, port=port, timeout=0.3)
        if health:
            return {"ok": True, "status": "started", "url": report_action_service_url(host, port), "pid": process.pid, "health": health}
        if process.poll() is not None:
            return {"ok": False, "status": "exited", "url": report_action_service_url(host, port), "pid": process.pid, "returncode": process.returncode}
        time.sleep(0.1)
    return {"ok": True, "status": "starting", "url": report_action_service_url(host, port), "pid": process.pid}


def open_html_artifact(html_path: Path):
    html_abs = str(Path(html_path).resolve())
    webbrowser.open(f"file:///{html_abs}" if platform.system() == "Windows" else f"file://{html_abs}")


def _report_action_api_response(route, payload):
    """Return the shared response payload for local report action endpoints."""
    context = payload.get("context") or {}
    language = normalize_followup_language(payload.get("language"))
    if route == "/followups":
        return load_existing_followups(context, language=language)
    kind = payload.get("kind")
    result = generate_and_save_followup_draft(
        kind,
        context,
        language=language,
        identity=payload.get("identity"),
        selected_issues=payload.get("selected_issues"),
        custom_concerns=payload.get("custom_concerns"),
        tone=payload.get("tone"),
        disclaimer_confirmed=bool(payload.get("disclaimer_confirmed")),
        timeout=LLM_TIMEOUT,
    )
    return {
        "ok": True,
        "kind": result.get("kind"),
        "language": result.get("language"),
        "tone": result.get("tone"),
        "model": result.get("model"),
        "text": result.get("text"),
        "paths": result.get("paths"),
    }


def _read_json_request_body(handler, max_bytes=2_000_000):
    length = int(handler.headers.get("Content-Length", "0") or "0")
    if length > max_bytes:
        raise ValueError("request_too_large")
    body = handler.rfile.read(length).decode("utf-8", errors="replace")
    return json.loads(body or "{}")


def serve_report_actions(host="127.0.0.1", port=8765):
    from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

    class Handler(BaseHTTPRequestHandler):
        server_version = "PaperAuditActions/1.0"

        def _send_json(self, payload, status=200):
            data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
            self.send_response(status)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
            self.send_header("Access-Control-Allow-Headers", "Content-Type")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)

        def do_OPTIONS(self):
            self._send_json({"ok": True})

        def do_GET(self):
            if self.path.rstrip("/") == "/health":
                self._send_json({"ok": True, "model": LLM_MODEL})
            else:
                self._send_json({"ok": False, "error": "not_found"}, 404)

        def do_POST(self):
            route = self.path.rstrip("/")
            if route not in {"/generate", "/followups"}:
                self._send_json({"ok": False, "error": "not_found"}, 404)
                return
            try:
                payload = _read_json_request_body(self)
                self._send_json(_report_action_api_response(route, payload))
            except ValueError as e:
                status = 413 if str(e) == "request_too_large" else 400
                self._send_json({"ok": False, "error": str(e)}, status)
            except Exception as e:
                self._send_json({"ok": False, "error": f"{type(e).__name__}: {_brief_text(str(e), 300)}"}, 500)

        def log_message(self, fmt, *args):
            print(f"[report-actions] {self.address_string()} {fmt % args}")

    httpd = ThreadingHTTPServer((host, int(port)), Handler)
    print(f"🌐 报告动作服务已启动: http://{host}:{port}")
    print("   在HTML报告中点击“生成 PubPeer Comment”或“生成期刊 Letter”即可调用已配置的LLM。按 Ctrl+C 停止。")
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        print("\n⏹️ 报告动作服务已停止")
    finally:
        httpd.server_close()
    return 0


def web_runner_default_output_stem(input_path, timestamp=None):
    return web_runner_default_output_stem_from_namespace(globals(), input_path, timestamp=timestamp)


def web_runner_config_status():
    return web_runner_config_status_from_namespace(globals())


class WebRunnerState:
    """State boundary for the local browser runner."""

    def __init__(self, history_path=None):
        self.history_path = _web_runner_history_path(history_path)
        self.lock = threading.Lock()
        self.runs = {}
        self.active_run_id = None
        self._load_history()

    def _load_history(self):
        try:
            payload = json.loads(self.history_path.read_text(encoding="utf-8"))
            items = payload.get("runs") if isinstance(payload, dict) else payload
            if isinstance(items, list):
                self.runs = {str(item.get("id")): dict(item) for item in items if item.get("id")}
                for run in self.runs.values():
                    if run.get("status") == "running":
                        run["status"] = "failed"
                        run["finished_at"] = run.get("finished_at") or _web_runner_now()
                        run["message"] = "上次本地服务退出时该任务仍在运行。"
        except FileNotFoundError:
            self.runs = {}
        except Exception:
            self.runs = {}

    def _save_history(self):
        self.history_path.parent.mkdir(parents=True, exist_ok=True)
        items = [_web_runner_safe_run(run) for run in self.runs.values()]
        items.sort(key=lambda item: item.get("started_at") or "", reverse=True)
        tmp_path = self.history_path.with_suffix(".tmp")
        tmp_path.write_text(json.dumps({"runs": items[:100]}, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp_path.replace(self.history_path)

    def list_runs(self):
        with self.lock:
            items = [_web_runner_safe_run(run) for run in self.runs.values()]
        items.sort(key=lambda item: item.get("started_at") or "", reverse=True)
        return items[:100]

    def get_run(self, run_id):
        with self.lock:
            run = self.runs.get(str(run_id))
            return _web_runner_safe_run(run) if run else None

    def _append_log(self, run_id, text):
        if text is None:
            return
        line = str(text).rstrip("\n")
        with self.lock:
            run = self.runs.get(run_id)
            if not run:
                return
            logs = run.setdefault("logs", [])
            logs.append(line)
            if len(logs) > 2000:
                del logs[: len(logs) - 2000]
            run["log_count"] = len(logs)

    def logs_since(self, run_id, offset=0):
        with self.lock:
            run = self.runs.get(str(run_id))
            if not run:
                return None
            logs = list(run.get("logs") or [])
        start = max(0, int(offset or 0))
        return {"ok": True, "run_id": str(run_id), "offset": len(logs), "lines": logs[start:]}

    def start_run(self, input_path, output=None, fresh=False):
        input_text = str(input_path or "").strip()
        if not input_text:
            return {"ok": False, "error": "input_path_required", "message": "请输入文件或目录路径。"}, 400
        resolved = resolve_web_runner_input_path(input_text, search_roots=_web_runner_common_search_roots())
        if not resolved.get("ok"):
            status = 409 if resolved.get("error") == "ambiguous_input_path" else 400
            return resolved, status
        resolved_input = str(Path(resolved.get("path")).expanduser())
        command = [
            sys.executable,
            str(_report_action_entrypoint()),
            resolved_input,
            "--json",
            "--no-open",
        ]
        output_text = str(output or "").strip()
        if not output_text:
            output_text = web_runner_default_output_stem(resolved_input)
        if output_text:
            try:
                Path(output_text).expanduser().parent.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                return {"ok": False, "error": "output_prepare_failed", "message": f"{type(e).__name__}: {_brief_text(str(e), 240)}"}, 500
            command.extend(["-o", output_text])
        if fresh:
            command.append("--fresh")

        with self.lock:
            if self.active_run_id:
                active = self.runs.get(self.active_run_id) or {}
                if active.get("status") == "running":
                    return {"ok": False, "error": "busy", "active_run": _web_runner_safe_run(active)}, 409
                self.active_run_id = None
            run_id = _web_runner_run_id(resolved_input)
            run = {
                "id": run_id,
                "input_path": resolved_input,
                "output": output_text,
                "fresh": bool(fresh),
                "command": command,
                "started_at": _web_runner_now(),
                "finished_at": None,
                "status": "running",
                "message": "审查运行中",
                "logs": [],
                "log_count": 0,
                "artifacts": {},
                "report_type": "",
            }
            self.runs[run_id] = run
            self.active_run_id = run_id
            self._save_history()

        try:
            process = subprocess.Popen(
                command,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
        except Exception as e:
            with self.lock:
                run = self.runs[run_id]
                run["status"] = "failed"
                run["finished_at"] = _web_runner_now()
                run["message"] = f"{type(e).__name__}: {_brief_text(str(e), 240)}"
                self.active_run_id = None
                self._save_history()
            return {"ok": False, "error": "start_failed", "run": self.get_run(run_id)}, 500

        with self.lock:
            run = self.runs[run_id]
            run["_process"] = process
            run["pid"] = getattr(process, "pid", None)
            self._save_history()

        thread = threading.Thread(target=self._watch_process, args=(run_id, process), daemon=True)
        thread.start()
        return {"ok": True, "run": self.get_run(run_id)}, 200

    def _watch_process(self, run_id, process):
        try:
            stdout = getattr(process, "stdout", None)
            if stdout is not None:
                for line in stdout:
                    self._append_log(run_id, line)
            returncode = process.wait()
        except Exception as e:
            self._append_log(run_id, f"[web-runner] {type(e).__name__}: {_brief_text(str(e), 240)}")
            returncode = getattr(process, "returncode", 1)
        self._finish_run(run_id, returncode)

    def _finish_run(self, run_id, returncode):
        with self.lock:
            run = self.runs.get(run_id)
            if not run:
                return
            canceled = bool(run.get("_cancel_requested"))
            run["returncode"] = returncode
            run["finished_at"] = _web_runner_now()
            if canceled:
                run["status"] = "canceled"
                run["message"] = "已取消。本输入的断点续作缓存会在下次重新运行时继续复用。"
            elif int(returncode or 0) == 0:
                run["status"] = "succeeded"
                run["message"] = "审查完成"
            else:
                run["status"] = "failed"
                run["message"] = f"审查失败，退出码 {returncode}"
            run.pop("_process", None)
            if self.active_run_id == run_id:
                self.active_run_id = None
        self.discover_artifacts(run_id)

    def cancel_run(self, run_id):
        with self.lock:
            run = self.runs.get(str(run_id))
            if not run:
                return {"ok": False, "error": "not_found"}, 404
            process = run.get("_process")
            if run.get("status") != "running" or process is None:
                return {"ok": True, "run": _web_runner_safe_run(run)}, 200
            run["_cancel_requested"] = True
            run["message"] = "正在取消"
            self._save_history()
        try:
            process.terminate()
            try:
                process.wait(timeout=5)
            except TypeError:
                process.wait()
            except subprocess.TimeoutExpired:
                process.kill()
        except Exception as e:
            self._append_log(str(run_id), f"[web-runner] 取消失败: {type(e).__name__}: {_brief_text(str(e), 240)}")
        return {"ok": True, "run": self.get_run(run_id)}, 200

    def _artifact_candidates(self, run):
        input_path = Path(run.get("input_path") or "")
        output_base = _web_runner_output_base(run.get("output"))
        candidates = []
        for report_type in ("complete", "limited"):
            md_path, html_path, json_path = audit_artifact_paths(input_path, report_type, output_path=output_base)
            candidates.append((report_type, {"markdown": md_path, "html": html_path, "json": json_path}))
        if output_base:
            failed_paths = failed_audit_artifact_paths(input_path, output_dir=output_base.parent, output_stem=output_base.name)
        else:
            failed_paths = failed_audit_artifact_paths(input_path)
        candidates.append(("failed", {"markdown": failed_paths[0], "html": failed_paths[1], "json": failed_paths[2]}))
        return candidates

    def discover_artifacts(self, run_id):
        with self.lock:
            run = self.runs.get(str(run_id))
            if not run:
                return None
            run_copy = dict(run)
        discovered = {}
        report_type = ""
        summary = {}
        for candidate_type, paths in self._artifact_candidates(run_copy):
            existing = {kind: str(Path(path).resolve()) for kind, path in paths.items() if path and Path(path).exists()}
            if existing:
                discovered.update(existing)
                report_type = candidate_type
                json_path = existing.get("json")
                if json_path:
                    try:
                        payload = json.loads(Path(json_path).read_text(encoding="utf-8"))
                        summary = _web_runner_report_summary_from_payload(payload, candidate_type)
                    except Exception:
                        summary = {}
                break
        folder = None
        if discovered:
            first_path = Path(next(iter(discovered.values())))
            folder = str(first_path.parent.resolve())
        elif run_copy.get("output"):
            base = _web_runner_output_base(run_copy.get("output"))
            if base:
                folder = str(base.parent.resolve())
        else:
            input_path = Path(run_copy.get("input_path") or "")
            if input_path:
                folder = str((input_path if input_path.is_dir() else input_path.parent).resolve())
        if folder:
            discovered["folder"] = folder
        with self.lock:
            run = self.runs.get(str(run_id))
            if run:
                run["artifacts"] = discovered
                run["report_type"] = report_type
                run["summary"] = summary
                self._save_history()
                return _web_runner_safe_run(run)
        return None

    def artifact_target(self, run_id, kind):
        if kind not in {"html", "markdown", "json", "folder"}:
            return None, "unknown_artifact"
        with self.lock:
            run = self.runs.get(str(run_id))
            if not run:
                return None, "not_found"
            artifacts = dict(run.get("artifacts") or {})
        if kind not in artifacts:
            self.discover_artifacts(run_id)
            with self.lock:
                run = self.runs.get(str(run_id)) or {}
                artifacts = dict(run.get("artifacts") or {})
        raw_path = artifacts.get(kind)
        if not raw_path:
            return None, "not_recorded"
        path = Path(raw_path)
        if kind != "folder" and not path.exists():
            return None, "missing"
        return path, ""


def desktop_gui_write_llm_config(api_key, api_url, model, config_path=None):
    return _desktop_gui_write_llm_config(
        api_key,
        api_url,
        model,
        config_path=config_path,
        default_api_url=LLM_API_URL,
        default_model=LLM_MODEL,
    )


def desktop_gui_checked_config_snapshot(llm_preflight_runner=None, mineru_preflight_runner=None, timeout=6):
    return desktop_gui_checked_config_snapshot_from_namespace(
        globals(),
        llm_preflight_runner=llm_preflight_runner,
        mineru_preflight_runner=mineru_preflight_runner,
        timeout=timeout,
    )


def desktop_gui_followup_context(run):
    return desktop_gui_followup_context_from_namespace(globals(), run)


def desktop_gui_generate_followup_draft(kind, run, language="zh", tone="conservative", timeout=None):
    return desktop_gui_generate_followup_draft_from_namespace(
        globals(),
        kind,
        run,
        language=language,
        tone=tone,
        timeout=timeout,
    )


def create_desktop_root(tk_module):
    try:
        from tkinterdnd2 import TkinterDnD

        return TkinterDnD.Tk()
    except Exception:
        return tk_module.Tk()


class DesktopGuiApp:
    """Native tkinter shell around the same local run state used by Web Runner."""

    def __init__(self, root, state=None, tk_module=None, ttk_module=None, filedialog_module=None, messagebox_module=None, opener=None):
        self.root = root
        self.state = state or WebRunnerState()
        self.tk = tk_module
        self.ttk = ttk_module
        self.filedialog = filedialog_module
        self.messagebox = messagebox_module
        self.opener = opener or open_desktop_path
        self.active_run_id = None
        self.last_run = None
        self.log_offset = 0
        self._config_refresh_serial = 0
        self.artifact_paths = {}
        self.auto_opened_run_ids = set()
        self.config_path = desktop_gui_config_file_path()
        self._build()
        self.refresh_config()
        self.refresh_runs()

    def _build(self):
        tk = self.tk
        ttk = self.ttk
        self.root.title("Veritas Report Studio")
        self.root.geometry("1220x760")
        self.root.minsize(1040, 640)
        self._configure_style()

        self.input_var = tk.StringVar()
        self.output_var = tk.StringVar()
        self.input_display_var = tk.StringVar(value="拖入论文或项目目录")
        self.output_display_var = tk.StringVar(value="选择报告目录")
        self.fresh_var = tk.BooleanVar(value=False)
        self.auto_open_var = tk.BooleanVar(value=True)
        self.status_var = tk.StringVar(value="待命")
        self.report_type_var = tk.StringVar(value="待生成")
        self.risk_var = tk.StringVar(value="待评估")
        self.summary_var = tk.StringVar(value="选择材料后开始分析。")
        self.stage_var = tk.StringVar(value="待命")
        self.progress_var = tk.DoubleVar(value=0.0)
        self.log_title_var = tk.StringVar(value="运行日志")
        self.config_summary_var = tk.StringVar(value="状态")
        self.config_row_vars = []

        main = ttk.Frame(self.root, padding=14, style="App.TFrame")
        main.pack(fill=tk.BOTH, expand=True)
        main.columnconfigure(0, weight=0)
        main.columnconfigure(1, weight=1)
        main.rowconfigure(0, weight=1)

        left = ttk.Frame(main, width=330, padding=(16, 14), style="Sidebar.TFrame")
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 16))
        left.grid_propagate(False)
        right = ttk.Frame(main, style="App.TFrame")
        right.grid(row=0, column=1, sticky="nsew")
        right.rowconfigure(1, weight=1)
        right.columnconfigure(0, weight=1)

        self.drag_drop_available = self._enable_drag_drop_package()

        ttk.Label(left, text="Veritas", style="Brand.TLabel").pack(anchor=tk.W, pady=(0, 16))

        input_card = ttk.Frame(left, padding=12, style="SidebarCard.TFrame")
        input_card.pack(fill=tk.X, pady=(0, 10))
        input_pick = ttk.Button(input_card, textvariable=self.input_display_var, command=self.choose_file, style="Picker.TButton")
        input_pick.pack(fill=tk.X)
        self._register_drop_target(input_card, self._handle_input_drop)
        self._register_drop_target(input_pick, self._handle_input_drop)

        output_card = ttk.Frame(left, padding=12, style="SidebarCard.TFrame")
        output_card.pack(fill=tk.X, pady=(0, 10))
        output_pick = ttk.Button(output_card, textvariable=self.output_display_var, command=self.choose_output_directory, style="Picker.TButton")
        output_pick.pack(fill=tk.X, pady=(0, 8))
        output_buttons = ttk.Frame(output_card, style="SidebarCard.TFrame")
        output_buttons.pack(fill=tk.X)
        ttk.Button(output_buttons, text="更改", command=self.choose_output_directory, style="Secondary.TButton").pack(side=tk.LEFT)
        ttk.Button(output_buttons, text="清空", command=self.clear_output, style="Ghost.TButton").pack(side=tk.LEFT, padx=(8, 0))
        self._register_drop_target(output_card, self._handle_output_drop)
        self._register_drop_target(output_pick, self._handle_output_drop)

        options_card = ttk.Frame(left, padding=12, style="SidebarCard.TFrame")
        options_card.pack(fill=tk.X, pady=(0, 12))
        self._sidebar_checkbutton(options_card, "重新开始", self.fresh_var).pack(side=tk.LEFT)
        self._sidebar_checkbutton(options_card, "完成后打开报告", self.auto_open_var).pack(side=tk.LEFT, padx=(12, 0))

        action_buttons = ttk.Frame(left, style="Sidebar.TFrame")
        action_buttons.pack(fill=tk.X, pady=(0, 12))
        self.start_button = ttk.Button(action_buttons, text="开始分析", command=self.start_run, style="Primary.TButton")
        self.start_button.pack(fill=tk.X, pady=(0, 6))
        self.cancel_button = ttk.Button(action_buttons, text="停止任务", command=self.cancel_run, state=tk.DISABLED, style="Danger.TButton")
        self.cancel_button.pack(fill=tk.X, pady=(0, 6))
        self.retry_button = ttk.Button(action_buttons, text="重新运行", command=self.retry_run, state=tk.DISABLED, style="Secondary.TButton")
        self.retry_button.pack(fill=tk.X)

        config_card = ttk.Frame(left, padding=10, style="Config.TFrame")
        config_card.pack(fill=tk.BOTH, expand=True)
        config_head = ttk.Frame(config_card, style="Config.TFrame")
        config_head.pack(fill=tk.X, pady=(0, 6))
        ttk.Label(config_head, text="能力", style="ConfigSection.TLabel").pack(side=tk.LEFT)
        ttk.Label(config_head, textvariable=self.config_summary_var, style="ConfigSummary.TLabel").pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(config_head, text="刷新", command=self.refresh_config, style="Tiny.TButton").pack(side=tk.RIGHT)
        ttk.Button(config_head, text="LLM 设置", command=self.open_llm_settings, style="Tiny.TButton").pack(side=tk.RIGHT, padx=(0, 6))
        self.config_rows_frame = ttk.Frame(config_card, style="Config.TFrame")
        self.config_rows_frame.pack(fill=tk.X)
        for index in range(len(DESKTOP_GUI_CONFIG_CAPABILITIES) + len(DESKTOP_GUI_CONFIG_DEPENDENCIES)):
            row = ttk.Frame(self.config_rows_frame, padding=(6, 4), style="ConfigChip.TFrame")
            row.grid(row=index // 2, column=index % 2, sticky="ew", padx=(0, 6 if index % 2 == 0 else 0), pady=(0, 4))
            self.config_rows_frame.columnconfigure(index % 2, weight=1)
            name_var = tk.StringVar()
            status_var = tk.StringVar()
            ttk.Label(row, textvariable=name_var, style="ConfigName.TLabel").pack(side=tk.LEFT)
            status_label = ttk.Label(row, textvariable=status_var, style="ConfigOk.TLabel")
            status_label.pack(side=tk.RIGHT)
            self.config_row_vars.append((name_var, status_var, status_label))

        report = ttk.Frame(right, padding=14, style="ReportCard.TFrame")
        report.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        report.columnconfigure(0, weight=1)
        metrics = ttk.Frame(report, style="ReportCard.TFrame")
        metrics.grid(row=0, column=0, sticky="ew")
        metrics.columnconfigure(0, weight=1)
        metrics.columnconfigure(1, weight=1)
        metrics.columnconfigure(2, weight=1)
        status_card = ttk.Frame(metrics, padding=10, style="Metric.TFrame")
        status_card.grid(row=0, column=0, sticky="ew", padx=(0, 8))
        type_card = ttk.Frame(metrics, padding=10, style="Metric.TFrame")
        type_card.grid(row=0, column=1, sticky="ew", padx=(0, 8))
        risk_card = ttk.Frame(metrics, padding=10, style="Metric.TFrame")
        risk_card.grid(row=0, column=2, sticky="ew")
        ttk.Label(status_card, text="状态", style="MetricLabel.TLabel").pack(side=tk.LEFT)
        ttk.Label(status_card, textvariable=self.status_var, style="MetricValue.TLabel").pack(side=tk.LEFT, padx=(8, 0))
        ttk.Label(type_card, text="产物", style="MetricLabel.TLabel").pack(side=tk.LEFT)
        ttk.Label(type_card, textvariable=self.report_type_var, style="MetricValue.TLabel").pack(side=tk.LEFT, padx=(8, 0))
        ttk.Label(risk_card, text="信号", style="MetricLabel.TLabel").pack(side=tk.LEFT)
        ttk.Label(risk_card, textvariable=self.risk_var, style="MetricValue.TLabel").pack(side=tk.LEFT, padx=(8, 0))

        summary_row = ttk.Frame(report, style="ReportCard.TFrame")
        summary_row.grid(row=1, column=0, sticky="ew", pady=(10, 0))
        summary_row.columnconfigure(1, weight=1)
        ttk.Label(summary_row, text="摘要", style="CardTitle.TLabel").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        ttk.Label(summary_row, textvariable=self.summary_var, wraplength=760, style="Summary.TLabel").grid(row=0, column=1, sticky="ew")
        progress_row = ttk.Frame(report, style="ReportCard.TFrame")
        progress_row.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        progress_row.columnconfigure(0, weight=1)
        self.progress_bar = ttk.Progressbar(progress_row, variable=self.progress_var, maximum=100, mode="determinate", style="Audit.Horizontal.TProgressbar")
        self.progress_bar.grid(row=0, column=0, sticky="ew")
        ttk.Label(progress_row, textvariable=self.stage_var, style="Stage.TLabel").grid(row=0, column=1, sticky=tk.E, padx=(12, 0))

        artifact_frame = ttk.Frame(report, style="ReportCard.TFrame")
        artifact_frame.grid(row=3, column=0, sticky=tk.W, pady=(12, 0))
        self.artifact_buttons = {}
        for kind, label in DESKTOP_GUI_ARTIFACT_LABELS.items():
            button = ttk.Button(artifact_frame, text=label, command=lambda k=kind: self.open_artifact(k), state=tk.DISABLED, style="Artifact.TButton")
            button.pack(side=tk.LEFT, padx=(0, 8))
            self.artifact_buttons[kind] = button
        self.followup_buttons = {}
        for kind, label in DESKTOP_GUI_FOLLOWUP_LABELS.items():
            button = ttk.Button(artifact_frame, text=label, command=lambda k=kind: self.generate_followup(k), state=tk.DISABLED, style="Artifact.TButton")
            button.pack(side=tk.LEFT, padx=(0, 8))
            self.followup_buttons[kind] = button

        log_frame = ttk.Frame(right, padding=12, style="LogCard.TFrame")
        log_frame.grid(row=1, column=0, sticky="nsew")
        log_frame.rowconfigure(1, weight=1)
        log_frame.columnconfigure(0, weight=1)
        ttk.Label(log_frame, textvariable=self.log_title_var, style="LogTitle.TLabel").grid(row=0, column=0, sticky=tk.W, pady=(0, 8))
        self.log_text = tk.Text(log_frame, height=28, wrap=tk.WORD, bd=0, relief=tk.FLAT)
        self.log_text.grid(row=1, column=0, sticky="nsew")
        self.log_text.configure(
            bg="#101317",
            fg="#d9e0ea",
            insertbackground="#d9e0ea",
            selectbackground="#30445f",
            font=("SF Mono", 10),
            padx=12,
            pady=12,
        )
        self._set_log_editable(False)

    def _configure_style(self):
        ttk = self.ttk
        try:
            ttk.Style().theme_use("clam")
        except Exception:
            pass
        style = ttk.Style()
        colors = {
            "app": "#f0f2f6",
            "sidebar": "#f8f9fb",
            "sidebar_card": "#ffffff",
            "config": "#ffffff",
            "text": "#111827",
            "muted": "#7b8190",
            "card": "#ffffff",
            "metric": "#f7f8fc",
            "accent": "#0071e3",
            "accent_dark": "#005bb5",
            "danger": "#ff453a",
            "danger_dark": "#d92d25",
            "log": "#101317",
        }
        self.root.configure(bg=colors["app"])
        style.configure("App.TFrame", background=colors["app"])
        style.configure("Sidebar.TFrame", background=colors["sidebar"])
        style.configure("SidebarCard.TFrame", background=colors["sidebar_card"], relief="flat", borderwidth=0)
        style.configure("Config.TFrame", background=colors["config"], relief="flat", borderwidth=0)
        style.configure("ConfigChip.TFrame", background="#f8f9fc", relief="flat", borderwidth=0)
        style.configure("ReportCard.TFrame", background=colors["card"], relief="flat", borderwidth=0)
        style.configure("Metric.TFrame", background=colors["metric"], relief="flat")
        style.configure("LogCard.TFrame", background=colors["log"], relief="flat")
        style.configure("CloseDot.TLabel", background=colors["sidebar"], foreground="#ff5f57", font=("Avenir Next", 12, "bold"))
        style.configure("MinDot.TLabel", background=colors["sidebar"], foreground="#ffbd2e", font=("Avenir Next", 12, "bold"))
        style.configure("ZoomDot.TLabel", background=colors["sidebar"], foreground="#28c840", font=("Avenir Next", 12, "bold"))
        style.configure("Brand.TLabel", background=colors["sidebar"], foreground=colors["text"], font=("Avenir Next", 23, "bold"))
        style.configure("BrandSub.TLabel", background=colors["sidebar"], foreground=colors["accent"], font=("Avenir Next", 13, "bold"))
        style.configure("SidebarHint.TLabel", background=colors["sidebar"], foreground=colors["muted"], font=("Avenir Next", 10))
        style.configure("SidebarSection.TLabel", background=colors["sidebar_card"], foreground=colors["muted"], font=("Avenir Next", 8, "bold"))
        style.configure("ConfigSection.TLabel", background=colors["config"], foreground=colors["muted"], font=("Avenir Next", 8, "bold"))
        style.configure("ConfigSummary.TLabel", background=colors["config"], foreground=colors["text"], font=("Avenir Next", 8, "bold"))
        style.configure("ConfigName.TLabel", background="#f8f9fc", foreground=colors["text"], font=("Avenir Next", 8))
        style.configure("ConfigOk.TLabel", background="#f8f9fc", foreground="#2e7d32", font=("Avenir Next", 8, "bold"))
        style.configure("ConfigWarn.TLabel", background="#f8f9fc", foreground="#b45309", font=("Avenir Next", 8, "bold"))
        style.configure("PageTitle.TLabel", background=colors["app"], foreground=colors["text"], font=("Avenir Next", 25, "bold"))
        style.configure("PageSub.TLabel", background=colors["app"], foreground=colors["muted"], font=("Avenir Next", 10))
        style.configure("StatusPill.TLabel", background="#eaf2ff", foreground="#0f5fb8", font=("Avenir Next", 10, "bold"), padding=(12, 6))
        style.configure("CardTitle.TLabel", background=colors["card"], foreground=colors["text"], font=("Avenir Next", 12, "bold"))
        style.configure("Summary.TLabel", background=colors["card"], foreground="#33363d", font=("Avenir Next", 11))
        style.configure("MetricLabel.TLabel", background=colors["metric"], foreground=colors["muted"], font=("Avenir Next", 8, "bold"))
        style.configure("MetricValue.TLabel", background=colors["metric"], foreground=colors["text"], font=("Avenir Next", 14, "bold"))
        style.configure("Stage.TLabel", background=colors["card"], foreground=colors["muted"], font=("Avenir Next", 10, "bold"))
        style.configure("Audit.Horizontal.TProgressbar", troughcolor="#e6e8ed", background=colors["accent"], bordercolor="#e6e8ed", lightcolor=colors["accent"], darkcolor=colors["accent"], thickness=9)
        style.configure("LogTitle.TLabel", background=colors["log"], foreground="#f5f5f7", font=("Avenir Next", 12, "bold"))
        style.configure("Sidebar.TCheckbutton", background=colors["sidebar_card"], foreground=colors["text"], font=("Avenir Next", 9))
        style.map("Sidebar.TCheckbutton", background=[("active", colors["sidebar_card"])], foreground=[("active", colors["text"])])
        style.configure("Path.TEntry", fieldbackground="#ffffff", foreground=colors["text"], padding=6)
        style.configure("Picker.TButton", background="#f8f9fc", foreground="#252b36", font=("Avenir Next", 10, "bold"), padding=(10, 16), borderwidth=0)
        style.map("Picker.TButton", background=[("active", "#eef6ff"), ("disabled", "#e4e5e8")], foreground=[("disabled", "#8c8c91")])
        style.configure("Primary.TButton", background=colors["accent"], foreground="#ffffff", font=("Avenir Next", 11, "bold"), padding=(12, 9), borderwidth=0)
        style.map("Primary.TButton", background=[("active", colors["accent_dark"]), ("disabled", "#8a9a93")])
        style.configure("Secondary.TButton", background="#f8f9fc", foreground=colors["text"], font=("Avenir Next", 9, "bold"), padding=(10, 7), borderwidth=0)
        style.map("Secondary.TButton", background=[("active", "#e9eef7"), ("disabled", "#d6d7d9")], foreground=[("disabled", "#8c8c91")])
        style.configure("Ghost.TButton", background=colors["sidebar_card"], foreground=colors["accent"], font=("Avenir Next", 9, "bold"), padding=(9, 7), borderwidth=0)
        style.map("Ghost.TButton", background=[("active", "#edf4ff"), ("disabled", "#d6d7d9")])
        style.configure("Tiny.TButton", background="#f8f9fc", foreground=colors["accent"], font=("Avenir Next", 8, "bold"), padding=(6, 3), borderwidth=0)
        style.map("Tiny.TButton", background=[("active", "#edf4ff"), ("disabled", "#d6d7d9")])
        style.configure("Danger.TButton", background=colors["danger"], foreground="#ffffff", font=("Avenir Next", 9, "bold"), padding=(10, 7), borderwidth=0)
        style.map("Danger.TButton", background=[("active", colors["danger_dark"]), ("disabled", "#d6d7d9")])
        style.configure("Artifact.TButton", background="#eef6ff", foreground="#1f4f7a", font=("Avenir Next", 9, "bold"), padding=(10, 7), borderwidth=0)
        style.map("Artifact.TButton", background=[("active", "#e6f0ff"), ("disabled", "#e4e5e8")], foreground=[("disabled", "#8c8c91")])

    def _sidebar_checkbutton(self, parent, text, variable):
        tk = self.tk
        button = tk.Checkbutton(
            parent,
            text=text,
            variable=variable,
            indicatoron=False,
            onvalue=True,
            offvalue=False,
            bg="#f3f6fb",
            activebackground="#eaf2ff",
            fg="#111827",
            activeforeground="#111827",
            selectcolor="#0071e3",
            font=("Avenir Next", 9, "bold"),
            bd=0,
            highlightthickness=0,
            relief=tk.FLAT,
            overrelief=tk.FLAT,
            padx=12,
            pady=7,
        )

        def sync(*_args):
            selected = bool(variable.get())
            button.configure(
                bg="#0071e3" if selected else "#f3f6fb",
                activebackground="#005bb5" if selected else "#eaf2ff",
                fg="#ffffff" if selected else "#111827",
                activeforeground="#ffffff" if selected else "#111827",
            )

        try:
            trace_id = variable.trace_add("write", sync)
            button._veritas_trace_id = trace_id
        except Exception:
            pass
        sync()
        return button

    def _render_config_snapshot(self, snapshot):
        if hasattr(self, "config_summary_var"):
            self.config_summary_var.set(snapshot.get("summary") or "不可用")
        rows = snapshot.get("rows") or []
        for index, row_vars in enumerate(getattr(self, "config_row_vars", [])):
            name_var, status_var, status_label = row_vars
            row = rows[index] if index < len(rows) else None
            if not row:
                name_var.set("")
                status_var.set("")
                status_label.configure(style="ConfigOk.TLabel")
                continue
            name_var.set(row["label"])
            status_var.set(row["status"])
            status_label.configure(style="ConfigOk.TLabel" if row.get("ok") else "ConfigWarn.TLabel")

    def _current_llm_config(self):
        return load_runtime_config(verbose=False).text_llm

    def _save_llm_settings(self, window, api_key_var, api_url_var, model_var):
        api_key = api_key_var.get().strip()
        api_url = api_url_var.get().strip()
        model = model_var.get().strip()
        if not api_key or not api_url or not model:
            self.messagebox.showerror("配置不完整", "请填写 API Key、API URL 和模型名称。")
            return
        try:
            path = desktop_gui_write_llm_config(api_key, api_url, model, config_path=getattr(self, "config_path", None))
            runtime_config = load_runtime_config(verbose=False)
            apply_runtime_config(runtime_config)
            self.refresh_config()
            window.destroy()
            self.messagebox.showinfo("已保存", f"LLM 设置已保存到 {path}。")
        except Exception as e:
            self.messagebox.showerror("保存失败", f"{type(e).__name__}: {_brief_text(str(e), 240)}")

    def open_llm_settings(self):
        tk = self.tk
        ttk = self.ttk
        try:
            llm = self._current_llm_config()
        except Exception:
            llm = CapabilityConfig("text_llm", api_url=LLM_API_URL, model=LLM_MODEL)
        window = tk.Toplevel(self.root)
        window.title("LLM 设置")
        window.transient(self.root)
        try:
            window.resizable(False, False)
        except Exception:
            pass
        frame = ttk.Frame(window, padding=18, style="App.TFrame")
        frame.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frame, text="文本 LLM", style="CardTitle.TLabel").grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 10))

        api_key_var = tk.StringVar(value=llm.api_key or "")
        api_url_var = tk.StringVar(value=llm.api_url or LLM_API_URL)
        model_var = tk.StringVar(value=llm.model or LLM_MODEL)

        fields = [
            ("API Key", api_key_var, "*"),
            ("API URL", api_url_var, ""),
            ("模型", model_var, ""),
        ]
        for row_index, (label, variable, show) in enumerate(fields, start=1):
            ttk.Label(frame, text=label, style="MetricLabel.TLabel").grid(row=row_index, column=0, sticky=tk.W, pady=(0, 8), padx=(0, 10))
            entry = ttk.Entry(frame, textvariable=variable, width=46, show=show)
            entry.grid(row=row_index, column=1, sticky="ew", pady=(0, 8))
            if row_index == 1:
                entry.focus_set()
        buttons = ttk.Frame(frame, style="App.TFrame")
        buttons.grid(row=4, column=0, columnspan=2, sticky=tk.E, pady=(8, 0))
        ttk.Button(buttons, text="取消", command=window.destroy, style="Secondary.TButton").pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(
            buttons,
            text="保存",
            command=lambda: self._save_llm_settings(window, api_key_var, api_url_var, model_var),
            style="Primary.TButton",
        ).pack(side=tk.LEFT)

    def _set_log_editable(self, editable):
        try:
            self.log_text.configure(state=self.tk.NORMAL if editable else self.tk.DISABLED)
        except Exception:
            pass

    def _reset_progress(self, label="待命"):
        if hasattr(self, "progress_var"):
            self.progress_var.set(0.0)
        if hasattr(self, "stage_var"):
            self.stage_var.set(label)

    def _update_progress_from_line(self, line):
        progress = desktop_gui_progress_from_log_line(line)
        if not progress:
            return False
        if hasattr(self, "progress_var"):
            self.progress_var.set(progress["percent"])
        if hasattr(self, "stage_var"):
            self.stage_var.set(progress["label"])
        return True

    def _clear_log_text(self):
        if hasattr(self, "log_title_var"):
            self.log_title_var.set("运行日志")
        self._set_log_editable(True)
        self.log_text.delete("1.0", self.tk.END)
        self._set_log_editable(False)

    def _replace_log_text(self, title, content):
        if hasattr(self, "log_title_var"):
            self.log_title_var.set(title)
        self._set_log_editable(True)
        try:
            self.log_text.delete("1.0", self.tk.END)
            self.log_text.insert(self.tk.END, str(content or ""))
            self.log_text.see("1.0")
        finally:
            self._set_log_editable(False)

    def _render_artifact_in_log(self, kind, path):
        label = DESKTOP_GUI_ARTIFACT_LABELS.get(kind, kind)
        content = desktop_gui_artifact_preview(path, kind)
        self._replace_log_text(f"报告预览 · {label}", content)

    def _display_path_text(self, raw_path, placeholder):
        text = str(raw_path or "").strip()
        if not text:
            return placeholder
        try:
            path = Path(text).expanduser()
            if path.name == "audit_report":
                path = path.parent
            label = str(path)
        except Exception:
            label = text
        if len(label) <= 36:
            return label
        return "..." + label[-33:]

    def _set_input_path(self, path):
        path_text = str(Path(path).expanduser())
        self.input_var.set(path_text)
        if hasattr(self, "input_display_var"):
            self.input_display_var.set(self._display_path_text(path_text, "拖入论文或项目目录"))

    def _set_output_path(self, output_stem):
        output_text = str(output_stem or "").strip()
        self.output_var.set(output_text)
        if hasattr(self, "output_display_var"):
            self.output_display_var.set(self._display_path_text(output_text, "选择报告目录"))

    def clear_output(self):
        self._set_output_path("")

    def _enable_drag_drop_package(self):
        try:
            self.root.tk.call("package", "require", "tkdnd")
            return True
        except Exception:
            return False

    def _register_drop_target(self, widget, handler):
        if not getattr(self, "drag_drop_available", False):
            return
        try:
            self.root.tk.call("tkdnd::drop_target", "register", widget._w, "DND_Files")
            callback = widget.register(lambda data: handler(data))
            self.root.tk.call("bind", widget._w, "<<Drop>>", f"{callback} %D")
        except Exception:
            pass

    def _path_from_drop_data(self, data):
        raw = str(data or "").strip()
        if not raw:
            return ""
        try:
            items = self.root.tk.splitlist(raw)
        except Exception:
            items = [raw]
        if not items:
            return ""
        item = str(items[0]).strip().strip("{}")
        if item.lower().startswith("file://"):
            return dropped_local_path_from_uri_text(item)
        return item

    def _handle_input_drop(self, data):
        path = self._path_from_drop_data(data)
        if path:
            self._set_input_path(path)

    def _handle_output_drop(self, data):
        path = self._path_from_drop_data(data)
        if not path:
            return
        dropped = Path(path).expanduser()
        output_dir = dropped.parent if dropped.is_file() else dropped
        self._set_output_path(output_dir / "audit_report")

    def choose_file(self):
        selected = self.filedialog.askopenfilename(title="选择材料文件")
        if selected:
            self._set_input_path(selected)

    def choose_directory(self):
        selected = self.filedialog.askdirectory(title="选择材料目录", mustexist=True)
        if selected:
            self._set_input_path(selected)

    def choose_output_directory(self):
        selected = self.filedialog.askdirectory(title="选择报告目录", mustexist=False)
        if selected:
            self._set_output_path(Path(selected).expanduser() / "audit_report")

    def refresh_config(self):
        self._config_refresh_serial = getattr(self, "_config_refresh_serial", 0) + 1
        serial = self._config_refresh_serial
        try:
            config = web_runner_config_status()
            self._render_config_snapshot(desktop_gui_config_snapshot(config, preflight_results={"text_llm": {"status": "检测中", "ok": False}}))
        except Exception as e:
            self._render_config_snapshot({"summary": f"{type(e).__name__}: {_brief_text(str(e), 120)}", "rows": []})
            return

        def worker():
            try:
                snapshot = desktop_gui_checked_config_snapshot()
            except Exception as e:
                snapshot = {"summary": f"{type(e).__name__}: {_brief_text(str(e), 120)}", "rows": []}

            def apply_snapshot():
                if serial == getattr(self, "_config_refresh_serial", 0):
                    self._render_config_snapshot(snapshot)

            try:
                self.root.after(0, apply_snapshot)
            except Exception:
                pass

        threading.Thread(target=worker, daemon=True).start()

    def refresh_runs(self):
        runs = self.state.list_runs()
        running = next((run for run in runs if run.get("status") == "running"), None)
        if running and not self.active_run_id:
            self.active_run_id = running.get("id")
            self.log_offset = 0
            self._clear_log_text()
            self.poll_run()
            return
        if runs and not self.active_run_id:
            latest = runs[0]
            self.last_run = latest
            self.log_offset = 0
            self._clear_log_text()
            self._reset_progress()
            self.render_run(latest)
            self._append_logs_since(latest.get("id"))
            self.retry_button.configure(state=self.tk.NORMAL if latest.get("input_path") else self.tk.DISABLED)

    def start_run(self):
        input_path = self.input_var.get().strip()
        if not input_path:
            self.messagebox.showerror("缺少材料", "请选择文件或项目目录。")
            return
        result, status = desktop_gui_start_run(self.state, input_path, self.output_var.get(), self.fresh_var.get())
        run = result.get("run") if isinstance(result, dict) else None
        if status != 200 or not run:
            message = result.get("message") or result.get("error") or "启动失败"
            self.status_var.set(message)
            self.messagebox.showerror("启动失败", message)
            return
        self.active_run_id = run.get("id")
        self.last_run = run
        self.log_offset = 0
        self.start_button.configure(state=self.tk.DISABLED)
        self.cancel_button.configure(state=self.tk.NORMAL)
        self.retry_button.configure(state=self.tk.DISABLED)
        self._clear_log_text()
        self._reset_progress("启动中")
        self.render_run(run)
        self.poll_run()

    def cancel_run(self):
        if not self.active_run_id:
            return
        self.state.cancel_run(self.active_run_id)
        self.poll_run()

    def retry_run(self):
        run = self.last_run or {}
        input_path = run.get("input_path") or self.input_var.get()
        if not input_path:
            return
        self.input_var.set(input_path)
        if hasattr(self, "input_display_var"):
            self.input_display_var.set(self._display_path_text(input_path, "拖入论文或项目目录"))
        self._set_output_path(run.get("output") or self.output_var.get())
        self.fresh_var.set(bool(run.get("fresh")))
        self.start_run()

    def poll_run(self):
        if not self.active_run_id:
            return
        self._append_logs_since(self.active_run_id)
        run = self.state.get_run(self.active_run_id)
        if run:
            self.render_run(run)
            if run.get("status") != "running":
                refreshed = self.state.discover_artifacts(self.active_run_id) if self.active_run_id else None
                if refreshed:
                    run = refreshed
                    self.render_run(run)
                self._append_logs_since(run.get("id"))
                self._maybe_auto_open_completed_report(run)
                self.active_run_id = None
                self.start_button.configure(state=self.tk.NORMAL)
                self.cancel_button.configure(state=self.tk.DISABLED)
                self.retry_button.configure(state=self.tk.NORMAL if run.get("input_path") else self.tk.DISABLED)
                return
        self.root.after(1000, self.poll_run)

    def _append_logs_since(self, run_id):
        if not run_id:
            return
        logs = self.state.logs_since(run_id, self.log_offset)
        if not logs:
            return
        self.log_offset = logs.get("offset", self.log_offset)
        lines = logs.get("lines") or []
        if not lines:
            return
        self._set_log_editable(True)
        try:
            for line in lines:
                self._update_progress_from_line(line)
                self.log_text.insert(self.tk.END, str(line) + "\n")
            self.log_text.see(self.tk.END)
        finally:
            self._set_log_editable(False)

    def render_run(self, run):
        self.last_run = run or self.last_run
        view = desktop_gui_run_summary(run)
        self.status_var.set(view["status_label"])
        self.report_type_var.set(view["report_type_label"])
        self.risk_var.set(view["risk_label"])
        self.summary_var.set(view["summary"] or "完成后显示摘要。")
        if run and run.get("status") == "succeeded" and hasattr(self, "progress_var"):
            self.progress_var.set(100.0)
            self.stage_var.set("已完成")
        elif run and run.get("status") in {"failed", "canceled"} and hasattr(self, "stage_var"):
            if not self.stage_var.get() or self.stage_var.get() == "待命":
                self.stage_var.set(view["status_label"])
        self.artifact_paths = view["artifacts"]
        for kind, button in self.artifact_buttons.items():
            button.configure(state=self.tk.NORMAL if kind in self.artifact_paths else self.tk.DISABLED)
        followup_enabled = bool(run and run.get("status") == "succeeded" and self.artifact_paths.get("json"))
        for button in getattr(self, "followup_buttons", {}).values():
            button.configure(state=self.tk.NORMAL if followup_enabled else self.tk.DISABLED)

    def open_artifact(self, kind):
        path = self.artifact_paths.get(kind)
        if not path:
            return
        try:
            if kind in {"html", "markdown", "json"}:
                self._render_artifact_in_log(kind, path)
            else:
                self.opener(path)
        except Exception as e:
            self.messagebox.showerror("打开失败", f"{type(e).__name__}: {_brief_text(str(e), 240)}")

    def _current_followup_run(self):
        run = dict(self.last_run or {})
        artifacts = dict(run.get("artifacts") or {})
        artifacts.update(getattr(self, "artifact_paths", {}) or {})
        run["artifacts"] = artifacts
        return run

    def _set_followup_buttons_state(self, state):
        for button in getattr(self, "followup_buttons", {}).values():
            button.configure(state=state)

    def generate_followup(self, kind):
        label = DESKTOP_GUI_FOLLOWUP_LABELS.get(kind, kind)
        run = self._current_followup_run()
        try:
            desktop_gui_followup_context(run)
        except Exception as e:
            self.messagebox.showerror("无法生成草稿", f"{type(e).__name__}: {_brief_text(str(e), 240)}")
            return
        self._set_followup_buttons_state(self.tk.DISABLED)
        if hasattr(self, "stage_var"):
            self.stage_var.set(f"{label} 生成中")

        def worker():
            try:
                result = desktop_gui_generate_followup_draft(kind, run)
                error = None
            except Exception as e:
                result = None
                error = e

            def apply_result():
                self._set_followup_buttons_state(self.tk.NORMAL)
                if error is not None:
                    if hasattr(self, "stage_var"):
                        self.stage_var.set("草稿生成失败")
                    self.messagebox.showerror("生成失败", f"{type(error).__name__}: {_brief_text(str(error), 240)}")
                    return
                self._replace_log_text(f"草稿 · {label}", result.get("text") or "")
                if hasattr(self, "stage_var"):
                    self.stage_var.set("草稿已生成")
                draft_path = (result.get("paths") or {}).get("draft_path")
                if self.messagebox:
                    self.messagebox.showinfo("已生成", f"草稿已保存: {draft_path}")

            try:
                self.root.after(0, apply_result)
            except Exception:
                apply_result()

        threading.Thread(target=worker, daemon=True).start()

    def _maybe_auto_open_completed_report(self, run):
        if not run or run.get("status") != "succeeded" or not self.auto_open_var.get():
            return
        run_id = str(run.get("id") or "")
        html_path = (run.get("artifacts") or {}).get("html")
        if not run_id or not html_path or run_id in self.auto_opened_run_ids:
            return
        self.auto_opened_run_ids.add(run_id)
        try:
            self._render_artifact_in_log("html", html_path)
        except Exception as e:
            self.messagebox.showerror("打开失败", f"{type(e).__name__}: {_brief_text(str(e), 240)}")


def run_desktop_gui(history_path=None):
    try:
        import tkinter as tk
        from tkinter import filedialog, messagebox, ttk
    except Exception as e:
        print(f"无法启动桌面 GUI：tkinter 不可用: {type(e).__name__}: {_brief_text(str(e), 240)}")
        return 1
    try:
        root = create_desktop_root(tk)
        DesktopGuiApp(root, state=WebRunnerState(history_path=history_path), tk_module=tk, ttk_module=ttk, filedialog_module=filedialog, messagebox_module=messagebox)
        root.mainloop()
        return 0
    except Exception as e:
        print(f"无法启动桌面 GUI：{type(e).__name__}: {_brief_text(str(e), 240)}")
        return 1


def gui_main():
    apply_runtime_config(load_runtime_config())
    return run_desktop_gui()


def render_web_runner_page():
    return """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Veritas Web Runner</title>
<style>
:root { color-scheme: light; --bg:#f5f5f3; --panel:#ffffff; --line:#d8d8d3; --text:#191919; --muted:#6a6a64; --accent:#155e75; --danger:#b42318; --ok:#237a4b; --warn:#9a5b00; --soft:#f0f7f8; }
* { box-sizing:border-box; }
body { margin:0; background:var(--bg); color:var(--text); font-family:ui-sans-serif,-apple-system,BlinkMacSystemFont,"Segoe UI","Noto Sans SC",sans-serif; font-size:14px; letter-spacing:0; }
header { height:54px; display:flex; align-items:center; justify-content:space-between; padding:0 20px; border-bottom:1px solid var(--line); background:#fff; }
h1 { font-size:18px; margin:0; font-weight:750; }
main { display:grid; grid-template-columns:minmax(360px, 520px) minmax(420px, 1fr); gap:16px; padding:16px; max-width:1440px; margin:0 auto; }
section { background:var(--panel); border:1px solid var(--line); border-radius:8px; padding:14px; }
h2 { font-size:14px; margin:0 0 12px; font-weight:750; }
label { display:block; color:var(--muted); font-size:12px; margin:10px 0 5px; }
input[type="text"] { width:100%; min-height:38px; border:1px solid var(--line); border-radius:6px; padding:8px 10px; font:inherit; background:#fff; color:var(--text); }
input[readonly] { background:#f7f7f5; color:#2f2f2c; cursor:default; }
.drop-zone { border:1px dashed var(--line); border-radius:8px; padding:10px; background:#fbfbfa; transition:border-color .12s ease, background .12s ease; }
.drop-zone.dragover { border-color:var(--accent); background:#edf7f8; }
.drop-zone input[type="text"] { margin-top:5px; }
.drop-hint { min-height:18px; margin-top:6px; }
.path-row { display:grid; grid-template-columns:minmax(0,1fr) auto auto; gap:8px; align-items:end; }
.path-row.output { grid-template-columns:minmax(0,1fr) auto; margin-top:5px; }
.row { display:flex; gap:8px; align-items:center; flex-wrap:wrap; }
.actions { margin-top:12px; display:flex; gap:8px; }
button, .linkbtn { min-height:36px; border:1px solid var(--line); border-radius:6px; background:#fff; color:var(--text); padding:7px 12px; font:inherit; font-weight:650; cursor:pointer; text-decoration:none; display:inline-flex; align-items:center; justify-content:center; }
button.primary { background:var(--accent); border-color:var(--accent); color:#fff; }
button.danger { border-color:#e3aaa5; color:var(--danger); }
button:disabled { opacity:.55; cursor:not-allowed; }
.status { display:inline-flex; min-height:26px; align-items:center; padding:3px 8px; border-radius:999px; border:1px solid var(--line); color:var(--muted); font-size:12px; }
.status.succeeded { color:var(--ok); border-color:#9bc7ad; }
.status.failed, .status.canceled { color:var(--danger); border-color:#e3aaa5; }
.hero-strip { display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:8px; margin-bottom:12px; }
.step { border:1px solid var(--line); border-radius:6px; padding:8px; background:#fbfbfa; min-height:58px; }
.step span { display:block; color:var(--muted); font-size:12px; }
.step strong { display:block; margin-top:3px; font-size:13px; }
.grid { display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:8px; }
.cell { border:1px solid var(--line); border-radius:6px; padding:8px; min-height:54px; overflow-wrap:anywhere; }
.cell span { display:block; color:var(--muted); font-size:12px; }
.cell strong { display:block; margin-top:3px; font-size:13px; }
#log { height:430px; overflow:auto; white-space:pre-wrap; word-break:break-word; background:#161616; color:#f4f4f0; border-radius:6px; padding:12px; font:12px/1.5 ui-monospace,SFMono-Regular,Menlo,Consolas,monospace; }
.run { border-top:1px solid var(--line); padding:10px 0; }
.run:first-child { border-top:0; padding-top:0; }
.run-title { display:flex; justify-content:space-between; gap:8px; align-items:center; }
.run path, code { font-family:ui-monospace,SFMono-Regular,Menlo,Consolas,monospace; }
.muted { color:var(--muted); font-size:12px; overflow-wrap:anywhere; }
.links { display:flex; gap:6px; flex-wrap:wrap; margin-top:8px; }
.check { display:flex; align-items:center; gap:7px; margin-top:10px; color:var(--text); }
.current-card { border:1px solid var(--line); border-radius:6px; padding:9px; margin-bottom:10px; min-height:58px; }
.current-card span { display:block; color:var(--muted); font-size:12px; }
.current-card strong { display:block; margin-top:3px; overflow-wrap:anywhere; }
.current-actions { display:flex; gap:6px; flex-wrap:wrap; margin:8px 0 10px; min-height:36px; }
.report-panel { border:1px solid var(--line); border-radius:8px; padding:10px; margin:0 0 10px; background:var(--soft); }
.report-panel.failed, .report-panel.canceled { background:#fff7f6; }
.report-panel.succeeded, .report-panel.complete, .report-panel.limited { background:#f3faf6; }
.report-head { display:flex; align-items:center; justify-content:space-between; gap:8px; margin-bottom:8px; }
.report-head h3 { margin:0; font-size:14px; }
.report-grid { display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:8px; }
.report-cell { border:1px solid rgba(0,0,0,.08); border-radius:6px; padding:8px; background:rgba(255,255,255,.72); overflow-wrap:anywhere; }
.report-cell span { display:block; color:var(--muted); font-size:12px; }
.report-cell strong { display:block; margin-top:3px; font-size:13px; }
.report-summary { margin:8px 0 0; color:#2f2f2c; line-height:1.55; overflow-wrap:anywhere; }
.feedback { min-height:34px; border:1px solid var(--line); border-radius:6px; padding:8px; margin-bottom:10px; background:#fbfbfa; color:#2f2f2c; overflow-wrap:anywhere; }
.feedback.failed { border-color:#e3aaa5; color:var(--danger); background:#fff7f6; }
.feedback.succeeded { border-color:#9bc7ad; color:var(--ok); background:#f3faf6; }
@media (max-width:900px) { main { grid-template-columns:1fr; padding:10px; } header { padding:0 12px; } #log { height:320px; } .hero-strip, .report-grid { grid-template-columns:1fr; } }
</style>
</head>
<body>
<header><h1>Veritas Web Runner</h1><span id="topStatus" class="status">local</span></header>
<main>
  <div>
    <section>
      <h2>审查任务</h2>
      <div class="hero-strip" aria-label="workflow">
        <div class="step"><span>1. 输入</span><strong>选择论文文件或项目目录</strong></div>
        <div class="step"><span>2. 审查</span><strong>本地启动正式 CLI 流程</strong></div>
        <div class="step"><span>3. 输出</span><strong>直接打开生成报告</strong></div>
      </div>
      <div id="inputDropZone" class="drop-zone">
        <label for="inputPath">输入路径 / 拖拽区域</label>
        <div class="path-row">
          <input id="inputPath" type="text" autocomplete="off" readonly placeholder="/path/to/paper-or-folder">
          <button id="pickFileBtn" type="button">文件</button>
          <button id="pickDirectoryBtn" type="button">目录</button>
        </div>
        <div id="dropHint" class="muted drop-hint">拖拽文件或目录到这里</div>
      </div>
      <label for="outputPath">输出路径</label>
      <div class="path-row output">
        <input id="outputPath" type="text" autocomplete="off" readonly placeholder="">
        <button id="pickOutputBtn" type="button">输出</button>
      </div>
      <label class="check"><input id="fresh" type="checkbox"> 从头重跑</label>
      <div class="actions">
        <button id="startBtn" class="primary">Start</button>
        <button id="cancelBtn" class="danger" disabled>Cancel</button>
      </div>
    </section>
    <section style="margin-top:16px">
      <h2>配置状态</h2>
      <div id="config" class="grid"></div>
    </section>
  </div>
  <div>
    <section>
      <div class="row" style="justify-content:space-between;margin-bottom:10px"><h2 style="margin:0">当前运行</h2><span id="runStatus" class="status">idle</span></div>
      <div id="currentRun" class="current-card"><span>输入</span><strong>No active run</strong></div>
      <div id="currentOutput" class="current-card"><span>输出</span><strong></strong></div>
      <div id="runFeedback" class="feedback">选择输入后点击 Start，生成的报告会在这里显示。</div>
      <div id="reportPanel" class="report-panel">
        <div class="report-head"><h3>报告输出</h3><span id="reportState" class="status">pending</span></div>
        <div class="report-grid">
          <div class="report-cell"><span>报告类型</span><strong id="reportType">待生成</strong></div>
          <div class="report-cell"><span>风险级别</span><strong id="reportRisk">待生成</strong></div>
          <div class="report-cell"><span>输出目录</span><strong id="reportFolder">待生成</strong></div>
        </div>
        <p id="reportSummary" class="report-summary">报告生成后会显示摘要和打开入口。</p>
        <div id="currentActions" class="current-actions"></div>
      </div>
      <div id="log" aria-label="live log"></div>
    </section>
    <section style="margin-top:16px">
      <h2>最近运行</h2>
      <div id="runs"></div>
    </section>
  </div>
</main>
<script>
const $ = (id) => document.getElementById(id);
let activeRunId = null;
let logOffset = 0;
let timer = null;
let selectedInputKind = '';
let lastRun = null;
async function api(path, options={}) {
  const res = await fetch(path, {headers: {'Content-Type': 'application/json'}, ...options});
  const data = await res.json();
  if (!res.ok) throw data;
  return data;
}
function setStatus(text, cls='') {
  const el = $('runStatus');
  el.className = 'status ' + cls;
  el.textContent = text;
}
function setFeedback(text, cls='') {
  const el = $('runFeedback');
  if (!el) return;
  el.className = 'feedback ' + cls;
  el.textContent = text || '选择输入后点击 Start，生成的报告会在这里显示。';
}
function escapeHtml(value) {
  return String(value || '').replace(/[&<>"']/g, (ch) => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[ch]));
}
function pathSeparator(path) {
  return String(path || '').includes('\\\\') ? '\\\\' : '/';
}
function trimPath(path) {
  return String(path || '').trim().replace(/[\\\\/]+$/, '');
}
function pathBaseName(path) {
  const cleaned = trimPath(path);
  const sep = pathSeparator(cleaned);
  return cleaned.split(sep).filter(Boolean).pop() || 'audit_project';
}
function pathParent(path) {
  const cleaned = trimPath(path);
  const sep = pathSeparator(cleaned);
  const parts = cleaned.split(sep);
  parts.pop();
  const parent = parts.join(sep);
  if (parent) return parent;
  return cleaned.startsWith(sep) ? sep : '.';
}
function joinPath(...parts) {
  const filtered = parts.map(p => String(p || '')).filter(Boolean);
  const sep = filtered.some(p => p.includes('\\\\')) ? '\\\\' : '/';
  const joined = filtered.join(sep);
  return sep === '\\\\' ? joined.replace(/\\\\{2,}/g, '\\\\') : joined.replace(/\\/{2,}/g, '/');
}
function stripExtension(name) {
  const dot = String(name || '').lastIndexOf('.');
  return dot > 0 ? name.slice(0, dot) : name;
}
function safeProjectName(name) {
  return String(name || 'audit_project').replace(/[<>:"/\\\\|?*\\x00-\\x1f]+/g, '_').replace(/[ .]+$/g, '') || 'audit_project';
}
function timestampForOutput(date = new Date()) {
  const pad = (value) => String(value).padStart(2, '0');
  return `${date.getFullYear()}${pad(date.getMonth() + 1)}${pad(date.getDate())}-${pad(date.getHours())}${pad(date.getMinutes())}${pad(date.getSeconds())}`;
}
function defaultOutputStemForInput(inputPath, kind = selectedInputKind, date = new Date()) {
  const cleaned = trimPath(inputPath);
  if (!cleaned) return '';
  const base = pathBaseName(cleaned);
  const project = safeProjectName(kind === 'directory' ? base : stripExtension(base));
  return joinPath(pathParent(cleaned), `${project}_${timestampForOutput(date)}`, 'audit_report');
}
function localPathFromFileUri(rawUri) {
  const uri = String(rawUri || '').trim();
  if (!uri.toLowerCase().startsWith('file://')) return '';
  try {
    const parsed = new URL(uri);
    if (parsed.protocol !== 'file:') return '';
    return decodeURIComponent(parsed.pathname || '').replace(/^\\/([A-Za-z]:\\/)/, '$1');
  } catch (_e) {
    return decodeURIComponent(uri.replace(/^file:\\/\\//i, ''));
  }
}
function droppedPathFromUriText(text) {
  const lines = String(text || '').split(/\\r?\\n/).map(line => line.trim()).filter(Boolean);
  for (const line of lines) {
    if (line.startsWith('#')) continue;
    const path = localPathFromFileUri(line);
    if (path) return path;
  }
  return '';
}
function droppedPathFromTransferText(dataTransfer) {
  if (!dataTransfer || !dataTransfer.getData) return '';
  for (const type of ['text/uri-list', 'text/plain']) {
    const path = droppedPathFromUriText(dataTransfer.getData(type));
    if (path) return path;
  }
  return '';
}
function setSelectedInputPath(path, kind = '') {
  selectedInputKind = kind;
  $('inputPath').value = path;
  $('dropHint').textContent = path;
  $('outputPath').value = defaultOutputStemForInput(path, kind);
  $('outputPath').dataset.userSelected = '';
  $('inputPath').dispatchEvent(new Event('change', {bubbles: true}));
}
function droppedPathInfoFromDataTransfer(dataTransfer) {
  const uriPath = droppedPathFromTransferText(dataTransfer);
  if (uriPath) return {path: uriPath, kind: 'file'};
  const items = Array.from((dataTransfer && dataTransfer.items) || []);
  for (const item of items) {
    const entry = item.webkitGetAsEntry ? item.webkitGetAsEntry() : null;
    if (entry) {
      const entryPath = entry.fullPath && entry.fullPath !== '/' ? entry.fullPath.replace(/^\\//, '') : entry.name;
      if (entryPath) return {path: entryPath, kind: entry.isDirectory ? 'directory' : 'file'};
    }
  }
  const files = Array.from((dataTransfer && dataTransfer.files) || []);
  if (!files.length) return {path: '', kind: ''};
  const file = files[0];
  return {path: file.path || file.webkitRelativePath || file.name || '', kind: 'file'};
}
function droppedPathFromDataTransfer(dataTransfer) {
  return droppedPathInfoFromDataTransfer(dataTransfer).path;
}
function applyDroppedPath(dataTransfer) {
  const info = droppedPathInfoFromDataTransfer(dataTransfer);
  if (!info.path) return false;
  setSelectedInputPath(info.path, info.kind);
  return true;
}
function dragHasFiles(event) {
  const types = Array.from((event.dataTransfer && event.dataTransfer.types) || []);
  return types.includes('Files');
}
async function chooseLocalPath(mode) {
  try {
    const data = await api('/api/pick-path', {method:'POST', body: JSON.stringify({mode})});
    if (mode === 'input_file') setSelectedInputPath(data.path, 'file');
    if (mode === 'input_directory') setSelectedInputPath(data.path, 'directory');
    if (mode === 'output_directory') {
      $('outputPath').value = joinPath(data.path, 'audit_report');
      $('outputPath').dataset.userSelected = 'true';
    }
  } catch (e) {
    setStatus(e.error || 'picker failed', 'failed');
    setFeedback(e.message || e.error || JSON.stringify(e), 'failed');
    renderReportPanel({status: 'failed', message: e.message || e.error || JSON.stringify(e)});
  }
}
function artifactLinks(run) {
  const arts = run.artifacts || {};
  return ['html','markdown','json','folder'].filter(k => arts[k]).map(k => `<a class="linkbtn" target="_blank" href="/artifact/${encodeURIComponent(run.id)}/${k}">${k}</a>`).join('');
}
function currentArtifactActions(run) {
  const links = artifactLinks(run);
  const retry = run && run.input_path && ['failed', 'canceled'].includes(run.status || '') ? '<button id="retryRunBtn" type="button">Retry</button>' : '';
  return `${links}${retry}`;
}
function reportLabel(reportType, status) {
  const value = reportType || status || '';
  const labels = {
    complete: '完整审查',
    limited: '范围受限审查',
    failed: '失败诊断',
    succeeded: '审查完成',
    running: '审查中',
    canceled: '已取消'
  };
  return labels[value] || value || '待生成';
}
function reportFolderForRun(run) {
  const arts = (run && run.artifacts) || {};
  if (arts.folder) return arts.folder;
  if (run && run.output) return pathParent(run.output);
  return '待生成';
}
function reportSummaryFallback(run) {
  const status = (run && run.status) || '';
  if (status === 'running') return '审查正在运行，日志会持续刷新。';
  if (status === 'succeeded') return '审查完成，报告入口已在下方列出。';
  if (status === 'failed') return (run && run.message) || '审查失败，请查看日志和失败诊断产物。';
  if (status === 'canceled') return (run && run.message) || '运行已取消，可在需要时重试。';
  return '报告生成后会显示摘要和打开入口。';
}
function safePanelClass(value) {
  const allowed = ['complete', 'limited', 'failed', 'succeeded', 'running', 'canceled'];
  return allowed.includes(value || '') ? value : '';
}
function renderReportPanel(run) {
  const active = run || {};
  const summary = active.summary || {};
  const reportType = summary.report_type || active.report_type || '';
  const status = active.status || '';
  const panelClass = safePanelClass(reportType) || safePanelClass(status);
  $('reportPanel').className = `report-panel ${panelClass}`;
  $('reportState').className = `status ${panelClass}`;
  $('reportState').textContent = reportLabel(reportType, status);
  $('reportType').textContent = reportLabel(reportType, status);
  $('reportRisk').textContent = summary.risk_level || (reportType || status ? '未标注' : '待生成');
  $('reportFolder').textContent = reportFolderForRun(active);
  $('reportSummary').textContent = summary.summary || reportSummaryFallback(active);
  $('currentActions').innerHTML = currentArtifactActions(active);
  const retry = $('retryRunBtn');
  if (retry) retry.addEventListener('click', () => retryRun(active));
}
function renderCurrentRun(run) {
  lastRun = run || lastRun;
  const active = run || {};
  const input = active.input_path || $('inputPath').value || 'No active run';
  const output = active.output || $('outputPath').value || '';
  $('currentRun').innerHTML = `<span>输入</span><strong>${escapeHtml(input)}</strong>`;
  $('currentOutput').innerHTML = `<span>输出</span><strong>${escapeHtml(output)}</strong>`;
  renderReportPanel(active);
  setFeedback(active.message || '', active.status || '');
}
function renderRun(run) {
  const summary = run.summary || {};
  const reportType = summary.report_type || run.report_type || run.status || '';
  const summaryLine = summary.summary ? `<div class="muted">${escapeHtml(summary.summary)}</div>` : '';
  const risk = summary.risk_level ? `<span class="status">${escapeHtml(summary.risk_level)}</span>` : '';
  return `<div class="run"><div class="run-title"><strong>${escapeHtml(reportLabel(reportType, run.status))}</strong><span class="status ${escapeHtml(run.status || '')}">${escapeHtml(run.status || '')}</span></div><div class="muted">${escapeHtml(run.input_path || '')}</div><div class="muted">${escapeHtml(run.started_at || '')}</div>${summaryLine}<div class="links">${risk}${artifactLinks(run)}</div></div>`;
}
async function startRunWithPayload(payload) {
  $('log').textContent = '';
  logOffset = 0;
  try {
    const data = await api('/api/runs', {method:'POST', body: JSON.stringify(payload)});
    activeRunId = data.run.id;
    lastRun = data.run;
    $('cancelBtn').disabled = false;
    setStatus('running');
    renderCurrentRun(data.run);
    if (!timer) timer = setInterval(pollLogs, 1200);
    pollLogs();
    refreshRuns();
  } catch (e) {
    setStatus(e.error || 'start failed', 'failed');
    setFeedback(e.message || e.error || JSON.stringify(e), 'failed');
    renderReportPanel({status: 'failed', message: e.message || e.error || JSON.stringify(e)});
  }
}
function startPayloadFromForm() {
  const payload = {input_path: $('inputPath').value, fresh: $('fresh').checked};
  if ($('outputPath').dataset.userSelected === 'true' && $('outputPath').value) {
    payload.output = $('outputPath').value;
  }
  return payload;
}
function retryRun(run) {
  if (!run || !run.input_path) return;
  if (activeRunId) return;
  $('inputPath').value = run.input_path || '';
  $('outputPath').value = run.output || '';
  $('fresh').checked = !!run.fresh;
  startRunWithPayload({input_path: run.input_path, output: run.output, fresh: !!run.fresh});
}
async function refreshRuns() {
  const data = await api('/api/runs');
  $('runs').innerHTML = (data.runs || []).map(renderRun).join('') || '<div class="muted">No runs yet</div>';
  const active = (data.runs || []).find(r => r.status === 'running');
  if (active && !activeRunId) {
    activeRunId = active.id; logOffset = 0; pollLogs();
  }
}
async function refreshConfig() {
  const data = await api('/api/config');
  const caps = data.capabilities || {};
  $('config').innerHTML = Object.keys(caps).map(k => `<div class="cell"><span>${k}</span><strong>${caps[k].ok ? 'ready' : 'needs config'}</strong><div class="muted">${(caps[k].missing || []).join(', ')}</div></div>`).join('') +
    `<div class="cell"><span>python-docx</span><strong>${data.optional_dependencies.python_docx ? 'available' : 'missing'}</strong></div><div class="cell"><span>openpyxl</span><strong>${data.optional_dependencies.openpyxl ? 'available' : 'missing'}</strong></div>`;
}
async function pollLogs() {
  if (!activeRunId) return;
  try {
    const data = await api(`/api/runs/${encodeURIComponent(activeRunId)}/logs?offset=${logOffset}`);
    logOffset = data.offset;
    if (data.lines && data.lines.length) {
      $('log').textContent += data.lines.join('\\n') + '\\n';
      $('log').scrollTop = $('log').scrollHeight;
    }
    const runData = await api(`/api/runs/${encodeURIComponent(activeRunId)}`);
    const run = runData.run;
    renderCurrentRun(run);
    setStatus(run.status || 'unknown', run.status || '');
    $('cancelBtn').disabled = run.status !== 'running';
    if (run.status !== 'running') {
      activeRunId = null;
      clearInterval(timer);
      timer = null;
      refreshRuns();
    }
  } catch (e) {
    setStatus(e.error || 'error', 'failed');
  }
}
$('startBtn').addEventListener('click', async () => {
  if (!$('outputPath').value && $('inputPath').value) {
    $('outputPath').value = defaultOutputStemForInput($('inputPath').value);
  }
  await startRunWithPayload(startPayloadFromForm());
});
$('pickFileBtn').addEventListener('click', () => chooseLocalPath('input_file'));
$('pickDirectoryBtn').addEventListener('click', () => chooseLocalPath('input_directory'));
$('pickOutputBtn').addEventListener('click', () => chooseLocalPath('output_directory'));
$('cancelBtn').addEventListener('click', async () => {
  if (!activeRunId) return;
  await api(`/api/runs/${encodeURIComponent(activeRunId)}/cancel`, {method:'POST', body:'{}'});
  pollLogs();
});
const dropZone = $('inputDropZone');
['dragenter', 'dragover'].forEach(name => dropZone.addEventListener(name, (event) => {
  if (!dragHasFiles(event)) return;
  event.preventDefault();
  event.stopPropagation();
  dropZone.classList.add('dragover');
}));
['dragleave', 'drop'].forEach(name => dropZone.addEventListener(name, (event) => {
  if (!dragHasFiles(event)) return;
  event.preventDefault();
  event.stopPropagation();
  dropZone.classList.remove('dragover');
  if (name === 'drop') applyDroppedPath(event.dataTransfer);
}));
['dragover', 'drop'].forEach(name => document.addEventListener(name, (event) => {
  if (!dragHasFiles(event)) return;
  event.preventDefault();
}));
refreshConfig();
renderCurrentRun(null);
refreshRuns();
</script>
</body>
</html>"""


def web_runner_cors_headers():
    return {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type",
    }


def serve_web_runner(host="127.0.0.1", port=8765, open_browser=True, history_path=None):
    from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

    state = WebRunnerState(history_path=history_path)

    class Handler(BaseHTTPRequestHandler):
        server_version = "VeritasWebRunner/1.0"

        def _send_bytes(self, data, content_type="application/octet-stream", status=200):
            self.send_response(status)
            self.send_header("Content-Type", content_type)
            for key, value in web_runner_cors_headers().items():
                self.send_header(key, value)
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)

        def _send_json(self, payload, status=200):
            data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
            self._send_bytes(data, "application/json; charset=utf-8", status=status)

        def _route(self):
            parsed = urllib.parse.urlparse(self.path)
            path = parsed.path.rstrip("/") or "/"
            query = urllib.parse.parse_qs(parsed.query)
            return path, query

        def do_OPTIONS(self):
            self._send_json({"ok": True})

        def do_GET(self):
            path, query = self._route()
            if path == "/":
                self._send_bytes(render_web_runner_page().encode("utf-8"), "text/html; charset=utf-8")
                return
            if path == "/health":
                self._send_json({"ok": True, "model": LLM_MODEL, "web_runner": True})
                return
            if path == "/api/config":
                self._send_json(web_runner_config_status())
                return
            if path == "/api/runs":
                self._send_json({"ok": True, "runs": state.list_runs()})
                return
            match = re.match(r"^/api/runs/([^/]+)$", path)
            if match:
                run = state.get_run(urllib.parse.unquote(match.group(1)))
                self._send_json({"ok": bool(run), "run": run} if run else {"ok": False, "error": "not_found"}, 200 if run else 404)
                return
            match = re.match(r"^/api/runs/([^/]+)/logs$", path)
            if match:
                payload = state.logs_since(urllib.parse.unquote(match.group(1)), (query.get("offset") or ["0"])[0])
                self._send_json(payload if payload else {"ok": False, "error": "not_found"}, 200 if payload else 404)
                return
            match = re.match(r"^/api/runs/([^/]+)/artifacts$", path)
            if match:
                run = state.discover_artifacts(urllib.parse.unquote(match.group(1)))
                self._send_json({"ok": bool(run), "run": run, "artifacts": (run or {}).get("artifacts", {})}, 200 if run else 404)
                return
            match = re.match(r"^/artifact/([^/]+)/([^/]+)$", path)
            if match:
                run_id = urllib.parse.unquote(match.group(1))
                kind = urllib.parse.unquote(match.group(2))
                target, error = state.artifact_target(run_id, kind)
                if error:
                    self._send_json({"ok": False, "error": error}, 404)
                    return
                if kind == "folder":
                    payload = json.dumps({"ok": True, "path": str(target)}, ensure_ascii=False, indent=2).encode("utf-8")
                    self._send_bytes(payload, "application/json; charset=utf-8")
                    return
                content_type = mimetypes.guess_type(str(target))[0] or "application/octet-stream"
                self._send_bytes(Path(target).read_bytes(), content_type)
                return
            self._send_json({"ok": False, "error": "not_found"}, 404)

        def do_POST(self):
            path, _query = self._route()
            if path in {"/generate", "/followups"}:
                try:
                    payload = _read_json_request_body(self)
                    self._send_json(_report_action_api_response(path, payload))
                except ValueError as e:
                    status = 413 if str(e) == "request_too_large" else 400
                    self._send_json({"ok": False, "error": str(e)}, status)
                except Exception as e:
                    self._send_json({"ok": False, "error": f"{type(e).__name__}: {_brief_text(str(e), 300)}"}, 500)
                return
            if path == "/api/runs":
                try:
                    payload = _read_json_request_body(self, max_bytes=20_000)
                    response, status = state.start_run(payload.get("input_path"), output=payload.get("output"), fresh=bool(payload.get("fresh")))
                    self._send_json(response, status)
                except Exception as e:
                    self._send_json({"ok": False, "error": f"{type(e).__name__}: {_brief_text(str(e), 300)}"}, 500)
                return
            if path == "/api/pick-path":
                try:
                    payload = _read_json_request_body(self, max_bytes=20_000)
                    response = pick_local_path(payload.get("mode"))
                    self._send_json(response, 200 if response.get("ok") else 400)
                except Exception as e:
                    self._send_json({"ok": False, "error": f"{type(e).__name__}: {_brief_text(str(e), 300)}"}, 500)
                return
            match = re.match(r"^/api/runs/([^/]+)/cancel$", path)
            if match:
                response, status = state.cancel_run(urllib.parse.unquote(match.group(1)))
                self._send_json(response, status)
                return
            self._send_json({"ok": False, "error": "not_found"}, 404)

        def log_message(self, fmt, *args):
            print(f"[web-runner] {self.address_string()} {fmt % args}")

    try:
        httpd = ThreadingHTTPServer((host, int(port)), Handler)
    except OSError as e:
        print(f"无法启动Web Runner: {e}")
        print(f"端口 {port} 可能已被占用；请改用 --web-port <port>。")
        return 1
    url = f"http://{host}:{int(port)}"
    print(f"Veritas Web Runner 已启动: {url}")
    print("仅监听 127.0.0.1。按 Ctrl+C 停止。")
    if open_browser:
        try:
            webbrowser.open(url)
        except Exception as e:
            print(f"自动打开浏览器失败: {e}")
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        print("\nWeb Runner 已停止")
    finally:
        httpd.server_close()
    return 0


# ══════════════════════════════════════════════════════════════
# 报告解析与格式化
# ══════════════════════════════════════════════════════════════

def split_references_from_text(text):
    """Remove reference sections from main audit text and return parsed tail text."""
    text = str(text or "")
    pattern = re.compile(
        r"(?im)^(?:\[\[BLOCK[^\]]*\]\]\s*)?(?:#+\s*)?(?:references?|bibliography|参考文献|参考资料|works cited)\s*(?:\[\[/BLOCK\]\]\s*)?$"
    )
    matches = list(pattern.finditer(text))
    if not matches:
        return text, ""
    start = matches[-1].start()
    main_text = text[:start].rstrip()
    references_text = text[start:].strip()
    return main_text, references_text


def audit_references(references_text, online=False, online_limit=50, timeout=10, cache=None):
    """Reference plausibility check with optional online scholarly database verification."""
    refs = parse_references(references_text)
    effective_online_limit = _effective_limit(online_limit, len(refs))
    online_checked = 0
    online_cache = cache if isinstance(cache, dict) else {}

    def base_issues_for(ref):
        ref_issues = []
        year = ref.get("year")
        if not year:
            ref_issues.append("missing_year")
        else:
            try:
                if int(year) > runtime_utc_year():
                    ref_issues.append("future_year")
            except (TypeError, ValueError):
                pass
        if not ref.get("doi"):
            ref_issues.append("missing_doi")
        if not ref.get("has_journal_hint"):
            ref_issues.append("missing_journal_or_source")
        if len(ref.get("text", "")) < 25:
            ref_issues.append("too_short")
        return ref_issues

    if online:
        fetch_jobs = []
        for idx, ref in enumerate(refs, 1):
            if idx <= effective_online_limit:
                cache_key = reference_cache_key(ref)
                online_result = online_cache.get(cache_key)
                if online_result:
                    ref["online"] = online_result
                else:
                    fetch_jobs.append((idx, ref, cache_key))
            else:
                ref["online"] = {
                    "online_status": "skipped",
                    "confidence": 0.0,
                    "problems": ["online_limit_reached"],
                    "matched_sources": [],
                    "query": build_reference_query(ref),
                }

        if fetch_jobs:
            workers = min(4, len(fetch_jobs))
            with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
                future_map = {
                    executor.submit(verify_reference_online, ref, timeout=timeout): (idx, ref, cache_key)
                    for idx, ref, cache_key in fetch_jobs
                }
                for future in concurrent.futures.as_completed(future_map):
                    idx, ref, cache_key = future_map[future]
                    try:
                        online_result = future.result()
                    except Exception as e:
                        online_result = {
                            "online_status": "error",
                            "confidence": 0.0,
                            "query": build_reference_query(ref),
                            "matched_sources": [],
                            "problems": ["all_sources_error"],
                            "source_errors": [f"verify_reference_online: {type(e).__name__}"],
                            "error_message": _brief_text(str(e), 240),
                        }
                    online_cache[cache_key] = online_result
                    ref["online"] = online_result

        online_checked = sum(
            1 for idx, ref in enumerate(refs, 1)
            if idx <= effective_online_limit and (ref.get("online") or {}).get("online_status") != "skipped"
        )

    issues = []
    for idx, ref in enumerate(refs, 1):
        ref_issues = base_issues_for(ref)

        if online and idx <= effective_online_limit:
            cache_key = reference_cache_key(ref)
            online_result = ref.get("online") or online_cache.get(cache_key) or {}
            ref["online"] = online_result
            online_status = online_result.get("online_status")
            if online_status in {"not_found", "weak", "error"}:
                ref_issues.append(f"online_{online_status}")
            ref_issues.extend(online_result.get("problems") or [])

        if ref_issues:
            issues.append({"index": idx, "issues": ref_issues, "text": ref.get("text", "")})
    status = "ok"
    if issues:
        status = "needs_review" if len(issues) < max(3, len(refs) // 3) else "weak"
    if online and refs:
        hard_online_issues = [
            item for item in issues
            if any(str(issue).startswith("online_") or issue in {"doi_not_found", "no_online_match"} for issue in item.get("issues", []))
        ]
        if hard_online_issues:
            status = "online_needs_review" if len(hard_online_issues) < max(3, len(refs) // 3) else "online_weak"
        elif effective_online_limit >= len(refs):
            online_statuses = [
                (ref.get("online") or {}).get("online_status")
                for ref in refs
            ]
            if online_statuses and all(item == "verified" for item in online_statuses):
                status = "ok"
            elif online_statuses and all(item in {"verified", "likely"} for item in online_statuses):
                status = "needs_review"
    return {
        "status": status,
        "reference_count": len(refs),
        "doi_count": sum(1 for r in refs if r.get("doi")),
        "year_count": sum(1 for r in refs if r.get("year")),
        "online_enabled": bool(online),
        "online_checked": online_checked,
        "issues": issues,
        "references": refs[:200],
        "note": (
            "在线真实性校检：优先用DOI精确检索，再用题名/年份在Crossref、OpenAlex和PubMed进行多源核验；"
            "结果为尽力检索证据，不等同于绝对证明。"
            if online else
            "离线格式/可核验性校检：检查DOI、年份、来源字段等基本信息；不代表已联网验证引用真实存在。"
        ),
    }


def parse_references(references_text):
    text = _clean_reference_text(references_text)
    text = _truncate_reference_suffix(text)
    text = re.sub(r"(?im)^(?:#+\s*)?(?:references?|bibliography|参考文献|参考资料|works cited)\s*$", "", text).strip()
    if not text:
        return []
    raw_items = _reference_items_from_numbered_lines(text)
    if len(raw_items) <= 1:
        raw_items = re.split(r"\n{2,}", text)
    refs = []
    for item in raw_items:
        item = re.sub(r"\s+", " ", item).strip()
        item = re.sub(r"^(?:\[\d+\]|\d+\.|\(\d+\))\s*", "", item)
        item = re.sub(r"^\d+\.\s*", "", item)
        if len(item) < 8:
            continue
        if _looks_like_reference_table_noise(item):
            continue
        doi_match = re.search(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", item, re.IGNORECASE)
        year = extract_reference_year_hint(item)
        has_journal_hint = bool(re.search(r"\b(?:journal|j\.|proc\.|nature|science|cell|ieee|acm|springer|elsevier|frontiers|plos|bmc|lancet)\b", item, re.IGNORECASE))
        doi = _normalize_doi(doi_match.group(0)) if doi_match else ""
        refs.append({
            "text": item,
            "doi": doi,
            "year": year,
            "has_journal_hint": has_journal_hint,
            "title_hint": extract_reference_title(item),
            "author_hint": extract_reference_author_hint(item),
            "container_hint": extract_reference_container_hint(item),
        })
    return refs


def _truncate_reference_suffix(text):
    """Drop non-reference sections accidentally captured after References."""
    text = str(text or "")
    suffix_heading = re.search(
        r"(?im)^\s*(?:figure\s+legends?|figures?|tables?|supplementary\s+(?:material|information)|acknowledg(?:e)?ments?)\s*$",
        text,
    )
    if suffix_heading:
        return text[:suffix_heading.start()].rstrip()
    return text


def _reference_items_from_numbered_lines(text):
    """Build reference items while tolerating MinerU per-page list numbering.

    MinerU can emit each page's reference_list with local numbering that restarts
    at 1, while the actual global reference number remains in the text, e.g.
    "2. 15.Liu...". Continuation lines can therefore look like "1. black-box...".
    Treat a numbered line as a new reference only when its visible/global number
    matches the next expected reference number; otherwise append it to the current
    item.
    """
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    items = []
    current = []
    expected_number = 1

    def flush():
        nonlocal current
        if current:
            items.append(" ".join(current).strip())
            current = []

    for line in lines:
        if re.fullmatch(r"(?i)(?:article\s+in\s+press|references?|bibliography|参考文献|参考资料|works cited)", line):
            continue
        bracketed = re.match(r"^\[(\d+)\]\s*(.+)$", line)
        if bracketed:
            visible_number = int(bracketed.group(1))
            body = bracketed.group(2).strip()
        else:
            numbered = re.match(r"^(\d+)\.\s*(.+)$", line)
            if not numbered:
                if current:
                    current.append(line)
                continue
            local_number = int(numbered.group(1))
            rest = numbered.group(2).strip()
            nested = re.match(r"^(\d+)\.?\s*(.+)$", rest)
            if nested:
                visible_number = int(nested.group(1))
                body = nested.group(2).strip()
            else:
                visible_number = local_number
                body = rest

        if not current or visible_number == expected_number:
            flush()
            current = [body]
            expected_number = visible_number + 1
        else:
            current.append(body)
    flush()
    return items


def _looks_like_reference_table_noise(item):
    """Avoid treating extracted tables as reference entries."""
    decoded = html.unescape(str(item or "")).strip()
    lowered = decoded.lower()
    if "[[table_start" in lowered or "[[table_continuation" in lowered:
        return True
    td_count = len(re.findall(r"</?t[dh]\b", lowered))
    tr_count = len(re.findall(r"</?tr\b", lowered))
    if "<table" in lowered and (td_count >= 4 or tr_count >= 2):
        return True
    pipe_cells = sum(1 for line in decoded.splitlines() if line.count("|") >= 3)
    if pipe_cells >= 2 and not re.search(r"\b(?:doi|pmid|arxiv|journal|volume|issue)\b", decoded, re.I):
        return True
    return False


def _clean_reference_text(text):
    text = str(text or "")
    text = re.sub(r"\[\[EXTRACTION_NOTE\]\].*?\[\[/EXTRACTION_NOTE\]\]", "\n", text, flags=re.S)
    text = re.sub(r"\[\[/?(?:BLOCK|FIGURE)[^\]]*\]\]", "\n", text, flags=re.I)
    text = re.sub(r"\[\[TABLE_START[^\]]*\]\]|\[\[TABLE_END\]\]|\[\[TABLE_CONTINUATION[^\]]*\]\]", "\n", text)
    text = re.sub(r"(?m)^===\s*文件:.*?===\s*$", "\n", text)
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _normalize_doi(value):
    value = html.unescape(str(value or "")).strip()
    value = re.sub(r"(?i)^https?://(?:dx\.)?doi\.org/", "", value)
    value = re.sub(r"(?i)^doi\s*[:：]\s*", "", value)
    value = value.strip().rstrip(".,;)]}")
    return value.lower()


REFERENCE_CONTAINER_WORD_RE = re.compile(
    r"\b(?:journal|jclin|proc\.?|proceedings|nature|science|cell|frontiers|plos|bmc|"
    r"lancet|thyroid|oncology|endocrinology|communications?|commun|annals|cancers|"
    r"cancer\s+letters?|cancer\s+lett|mol\s+cancer|jama|esmo)\b",
    re.I,
)


def _looks_like_reference_container_part(part):
    part = str(part or "").strip()
    if not REFERENCE_CONTAINER_WORD_RE.search(part):
        return False
    normalized = _normalize_title(part)
    known_short = {
        "ca cancer jclin",
        "nat commun",
        "mol cancer",
        "cancer lett",
        "cancer letters",
        "jama",
        "the lancet",
        "thyroid",
        "cancers basel",
        "esmo open",
    }
    if normalized in known_short:
        return True
    if len(_title_tokens(part)) > 6:
        return False
    return bool(re.search(
        r"\b(?:vol\.?|volume)\b|\b\d+\s*,|\b\d+\s+\d+|\(\d{4}\)|\b\d{1,5}\s*[-–]\s*\d{1,5}\b",
        part,
        re.I,
    ))


def _looks_like_reference_author_fragment(part):
    part = str(part or "").strip()
    if not part:
        return False
    if re.fullmatch(r"(?:[A-Z]\.?){1,4}", part):
        return True
    if re.search(r"\bet\s+al\b", part, re.I):
        return True
    if re.search(r"(?:^|[\s,&])(?:[A-Z]\.){1,3}(?:,|\s|$)", part):
        return True
    if re.search(r"\b[A-Z][A-Za-z'’-]+,\s*[A-Z](?:\.|$)", part):
        return True
    return False


def _name_tokens(value):
    return {
        token
        for token in re.findall(r"[a-z\u4e00-\u9fff]+", _normalize_title(value))
        if len(token) >= 3
    }


def _author_similarity(query_author, match_authors):
    left = _name_tokens(query_author)
    if not left:
        return 0.0
    right = set()
    for author in match_authors or []:
        right.update(_name_tokens(author))
    if not right:
        return 0.0
    return len(left & right) / max(len(left), 1)


def _reference_year(ref):
    if isinstance(ref, dict):
        value = ref.get("year") or ref.get("publication_year") or ""
    else:
        value = str(ref or "")
    match = re.search(r"\b(19|20)\d{2}\b", str(value))
    return match.group(0) if match else ""


def extract_reference_year_hint(text):
    text = str(text or "")
    parenthetical = re.findall(r"\(((?:19|20)\d{2})\)", text)
    if parenthetical:
        return parenthetical[-1]
    years = re.findall(r"\b((?:19|20)\d{2})\b", text)
    return years[-1] if years else ""


def extract_reference_author_hint(text):
    text = re.sub(r"^(?:\[\d+\]|\d+\.|\(\d+\))\s*", "", str(text or "")).strip()
    before_year = re.split(r"\b(?:19|20)\d{2}\b", text, maxsplit=1)[0]
    before_title = before_year.split(".")[0] if "." in before_year else before_year
    names = re.findall(r"\b[A-Z][A-Za-z'’-]{2,}\b", before_title)
    return " ".join(names[:3])


def extract_reference_container_hint(text):
    text = re.sub(r"^(?:\[\d+\]|\d+\.|\(\d+\))\s*", "", str(text or "")).strip()
    text = re.sub(r"https?://\S+", "", text, flags=re.I)
    parts = [p.strip(" .;:") for p in re.split(r"\.\s+", text) if p.strip(" .;:")]
    for part in parts:
        if _looks_like_reference_author_fragment(part):
            continue
        if len(_title_tokens(part)) > 8:
            continue
        if _looks_like_reference_container_part(part):
            container = re.split(r"\b(?:vol\.?|volume|\d+\s*,|\d+\s+\d|\(\d{4}\))", part, maxsplit=1, flags=re.I)[0]
            return container.strip(" .;:")[:160]
    return ""


def extract_reference_title(text):
    text = re.sub(r"^(?:\[\d+\]|\d+\.|\(\d+\))\s*", "", str(text or "")).strip()
    text = re.sub(r"\bdoi\s*[:：]?\s*10\.\S+", "", text, flags=re.I).strip()
    text = re.sub(r"https?://\S+", "", text, flags=re.I).strip()
    parts = [p.strip(" .;:") for p in re.split(r"\.\s+", text) if p.strip(" .;:")]
    if not parts:
        return _brief_text(text, 160)
    candidates = []
    for part in parts:
        if _looks_like_reference_author_fragment(part):
            continue
        if _looks_like_reference_container_part(part):
            continue
        if len(_title_tokens(part)) >= 2:
            candidates.append(part)
    if candidates:
        return candidates[0][:360]
    return parts[min(1, len(parts) - 1)][:360]


def build_reference_query(ref):
    title = ref.get("title_hint") or extract_reference_title(ref.get("text", ""))
    author = ref.get("author_hint") or extract_reference_author_hint(ref.get("text", ""))
    year = _reference_year(ref)
    doi = _normalize_doi(ref.get("doi", ""))
    container = ref.get("container_hint") or extract_reference_container_hint(ref.get("text", ""))
    query_parts = [p for p in (title, author, year, container) if p]
    bibliographic = " ".join(query_parts) or ref.get("text", "")[:240]
    if not doi and ref.get("text"):
        bibliographic = ref.get("text", "")[:600]
    return {
        "doi": doi,
        "title": title,
        "author": author,
        "container": container,
        "year": year,
        "bibliographic": bibliographic,
    }


def reference_cache_key(ref):
    query = build_reference_query(ref)
    if query.get("doi"):
        return "doi:" + query["doi"]
    key = f"{_normalize_title(query.get('title'))}|{query.get('year', '')}"
    return "title:" + hashlib.sha256(key.encode("utf-8", errors="ignore")).hexdigest()[:24]


def _reference_get_json(url, timeout=10, headers=None):
    # Reference verification already fans out across Crossref, OpenAlex, PubMed,
    # DOI landing pages, and official publisher sites. Keep each individual
    # source fast-fail so a full bibliography cannot hang on one slow provider.
    data, _ = _http_request(url, "GET", headers=headers or {}, timeout=timeout)
    return json.loads(data.decode("utf-8", errors="replace"))


def _crossref_work_to_match(work):
    title = " ".join(work.get("title") or []).strip()
    year = ""
    for key in ("published-print", "published-online", "published", "created"):
        parts = ((work.get(key) or {}).get("date-parts") or [[]])[0]
        if parts:
            year = str(parts[0])
            break
    authors = []
    for author in work.get("author") or []:
        name = " ".join(p for p in (author.get("given"), author.get("family")) if p)
        if name:
            authors.append(name)
    container = " ".join(work.get("container-title") or work.get("short-container-title") or [])
    return {
        "source": "Crossref",
        "title": title,
        "year": year,
        "doi": _normalize_doi(work.get("DOI", "")),
        "authors": authors[:5],
        "container": container,
        "url": work.get("URL", ""),
        "retracted": bool(work.get("relation", {}).get("is-retracted-by")),
    }


def _openalex_work_to_match(work):
    title = work.get("display_name") or work.get("title") or ""
    authors = []
    for authorship in work.get("authorships") or []:
        name = ((authorship.get("author") or {}).get("display_name") or "").strip()
        if name:
            authors.append(name)
    location = work.get("primary_location") or {}
    source = location.get("source") or {}
    return {
        "source": "OpenAlex",
        "title": title,
        "year": str(work.get("publication_year") or ""),
        "doi": _normalize_doi(work.get("doi", "")),
        "authors": authors[:5],
        "container": source.get("display_name", ""),
        "url": work.get("doi") or work.get("id") or "",
        "retracted": bool(work.get("is_retracted")),
    }


def _pubmed_summary_to_match(uid, item):
    title = item.get("title") or ""
    authors = []
    for author in item.get("authors") or []:
        name = (author.get("name") or "").strip()
        if name:
            authors.append(name)
    pubdate = item.get("pubdate") or ""
    year = _reference_year(pubdate)
    doi = ""
    for article_id in item.get("articleids") or []:
        if str(article_id.get("idtype", "")).lower() == "doi":
            doi = _normalize_doi(article_id.get("value", ""))
            break
    return {
        "source": "PubMed",
        "title": title,
        "year": year,
        "doi": doi,
        "authors": authors[:5],
        "container": item.get("fulljournalname") or item.get("source") or "",
        "url": f"https://pubmed.ncbi.nlm.nih.gov/{uid}/",
        "retracted": "retracted publication" in " ".join(item.get("pubtype") or []).lower(),
    }


def lookup_crossref_reference(ref, timeout=10):
    query = build_reference_query(ref)
    matches = []
    last_error = None
    if query.get("doi"):
        try:
            url = "https://api.crossref.org/works/" + urllib.parse.quote(query["doi"], safe="")
            data = _reference_get_json(url, timeout=timeout)
            work = data.get("message") or {}
            if work:
                matches.append(_crossref_work_to_match(work))
                return matches
        except Exception as e:
            last_error = e
    if query.get("title"):
        try:
            title = urllib.parse.quote(query["title"])
            url = f"https://api.crossref.org/works?query.title={title}&rows=5"
            data = _reference_get_json(url, timeout=timeout)
            for work in (data.get("message") or {}).get("items") or []:
                matches.append(_crossref_work_to_match(work))
        except Exception as e:
            last_error = e
    bibliographic = urllib.parse.quote(query.get("bibliographic") or "")
    if not bibliographic:
        return matches
    try:
        url = f"https://api.crossref.org/works?query.bibliographic={bibliographic}&rows=5"
        data = _reference_get_json(url, timeout=timeout)
        for work in (data.get("message") or {}).get("items") or []:
            matches.append(_crossref_work_to_match(work))
    except Exception as e:
        last_error = e
    if last_error and not matches:
        raise last_error
    return matches


def lookup_openalex_reference(ref, timeout=10):
    query = build_reference_query(ref)
    matches = []
    last_error = None
    if query.get("doi"):
        try:
            url = "https://api.openalex.org/works/doi:" + urllib.parse.quote(query["doi"], safe="")
            data = _reference_get_json(url, timeout=timeout)
            if data:
                matches.append(_openalex_work_to_match(data))
                return matches
        except Exception as e:
            last_error = e
    if query.get("title"):
        try:
            title = urllib.parse.quote(query["title"])
            url = f"https://api.openalex.org/works?filter=title.search:{title}&per-page=5"
            data = _reference_get_json(url, timeout=timeout)
            for work in (data.get("results") or []):
                matches.append(_openalex_work_to_match(work))
        except Exception as e:
            last_error = e
    search = urllib.parse.quote(query.get("bibliographic") or "")
    if not search:
        return matches
    try:
        url = f"https://api.openalex.org/works?search={search}&per-page=5"
        data = _reference_get_json(url, timeout=timeout)
        for work in (data.get("results") or []):
            matches.append(_openalex_work_to_match(work))
    except Exception as e:
        last_error = e
    if last_error and not matches:
        raise last_error
    return matches


def lookup_pubmed_reference(ref, timeout=10):
    query = build_reference_query(ref)
    term = query.get("doi") or query.get("bibliographic") or ""
    if not term:
        return []
    search_url = (
        "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
        f"?db=pubmed&retmode=json&retmax=5&term={urllib.parse.quote(term)}"
    )
    search = _reference_get_json(search_url, timeout=timeout)
    ids = ((search.get("esearchresult") or {}).get("idlist") or [])[:5]
    if not ids:
        return []
    summary_url = (
        "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"
        f"?db=pubmed&retmode=json&id={','.join(ids)}"
    )
    summary = _reference_get_json(summary_url, timeout=timeout)
    result = summary.get("result") or {}
    return [_pubmed_summary_to_match(uid, result.get(uid) or {}) for uid in ids if result.get(uid)]


REFERENCE_OFFICIAL_SITE_RULES = [
    (("ca cancer", "international journal of cancer"), "Wiley Online Library", "https://onlinelibrary.wiley.com/action/doSearch?AllField={query}"),
    (("thyroid",), "Mary Ann Liebert", "https://www.liebertpub.com/action/doSearch?AllField={query}"),
    (("nature reviews", "nat commun", "nature communications"), "Nature", "https://www.nature.com/search?q={query}"),
    (("current opinion in oncology", "lww",), "LWW Journals", "https://journals.lww.com/pages/results.aspx?txtKeywords={query}"),
    (("journal of clinical endocrinology", "endocrinology and metabolism"), "Oxford Academic", "https://academic.oup.com/search-results?page=1&q={query}"),
    (("annals of oncology", "esmo open"), "Elsevier ClinicalKey", "https://www.annalsofoncology.org/action/doSearch?AllField={query}"),
    (("lancet",), "The Lancet", "https://www.thelancet.com/action/doSearch?AllField={query}"),
    (("cancers basel", "mdpi",), "MDPI", "https://www.mdpi.com/search?q={query}"),
    (("proceedings of the national academy", "pnas"), "PNAS", "https://www.pnas.org/action/doSearch?AllField={query}"),
    (("mol cancer", "molecular cancer"), "BMC Molecular Cancer", "https://molecular-cancer.biomedcentral.com/search?query={query}"),
    (("jama",), "JAMA Network", "https://jamanetwork.com/searchresults?q={query}"),
    (("endocrine", "hashimoto", "papillary thyroid carcinoma"), "Springer Link", "https://link.springer.com/search?query={query}"),
    (("kolmogorov arnold networks", "arxiv"), "arXiv", "https://arxiv.org/search/?query={query}&searchtype=all&source=header"),
]


def _html_to_searchable_text(content):
    raw = content.decode("utf-8", errors="replace") if isinstance(content, (bytes, bytearray)) else str(content or "")
    raw = re.sub(r"(?is)<(script|style|noscript).*?</\1>", " ", raw)
    raw = re.sub(r"(?is)<[^>]+>", " ", raw)
    return re.sub(r"\s+", " ", html.unescape(raw)).strip()


def _html_title(content):
    raw = content.decode("utf-8", errors="replace") if isinstance(content, (bytes, bytearray)) else str(content or "")
    match = re.search(r"(?is)<title[^>]*>(.*?)</title>", raw)
    if not match:
        return ""
    return re.sub(r"\s+", " ", html.unescape(re.sub(r"<[^>]+>", " ", match.group(1)))).strip()


def _official_page_matches_reference(ref, page_text):
    query = build_reference_query(ref)
    title_tokens = _title_tokens(query.get("title") or "")
    page_tokens = _title_tokens(page_text)
    if not title_tokens or not page_tokens:
        return False
    coverage = len(title_tokens & page_tokens) / max(len(title_tokens), 1)
    year = query.get("year")
    if not year:
        return coverage >= 0.82
    years = {_reference_year(token) for token in re.findall(r"\b(?:19|20)\d{2}\b", page_text)}
    years.discard("")
    year_ok = year in years or any(abs(int(year) - int(item)) <= 1 for item in years)
    return (coverage >= 0.72 and year_ok) or coverage >= 0.9


def _official_site_search_urls(ref):
    query = build_reference_query(ref)
    probe = _normalize_title(" ".join([
        query.get("container", ""),
        query.get("title", ""),
        ref.get("text", ""),
    ]))
    title = query.get("title") or ""
    words = re.findall(r"[A-Za-z0-9][A-Za-z0-9'’-]*", title)
    search_terms = [title or query.get("bibliographic") or ref.get("text", "")[:180]]
    if len(words) >= 7:
        # OCR can damage the first word of a title; retry with the distinctive
        # title tail before declaring that the publisher site cannot find it.
        search_terms.append(" ".join(words[1:]))
    if len(words) >= 11:
        distinctive = [word for word in words if len(word) >= 4][:12]
        if len(distinctive) >= 4:
            search_terms.append(" ".join(distinctive))
    seen = set()
    urls = []
    for needles, label, template in REFERENCE_OFFICIAL_SITE_RULES:
        if any(needle in probe for needle in needles):
            for term in search_terms:
                search = urllib.parse.quote(term)
                url = template.format(query=search)
                if url not in seen:
                    urls.append((label, url))
                    seen.add(url)
    return urls


def lookup_official_site_reference(ref, timeout=10):
    """Verify references from DOI landing pages and publisher/official site searches."""
    query = build_reference_query(ref)
    matches = []
    if query.get("doi"):
        doi_url = "https://doi.org/" + urllib.parse.quote(query["doi"], safe="/")
        data, _ = _http_request(doi_url, "GET", headers={"Accept": "text/html,*/*;q=0.8"}, timeout=timeout)
        page_text = _html_to_searchable_text(data)
        title = _html_title(data) or query.get("title") or query.get("doi")
        if query.get("doi") or _official_page_matches_reference(ref, page_text):
            matches.append({
                "source": "DOI landing page",
                "title": query.get("title") or title,
                "year": query.get("year"),
                "doi": query.get("doi"),
                "authors": [query.get("author")] if query.get("author") else [],
                "container": query.get("container"),
                "url": doi_url,
                "retracted": False,
                "official_site": True,
            })

    for label, url in _official_site_search_urls(ref):
        data, _ = _http_request(url, "GET", headers={"Accept": "text/html,*/*;q=0.8"}, timeout=timeout)
        page_text = _html_to_searchable_text(data)
        if not _official_page_matches_reference(ref, page_text):
            continue
        matches.append({
            "source": f"Official site: {label}",
            "title": query.get("title") or _html_title(data),
            "year": query.get("year"),
            "doi": query.get("doi"),
            "authors": [query.get("author")] if query.get("author") else [],
            "container": query.get("container") or label,
            "url": url,
            "retracted": False,
            "official_site": True,
        })
    return matches


def _score_reference_match(ref, match):
    query = build_reference_query(ref)
    problems = []
    score = 0.0
    ref_doi = query.get("doi")
    match_doi = _normalize_doi(match.get("doi", ""))
    title_sim = max(
        _token_similarity(query.get("title") or "", match.get("title", "")),
        _token_similarity(ref.get("text", ""), match.get("title", "")),
    )
    author_sim = _author_similarity(query.get("author"), match.get("authors") or [])
    container_sim = _token_similarity(query.get("container") or "", match.get("container", ""))
    if ref_doi:
        if match_doi and ref_doi == match_doi:
            score += 0.72
        elif match_doi:
            problems.append("doi_mismatch")
            score -= 0.2
        else:
            problems.append("doi_missing_in_source")
        score += min(title_sim, 1.0) * 0.18
        score += min(author_sim, 1.0) * 0.04
    else:
        score += min(title_sim, 1.0) * 0.62
        score += min(author_sim, 1.0) * 0.14
        score += min(container_sim, 1.0) * 0.08
    if title_sim < 0.45 and not ref_doi:
        problems.append("title_low_similarity")
    ref_year = query.get("year")
    match_year = _reference_year(match.get("year", ""))
    if ref_year and match_year:
        if ref_year == match_year:
            score += 0.06 if ref_doi else 0.16
        elif abs(int(ref_year) - int(match_year)) <= 1:
            score += 0.03 if ref_doi else 0.08
            problems.append("year_near_mismatch")
        else:
            problems.append("year_mismatch")
            score -= 0.1
    elif ref_year and not match_year:
        problems.append("year_missing_in_source")
    if (
        not ref_doi
        and match_doi
        and ref_year
        and match_year
        and (ref_year == match_year or abs(int(ref_year) - int(match_year)) <= 1)
        and title_sim >= 0.78
    ):
        score = max(score, 0.93)
    if match.get("source") == "DOI landing page" and ref_doi and ref_doi == match_doi:
        score = max(score, 0.95)
    if match.get("official_site") and title_sim >= 0.82 and (not ref_year or not match_year or abs(int(ref_year) - int(match_year)) <= 1):
        score = max(score, 0.94)
    if match.get("retracted"):
        problems.append("source_marks_retracted")
    return max(0.0, min(1.0, score)), problems


def _score_reference_matches(ref, raw_matches):
    scored = []
    for match in raw_matches:
        score, problems = _score_reference_match(ref, match)
        enriched = dict(match)
        enriched["match_score"] = round(score, 3)
        enriched["_match_problems"] = problems
        scored.append(enriched)
    scored.sort(key=lambda m: m.get("match_score", 0), reverse=True)
    return scored


def verify_reference_online(ref, timeout=10):
    query = build_reference_query(ref)
    source_errors = []
    raw_matches = []
    standard_sources_ok = 0
    for lookup in (lookup_crossref_reference, lookup_openalex_reference, lookup_pubmed_reference):
        try:
            raw_matches.extend(lookup(ref, timeout=timeout))
            standard_sources_ok += 1
        except Exception as e:
            status = getattr(getattr(e, "response", None), "status_code", None)
            if status not in {404, 410}:
                source_errors.append(f"{lookup.__name__.replace('lookup_', '').replace('_reference', '')}: {type(e).__name__}")

    scored = _score_reference_matches(ref, raw_matches)
    if standard_sources_ok and (not scored or scored[0].get("match_score", 0) < 0.92):
        try:
            official_matches = lookup_official_site_reference(ref, timeout=timeout)
            if official_matches:
                raw_matches.extend(official_matches)
                scored = _score_reference_matches(ref, raw_matches)
        except Exception as e:
            source_errors.append(f"official_site: {type(e).__name__}")

    best = scored[0] if scored else None
    confidence = float(best.get("match_score", 0.0)) if best else 0.0
    problems = []
    if not scored:
        problems.append("doi_not_found" if query.get("doi") else "no_online_match")
    if source_errors and not scored and not standard_sources_ok:
        problems.append("all_sources_error")
    elif source_errors:
        problems.append("partial_source_error")
    if best:
        problems.extend(best.get("_match_problems") or [])
    problems = list(dict.fromkeys(problems[:8]))

    if confidence >= 0.92:
        status = "verified"
    elif confidence >= 0.68:
        status = "likely"
    elif confidence >= 0.38:
        status = "weak"
    elif source_errors and not scored:
        status = "error"
    else:
        status = "not_found"

    return {
        "online_status": status,
        "confidence": round(confidence, 3),
        "query": query,
        "matched_sources": [{k: v for k, v in match.items() if k != "_match_problems"} for match in scored[:5]],
        "problems": problems,
        "source_errors": source_errors,
    }


def _reference_online_summary(online):
    if not online:
        return "未联网检索"
    status = online.get("online_status", "N/A")
    confidence = online.get("confidence", 0)
    sources = []
    for match in online.get("matched_sources") or []:
        label = match.get("source", "")
        if label and label not in sources:
            sources.append(label)
    source_text = "/".join(sources) if sources else "无命中"
    return f"{status} / {confidence} / {source_text}"


REFERENCE_ISSUE_LABELS = {
    "missing_year": "缺少年份",
    "future_year": "年份晚于运行时当前年份",
    "missing_doi": "缺少DOI",
    "missing_journal_or_source": "缺少期刊/来源",
    "too_short": "引用过短",
    "online_not_found": "在线未检索到",
    "online_weak": "在线弱匹配",
    "online_error": "在线检索异常",
    "doi_not_found": "DOI未命中",
    "no_online_match": "无在线命中",
    "all_sources_error": "外部源均异常",
    "partial_source_error": "部分外部源异常",
    "doi_mismatch": "DOI不一致",
    "title_low_similarity": "题名相似度低",
    "year_mismatch": "年份不一致",
    "year_near_mismatch": "年份接近但不一致",
    "source_marks_retracted": "来源标记撤稿",
}


def _reference_issue_text(issues):
    labels = [REFERENCE_ISSUE_LABELS.get(str(issue), str(issue)) for issue in (issues or [])]
    return ", ".join(dict.fromkeys(labels))


def _reference_display_title(ref):
    raw = ref.get("title_hint") or ref.get("text", "")
    if _looks_like_reference_table_noise(raw):
        return "疑似表格片段，未作为有效参考文献"
    return _brief_text(raw, 140)


def _reference_text_html(text):
    if _looks_like_reference_table_noise(text):
        table_html = _escaped_html_table_fragment_to_html(text) or render_evidence_summary_html(text)
        return (
            '<div class="reference-table-note">该条内容看起来是表格提取噪声，不应作为参考文献核验对象。</div>'
            f"{table_html}"
        )
    return _html_escape(text)


def _reference_query_text(ref, query):
    raw = ref.get("text", "")
    if _looks_like_reference_table_noise(raw):
        return "疑似表格提取噪声，未生成参考文献检索式"
    return query.get("doi") or query.get("bibliographic") or ""


def format_reference_audit_markdown(reference_audit):
    if reference_audit is None:
        return []
    lines = [
        '<a id="reference-audit"></a>',
        "## 📚 参考文献真实性/可核验性校检",
        "",
        f"**状态**: {reference_audit.get('status', 'N/A')}",
        f"**参考文献数量**: {reference_audit.get('reference_count', 0)}",
        f"**含 DOI 数量**: {reference_audit.get('doi_count', 0)}",
        f"**含年份数量**: {reference_audit.get('year_count', 0)}",
        f"**在线检索**: {'启用' if reference_audit.get('online_enabled') else '未启用'}"
        + (f"（已检索 {reference_audit.get('online_checked', 0)} 条）" if reference_audit.get("online_enabled") else ""),
        f"> {reference_audit.get('note', '')}",
        "",
    ]
    issues = reference_audit.get("issues", [])
    if issues:
        lines.append("| # | 问题 | 在线证据 | 引用摘录 |")
        lines.append("|---|------|----------|----------|")
        refs_by_index = {i + 1: ref for i, ref in enumerate(reference_audit.get("references", []))}
        for item in issues[:30]:
            ref = refs_by_index.get(item.get("index"), {})
            lines.append(
                f"| {item.get('index')} | {_md_escape_cell(_reference_issue_text(item.get('issues', [])))} | "
                f"{_md_escape_cell(_reference_online_summary(ref.get('online')))} | "
                f"{_md_escape_cell(_brief_text(item.get('text', ''), 220))} |"
            )
    else:
        lines.append("> 未发现明显格式缺失；仍建议对关键引用进行数据库/DOI人工核验。")
    lines.append("")
    return lines


def format_reference_audit_html(reference_audit):
    if reference_audit is None:
        return ""
    status_labels = {
        "verified": "已验证",
        "likely": "较可能真实",
        "weak": "弱匹配",
        "not_found": "未检索到",
        "error": "检索异常",
        "skipped": "未检索",
    }
    issue_map = {item.get("index"): item for item in reference_audit.get("issues", [])}
    cards = ""
    for idx, ref in enumerate(reference_audit.get("references", [])[:60], 1):
        online = ref.get("online") or {}
        status = online.get("online_status", "offline")
        status_text = status_labels.get(status, status if status != "offline" else "离线检查")
        confidence = online.get("confidence", 0)
        issue = issue_map.get(idx, {})
        issue_text = _reference_issue_text(issue.get("issues", [])) if issue else "未发现明显问题"
        matches = ""
        for match in (online.get("matched_sources") or [])[:3]:
            title = _html_escape(_brief_text(match.get("title", ""), 150))
            meta = _html_escape(" | ".join(str(x) for x in (match.get("source"), match.get("year"), match.get("doi")) if x))
            url = _html_escape(match.get("url", ""))
            link = f' <a href="{url}" target="_blank" rel="noopener">来源</a>' if url else ""
            matches += f"<li><strong>{title or '无题名'}</strong><br><span>{meta}</span>{link}</li>"
        if not matches:
            matches = "<li>无在线命中记录</li>"
        query = online.get("query") or build_reference_query(ref)
        cards += f"""
        <details class="reference-card" id="reference-{idx}">
          <summary class="reference-summary">
            <span class="reference-index">#{idx}</span>
            <span class="reference-status reference-{_html_escape(status)}">{_html_escape(status_text)}</span>
            <span class="reference-title">{_html_escape(_reference_display_title(ref))}</span>
            <span class="reference-confidence">置信度 {confidence}</span>
            <span class="reference-issues">{_html_escape(_brief_text(issue_text, 140))}</span>
          </summary>
          <div class="reference-body">
            <div><strong>引用原文</strong>: {_reference_text_html(ref.get('text', ''))}</div>
            <p><strong>检索式</strong>: {_html_escape(_reference_query_text(ref, query))}</p>
            <p><strong>问题</strong>: {_html_escape(issue_text)}</p>
            <ul class="reference-matches">{matches}</ul>
          </div>
        </details>"""

    if not cards:
        cards = '<div class="muted">未发现可解析参考文献。</div>'
    return f"""
  <div class="section reference-section" id="reference-audit">
    <h2>参考文献真实性/可核验性校检</h2>
    <p><strong>状态</strong>: {_html_escape(reference_audit.get('status', 'N/A'))} | <strong>数量</strong>: {reference_audit.get('reference_count', 0)} | <strong>DOI</strong>: {reference_audit.get('doi_count', 0)} | <strong>年份</strong>: {reference_audit.get('year_count', 0)} | <strong>在线检索</strong>: {'启用' if reference_audit.get('online_enabled') else '未启用'}（{reference_audit.get('online_checked', 0)}条）</p>
    <p class="section-hint">{_html_escape(reference_audit.get('note', ''))}</p>
    <div class="reference-list">{cards}</div>
  </div>"""


RESOURCE_STATUS_LABELS = {
    "available": "可访问",
    "unavailable": "不可访问",
    "access_restricted": "访问受限",
    "malformed": "链接格式错误",
    "error": "检测异常",
    "skipped": "未检测",
}


def extract_paper_resources(text):
    """Extract code repositories and deployed online resources mentioned by the paper text."""
    raw_text = str(text or "")
    resources = []
    seen = set()
    url_pattern = re.compile(r"\b(?:https?|htps)://[^\s<>'\"\]\)）}]+", re.I)
    for match in url_pattern.finditer(raw_text):
        raw_url = _clean_resource_url(match.group(0))
        if not raw_url:
            continue
        start, end = match.span()
        context = _resource_context(raw_text, start, end)
        resource_type = _classify_resource(raw_url, context)
        if resource_type == "other":
            continue
        key = raw_url.lower()
        if key in seen:
            continue
        seen.add(key)
        resources.append({
            "url": raw_url,
            "type": resource_type,
            "context": context,
        })
    return resources


def _clean_resource_url(url):
    url = html.unescape(str(url or "")).strip()
    url = url.replace("\\_", "_")
    url = re.sub(r"[`*_]+$", "", url)
    url = url.rstrip(".,;:，。；：")
    while url.endswith((")", "）", "]", "}")) and url.count("(") < url.count(")"):
        url = url[:-1].rstrip()
    return url


def _resource_context(text, start, end, radius=180):
    snippet = str(text or "")[max(0, start - radius):min(len(text), end + radius)]
    snippet = re.sub(r"\[\[/?(?:BLOCK|FIGURE|TABLE_START|TABLE_END|TABLE_CONTINUATION|EXTRACTION_NOTE)[^\]]*\]\]", " ", snippet, flags=re.I)
    snippet = re.sub(r"\s+", " ", snippet).strip()
    return _brief_text(snippet, 420)


def _classify_resource(url, context=""):
    lowered = (str(url or "") + " " + str(context or "")).lower()
    host = urllib.parse.urlparse(url if not url.lower().startswith("htps://") else "https://" + url[7:]).netloc.lower()
    if any(domain in host for domain in ("github.com", "gitlab.com", "bitbucket.org")):
        return "code_repository"
    if any(domain in host for domain in ("zenodo.org", "figshare.com", "osf.io", "ncbi.nlm.nih.gov", "portal.gdc.cancer.gov")):
        return "data_repository"
    if any(domain in host for domain in (
        "streamlit.app", "huggingface.co", "shinyapps.io", "herokuapp.com",
        "vercel.app", "netlify.app", "github.io",
    )):
        return "deployed_resource"
    if re.search(r"\b(code availability|code available|source code|github|repository|repo)\b", lowered):
        return "code_repository"
    if re.search(r"\b(streamlit|web calculator|web-based calculator|online platform|online predictive|publicly accessible|deployed)\b", lowered):
        return "deployed_resource"
    return "other"


def verify_resource_availability(resource, timeout=10):
    url = _clean_resource_url((resource or {}).get("url", ""))
    if not re.match(r"^https?://", url, flags=re.I):
        return {
            "status": "malformed",
            "http_status": None,
            "problem": "malformed_url",
            "message": "URL scheme is malformed or unsupported.",
        }
    headers = {"Accept": "text/html,application/json,*/*;q=0.8"}
    try:
        _, status = _http_request(url, "GET", headers=headers, timeout=timeout)
        return {
            "status": "available" if 200 <= int(status) < 400 else "unavailable",
            "http_status": int(status),
            "problem": "",
            "message": "reachable",
        }
    except Exception as e:
        response = getattr(e, "response", None)
        status = getattr(response, "status_code", None)
        if status in {401, 403}:
            availability = "access_restricted"
            problem = "access_restricted"
        elif status in {404, 410}:
            availability = "unavailable"
            problem = "not_found"
        elif status:
            availability = "error"
            problem = f"http_{status}"
        else:
            availability = "error"
            problem = type(e).__name__
        return {
            "status": availability,
            "http_status": status,
            "problem": problem,
            "message": _brief_text(str(e), 240),
        }


def audit_resources(text, online=True, timeout=10, cache=None):
    resources = extract_paper_resources(text)
    resource_cache = cache if isinstance(cache, dict) else {}
    checked = 0
    issues = []
    for idx, resource in enumerate(resources, 1):
        if online:
            cache_key = resource["url"].lower()
            result = resource_cache.get(cache_key)
            if not result:
                result = verify_resource_availability(resource, timeout=timeout)
                resource_cache[cache_key] = result
            resource["availability"] = result
            checked += 1
            if result.get("status") in {"unavailable", "access_restricted", "malformed", "error"}:
                issues.append({
                    "index": idx,
                    "url": resource.get("url", ""),
                    "type": resource.get("type", ""),
                    "status": result.get("status", ""),
                    "problem": result.get("problem", ""),
                    "context": resource.get("context", ""),
                })
        else:
            resource["availability"] = {"status": "skipped", "problem": "online_disabled"}

    status = "ok"
    if issues:
        error_count = sum(1 for item in issues if item.get("status") == "error")
        status = "error" if resources and online and error_count == len(resources) else "needs_review"
    return {
        "status": status,
        "resource_count": len(resources),
        "online_enabled": bool(online),
        "online_checked": checked,
        "issues": issues,
        "resources": resources[:200],
        "note": "校检论文声明的代码仓库、在线计算器、部署平台等资源是否可访问；URL格式错误会单独标记。",
    }


def _resource_status_text(status):
    return RESOURCE_STATUS_LABELS.get(str(status or ""), str(status or "N/A"))


def _resource_type_text(resource_type):
    mapping = {
        "code_repository": "代码仓库",
        "data_repository": "数据资源库",
        "deployed_resource": "在线资源/部署平台",
    }
    return mapping.get(str(resource_type or ""), str(resource_type or "N/A"))


def format_resource_audit_markdown(resource_audit):
    if resource_audit is None:
        return []
    lines = [
        '<a id="resource-audit"></a>',
        "## 🔗 代码仓库与在线资源可用性校检",
        "",
        f"**状态**: {resource_audit.get('status', 'N/A')}",
        f"**资源数量**: {resource_audit.get('resource_count', 0)}",
        f"**在线检测**: {'启用' if resource_audit.get('online_enabled') else '未启用'}"
        + (f"（已检测 {resource_audit.get('online_checked', 0)} 项）" if resource_audit.get("online_enabled") else ""),
        f"> {resource_audit.get('note', '')}",
        "",
    ]
    resources = resource_audit.get("resources") or []
    if resources:
        lines.append("| # | 类型 | URL | 可用性 | 问题 | 上下文 |")
        lines.append("|---|------|-----|--------|------|--------|")
        for idx, resource in enumerate(resources[:40], 1):
            availability = resource.get("availability") or {}
            lines.append(
                f"| {idx} | {_md_escape_cell(_resource_type_text(resource.get('type')))} | "
                f"{_md_escape_cell(resource.get('url', ''))} | "
                f"{_md_escape_cell(_resource_status_text(availability.get('status')))} | "
                f"{_md_escape_cell(availability.get('problem', '') or '-')} | "
                f"{_md_escape_cell(_brief_text(resource.get('context', ''), 180))} |"
            )
    else:
        lines.append("> 未识别到代码仓库或论文部署的在线资源链接。")
    lines.append("")
    return lines


def format_resource_audit_html(resource_audit):
    if resource_audit is None:
        return ""
    rows = ""
    for idx, resource in enumerate((resource_audit.get("resources") or [])[:80], 1):
        availability = resource.get("availability") or {}
        url = resource.get("url", "")
        href = _html_escape(url) if re.match(r"^https?://", url, flags=re.I) else ""
        url_html = f'<a href="{href}" target="_blank" rel="noopener">{_html_escape(url)}</a>' if href else _html_escape(url)
        rows += (
            "<tr>"
            f"<td>{idx}</td>"
            f"<td>{_html_escape(_resource_type_text(resource.get('type')))}</td>"
            f"<td>{url_html}</td>"
            f"<td>{_html_escape(_resource_status_text(availability.get('status')))}</td>"
            f"<td>{_html_escape(availability.get('problem', '') or '-')}</td>"
            f"<td>{_html_escape(_brief_text(resource.get('context', ''), 220))}</td>"
            "</tr>"
        )
    if rows:
        body = f"""
    <table>
      <thead><tr><th>#</th><th>类型</th><th>URL</th><th>可用性</th><th>问题</th><th>上下文</th></tr></thead>
      <tbody>{rows}</tbody>
    </table>"""
    else:
        body = '<p class="section-hint">未识别到代码仓库或论文部署的在线资源链接。</p>'
    return f"""
  <div class="section resource-section" id="resource-audit">
    <h2>代码仓库与在线资源可用性校检</h2>
    <p><strong>状态</strong>: {_html_escape(resource_audit.get('status', 'N/A'))} | <strong>数量</strong>: {resource_audit.get('resource_count', 0)} | <strong>在线检测</strong>: {'启用' if resource_audit.get('online_enabled') else '未启用'}（{resource_audit.get('online_checked', 0)}项）</p>
    <p class="section-hint">{_html_escape(resource_audit.get('note', ''))}</p>
    {body}
  </div>"""


def _md_escape_cell(text):
    """Markdown表格单元格转义与压缩。"""
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    return text.replace("|", "\\|")


def _cross_file_source_label(category):
    return {
        "main_text": "正文",
        "supplement": "补充材料",
        "data_file": "数据文件",
        "other": "其他材料",
    }.get(category or "", category or "未知来源")


def _cross_file_source_rank(category):
    return {
        "main_text": 0,
        "supplement": 1,
        "data_file": 2,
        "other": 3,
    }.get(category or "", 9)


def _cross_file_segment_text(text):
    raw = _clean_mineru_table_block(str(text or ""))
    raw = re.sub(r"\[\[/?(?:BLOCK|FIGURE)[^\]]*\]\]", " ", raw, flags=re.I)
    segments = []
    for line in raw.splitlines():
        line = re.sub(r"\s+", " ", line).strip()
        if not line:
            continue
        for part in re.split(r"(?<=[。.!?])\s+|\s{2,}", line):
            part = part.strip()
            if 12 <= len(part) <= 800:
                segments.append(part)
    return segments[:1200]


_CROSS_FILE_STOPWORDS = {
    "the", "and", "for", "with", "from", "that", "this", "were", "was", "are",
    "into", "have", "has", "had", "not", "all", "table", "figure", "supplement",
    "supplementary", "group", "groups", "cohort", "sample", "samples", "patients",
    "subjects", "mice", "cells", "results", "method", "methods", "study", "data",
}


def _cross_file_terms(text):
    lowered = str(text or "").lower()
    terms = set()
    for token in re.findall(r"[a-z][a-z0-9_-]{2,}|[A-Za-z]*\d+[A-Za-z]*", lowered):
        token = token.strip("_-")
        if token and token not in _CROSS_FILE_STOPWORDS:
            terms.add(token)
    return terms


def _cross_file_is_noisy(text):
    raw = str(text or "")
    if "[文件解析失败" in raw or "[文本过长已截断]" in raw:
        return True
    pipe_count = raw.count("|")
    return pipe_count >= 8 and pipe_count > max(2, len(raw) // 80)


def _extract_cross_file_sample_records(entry):
    records = []
    patterns = [
        re.compile(r"\b[nN]\s*[=:：]\s*(\d{1,5})\b"),
        re.compile(r"\b(\d{1,5})\s+(?:patients?|subjects?|participants?|samples?|mice|cells|cases)\b", re.I),
    ]
    for segment in _cross_file_segment_text(entry.get("text", "")):
        if not re.search(r"\b(?:n\s*[=:：]|\d+\s+(?:patients?|subjects?|participants?|samples?|mice|cells|cases))", segment, re.I):
            continue
        for pattern in patterns:
            for match in pattern.finditer(segment):
                try:
                    value = int(match.group(1))
                except Exception:
                    continue
                if value <= 0:
                    continue
                records.append({
                    "value": value,
                    "terms": _cross_file_terms(segment),
                    "excerpt": segment,
                    "file": entry.get("file", ""),
                    "path": entry.get("path", ""),
                    "category": entry.get("category", ""),
                    "noisy": _cross_file_is_noisy(segment),
                })
    return records


def _cross_file_shared_terms(a, b):
    return sorted((a.get("terms") or set()) & (b.get("terms") or set()))


def _cross_file_context_match(a, b):
    shared = _cross_file_shared_terms(a, b)
    if shared:
        return shared
    a_text = str(a.get("excerpt") or "").lower()
    b_text = str(b.get("excerpt") or "").lower()
    figure_tokens_a = set(re.findall(r"\b(?:fig(?:ure)?|table)\s*s?\d+[a-z]?\b", a_text, flags=re.I))
    figure_tokens_b = set(re.findall(r"\b(?:fig(?:ure)?|table)\s*s?\d+[a-z]?\b", b_text, flags=re.I))
    return sorted(figure_tokens_a & figure_tokens_b)


def _cross_file_finding(conflict_type, severity, claim, counter, reason, manual_check):
    return {
        "conflict_type": conflict_type,
        "severity": severity,
        "claim": claim.get("text", ""),
        "claim_source": claim.get("category", ""),
        "claim_source_label": _cross_file_source_label(claim.get("category", "")),
        "claim_file": claim.get("file", ""),
        "claim_excerpt": _brief_text(claim.get("excerpt", ""), 420),
        "counter_evidence": counter.get("text", ""),
        "counter_source": counter.get("category", ""),
        "counter_source_label": _cross_file_source_label(counter.get("category", "")),
        "counter_file": counter.get("file", ""),
        "counter_excerpt": _brief_text(counter.get("excerpt", ""), 420),
        "reason": reason,
        "manual_check": manual_check,
    }


def _cross_file_sample_findings(entries):
    records = []
    for entry in entries:
        records.extend(_extract_cross_file_sample_records(entry))
    findings = []
    seen = set()
    for idx, a in enumerate(records):
        for b in records[idx + 1:]:
            if a.get("value") == b.get("value"):
                continue
            if a.get("category") == b.get("category") and a.get("file") == b.get("file"):
                continue
            shared = _cross_file_context_match(a, b)
            if not shared:
                continue
            severity = "weak" if a.get("noisy") or b.get("noisy") else "strong"
            key = (
                "sample_size_mismatch",
                tuple(sorted([a.get("file", ""), b.get("file", "")])),
                tuple(sorted([a.get("value"), b.get("value")])),
                tuple(shared[:4]),
            )
            if key in seen:
                continue
            seen.add(key)
            first, second = sorted([a, b], key=lambda item: _cross_file_source_rank(item.get("category")))
            findings.append(_cross_file_finding(
                "sample_size_mismatch",
                severity,
                {
                    **first,
                    "text": f"{_cross_file_source_label(first.get('category'))}报告样本量 n={first.get('value')}",
                },
                {
                    **second,
                    "text": f"{_cross_file_source_label(second.get('category'))}报告样本量 n={second.get('value')}",
                },
                f"相近上下文共享关键词 {', '.join(shared[:6])}，但样本量分别为 {first.get('value')} 和 {second.get('value')}。",
                "核对同一实验/队列/分组的最终纳入样本数、排除标准和表格版本是否一致。",
            ))
    return findings


def _normalize_group_label(label):
    label = re.sub(r"\s+", " ", str(label or "").strip().lower())
    aliases = {
        "wt": "wildtype",
        "wild-type": "wildtype",
        "ko": "knockout",
    }
    return aliases.get(label, label)


def _extract_cross_file_group_labels(entry):
    labels = {}
    patterns = [
        re.compile(r"\b(control|vehicle|placebo|treatment|treated|case|experimental|sham|wildtype|wild-type|wt|knockout|ko|disease)\s+(?:group|arm|cohort)\b", re.I),
        re.compile(r"\b(?:group|arm|cohort)\s+(?:of\s+)?(control|vehicle|placebo|treatment|treated|case|experimental|sham|wildtype|wild-type|wt|knockout|ko|disease)\b", re.I),
    ]
    for segment in _cross_file_segment_text(entry.get("text", "")):
        for pattern in patterns:
            for match in pattern.finditer(segment):
                label = _normalize_group_label(match.group(1))
                labels.setdefault(label, {
                    "label": label,
                    "excerpt": segment,
                    "file": entry.get("file", ""),
                    "path": entry.get("path", ""),
                    "category": entry.get("category", ""),
                })
    return labels


def _cross_file_group_findings(entries):
    by_category = collections.defaultdict(dict)
    for entry in entries:
        labels = _extract_cross_file_group_labels(entry)
        by_category[entry.get("category", "")].update(labels)
    main_labels = by_category.get("main_text") or {}
    other_labels = {}
    for category, labels in by_category.items():
        if category != "main_text":
            other_labels.update(labels)
    findings = []
    if "control" in main_labels and "vehicle" in other_labels and "vehicle" not in main_labels:
        findings.append(_cross_file_finding(
            "group_label_mismatch",
            "medium",
            {**main_labels["control"], "text": "正文使用 Control group"},
            {**other_labels["vehicle"], "text": "补充/数据材料使用 Vehicle group"},
            "正文与补充/数据材料使用了不同的对照组标签；两者可能是同义设计，也可能代表分组命名不一致。",
            "核对方法学定义、图表标签和原始分组编码，确认 Control 与 Vehicle 是否为同一组。",
        ))
    return findings


def _extract_supplementary_refs(text):
    refs = []
    pattern = re.compile(r"\b(?:Supplementary|Supplemental|附表|补充图|补充表)\s*(?:Fig(?:ure)?|Table)?\s*S?(\d+[A-Za-z]?)\b|\b(?:Fig(?:ure)?|Table)\s*S(\d+[A-Za-z]?)\b", re.I)
    for segment in _cross_file_segment_text(text):
        for match in pattern.finditer(segment):
            number = (match.group(1) or match.group(2) or "").lower()
            if number:
                refs.append((f"s{number}" if not number.startswith("s") else number, segment))
    return refs


def _cross_file_figure_table_findings(entries):
    main_text = "\n".join(entry.get("text", "") for entry in entries if entry.get("category") == "main_text")
    supplemental_text = "\n".join(entry.get("text", "") for entry in entries if entry.get("category") in {"supplement", "data_file"})
    if not main_text or not supplemental_text:
        return []
    supplement_lower = supplemental_text.lower()
    findings = []
    seen = set()
    for ref_id, excerpt in _extract_supplementary_refs(main_text):
        if ref_id in seen:
            continue
        seen.add(ref_id)
        compact = ref_id.replace("s", "")
        if ref_id in supplement_lower or f"table {compact}" in supplement_lower or f"figure {compact}" in supplement_lower:
            continue
        findings.append(_cross_file_finding(
            "supplement_reference_gap",
            "weak",
            {
                "category": "main_text",
                "file": "main_text",
                "excerpt": excerpt,
                "text": f"正文引用补充材料 {ref_id.upper()}",
            },
            {
                "category": "supplement",
                "file": "supplement/data files",
                "excerpt": f"未在已提取补充/数据文本中找到 {ref_id.upper()} 的直接标记。",
                "text": "补充材料标记覆盖不足",
            },
            "正文出现补充图表引用，但已提取补充材料中未找到对应编号标记。",
            "核对补充材料文件是否完整、编号是否被OCR/表格提取改写，或是否缺失对应补充图表。",
        ))
        if len(findings) >= 8:
            break
    return findings


def build_cross_file_consistency_audit(file_entries, root_path=None):
    entries = []
    for entry in file_entries or []:
        text = str(entry.get("text") or "")
        category = entry.get("category") or "other"
        if not text.strip() or category == "reference":
            continue
        entries.append({
            "file": entry.get("file") or Path(entry.get("path", "")).name,
            "path": entry.get("path") or entry.get("file") or "",
            "category": category,
            "text": text,
        })
    cross_categories = {entry.get("category") for entry in entries if entry.get("category") != "main_text"}
    if len(entries) < 2 or not cross_categories:
        return {
            "status": "skipped",
            "checked_files": len(entries),
            "finding_count": 0,
            "strong_count": 0,
            "medium_count": 0,
            "weak_count": 0,
            "findings": [],
            "note": "缺少可比较的跨文件材料；跨文件一致性审查已跳过。",
        }
    findings = []
    findings.extend(_cross_file_sample_findings(entries))
    findings.extend(_cross_file_group_findings(entries))
    findings.extend(_cross_file_figure_table_findings(entries))
    severity_rank = {"strong": 0, "medium": 1, "weak": 2}
    findings = sorted(findings, key=lambda item: (severity_rank.get(item.get("severity"), 9), item.get("conflict_type", ""), item.get("claim_file", "")))[:40]
    return {
        "status": "ok",
        "checked_files": len(entries),
        "finding_count": len(findings),
        "strong_count": sum(1 for item in findings if item.get("severity") == "strong"),
        "medium_count": sum(1 for item in findings if item.get("severity") == "medium"),
        "weak_count": sum(1 for item in findings if item.get("severity") == "weak"),
        "findings": findings,
        "note": "基于已提取文本的跨文件一致性审查；不等同于最终科研不端判断。",
    }


_EVIDENCE_SECTION_ALIASES = {
    "abstract": "abstract",
    "methods": "methods",
    "method": "methods",
    "materials and methods": "methods",
    "materials & methods": "methods",
    "results": "results",
    "result": "results",
    "discussion": "discussion",
    "conclusion": "conclusion",
    "conclusions": "conclusion",
}


_EVIDENCE_STRONG_CLAIM_TERMS = (
    "demonstrate", "demonstrated", "demonstrates", "prove", "proved", "proves",
    "confirm", "confirmed", "confirms", "establish", "established", "significant",
    "significantly", "markedly", "robustly", "definitive", "definitively",
    "证明", "证实", "确定", "显著", "明确", "强烈", "完全",
)


_EVIDENCE_SUPPORT_HINT_RE = re.compile(
    r"\b(?:p\s*[<=>]|n\s*[=:：]|\d+\s+(?:patients?|subjects?|participants?|samples?|mice|cells|cases)|fig(?:ure)?|table)\b|图\s*\d+|表\s*\d+",
    re.I,
)


def _evidence_section_name(raw):
    return _EVIDENCE_SECTION_ALIASES.get(re.sub(r"\s+", " ", str(raw or "").strip().lower()))


def _extract_evidence_sections(text):
    sections = collections.defaultdict(list)
    current = None
    heading_re = re.compile(
        r"^\s*(abstract|materials\s+(?:and|&)\s+methods|methods?|results?|discussion|conclusions?)\s*(?::\s*(.*)|$)",
        re.I,
    )
    for line in str(text or "").splitlines():
        stripped = line.strip()
        match = heading_re.match(stripped)
        if match and len((match.group(2) or "").strip()) <= 240:
            section_name = _evidence_section_name(match.group(1))
            trailing = (match.group(2) or "").strip()
            if section_name:
                current = section_name
                if trailing:
                    sections[current].append(trailing)
                continue
        if current:
            sections[current].append(line)
    return {key: "\n".join(value).strip() for key, value in sections.items()}


def _extract_section_sample_records(text, section):
    entry = {"text": text, "file": section, "path": section, "category": section}
    records = []
    for record in _extract_cross_file_sample_records(entry):
        record = dict(record)
        record["section"] = section
        records.append(record)
    return records


def _extract_section_group_records(text, section):
    entry = {"text": text, "file": section, "path": section, "category": section}
    records = []
    for label, record in _extract_cross_file_group_labels(entry).items():
        item = dict(record)
        item["label"] = label
        item["section"] = section
        records.append(item)
    return records


def _evidence_segment_has_strong_claim(segment):
    lowered = str(segment or "").lower()
    return any(term in lowered or term in segment for term in _EVIDENCE_STRONG_CLAIM_TERMS)


def _evidence_claim_keywords(text):
    terms = _cross_file_terms(text)
    generic = {
        "abstract", "conclusion", "conclusions", "paper", "article", "study",
        "studies", "result", "results", "finding", "findings", "show", "shows",
        "shown", "demonstrate", "demonstrated", "demonstrates", "significant",
        "significantly", "prove", "proved", "confirm", "confirmed",
    }
    return sorted(term for term in terms if term not in generic)[:12]


def _results_support_claim(result_text, claim_segment):
    result_lower = str(result_text or "").lower()
    keywords = _evidence_claim_keywords(claim_segment)
    shared = [term for term in keywords if term.lower() in result_lower]
    if len(shared) >= 2:
        return True
    if shared and _EVIDENCE_SUPPORT_HINT_RE.search(result_text or ""):
        return True
    return False


def _build_claim_chain_findings(full_text):
    sections = _extract_evidence_sections(full_text)
    methods_text = sections.get("methods", "")
    results_text = sections.get("results", "")
    findings = []

    if methods_text and results_text:
        method_samples = _extract_section_sample_records(methods_text, "methods")
        result_samples = _extract_section_sample_records(results_text, "results")
        seen = set()
        for method_record in method_samples:
            for result_record in result_samples:
                if method_record.get("value") == result_record.get("value"):
                    continue
                shared = _cross_file_context_match(method_record, result_record)
                if not shared:
                    continue
                key = ("sample_size_chain_mismatch", method_record.get("value"), result_record.get("value"), tuple(shared[:5]))
                if key in seen:
                    continue
                seen.add(key)
                findings.append({
                    "type": "methods_results_sample_size_mismatch",
                    "severity": "strong" if not method_record.get("noisy") and not result_record.get("noisy") else "medium",
                    "chain": "Methods -> Results",
                    "method_excerpt": _brief_text(method_record.get("excerpt", ""), 420),
                    "result_excerpt": _brief_text(result_record.get("excerpt", ""), 420),
                    "reason": f"Methods 与 Results 在相近上下文共享关键词 {', '.join(shared[:6])}，但样本量分别为 {method_record.get('value')} 和 {result_record.get('value')}。",
                    "manual_check": "核对同一实验/队列/分组的纳入样本量、排除标准和最终图表版本。",
                    "evidence_keys": sorted(set(shared) | {f"n={method_record.get('value')}", f"n={result_record.get('value')}"}),
                })
                if len(findings) >= 20:
                    break
            if len(findings) >= 20:
                break

        method_groups = {item.get("label"): item for item in _extract_section_group_records(methods_text, "methods")}
        result_groups = {item.get("label"): item for item in _extract_section_group_records(results_text, "results")}
        if "control" in method_groups and "vehicle" in result_groups and "vehicle" not in method_groups:
            findings.append({
                "type": "methods_results_group_label_mismatch",
                "severity": "medium",
                "chain": "Methods -> Results",
                "method_excerpt": _brief_text(method_groups["control"].get("excerpt", ""), 420),
                "result_excerpt": _brief_text(result_groups["vehicle"].get("excerpt", ""), 420),
                "reason": "Methods 使用 Control group，Results 使用 Vehicle group；两者可能是同义设计，也可能是分组命名不一致。",
                "manual_check": "核对方法学定义、图表标签和原始分组编码，确认 Control 与 Vehicle 是否为同一组。",
                "evidence_keys": ["control", "vehicle"],
            })

    conclusion_sources = []
    for section_name in ("abstract", "conclusion"):
        for segment in _cross_file_segment_text(sections.get(section_name, "")):
            if not _evidence_segment_has_strong_claim(segment):
                continue
            keywords = _evidence_claim_keywords(segment)
            if not keywords:
                continue
            if _results_support_claim(results_text, segment):
                continue
            conclusion_sources.append((section_name, segment, keywords))
    for section_name, segment, keywords in conclusion_sources[:12]:
        findings.append({
            "type": "strong_claim_without_result_support",
            "severity": "medium",
            "chain": f"{section_name.title()} -> Results",
            "claim_excerpt": _brief_text(segment, 420),
            "result_excerpt": _brief_text(results_text, 420) if results_text else "未识别到 Results 段落或结果证据不足。",
            "reason": "摘要/结论出现强结论措辞，但 Results 段落中未找到足够接近的结果、图表或统计支撑。",
            "manual_check": "回查 Results 中是否存在对应指标、图表编号、统计量和方向一致的结果描述。",
            "evidence_keys": keywords[:8],
        })

    severity_rank = {"strong": 0, "medium": 1, "weak": 2}
    return sorted(findings, key=lambda item: (severity_rank.get(item.get("severity"), 9), item.get("type", "")))[:40]


def _extract_evidence_refs(text):
    raw = str(text or "")
    refs = set()
    for prefix, number in re.findall(r"\b(fig(?:ure)?|table)\s*\.?\s*(s?\d+[a-z]?)\b", raw, flags=re.I):
        kind = "figure" if prefix.lower().startswith("fig") else "table"
        refs.add(f"{kind}:{number.lower()}")
    for prefix, number in re.findall(r"([图表])\s*\.?\s*(\d+[a-zA-Z]?)", raw):
        kind = "figure" if prefix == "图" else "table"
        refs.add(f"{kind}:{number.lower()}")
    return sorted(refs)


def _evidence_keys_from_text(text):
    keys = set(_extract_evidence_refs(text))
    for value in re.findall(r"\b[nN]\s*[=:：]\s*(\d{1,5})\b", str(text or "")):
        keys.add(f"n={value}")
    for term in _cross_file_terms(text):
        if term not in {"risk", "audit", "evidence", "reason", "manual"}:
            keys.add(term)
    return sorted(keys)[:24]


def _evidence_item(source_type, severity, title, detail="", excerpt="", keys=None, source_id=None):
    text = " ".join(str(part or "") for part in (title, detail, excerpt))
    return {
        "id": source_id or f"{source_type}-{_text_fingerprint(text)[:10]}",
        "source_type": source_type,
        "severity": severity if severity in {"strong", "medium", "weak"} else "weak",
        "title": _brief_text(title, 160),
        "detail": _brief_text(detail, 520),
        "excerpt": _brief_text(excerpt, 520),
        "keys": sorted(set(keys or []) | set(_evidence_keys_from_text(text)))[:32],
    }


def _evidence_items_from_chain_findings(findings):
    items = []
    for idx, finding in enumerate(findings or [], 1):
        excerpt = " || ".join(
            part for part in (
                finding.get("method_excerpt"),
                finding.get("result_excerpt"),
                finding.get("claim_excerpt"),
            )
            if part
        )
        items.append(_evidence_item(
            "claim_chain",
            finding.get("severity", "weak"),
            finding.get("type", "claim_chain_finding"),
            finding.get("reason") or finding.get("manual_check") or "",
            excerpt,
            keys=finding.get("evidence_keys") or [],
            source_id=f"chain-{idx}",
        ))
    return items


def _evidence_items_from_cross_file(meta):
    audit = (meta or {}).get("cross_file_consistency_audit") or {}
    items = []
    for idx, finding in enumerate((audit.get("findings") or [])[:30], 1):
        excerpt = (
            f"{finding.get('claim_source_label')} / {finding.get('claim_file')}: {finding.get('claim_excerpt')} "
            f"|| {finding.get('counter_source_label')} / {finding.get('counter_file')}: {finding.get('counter_excerpt')}"
        )
        items.append(_evidence_item(
            "cross_file_consistency",
            finding.get("severity", "weak"),
            finding.get("conflict_type", "cross_file_finding"),
            finding.get("reason") or finding.get("manual_check") or "",
            excerpt,
            source_id=f"cross-file-{idx}",
        ))
    return items


def _evidence_items_from_llm_report(report):
    checks = sorted(report.get("checks", []) if isinstance(report, dict) else [], key=_check_sort_key)
    items = []
    for idx, check in enumerate(checks[:80], 1):
        if not _is_suspicious_check(check):
            continue
        score = _check_suspicion_score(check)
        severity = "strong" if "红旗" in str(check.get("verdict", "")) or score >= 220 else "medium"
        items.append(_evidence_item(
            "llm_check",
            severity,
            f"{check.get('category', 'N/A')} / {check.get('item', 'N/A')}",
            _check_reason(check),
            _check_source_text(check),
            source_id=f"llm-{idx}",
        ))
    return items[:30]


def _evidence_items_from_stat(stat_result):
    stat_result = stat_result or {}
    items = []
    if stat_result.get("number_consistency"):
        items.append(_evidence_item(
            "local_stat",
            "medium",
            "数字自洽性需复核",
            stat_result.get("number_consistency"),
            source_id="stat-number-consistency",
        ))
    if stat_result.get("benford_status") and "高偏差" in str(stat_result.get("benford_status")):
        items.append(_evidence_item(
            "local_stat",
            "medium",
            "Benford分布偏差较高",
            f"偏差={round(stat_result.get('benford_deviation') or 0, 3)}；建议核对原始数值来源。",
            source_id="stat-benford",
        ))
    if (stat_result.get("p_value_abnormal") or 0) > 0:
        severity = "medium" if (stat_result.get("p_value_abnormal") or 0) >= 3 else "weak"
        items.append(_evidence_item(
            "local_stat",
            severity,
            "p值分布需复核",
            f"p值数量/异常: {stat_result.get('p_value_count', 0)} / {stat_result.get('p_value_abnormal', 0)}。",
            source_id="stat-p-values",
        ))
    return items


def _evidence_items_from_reference(meta):
    reference_audit = (meta or {}).get("reference_audit") or {}
    refs_by_index = {i + 1: ref for i, ref in enumerate(reference_audit.get("references") or [])}
    items = []
    for issue in (reference_audit.get("issues") or [])[:12]:
        ref = refs_by_index.get(issue.get("index"), {})
        online = ref.get("online") or {}
        status = online.get("online_status")
        if status and status not in {"not_found", "weak", "error"}:
            continue
        items.append(_evidence_item(
            "reference_audit",
            "medium" if status in {"not_found", "error"} else "weak",
            f"参考文献 #{issue.get('index')} 在线证据不足",
            "；".join(str(x) for x in issue.get("issues", []) if x) or status or "reference issue",
            issue.get("text", ""),
            source_id=f"reference-{issue.get('index')}",
        ))
    return items


def _evidence_items_from_resource(meta):
    resource_audit = (meta or {}).get("resource_audit") or {}
    items = []
    for idx, issue in enumerate((resource_audit.get("issues") or [])[:12], 1):
        status = issue.get("status")
        severity = "medium" if status in {"unavailable", "access_restricted", "malformed", "error"} else "weak"
        items.append(_evidence_item(
            "resource_audit",
            severity,
            f"论文资源需复核: {issue.get('type') or 'resource'}",
            issue.get("problem") or status or "",
            issue.get("url") or "",
            source_id=f"resource-{idx}",
        ))
    return items


def _evidence_items_from_image(meta):
    image_audit = (meta or {}).get("image_audit") or {}
    items = []
    for idx, img in enumerate((image_audit.get("images") or [])[:30], 1):
        sem = img.get("semantic") or {}
        detector = img.get("detector") or {}
        detector_score = detector.get("score")
        local_warning = img.get("risk") == "local_warning"
        semantic_warning = sem.get("reasonability") in {"需人工核对", "可疑"} or sem.get("status") == "error"
        detector_warning = detector.get("status") == "ok" and detector_score is not None and detector_score >= 50
        if not (local_warning or semantic_warning or detector_warning):
            continue
        severity = "medium" if sem.get("reasonability") == "可疑" or (detector_score or 0) >= 80 else "weak"
        details = []
        if local_warning:
            details.append("本地图像规则提示需复核")
        if semantic_warning:
            details.append(f"图像语义: {sem.get('reasonability') or sem.get('status')}")
        if detector_warning:
            details.append(f"imagedetector score={detector_score}")
        items.append(_evidence_item(
            "image_audit",
            severity,
            f"图像需复核: {img.get('file') or idx}",
            "；".join(details),
            " ".join(str(x or "") for x in (img.get("file"), sem.get("summary"), detector.get("message"))),
            source_id=f"image-{idx}",
        ))
    return items


def _cluster_severity(evidence):
    severities = [item.get("severity") for item in evidence]
    source_types = {item.get("source_type") for item in evidence}
    if "strong" in severities:
        return "strong"
    if len(source_types) >= 2 and "medium" in severities:
        return "strong"
    if "medium" in severities:
        return "medium"
    if len(source_types) >= 2:
        return "medium"
    return "weak"


def _build_evidence_clusters(items):
    clusters = []
    for item in items:
        keys = set(item.get("keys") or [])
        target = None
        if keys:
            for cluster in clusters:
                if keys & set(cluster.get("_keys") or []):
                    target = cluster
                    break
        if target is None:
            target = {"_keys": sorted(keys), "evidence": []}
            clusters.append(target)
        else:
            target["_keys"] = sorted(set(target.get("_keys") or []) | keys)[:48]
        target["evidence"].append(item)

    severity_rank = {"strong": 0, "medium": 1, "weak": 2}
    rendered = []
    for idx, cluster in enumerate(clusters, 1):
        evidence = cluster.get("evidence") or []
        severity = _cluster_severity(evidence)
        source_types = sorted({item.get("source_type") for item in evidence if item.get("source_type")})
        key_refs = [key for key in cluster.get("_keys") or [] if key.startswith(("figure:", "table:", "n="))]
        title = " / ".join(key_refs[:3]) if key_refs else (evidence[0].get("title") if evidence else "证据簇")
        rendered.append({
            "id": f"cluster-{idx}",
            "severity": severity,
            "title": _brief_text(title, 140),
            "source_types": source_types,
            "evidence_count": len(evidence),
            "keys": (cluster.get("_keys") or [])[:24],
            "summary": _brief_text("; ".join(item.get("title", "") for item in evidence[:4] if item.get("title")), 360),
            "evidence": evidence[:12],
        })
    rendered.sort(key=lambda item: (severity_rank.get(item.get("severity"), 9), -item.get("evidence_count", 0), item.get("title", "")))
    for idx, cluster in enumerate(rendered, 1):
        cluster["id"] = f"cluster-{idx}"
    return rendered[:40]


def build_evidence_chain_audit(full_text, file_entries, report, meta, stat_result):
    meta = meta or {}
    full_text = str(full_text or "")
    file_entries = file_entries or []
    claim_findings = _build_claim_chain_findings(full_text) if full_text.strip() else []
    items = []
    items.extend(_evidence_items_from_chain_findings(claim_findings))
    items.extend(_evidence_items_from_cross_file(meta))
    items.extend(_evidence_items_from_llm_report(report or {}))
    items.extend(_evidence_items_from_stat(stat_result or {}))
    items.extend(_evidence_items_from_reference(meta))
    items.extend(_evidence_items_from_resource(meta))
    items.extend(_evidence_items_from_image(meta))
    clusters = _build_evidence_clusters(items)
    note_scope = "单文件链条审查范围较窄；未发现可比较的跨文件材料。" if len(file_entries) <= 1 else "已结合全文、跨文件材料和现有审查信号进行证据簇聚合。"
    if not full_text.strip() and not items:
        status = "skipped"
        note = "缺少全文和可聚合审查信号；证据链审查已跳过。"
    else:
        status = "ok"
        note = f"{note_scope} 该结果用于人工复核排序，不等同于科研不端判断。"
    return {
        "status": status,
        "checked_files": len(file_entries or []),
        "cluster_count": len(clusters),
        "finding_count": len(claim_findings),
        "strong_count": sum(1 for item in clusters if item.get("severity") == "strong"),
        "medium_count": sum(1 for item in clusters if item.get("severity") == "medium"),
        "weak_count": sum(1 for item in clusters if item.get("severity") == "weak"),
        "clusters": clusters,
        "claim_chain_findings": claim_findings,
        "note": note,
    }


def _cross_file_severity_label(severity):
    return {
        "strong": "强证据冲突",
        "medium": "中等疑点",
        "weak": "弱信号/需人工核对",
    }.get(severity or "", severity or "未知")


def format_evidence_chain_audit_markdown(audit):
    if audit is None:
        return []
    lines = [
        '<a id="evidence-chain"></a>',
        "## 🔗 证据链与证据簇审查",
        "",
        f"**状态**: {audit.get('status', 'N/A')}",
        f"**证据簇**: {audit.get('cluster_count', 0)}（强 {audit.get('strong_count', 0)} / 中 {audit.get('medium_count', 0)} / 弱 {audit.get('weak_count', 0)}）",
        f"**链条发现**: {audit.get('finding_count', 0)}",
        f"> {audit.get('note', '')}",
        "",
    ]
    findings = audit.get("claim_chain_findings") or []
    if findings:
        lines.append('<a id="evidence-chain-findings"></a>')
        lines.append("### Methods → Results → Abstract/Conclusion 链条发现")
        lines.append("")
        lines.append("| # | 级别 | 链条 | 类型 | 证据摘要 | 复核建议 |")
        lines.append("|---|------|------|------|----------|----------|")
        for idx, finding in enumerate(findings[:20], 1):
            evidence = " || ".join(
                part for part in (
                    finding.get("method_excerpt"),
                    finding.get("result_excerpt"),
                    finding.get("claim_excerpt"),
                )
                if part
            )
            lines.append(
                f"| {idx} | {_md_escape_cell(_cross_file_severity_label(finding.get('severity')))} | "
                f"{_md_escape_cell(finding.get('chain', ''))} | {_md_escape_cell(finding.get('type', ''))} | "
                f"{_md_escape_cell(_brief_text(evidence or finding.get('reason', ''), 360))} | "
                f"{_md_escape_cell(finding.get('manual_check', ''))} |"
            )
        lines.append("")
    clusters = audit.get("clusters") or []
    if clusters:
        lines.append('<a id="evidence-chain-clusters"></a>')
        lines.append("### 证据簇")
        lines.append("")
        lines.append("| # | 级别 | 主题 | 来源 | 证据数 | 摘要 |")
        lines.append("|---|------|------|------|--------|------|")
        for idx, cluster in enumerate(clusters[:30], 1):
            lines.append(
                f"| {idx} | {_md_escape_cell(_cross_file_severity_label(cluster.get('severity')))} | "
                f"{_md_escape_cell(cluster.get('title', ''))} | {_md_escape_cell(', '.join(cluster.get('source_types') or []))} | "
                f"{cluster.get('evidence_count', 0)} | {_md_escape_cell(_brief_text(cluster.get('summary', ''), 260))} |"
            )
    else:
        lines.append("> 未形成明确证据簇；仍建议人工核对关键结果链条。")
    lines.append("")
    return lines


def format_evidence_chain_audit_html(audit):
    if not audit:
        return ""
    findings = audit.get("claim_chain_findings") or []
    finding_cards = ""
    for idx, finding in enumerate(findings[:20], 1):
        evidence = " || ".join(
            part for part in (
                finding.get("method_excerpt"),
                finding.get("result_excerpt"),
                finding.get("claim_excerpt"),
            )
            if part
        )
        finding_cards += f"""
        <details class="cross-file-card evidence-chain-card" id="evidence-chain-finding-{idx}">
          <summary class="cross-file-summary">
            <span class="cross-file-rank">#{idx}</span>
            <span class="cross-file-severity cross-file-{_html_escape(finding.get('severity', ''))}">{_html_escape(_cross_file_severity_label(finding.get('severity')))}</span>
            <span class="cross-file-title">{_html_escape(finding.get('chain', ''))}</span>
            <span class="cross-file-reason">{_html_escape(_brief_text(finding.get('reason', ''), 150))}</span>
          </summary>
          <div class="cross-file-body">
            <p>{_html_escape(evidence)}</p>
            <p><strong>复核建议</strong>: {_html_escape(finding.get('manual_check', ''))}</p>
          </div>
        </details>"""
    cluster_cards = ""
    for idx, cluster in enumerate((audit.get("clusters") or [])[:30], 1):
        evidence_rows = ""
        for item in (cluster.get("evidence") or [])[:8]:
            evidence_rows += f"""
            <li><strong>{_html_escape(item.get('source_type', ''))}</strong> · {_html_escape(item.get('title', ''))}<br><small>{_html_escape(item.get('detail') or item.get('excerpt') or '')}</small></li>"""
        cluster_cards += f"""
        <details class="cross-file-card evidence-cluster-card" id="evidence-chain-cluster-{idx}">
          <summary class="cross-file-summary">
            <span class="cross-file-rank">#{idx}</span>
            <span class="cross-file-severity cross-file-{_html_escape(cluster.get('severity', ''))}">{_html_escape(_cross_file_severity_label(cluster.get('severity')))}</span>
            <span class="cross-file-title">{_html_escape(cluster.get('title', ''))}</span>
            <span class="cross-file-reason">{_html_escape(_brief_text(cluster.get('summary', ''), 150))}</span>
          </summary>
          <div class="cross-file-body">
            <p><strong>来源</strong>: {_html_escape(', '.join(cluster.get('source_types') or []))} · <strong>证据数</strong>: {cluster.get('evidence_count', 0)}</p>
            <ul>{evidence_rows}</ul>
          </div>
        </details>"""
    if not finding_cards:
        finding_cards = '<div class="muted">未发现明确 Methods/Results/Abstract/Conclusion 链条问题。</div>'
    if not cluster_cards:
        cluster_cards = '<div class="muted">未形成明确证据簇；仍建议人工核对关键结果链条。</div>'
    return f"""
  <div class="section cross-file-section evidence-chain-section" id="evidence-chain">
    <h2>证据链与证据簇审查</h2>
    <p class="section-hint">{_html_escape(audit.get('note', ''))}</p>
    <p><strong>状态</strong>: {_html_escape(audit.get('status', 'N/A'))} | <strong>证据簇</strong>: {audit.get('cluster_count', 0)}（强 {audit.get('strong_count', 0)} / 中 {audit.get('medium_count', 0)} / 弱 {audit.get('weak_count', 0)}） | <strong>链条发现</strong>: {audit.get('finding_count', 0)}</p>
    <h3 id="evidence-chain-findings">链条发现</h3>
    {finding_cards}
    <h3 id="evidence-chain-clusters">证据簇</h3>
    {cluster_cards}
  </div>"""


def format_cross_file_consistency_markdown(audit):
    if audit is None:
        return []
    lines = [
        '<a id="cross-file-consistency"></a>',
        "## 🧩 跨文件一致性审查",
        "",
        f"**状态**: {audit.get('status', 'N/A')}",
        f"**检查文件数**: {audit.get('checked_files', 0)}",
        f"**发现数**: {audit.get('finding_count', 0)}（强 {audit.get('strong_count', 0)} / 中 {audit.get('medium_count', 0)} / 弱 {audit.get('weak_count', 0)}）",
        f"> {audit.get('note', '')}",
        "",
    ]
    findings = audit.get("findings") or []
    if findings:
        lines.append("| # | 级别 | 类型 | 证据A | 证据B | 复核建议 |")
        lines.append("|---|------|------|-------|-------|----------|")
        for idx, finding in enumerate(findings[:30], 1):
            claim = f"{finding.get('claim_source_label') or finding.get('claim_source')} / {finding.get('claim_file')}: {finding.get('claim_excerpt')}"
            counter = f"{finding.get('counter_source_label') or finding.get('counter_source')} / {finding.get('counter_file')}: {finding.get('counter_excerpt')}"
            lines.append(
                f"| {idx} | {_md_escape_cell(_cross_file_severity_label(finding.get('severity')))} | "
                f"{_md_escape_cell(finding.get('conflict_type', ''))} | {_md_escape_cell(_brief_text(claim, 260))} | "
                f"{_md_escape_cell(_brief_text(counter, 260))} | {_md_escape_cell(finding.get('manual_check', ''))} |"
            )
    else:
        lines.append("> 未发现明确跨文件不一致；仍建议人工抽查关键表格、补充材料和正文结论。")
    lines.append("")
    return lines


def format_cross_file_consistency_html(audit):
    if not audit:
        return ""
    findings = audit.get("findings") or []
    if findings:
        cards = ""
        for idx, finding in enumerate(findings[:40], 1):
            cards += f"""
        <details class="cross-file-card" id="cross-file-finding-{idx}">
          <summary class="cross-file-summary">
            <span class="cross-file-rank">#{idx}</span>
            <span class="cross-file-severity cross-file-{_html_escape(finding.get('severity', ''))}">{_html_escape(_cross_file_severity_label(finding.get('severity')))}</span>
            <span class="cross-file-title">{_html_escape(finding.get('conflict_type', ''))}</span>
            <span class="cross-file-reason">{_html_escape(_brief_text(finding.get('reason', ''), 140))}</span>
          </summary>
          <div class="cross-file-body">
            <div><strong>{_html_escape(finding.get('claim_source_label') or finding.get('claim_source'))} / {_html_escape(finding.get('claim_file', ''))}</strong><p>{_html_escape(finding.get('claim_excerpt', ''))}</p></div>
            <div><strong>{_html_escape(finding.get('counter_source_label') or finding.get('counter_source'))} / {_html_escape(finding.get('counter_file', ''))}</strong><p>{_html_escape(finding.get('counter_excerpt', ''))}</p></div>
            <p><strong>复核建议</strong>: {_html_escape(finding.get('manual_check', ''))}</p>
          </div>
        </details>"""
    else:
        cards = '<div class="muted">未发现明确跨文件不一致；仍建议人工抽查关键表格、补充材料和正文结论。</div>'
    return f"""
  <div class="section cross-file-section" id="cross-file-consistency">
    <h2>跨文件一致性审查</h2>
    <p><strong>状态</strong>: {_html_escape(audit.get('status', 'N/A'))} | <strong>文件</strong>: {audit.get('checked_files', 0)} | <strong>发现</strong>: {audit.get('finding_count', 0)}（强 {audit.get('strong_count', 0)} / 中 {audit.get('medium_count', 0)} / 弱 {audit.get('weak_count', 0)}）</p>
    <p class="section-hint">{_html_escape(audit.get('note', ''))}</p>
    <div class="cross-file-list">{cards}</div>
  </div>"""


def _is_suspicious_check(c):
    verdict = str(c.get("verdict", ""))
    return ("红旗" in verdict) or ("疑点" in verdict) or ("可疑" in verdict)


def _check_suspicion_score(c):
    """Higher score means the finding should be reviewed earlier."""
    verdict = str(c.get("verdict", ""))
    text = " ".join(
        str(c.get(k, "") or "")
        for k in ("category", "item", "source_text", "quote", "evidence", "reason", "detail", "analysis", "explanation")
    )

    if "红旗" in verdict:
        score = 300
    elif "疑点" in verdict or "可疑" in verdict:
        score = 200
    elif "通过" in verdict:
        score = 0
    else:
        score = 100

    high_terms = (
        "造假", "伪造", "篡改", "捏造", "复制粘贴", "严重", "重大", "直接矛盾",
        "明显矛盾", "无法复现", "不可靠", "否决", "致命", "必须公开", "数据真实性",
    )
    medium_terms = (
        "矛盾", "异常", "重复", "缺失", "不一致", "过拟合", "无验证", "样本量",
        "p值", "多重比较", "利益冲突", "方法论缺陷",
    )
    low_terms = (
        "OCR", "提取", "人工核对", "原PDF", "表格结构", "暂判", "无法判定",
        "不宜判定", "无理由认定", "可能", "需确认",
    )

    score += sum(18 for term in high_terms if term in text)
    score += sum(8 for term in medium_terms if term in text)
    score -= sum(10 for term in low_terms if term in text)

    return max(score, 0)


def _check_source_tags(c: Dict[str, Any]) -> List[str]:
    text = _check_text_blob(c).lower()
    tags = []
    if c.get("_runtime_year_check"):
        tags.append("本地规则")
    if any(term in text for term in ("crossref", "openalex", "pubmed", "doi", "在线", "元数据")):
        tags.append("在线核验")
    if any(term in text for term in ("imagedetector", "ai概率", "ai probability")):
        tags.append("imagedetector")
    if any(term in text for term in ("图像语义", "visible_text", "semantic", "多模态")):
        tags.append("图像语义")
    if any(term in text for term in ("benford", "p值", "p-value", "统计")):
        tags.append("统计线索")
    if _is_extraction_limited_check(c):
        tags.append("MinerU/OCR提取")
    if not tags:
        tags.append("LLM语义")
    return list(dict.fromkeys(tags))


def _merged_group_summary_text(c: Dict[str, Any]) -> str:
    group = c.get("merged_group") or {}
    if not group:
        return ""
    chunks = "/".join(str(item) for item in group.get("source_chunks") or [])
    items = "、".join(str(item) for item in (group.get("items") or [])[:5] if item)
    return f"已合并 {group.get('count', 0)} 条相近疑点；来源分块: {chunks or 'N/A'}；原始项: {items or 'N/A'}"


def _merged_group_html(c: Dict[str, Any]) -> str:
    group = c.get("merged_group") or {}
    if not group:
        return ""
    rows = ""
    for idx, member in enumerate(group.get("members") or [], 1):
        rows += f"""
        <tr>
          <td>{idx}</td>
          <td>{_html_escape(member.get('chunk', '-'))}</td>
          <td>{_html_escape(member.get('item', '-'))}</td>
          <td>{_html_escape(member.get('verdict', '-'))}</td>
          <td>{_html_escape(_brief_text(member.get('evidence') or member.get('source_text') or '', 160))}</td>
        </tr>"""
    return f"""
    <details class="merged-group">
      <summary>{_html_escape(_merged_group_summary_text(c))}</summary>
      <table>
        <thead><tr><th>#</th><th>分块</th><th>原始项</th><th>原判定</th><th>证据摘要</th></tr></thead>
        <tbody>{rows}</tbody>
      </table>
    </details>"""


def _check_sort_key(c):
    return (-_check_suspicion_score(c), str(c.get("category", "")), str(c.get("item", "")))


def _check_verdict_class(verdict):
    verdict = str(verdict or "")
    if "红旗" in verdict:
        return "verdict-red"
    if "疑点" in verdict or "可疑" in verdict:
        return "verdict-yellow"
    return "verdict-green"


def _check_source_text(c):
    """尽量提取LLM给出的原文支撑/证据字段。兼容不同JSON字段名。"""
    for k in ("source_text", "quote", "original_text", "原文", "原文摘录", "evidence"):
        v = c.get(k)
        if isinstance(v, (list, tuple)):
            v = "；".join(str(x) for x in v if x)
        if v:
            return str(v)
    return ""


def _sanitize_reason_text(text):
    """Reason/detail fields are prose only; strip nested JSON and table markup noise."""
    raw = str(text or "").strip()
    if not raw:
        return ""

    extracted = []
    try:
        parsed = json.loads(raw)
    except Exception:
        parsed = None

    if isinstance(parsed, dict):
        for key in ("summary", "reason", "detail", "analysis", "explanation", "conclusion"):
            value = parsed.get(key)
            if isinstance(value, str) and value.strip():
                extracted.append(value)
        for check in parsed.get("checks", []) if isinstance(parsed.get("checks"), list) else []:
            if not isinstance(check, dict):
                continue
            for key in ("reason", "detail", "analysis", "explanation"):
                value = check.get(key)
                if isinstance(value, str) and value.strip():
                    extracted.append(value)
    elif raw.lstrip().startswith("{") and ('"checks"' in raw or '"summary"' in raw):
        for key in ("summary", "reason", "detail", "analysis", "explanation", "conclusion"):
            for match in re.finditer(rf'"{key}"\s*:\s*"((?:\\.|[^"\\])*)"', raw):
                extracted.append(_json_string_unescape(match.group(1)))

    text = " ".join(extracted) if extracted else raw
    had_table_noise = bool(re.search(r"\[\[TABLE_|<\s*/?\s*t[rdh]\b|&lt;\s*/?\s*t[rdh]\b", text, flags=re.I))
    text = html.unescape(text)
    text = _clean_mineru_table_block(text)
    text = re.sub(r"\[\[/?(?:BLOCK|FIGURE)[^\]]*\]\]", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\b(?:source_text|evidence|checks|summary|risk_level|detection_score|verdict)\b\s*[:：]?", " ", text)
    text = re.sub(r"\s+", " ", text).strip(" ,;，；。")

    if had_table_noise:
        prefix = "表格原文已在证据区渲染；此处仅保留文字判断。"
        if text:
            return prefix + " " + _brief_text(text, 520)
        return prefix + " 请人工核对原PDF表格。"
    return _brief_text(text, 700)


def _check_reason(c):
    """提取可疑原因/细节，优先detail，其次reason/explanation。"""
    for k in ("detail", "reason", "analysis", "explanation", "说明"):
        v = c.get(k)
        if v:
            return _sanitize_reason_text(v)
    return ""


def build_audit_action_items(report, meta, stat_result, limit=8):
    items = []
    evidence_chain_audit = (meta or {}).get("evidence_chain_audit") or {}
    strong_clusters = [cluster for cluster in (evidence_chain_audit.get("clusters") or []) if cluster.get("severity") == "strong"]
    if strong_clusters:
        top = strong_clusters[0]
        items.append({
            "score": 285,
            "source": "证据链与证据簇审查",
            "title": f"{len(strong_clusters)}个强证据簇需优先复核",
            "detail": _brief_text(
                f"{top.get('title')}: {top.get('summary')}。优先核对 Methods、Results、图表、补充材料和现有审查信号是否指向同一证据链。",
                220,
            ),
            "anchor": "evidence-chain-clusters",
        })
    elif evidence_chain_audit.get("cluster_count"):
        items.append({
            "score": 238,
            "source": "证据链与证据簇审查",
            "title": f"{evidence_chain_audit.get('cluster_count')}个证据簇需复核",
            "detail": "证据簇用于把孤立疑点合并为可人工核查的问题组；建议优先查看中等以上证据簇。",
            "anchor": "evidence-chain-clusters",
        })
    checks = sorted(report.get("checks", []) if isinstance(report, dict) else [], key=_check_sort_key)
    for check_idx, c in enumerate(checks, 1):
        if not _is_suspicious_check(c):
            continue
        items.append({
            "score": _check_suspicion_score(c),
            "source": "LLM语义审查",
            "title": f"{c.get('category', 'N/A')} / {c.get('item', 'N/A')}",
            "detail": _brief_text(_check_reason(c) or _check_source_text(c) or "需人工复核。", 180),
            "anchor": f"check-{check_idx}",
        })
    if stat_result.get("benford_status") and "高偏差" in str(stat_result.get("benford_status")):
        items.append({
            "score": 260,
            "source": "本地统计",
            "title": "Benford分布偏差较高",
            "detail": f"偏差={round(stat_result.get('benford_deviation') or 0, 3)}，建议核对原始数值来源和批量生成痕迹。",
            "anchor": "local-statistics",
        })
    reference_audit = meta.get("reference_audit") or {}
    if reference_audit.get("online_enabled"):
        bad_refs = []
        refs_by_index = {i + 1: ref for i, ref in enumerate(reference_audit.get("references", []))}
        for issue in reference_audit.get("issues", []):
            ref = refs_by_index.get(issue.get("index"), {})
            online = ref.get("online") or {}
            if online.get("online_status") in {"not_found", "weak", "error"}:
                bad_refs.append(issue.get("index"))
        if bad_refs:
            items.append({
                "score": 240,
                "source": "参考文献在线检索",
                "title": f"{len(bad_refs)}条参考文献在线证据不足",
                "detail": f"优先核对编号: {bad_refs[:10]}；检查DOI、题名、年份是否与数据库命中一致。",
                "anchor": "reference-audit",
            })
    resource_audit = meta.get("resource_audit") or {}
    if resource_audit.get("issues"):
        issues = resource_audit.get("issues") or []
        items.append({
            "score": 235,
            "source": "资源可用性校检",
            "title": f"{len(issues)}项代码仓库/在线资源需复核",
            "detail": "优先核对不可访问、格式错误或访问受限的代码仓库、Streamlit等论文声明资源。",
            "anchor": "resource-audit",
        })
    cross_file_audit = meta.get("cross_file_consistency_audit") or {}
    cross_findings = cross_file_audit.get("findings") or []
    if cross_findings:
        strong = cross_file_audit.get("strong_count", 0)
        medium = cross_file_audit.get("medium_count", 0)
        score = 255 if strong else 225
        items.append({
            "score": score,
            "source": "跨文件一致性审查",
            "title": f"{len(cross_findings)}项跨文件一致性疑点",
            "detail": f"强证据冲突 {strong} 项，中等疑点 {medium} 项；优先核对正文、补充材料和数据表中的样本量/分组/图表编号。",
            "anchor": "cross-file-consistency",
        })
    image_audit = meta.get("image_audit") or {}
    image_warnings = [img for img in image_audit.get("images", []) if img.get("risk") == "local_warning"]
    semantic_warnings = []
    detector_warnings = []
    for img in image_audit.get("images", []):
        sem = img.get("semantic") or {}
        if sem.get("reasonability") in {"需人工核对", "可疑"} or sem.get("status") == "error":
            semantic_warnings.append(img)
        detector = img.get("detector") or {}
        score = detector.get("score")
        if detector.get("status") == "ok" and score is not None and score >= 50:
            detector_warnings.append(img)
    if image_warnings or semantic_warnings or detector_warnings:
        items.append({
            "score": 230,
            "source": "图像检测",
            "title": f"{len(image_warnings)}张本地异常 / {len(semantic_warnings)}张需语义复核 / {len(detector_warnings)}张AI概率偏高",
            "detail": "查看 image_ai_review_manifest.html 中的自动imagedetector结果，并核对图像语义分析描述是否与论文图注一致。",
            "anchor": "image-audit",
        })
    items.sort(key=lambda item: (-item["score"], item["source"], item["title"]))
    selected = []
    used_sources = set()
    for item in items:
        if item["source"] in used_sources:
            continue
        selected.append(item)
        used_sources.add(item["source"])
        if len(selected) >= limit:
            return selected
    for item in items:
        if item in selected:
            continue
        selected.append(item)
        if len(selected) >= limit:
            break
    return selected


def format_audit_action_summary_markdown(report, meta, stat_result):
    items = build_audit_action_items(report, meta, stat_result)
    lines = ["## 🎯 行动优先级摘要", ""]
    if not items:
        lines.append("> 未形成高优先级行动项；仍建议按报告逐项抽查原文、图表和引用。")
        lines.append("")
        return lines
    lines.append("| 优先级 | 来源 | 事项 | 复核建议 |")
    lines.append("|--------|------|------|----------|")
    for idx, item in enumerate(items, 1):
        title = item["title"]
        if item.get("anchor"):
            title = f"[{title}](#{item['anchor']})"
        lines.append(
            f"| {idx} | {_md_escape_cell(item['source'])} | {_md_escape_cell(title)} | "
            f"{_md_escape_cell(item['detail'])} |"
        )
    lines.append("")
    return lines


def format_audit_action_summary_html(report, meta, stat_result):
    items = build_audit_action_items(report, meta, stat_result)
    if not items:
        content = '<p class="section-hint">未形成高优先级行动项；仍建议按报告逐项抽查原文、图表和引用。</p>'
    else:
        cards = ""
        for idx, item in enumerate(items, 1):
            title_html = _html_escape(item["title"])
            if item.get("anchor"):
                title_html = f'<a href="#{_html_escape(item["anchor"])}">{title_html}</a>'
            cards += f"""
            <div class="action-card">
              <span class="action-rank">#{idx}</span>
              <div>
                <strong>{title_html}</strong>
                <p>{_html_escape(item['detail'])}</p>
              </div>
              <span class="action-source">{_html_escape(item['source'])}</span>
            </div>"""
        content = f'<div class="action-list">{cards}</div>'
    return f"""
  <div class="section action-section">
    <h2>行动优先级摘要</h2>
    {content}
  </div>"""


def build_review_overview(report, meta, stat_result, action_limit=3):
    meta = meta or {}
    report = report or {}
    stat_result = stat_result or {}
    breakdown = report.get("score_breakdown") or {}
    checks = report.get("checks", []) if isinstance(report, dict) else []
    suspicious_count = sum(1 for check in checks if _is_suspicious_check(check))
    evidence_warnings = breakdown.get("evidence_warnings")
    if evidence_warnings is None:
        evidence_warnings = suspicious_count
    extraction_warnings = breakdown.get("extraction_warnings", 0)
    artifact_type = meta.get("artifact_type") or "complete"
    llm_coverage = meta.get("llm_coverage") or (
        f"{meta.get('llm_success_chunks')}/{meta.get('chunk_count')}"
        if meta.get("chunk_count") and meta.get("llm_success_chunks") is not None
        else "单块/未分块"
    )
    artifacts = meta.get("artifact_paths") or {}
    return {
        "report_type": artifact_type,
        "report_type_label": "范围受限审查 (limited)" if artifact_type == "limited" else "完整审查 (complete)",
        "risk_level": report.get("risk_level", "未知"),
        "detection_score": report.get("detection_score", 0),
        "red_flags": breakdown.get("red_flags", suspicious_count),
        "evidence_warnings": evidence_warnings,
        "extraction_warnings": extraction_warnings,
        "p_value_warnings": stat_result.get("p_value_abnormal", 0),
        "llm_coverage": llm_coverage,
        "llm_failed_chunks": meta.get("llm_failed_chunks") or [],
        "top_actions": build_audit_action_items(report, meta, stat_result, limit=action_limit),
        "artifact_paths": artifacts,
        "limited_reasons": meta.get("limited_reasons") or [],
        "completed": [
            "文本提取",
            "本地统计",
            "LLM语义审查",
            *("参考文献校检" for _ in [1] if meta.get("reference_audit") is not None),
            *("资源可用性校检" for _ in [1] if meta.get("resource_audit") is not None),
            *("跨文件一致性审查" for _ in [1] if meta.get("cross_file_consistency_audit") is not None),
            *("证据链审查" for _ in [1] if meta.get("evidence_chain_audit") is not None),
            *("图像审查" for _ in [1] if meta.get("image_audit") is not None),
        ],
    }


def format_review_overview_markdown(report, meta, stat_result):
    overview = build_review_overview(report, meta, stat_result)
    lines = [
        '<a id="review-overview"></a>',
        "## 复核概览",
        "",
        "| 项目 | 内容 |",
        "|------|------|",
        f"| 报告类型 | {_md_escape_cell(overview['report_type_label'])} |",
        f"| 复核优先级 | {_md_escape_cell(overview['risk_level'])} |",
        f"| 证据风险分 | {overview['detection_score']} / 100 |",
        f"| 红旗/证据/提取警告 | {overview['red_flags']} / {overview['evidence_warnings']} / {overview['extraction_warnings']} |",
        f"| LLM覆盖 | {_md_escape_cell(str(overview['llm_coverage']))} |",
        f"| 失败分块 | {_md_escape_cell(str(overview['llm_failed_chunks'] or '无'))} |",
    ]
    if overview["artifact_paths"]:
        paths = "；".join(f"{key}: {value}" for key, value in overview["artifact_paths"].items())
        lines.append(f"| 产物路径 | {_md_escape_cell(paths)} |")
    lines.append("")
    if overview["top_actions"]:
        lines.append("**Top 3 复核动作**")
        lines.append("")
        for idx, item in enumerate(overview["top_actions"], 1):
            title = item["title"]
            if item.get("anchor"):
                title = f"[{title}](#{item['anchor']})"
            lines.append(f"{idx}. {title} - {item['detail']}")
        lines.append("")
    if overview["report_type"] == "limited":
        lines.append('<a id="limitation-panel"></a>')
        lines.append("### 范围限制面板")
        lines.append("")
        lines.append("**限制原因**: " + ("；".join(overview["limited_reasons"]) if overview["limited_reasons"] else "未记录具体限制原因"))
        lines.append("")
        lines.append("**已完成**: " + "；".join(overview["completed"]))
        lines.append("")
    return lines


def format_review_overview_html(report, meta, stat_result):
    overview = build_review_overview(report, meta, stat_result)
    actions = ""
    for idx, item in enumerate(overview["top_actions"], 1):
        title = _html_escape(item["title"])
        if item.get("anchor"):
            title = f'<a href="#{_html_escape(item["anchor"])}">{title}</a>'
        actions += f"""
        <li><span>#{idx}</span><strong>{title}</strong><p>{_html_escape(item.get('detail', ''))}</p></li>"""
    if not actions:
        actions = "<li><strong>暂无高优先级动作</strong><p>建议按报告章节抽查原文、图表和引用。</p></li>"
    artifacts = ""
    for key, value in (overview.get("artifact_paths") or {}).items():
        artifacts += f"<li><strong>{_html_escape(key)}</strong>: {_html_escape(value)}</li>"
    if not artifacts:
        artifacts = "<li>最终产物路径将在写入后记录。</li>"
    limitation_panel = ""
    if overview["report_type"] == "limited":
        reasons = overview["limited_reasons"] or ["未记录具体限制原因"]
        reason_items = "".join(f"<li>{_html_escape(reason)}</li>" for reason in reasons)
        completed_items = "".join(f"<li>{_html_escape(item)}</li>" for item in overview["completed"])
        limitation_panel = f"""
  <div class="section limitation-panel" id="limitation-panel">
    <h2>范围限制面板</h2>
    <div class="limitation-grid">
      <div><strong>限制来源</strong><ul>{reason_items}</ul></div>
      <div><strong>已完成审查</strong><ul>{completed_items}</ul></div>
    </div>
  </div>"""
    return f"""
  <div class="section review-overview" id="review-overview">
    <h2>复核概览</h2>
    <div class="overview-grid">
      <div><span>报告类型</span><strong>{_html_escape(overview['report_type_label'])}</strong></div>
      <div><span>复核优先级</span><strong>{_html_escape(overview['risk_level'])}</strong></div>
      <div><span>证据风险分</span><strong>{overview['detection_score']} / 100</strong></div>
      <div><span>红旗 / 证据警告 / 提取警告</span><strong>{overview['red_flags']} / {overview['evidence_warnings']} / {overview['extraction_warnings']}</strong></div>
      <div><span>LLM覆盖</span><strong>{_html_escape(overview['llm_coverage'])}</strong></div>
      <div><span>失败分块</span><strong>{_html_escape(overview['llm_failed_chunks'] or '无')}</strong></div>
    </div>
    <div class="overview-columns">
      <div>
        <h3>Top 3 复核动作</h3>
        <ol class="overview-actions">{actions}</ol>
      </div>
      <div>
        <h3>产物路径</h3>
        <ul class="artifact-list">{artifacts}</ul>
      </div>
    </div>
  </div>
  {limitation_panel}"""


def _report_action_context(report, pdf_path, meta, stat_result):
    meta = meta or {}
    paper_identity = meta.get("paper_identity") or {}
    checks = sorted(report.get("checks", []) if isinstance(report, dict) else [], key=_check_sort_key)
    suspicious = [c for c in checks if _is_suspicious_check(c)]
    selected = suspicious[:10] if suspicious else checks[:8]
    issues = []
    for idx, c in enumerate(selected, 1):
        issues.append({
            "id": f"issue-{idx}",
            "source": "audit",
            "category": c.get("category", ""),
            "item": c.get("item", ""),
            "verdict": c.get("verdict", ""),
            "evidence": _brief_text(_clean_mineru_table_block(_check_source_text(c)), 900),
            "reason": _brief_text(_check_reason(c), 900),
        })
    cross_file_audit = (meta or {}).get("cross_file_consistency_audit") or {}
    cross_file_issues = []
    for idx, finding in enumerate((cross_file_audit.get("findings") or [])[:8], 1):
        cross_file_issue = {
            "id": f"cross-file-{idx}",
            "source": "cross_file_consistency",
            "category": "跨文件一致性审查",
            "item": finding.get("conflict_type", ""),
            "verdict": _cross_file_severity_label(finding.get("severity")),
            "evidence": _brief_text(
                f"{finding.get('claim_source_label')} / {finding.get('claim_file')}: {finding.get('claim_excerpt')} "
                f"|| {finding.get('counter_source_label')} / {finding.get('counter_file')}: {finding.get('counter_excerpt')}",
                900,
            ),
            "reason": _brief_text(finding.get("reason") or finding.get("manual_check"), 900),
        }
        cross_file_issues.append(cross_file_issue)
    issues = cross_file_issues + issues
    evidence_chain_audit = (meta or {}).get("evidence_chain_audit") or {}
    evidence_chain_issues = []
    for idx, cluster in enumerate((evidence_chain_audit.get("clusters") or [])[:10], 1):
        evidence_chain_issues.append({
            "id": cluster.get("id") or f"evidence-cluster-{idx}",
            "source": "evidence_chain_audit",
            "category": "证据链与证据簇审查",
            "item": cluster.get("title", ""),
            "verdict": _cross_file_severity_label(cluster.get("severity")),
            "evidence": _brief_text(cluster.get("summary", ""), 900),
            "reason": _brief_text(
                f"来源: {', '.join(cluster.get('source_types') or [])}；证据数: {cluster.get('evidence_count', 0)}。",
                900,
            ),
            "default_selected": cluster.get("severity") == "strong",
        })
    if evidence_chain_issues:
        issues = evidence_chain_issues + issues
    reference_audit = (meta or {}).get("reference_audit") or {}
    ref_issues = []
    for issue in (reference_audit.get("issues") or [])[:8]:
        ref_issues.append({
            "index": issue.get("index"),
            "issues": issue.get("issues", []),
            "text": _brief_text(_clean_reference_text(issue.get("text", "")), 500),
        })
    image_audit = (meta or {}).get("image_audit") or {}
    image_issues = []
    for img in (image_audit.get("images") or [])[:8]:
        sem = img.get("semantic") or {}
        detector = img.get("detector") or {}
        if img.get("risk") == "local_warning" or sem.get("reasonability") in {"需人工核对", "可疑"} or (detector.get("score") or 0) >= 50:
            image_issues.append({
                "file": img.get("file"),
                "local_issues": img.get("issues", []),
                "semantic": _brief_text(sem.get("summary", ""), 360),
                "detector_score": detector.get("score"),
            })
    resource_audit = (meta or {}).get("resource_audit") or {}
    resource_issues = []
    for issue in (resource_audit.get("issues") or [])[:8]:
        resource_issues.append({
            "index": issue.get("index"),
            "url": issue.get("url"),
            "type": issue.get("type"),
            "status": issue.get("status"),
            "problem": issue.get("problem"),
        })
    return {
        "paper": str(pdf_path),
        "artifact_type": meta.get("artifact_type") or meta.get("report_type") or "complete",
        "limited_reasons": meta.get("limited_reasons") or [],
        "artifact_paths": meta.get("artifact_paths") or {},
        "followups_dir": meta.get("followups_dir") or str((Path((meta.get("artifact_paths") or {}).get("html") or pdf_path).parent / "followups")),
        "paper_identity": {
            "title": _brief_text(paper_identity.get("title", ""), 300),
            "journal": _brief_text(paper_identity.get("journal", ""), 220),
            "authors": [
                _brief_text(author, 120)
                for author in (paper_identity.get("authors") or [])
                if str(author or "").strip()
            ][:8],
            "doi": _brief_text(paper_identity.get("doi", ""), 120),
            "year": _brief_text(paper_identity.get("year", ""), 20),
        },
        "summary": _brief_text(report.get("summary", ""), 1200) if isinstance(report, dict) else "",
        "risk_level": report.get("risk_level", "") if isinstance(report, dict) else "",
        "detection_score": report.get("detection_score", "") if isinstance(report, dict) else "",
        "conclusion": _brief_text(report.get("conclusion", ""), 1200) if isinstance(report, dict) else "",
        "top_issues": issues,
        "cross_file_consistency": {
            "status": cross_file_audit.get("status"),
            "checked_files": cross_file_audit.get("checked_files"),
            "finding_count": cross_file_audit.get("finding_count"),
            "findings": cross_file_issues,
        },
        "evidence_chain_audit": {
            "status": evidence_chain_audit.get("status"),
            "cluster_count": evidence_chain_audit.get("cluster_count"),
            "finding_count": evidence_chain_audit.get("finding_count"),
            "strong_count": evidence_chain_audit.get("strong_count"),
            "clusters": evidence_chain_issues,
        },
        "stat": {
            "number_count": stat_result.get("number_count"),
            "p_value_count": stat_result.get("p_value_count"),
            "p_value_abnormal": stat_result.get("p_value_abnormal"),
            "number_consistency": stat_result.get("number_consistency"),
            "benford_status": stat_result.get("benford_status"),
        },
        "references": {
            "reference_count": reference_audit.get("reference_count"),
            "online_checked": reference_audit.get("online_checked"),
            "issues": ref_issues,
        },
        "resources": {
            "resource_count": resource_audit.get("resource_count"),
            "online_checked": resource_audit.get("online_checked"),
            "issues": resource_issues,
        },
        "images": image_issues,
    }


def format_web_action_panel_html(report, pdf_path, meta, stat_result):
    meta = meta or {}
    service = meta.get("report_actions") or {}
    host = service.get("host") or "127.0.0.1"
    port = int(service.get("port") or 8765)
    service_url = report_action_service_url(host, port)
    generate_url = f"{service_url}/generate"
    followups_url = f"{service_url}/followups"
    context = _report_action_context(report, pdf_path, meta, stat_result or {})
    context_json = _json_for_script_tag(context)
    generate_url_json = _json_for_script_tag(generate_url)
    followups_url_json = _json_for_script_tag(followups_url)
    service_url_json = _json_for_script_tag(service_url)
    startup_command_json = _json_for_script_tag(f"python paper_audit.py --serve-report-actions --report-actions-port {port}")
    return f"""
  <div class="section web-action-section">
    <h2>一键生成后续沟通草稿</h2>
    <p class="section-hint">草稿由本地配置的LLM生成。生成前请确认文章身份、证据范围和语气；生成后仍需人工核对。</p>
    <div class="identity-grid" aria-label="文章身份确认">
      <label>标题<input id="followup-title" type="text" placeholder="文章标题"></label>
      <label>期刊<input id="followup-journal" type="text" placeholder="期刊"></label>
      <label>作者<input id="followup-authors" type="text" placeholder="作者，逗号分隔"></label>
      <label>DOI<input id="followup-doi" type="text" placeholder="DOI"></label>
      <label>年份<input id="followup-year" type="text" placeholder="年份"></label>
    </div>
    <div class="web-action-toolbar">
      <label class="inline-control">语言
        <select id="draft-language" class="draft-language-select" aria-label="草稿语言">
          <option value="zh">中文</option>
          <option value="en">English</option>
        </select>
      </label>
      <label class="inline-control">语气
        <select id="draft-tone" class="draft-language-select" aria-label="草稿语气">
          <option value="conservative">保守</option>
          <option value="standard">标准</option>
          <option value="firm">强硬</option>
        </select>
      </label>
      <button type="button" class="action-button" data-action-kind="pubpeer_comment">生成 PubPeer Comment</button>
      <button type="button" class="action-button" data-action-kind="journal_letter">生成期刊 Letter</button>
      <button type="button" class="secondary-button" id="copy-generated-draft">复制草稿</button>
    </div>
    <div class="evidence-picker">
      <strong>写入草稿的证据</strong>
      <div id="followup-evidence-list" class="evidence-choice-list"></div>
    </div>
    <label class="custom-concern-label">自定义关注点
      <textarea id="custom-followup-concerns" class="custom-concern-input" placeholder="每行一个人工补充关注点，会标记为 user_added。"></textarea>
    </label>
    <label class="manual-confirmation">
      <input id="manual-review-confirmation" type="checkbox">
      我已确认文章身份、证据选择和语气设置；生成内容仅作为基于阅读和理解文章后的学术问题表达草稿，发送前仍需人工复核。
    </label>
    <div id="existing-followups" class="existing-followups"></div>
    <div id="web-action-status" class="web-action-status">动作服务: <code>{_html_escape(service_url)}</code></div>
    <textarea id="generated-draft" class="generated-draft" spellcheck="false" placeholder="生成的草稿会显示在这里，可直接编辑。"></textarea>
  </div>
  <script id="paper-audit-action-context" type="application/json">{context_json}</script>
  <script>
  (function() {{
    const statusEl = document.getElementById('web-action-status');
    const outputEl = document.getElementById('generated-draft');
    const contextEl = document.getElementById('paper-audit-action-context');
    const languageEl = document.getElementById('draft-language');
    const toneEl = document.getElementById('draft-tone');
    const evidenceEl = document.getElementById('followup-evidence-list');
    const existingEl = document.getElementById('existing-followups');
    const confirmationEl = document.getElementById('manual-review-confirmation');
    const actionLabels = {{
      pubpeer_comment: 'PubPeer comment',
      journal_letter: 'journal letter'
    }};
    const languageLabels = {{ zh: '中文', en: 'English' }};
    const generateUrl = {generate_url_json};
    const followupsUrl = {followups_url_json};
    const serviceUrl = {service_url_json};
    const startupCommand = {startup_command_json};
    let reportContext = {{}};
    function esc(value) {{
      return String(value == null ? '' : value).replace(/[&<>"']/g, function(ch) {{
        return {{ '&': '&amp;', '<': '&lt;', '>': '&gt;', '"': '&quot;', "'": '&#39;' }}[ch];
      }});
    }}
    function setStatus(text, isError) {{
      statusEl.textContent = text;
      statusEl.className = 'web-action-status' + (isError ? ' error' : '');
    }}
    function setStatusHtml(html, isError) {{
      statusEl.innerHTML = html;
      statusEl.className = 'web-action-status' + (isError ? ' error' : '');
    }}
    function serviceFailure(prefix, err) {{
      setStatusHtml(
        esc(prefix) + '：本机动作服务未响应。服务地址: <code>' + esc(serviceUrl) + '</code>。请运行: <code>' +
        esc(startupCommand) + '</code>，然后刷新本页面或重新点击生成。详情: ' + esc(err && err.message ? err.message : err),
        true
      );
    }}
    function readContext() {{
      try {{
        reportContext = JSON.parse(contextEl.textContent || '{{}}');
        return reportContext;
      }} catch (err) {{
        setStatus('无法读取报告上下文: ' + err.message, true);
        return {{}};
      }}
    }}
    function populateIdentity(identity) {{
      identity = identity || {{}};
      document.getElementById('followup-title').value = identity.title || '';
      document.getElementById('followup-journal').value = identity.journal || '';
      document.getElementById('followup-authors').value = Array.isArray(identity.authors) ? identity.authors.join(', ') : (identity.authors || '');
      document.getElementById('followup-doi').value = identity.doi || '';
      document.getElementById('followup-year').value = identity.year || '';
    }}
    function identityFromForm() {{
      return {{
        title: document.getElementById('followup-title').value.trim(),
        journal: document.getElementById('followup-journal').value.trim(),
        authors: document.getElementById('followup-authors').value.split(/[,;，；]/).map(function(x) {{ return x.trim(); }}).filter(Boolean),
        doi: document.getElementById('followup-doi').value.trim(),
        year: document.getElementById('followup-year').value.trim()
      }};
    }}
    function renderEvidence(context) {{
      const issues = Array.isArray(context.top_issues) ? context.top_issues : [];
      if (!issues.length) {{
        evidenceEl.innerHTML = '<p class="section-hint">没有可自动勾选的高优先级证据；可在自定义关注点中补充。</p>';
        return;
      }}
      evidenceEl.innerHTML = issues.map(function(issue, idx) {{
        const verdict = String(issue.verdict || '');
        const checked = (issue.default_selected || verdict.indexOf('红旗') >= 0 || verdict.indexOf('高') >= 0 || verdict.indexOf('强证据') >= 0 || idx < 3) ? ' checked' : '';
        const label = [issue.category, issue.item, issue.verdict].filter(Boolean).join(' · ');
        const detail = issue.reason || issue.evidence || '';
        return '<label class="evidence-choice"><input type="checkbox" data-issue-index="' + idx + '"' + checked + '> <span><strong>' +
          esc(label || ('证据 ' + (idx + 1))) + '</strong><small>' + esc(detail).slice(0, 220) + '</small></span></label>';
      }}).join('');
    }}
    function selectedIssues() {{
      const issues = Array.isArray(reportContext.top_issues) ? reportContext.top_issues : [];
      return Array.from(evidenceEl.querySelectorAll('input[type="checkbox"]:checked')).map(function(input) {{
        return issues[Number(input.getAttribute('data-issue-index'))];
      }}).filter(Boolean);
    }}
    function customConcerns() {{
      return document.getElementById('custom-followup-concerns').value.split(/\\n+/).map(function(x) {{ return x.trim(); }}).filter(Boolean);
    }}
    function renderExisting(data) {{
      const drafts = (data && data.drafts) || {{}};
      const kinds = Object.keys(drafts);
      if (!kinds.length) {{
        existingEl.textContent = '当前语言暂无已生成草稿。';
        return;
      }}
      existingEl.innerHTML = '已生成: ' + kinds.map(function(kind) {{
        return '<button type="button" class="secondary-button existing-draft-button" data-existing-kind="' + esc(kind) + '">' + esc(actionLabels[kind] || kind) + '</button>';
      }}).join(' ');
      existingEl.querySelectorAll('[data-existing-kind]').forEach(function(btn) {{
        btn.addEventListener('click', function() {{
          const kind = btn.getAttribute('data-existing-kind');
          outputEl.value = drafts[kind] && drafts[kind].text ? drafts[kind].text : '';
          setStatus('已载入已生成的 ' + (actionLabels[kind] || kind) + '。', false);
        }});
      }});
    }}
    async function loadExisting() {{
      const context = readContext();
      if (context.artifact_type === 'failed') {{
        setStatus('失败诊断报告不允许生成 PubPeer Comment 或期刊 Letter；请先修复关键服务后重新生成审查报告。', true);
        document.querySelectorAll('[data-action-kind]').forEach(function(btn) {{ btn.disabled = true; }});
        return;
      }}
      try {{
        const resp = await fetch(followupsUrl, {{
          method: 'POST',
          headers: {{ 'Content-Type': 'application/json' }},
          body: JSON.stringify({{ context: context, language: (languageEl && languageEl.value) || 'zh' }})
        }});
        const data = await resp.json();
        if (!resp.ok || !data.ok) {{ throw new Error(data.error || ('HTTP ' + resp.status)); }}
        if (data.identity && (data.identity.title || data.identity.journal || data.identity.authors)) {{
          populateIdentity(data.identity);
        }}
        renderExisting(data);
      }} catch (err) {{
        serviceFailure('读取已生成草稿失败', err);
      }}
    }}
    async function generate(kind) {{
      const context = readContext();
      if (context.artifact_type === 'failed') {{
        setStatus('失败诊断报告不允许生成 PubPeer Comment 或期刊 Letter。', true);
        return;
      }}
      if (!confirmationEl.checked) {{
        setStatus('请先勾选人工复核确认，再生成外部沟通草稿。', true);
        return;
      }}
      const language = (languageEl && languageEl.value) || 'zh';
      setStatus('正在生成 ' + languageLabels[language] + ' ' + actionLabels[kind] + ' ...', false);
      outputEl.value = '';
      try {{
        const resp = await fetch(generateUrl, {{
          method: 'POST',
          headers: {{ 'Content-Type': 'application/json' }},
          body: JSON.stringify({{
            kind,
            context,
            language,
            tone: (toneEl && toneEl.value) || 'conservative',
            identity: identityFromForm(),
            selected_issues: selectedIssues(),
            custom_concerns: customConcerns(),
            disclaimer_confirmed: true
          }})
        }});
        const data = await resp.json();
        if (!resp.ok || !data.ok) {{
          throw new Error(data.error || ('HTTP ' + resp.status));
        }}
        outputEl.value = data.text || '';
        const path = data.paths && data.paths.draft_path ? ' 已保存: ' + data.paths.draft_path : '';
        setStatus('已生成 ' + actionLabels[kind] + '。请人工核对后再使用。' + path, false);
        loadExisting();
      }} catch (err) {{
        serviceFailure('生成失败', err);
      }}
    }}
    document.querySelectorAll('[data-action-kind]').forEach((btn) => {{
      btn.addEventListener('click', () => generate(btn.getAttribute('data-action-kind')));
    }});
    document.getElementById('copy-generated-draft').addEventListener('click', async () => {{
      try {{
        await navigator.clipboard.writeText(outputEl.value || '');
        setStatus('草稿已复制到剪贴板。', false);
      }} catch (err) {{
        outputEl.select();
        setStatus('浏览器未允许自动复制，请手动复制文本框内容。', true);
      }}
    }});
    const initialContext = readContext();
    populateIdentity(initialContext.paper_identity || {{}});
    renderEvidence(initialContext);
    if (languageEl) {{ languageEl.addEventListener('change', loadExisting); }}
    loadExisting();
  }})();
  </script>"""


def format_report(report, pdf_path, meta, stat_result):
    """将审查结果格式化为Markdown报告"""
    meta = normalize_run_meta(meta, pdf_path)
    risk_icons = {"高": "🔴", "中": "🟡", "低": "🟢", "严重证据冲突": "⚫️"}
    artifact_type = meta.get("artifact_type") or "complete"
    artifact_label = "范围受限审查 (limited)" if artifact_type == "limited" else "完整审查 (complete)"
    runtime = meta.get("runtime") or {}
    lines = [
        f"# 📄 学术论文审查报告 [耿同学标准]",
        f"",
        f"**文件**: `{pdf_path}`",
        f"**产物类型**: {artifact_label}",
        f"**版本**: prompt={meta.get('prompt_version', PROMPT_VERSION)}；schema={meta.get('schema_version', SCHEMA_VERSION)}；adapter={meta.get('adapter_version', ADAPTER_VERSION)}；rules={meta.get('risk_rule_version', report.get('rule_version', RISK_RULE_VERSION))}",
        f"**文件大小**: {meta.get('size_mb', 'N/A')} MB",
        f"**提取字符数**: {meta.get('total_chars', meta.get('chars', 'N/A'))}",
        f"**提取方式**: {meta.get('extraction_method', meta.get('source', 'N/A'))}",
    ]
    if meta.get("limited_reasons"):
        lines.append(f"**范围限制**: {'；'.join(meta.get('limited_reasons') or [])}")
    # 显示分块信息（如果是分块审查）
    if meta.get("chunk_count") and meta["chunk_count"] > 1:
        lines.append(f"**审查方式**: 分块审查 | {meta['chunk_count']}块 | 单块上限{meta['chunk_size']}字符 | 重叠{meta['overlap']}字符")
    if meta.get("llm_coverage"):
        failed_chunks = meta.get("llm_failed_chunks") or []
        if meta.get("llm_partial_report") or failed_chunks:
            lines.append(f"**LLM覆盖率**: ⚠️ 部分报告，仅成功审查 {meta.get('llm_coverage')} 个分块；失败块: {failed_chunks or '无'}")
            lines.append("> ⚠️ 本报告只基于成功返回的LLM分块合并，未覆盖失败分块全文；结论只能作为阶段性结果，建议稍后用 `--llm-cache-only` 或更稳定API补跑。")
        else:
            lines.append(f"**LLM覆盖率**: ✅ {meta.get('llm_coverage')} 个分块全部成功")
    lines.append(f"**审查时间**: {runtime.get('local_time') or time.strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"**运行时UTC年份**: {runtime.get('utc_year', runtime_utc_year())}（用于未来发表年份等非LLM日期判断）")

    if not report.get("parse_error"):
        lines.extend([""])
        lines.extend(format_review_overview_markdown(report, meta, stat_result))

    lines.extend([
        f"",
        f'<a id="local-statistics"></a>',
        f"## 📊 本地统计检测结果",
        f"| 检测项 | 结果 | 状态 |",
        f"|--------|------|------|",
        f"| Benford分布偏差 | {round(stat_result['benford_deviation'],3) if stat_result['benford_deviation'] else '样本不足'} | {stat_result['benford_status'] or 'N/A'} |",
        f"| p值数量/异常 | {stat_result['p_value_count']} / {stat_result['p_value_abnormal']}个>0.05 | {'⚠️异常' if stat_result['p_value_abnormal'] else '✅正常'} |",
        f"| 标准差提及 | {stat_result['sd_count']}处 | N/A |",
    ])
    lines.append(f"| 提取数字数 | {stat_result['number_count']} | - |")

    if stat_result.get("number_consistency"):
        lines.append(f"| 数字自洽性 | {stat_result['number_consistency']} | ⚠️矛盾 |")

    lines.append("")

    if report.get("parse_error"):
        lines.append("## ⚠️ LLM报告解析失败（原始输出）")
        lines.append(f"```\n{report['raw_output']}\n```")
        return "\n".join(lines)

    lines.append(f"## 总评: {report.get('summary', 'N/A')}")
    risk = report.get('risk_level', '未知')
    lines.append(f"**复核优先级**: {risk_icons.get(risk, '⚪')} {risk}")
    lines.append(f"**证据风险分**: {report.get('detection_score', 0)} / 100 (辅助排序指标，越高表示越需要优先复核)")
    breakdown = report.get("score_breakdown") or {}
    if breakdown:
        lines.append(
            "**计分拆解**: "
            f"红旗 {breakdown.get('red_flags', 0)}；"
            f"证据型疑点 {breakdown.get('evidence_warnings', 0)}；"
            f"提取质量疑点 {breakdown.get('extraction_warnings', 0)}；"
            f"统计调整 {', '.join(breakdown.get('stat_adjustments') or []) or '无'}"
        )
    lines.append("")
    lines.extend(format_audit_action_summary_markdown(report, meta, stat_result))

    checks = sorted(report.get("checks", []), key=_check_sort_key)
    if checks:
        suspicious = [c for c in checks if _is_suspicious_check(c)]
        lines.append('<a id="suspicious-findings"></a>')
        lines.append("## 🚩 可疑点证据汇总表")
        lines.append("")
        if suspicious:
            lines.append("| # | 判定 | 来源类型 | 分类/检查项 | 原文证据摘录 | 可疑原因 |")
            lines.append("|---|------|----------|-------------|--------------|----------|")
            for i, c in enumerate(suspicious[:5], 1):
                cat_item = f"{c.get('category', 'N/A')} / {c.get('item', 'N/A')}"
                lines.append(
                    f"| {i} | {_md_escape_cell(c.get('verdict', 'N/A'))} | {_md_escape_cell(' + '.join(_check_source_tags(c)))} | {_md_escape_cell(cat_item)} | "
                    f"{_md_escape_cell(_brief_text(_check_source_text(c), 220) or '未提供明确原文摘录')} | "
                    f"{_md_escape_cell(_brief_text(_check_reason(c), 220) or '未提供详细原因')} |"
                )
            if len(suspicious) > 5:
                lines.append("")
                lines.append(f"> 仅显示 Top 5 可疑点；完整 {len(suspicious)} 条见下方全部检查项和 HTML/JSON。")
        else:
            lines.append("> 未发现红旗/疑点项；仍建议人工核验关键数据、图表和引用。")
        lines.append("")

        lines.append('<a id="all-checks"></a>')
        lines.append("## 🔍 全部检查项概览")
        lines.append("")
        lines.append("| # | 分类 | 检查项 | 判定 | 证据摘要 |")
        lines.append("|---|------|--------|------|----------|")
        for i, c in enumerate(checks, 1):
            lines.append(f"| {i} | {_md_escape_cell(c.get('category', 'N/A'))} | {_md_escape_cell(c.get('item', 'N/A'))} | {_md_escape_cell(c.get('verdict', 'N/A'))} | {_md_escape_cell(_brief_text(_check_source_text(c), 120) or '-')} |")
        lines.append("")

        lines.append('<a id="finding-details"></a>')
        lines.append("## 📋 逐条详细分析（含原文支撑）")
        lines.append("")
        for i, c in enumerate(checks, 1):
            lines.append(f'<a id="check-{i}"></a>')
            lines.append(f"### {i}. {c.get('category', 'N/A')} - {c.get('item', 'N/A')} — {c.get('verdict', 'N/A')}")
            source = _check_source_text(c)
            reason = _check_reason(c)
            if source:
                lines.append(f"> **原文/证据摘录**: {source}")
            else:
                lines.append("> **原文/证据摘录**: LLM未提供明确原文摘录，请人工回查对应段落。")
            if reason:
                lines.append(f"\n**可疑原因/详细说明**：{reason}")
            merged_summary = _merged_group_summary_text(c)
            if merged_summary:
                lines.append(f"\n**相近疑点统合**：{merged_summary}。完整成员见 HTML 展开区或 JSON `merged_group.members`。")
            lines.append("")

    if report.get("conclusion"):
        lines.append("## 📝 综合结论")
        lines.append(f"\n{report['conclusion']}")
        lines.append("")

    lines.extend(format_evidence_chain_audit_markdown(meta.get("evidence_chain_audit")))
    lines.extend(format_image_audit_markdown(meta.get("image_audit")))
    lines.extend(format_cross_file_consistency_markdown(meta.get("cross_file_consistency_audit")))
    lines.extend(format_resource_audit_markdown(meta.get("resource_audit")))
    lines.extend(format_reference_audit_markdown(meta.get("reference_audit")))

    return "\n".join(lines)


def format_html_report(report, pdf_path, meta, stat_result):
    """将审查结果格式化为紧凑、可审阅的HTML报告"""
    meta = normalize_run_meta(meta, pdf_path)
    risk_colors = {"高": "#b42318", "中": "#a16207", "低": "#166534", "严重证据冲突": "#111827"}
    risk_icons = {"高": "高复核优先级", "中": "中复核优先级", "低": "低复核优先级", "严重证据冲突": "严重证据冲突"}
    risk = report.get('risk_level', '未知')
    risk_color = risk_colors.get(risk, "#6b7280")
    risk_icon = risk_icons.get(risk, "未知风险")
    artifact_type = meta.get("artifact_type") or "complete"
    artifact_label = "范围受限审查 (limited)" if artifact_type == "limited" else "完整审查 (complete)"
    artifact_badge = "范围受限 limited" if artifact_type == "limited" else "完整审查 complete"
    artifact_color = "#8a5a00" if artifact_type == "limited" else "#166534"
    runtime = meta.get("runtime") or {}
    limited_notice = ""
    if meta.get("limited_reasons"):
        limited_notice = f"""
  <div class="section coverage-warning">
    <h2>范围受限审查</h2>
    <p>{_html_escape('；'.join(meta.get('limited_reasons') or []))}</p>
  </div>"""

    # 统计检测状态
    benford_val = round(stat_result['benford_deviation'], 3) if stat_result['benford_deviation'] else '样本不足'
    benford_status = stat_result.get('benford_status', 'N/A') or 'N/A'
    p_abnormal = stat_result['p_value_abnormal']
    p_status_class = "status-warn" if p_abnormal else "status-ok"

    # 分块信息
    chunk_info = ""
    if meta.get("chunk_count") and meta["chunk_count"] > 1:
        chunk_info = f"""
        <div class="meta-item">
            <span>审查方式</span>
            <strong>分块审查 · {meta['chunk_count']}块 · 单块{meta['chunk_size']}字符 · 重叠{meta['overlap']}字符</strong>
        </div>"""

    # 数字自洽性
    number_consistency = ""
    if stat_result.get("number_consistency"):
        number_consistency = f"""
        <tr>
            <td>数字自洽性</td>
            <td>{stat_result['number_consistency']}</td>
            <td><span class="status-warn">⚠️ 矛盾</span></td>
        </tr>"""

    # LLM覆盖率/部分报告提示
    coverage_banner = ""
    if meta.get("llm_coverage"):
        failed_chunks = meta.get("llm_failed_chunks") or []
        if meta.get("llm_partial_report") or failed_chunks:
            coverage_banner = f"""
  <div class="section coverage-warning">
    <h2>LLM覆盖不足</h2>
    <p><strong>成功审查分块</strong>: {_html_escape(meta.get('llm_coverage'))}</p>
    <p><strong>失败块</strong>: {_html_escape(failed_chunks or '无')}</p>
    <p>本报告只基于成功返回的LLM分块合并，未覆盖失败分块全文；结论只能作为阶段性结果。建议稍后使用 <code>--llm-cache-only</code> 复用成功缓存，或切换更稳定API补跑。</p>
  </div>"""
        else:
            coverage_banner = f"""
  <div class="section coverage-ok">
    <h2>LLM覆盖率</h2>
    <p>{_html_escape(meta.get('llm_coverage'))} 个分块全部成功。</p>
  </div>"""
    action_summary_html = format_audit_action_summary_html(report, meta, stat_result) if not report.get("parse_error") else ""
    review_overview_html = format_review_overview_html(report, meta, stat_result) if not report.get("parse_error") else ""
    resource_audit_html = format_resource_audit_html(meta.get("resource_audit"))
    reference_audit_html = format_reference_audit_html(meta.get("reference_audit"))
    image_audit_html = format_image_audit_html(meta.get("image_audit"))
    cross_file_audit_html = format_cross_file_consistency_html(meta.get("cross_file_consistency_audit"))
    evidence_chain_audit_html = format_evidence_chain_audit_html(meta.get("evidence_chain_audit"))
    web_action_panel_html = format_web_action_panel_html(report, pdf_path, meta, stat_result) if not report.get("parse_error") else ""
    breakdown = report.get("score_breakdown") or {}
    score_breakdown_html = ""
    if breakdown:
        score_breakdown_html = f"""
      <div class="score-breakdown">
        红旗 {breakdown.get('red_flags', 0)} · 证据型疑点 {breakdown.get('evidence_warnings', 0)} · 提取质量疑点 {breakdown.get('extraction_warnings', 0)} · 统计调整 {_html_escape(', '.join(breakdown.get('stat_adjustments') or []) or '无')}
      </div>"""
    summary_text = _html_escape(report.get('summary', 'N/A'))
    extracted_chars = meta.get('total_chars', meta.get('chars', 'N/A'))
    extraction_method = meta.get('extraction_method', meta.get('source', 'N/A'))

    # 解析失败
    if report.get("parse_error"):
        checks_html = f"""
        <div class="section">
            <h2>LLM报告解析失败（原始输出）</h2>
            <pre class="error-block">{_html_escape(report.get('raw_output', ''))}</pre>
        </div>"""
        conclusion_html = ""
    else:
        # 可疑点证据汇总表 + 全部检查概览 + 详细分析
        checks = sorted(report.get("checks", []), key=_check_sort_key)
        suspicious = [c for c in checks if _is_suspicious_check(c)]

        suspicious_items = ""
        for i, c in enumerate(suspicious[:5], 1):
            verdict = c.get('verdict', 'N/A')
            verdict_class = _check_verdict_class(verdict)
            cat_item = f"{c.get('category', 'N/A')} / {c.get('item', 'N/A')}"
            source = _check_source_text(c)
            reason = _check_reason(c)
            brief = _brief_text(reason or source or "未提供详细原因", 120)
            suspicion_score = _check_suspicion_score(c)
            source_tags = " + ".join(_check_source_tags(c))
            merged_html = _merged_group_html(c)
            suspicious_items += f"""
            <details class="suspicion-card" id="suspicious-finding-{i}">
                <summary class="suspicion-summary">
                    <span class="suspicion-rank">#{i}</span>
                    <span class="{verdict_class} suspicion-verdict">{_html_escape(verdict)}</span>
                    <span class="suspicion-title">{_html_escape(cat_item)}</span>
                    <span class="suspicion-score">复核分 {suspicion_score}</span>
                    <span class="suspicion-brief"><strong>{_html_escape(source_tags)}</strong> · {_html_escape(brief)}</span>
                    <span class="summary-action">查看详情</span>
                </summary>
                <div class="suspicion-body">
                    {merged_html}
                    <div class="detail-evidence"><strong>原文/证据摘录</strong>{render_evidence_html(source or 'LLM未提供明确原文摘录，请人工回查对应段落。')}</div>
                    <div class="detail-text"><strong>可疑原因/详细说明</strong><p>{_html_escape(reason or 'LLM未提供详细说明。')}</p></div>
                </div>
            </details>"""
        if len(suspicious) > 5:
            suspicious_items += f'<div class="muted">仅显示 Top 5；完整 {len(suspicious)} 条见下方全部检查项。</div>'
        if not suspicious_items:
            suspicious_items = '<div class="muted">未发现红旗/疑点项；仍建议人工核验关键数据、图表和引用。</div>'

        checks_table_rows = ""
        for i, c in enumerate(checks, 1):
            verdict = c.get('verdict', 'N/A')
            verdict_class = _check_verdict_class(verdict)
            checks_table_rows += f"""
            <tr>
                <td>{i}</td>
                <td>{_html_escape(c.get('category', 'N/A'))}</td>
                <td>{_html_escape(c.get('item', 'N/A'))}</td>
                <td><span class="{verdict_class}">{_html_escape(verdict)}</span></td>
                <td class="evidence-cell">{render_evidence_summary_html(_check_source_text(c), 120)}</td>
            </tr>"""

        detail_cards = ""
        for i, c in enumerate(checks, 1):
            verdict = c.get('verdict', 'N/A')
            verdict_class = _check_verdict_class(verdict)
            source = _check_source_text(c)
            reason = _check_reason(c)
            source_html = render_evidence_html(source or 'LLM未提供明确原文摘录，请人工回查对应段落。')
            merged_html = _merged_group_html(c)
            detail_cards += f"""
            <details class="detail-card" id="check-{i}">
                <summary class="detail-header detail-summary">
                    <span class="detail-num">#{i}</span>
                    <span class="detail-cat">{_html_escape(c.get('category', 'N/A'))}</span>
                    <span class="detail-item">{_html_escape(c.get('item', 'N/A'))}</span>
                    <span class="{verdict_class} detail-verdict">{_html_escape(verdict)}</span>
                    <span class="detail-brief">{_html_escape(_brief_text(reason or source or '无摘要', 120))}</span>
                    <span class="summary-action">查看详情</span>
                </summary>
                <div class="detail-body">
                    {merged_html}
                    <div class="detail-evidence"><strong>原文/证据摘录</strong>{source_html}</div>
                    <div class="detail-text"><strong>可疑原因/详细说明</strong><p>{_html_escape(reason or 'LLM未提供详细说明。')}</p></div>
                </div>
            </details>"""

        checks_html = f"""
        <div class="section evidence-summary" id="suspicious-findings">
            <h2>Top 可复核证据</h2>
            <p class="section-hint">按可复核性和证据强度排序；默认显示前5项，展开后查看原文证据、判断理由和相近疑点统合来源。</p>
            <div class="suspicion-list">{suspicious_items}</div>
        </div>
        <div class="section" id="all-checks">
            <h2>全部检查项概览</h2>
            <table class="checks-table">
                <thead><tr><th>#</th><th>分类</th><th>检查项</th><th>判定</th><th>证据摘要</th></tr></thead>
                <tbody>{checks_table_rows}</tbody>
            </table>
        </div>
        <div class="section" id="finding-details">
            <h2>逐条详细分析（含原文支撑）</h2>
            {detail_cards}
        </div>"""

        conclusion_html = ""
        if report.get("conclusion"):
            conclusion_html = f"""
            <div class="section conclusion-section">
                <h2>综合结论</h2>
                <p class="conclusion-text">{_html_escape(report['conclusion'])}</p>
            </div>"""

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>学术论文审查报告</title>
<style>
  :root {{
    --bg: #f7f3ec;
    --paper: #fffdf8;
    --surface: #fffaf1;
    --surface2: #f3eadc;
    --text: #2b241d;
    --text-muted: #7b7065;
    --accent: #c15f3c;
    --accent2: #2f6f73;
    --border: #e4d8c8;
    --red: #b42318;
    --yellow: #b7791f;
    --green: #2f7d50;
    --shadow: 0 22px 70px rgba(61, 45, 31, 0.10);
  }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: ui-sans-serif, -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Noto Sans SC', sans-serif;
    background: #f6f6f6;
    color: var(--text);
    line-height: 1.6;
    padding: 28px 20px;
  }}
  .container {{ max-width: 1120px; margin: 0 auto; }}
  .header {{
    background: rgba(255,253,248,0.86);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 36px;
    margin-bottom: 28px;
    text-align: center;
    box-shadow: none;
  }}
  .header h1 {{ font-size: 32px; margin-bottom: 18px; letter-spacing: 0; }}
  .risk-badge {{
    display: inline-block;
    font-size: 22px;
    font-weight: 700;
    padding: 8px 24px;
    border-radius: 999px;
    color: #fff;
    background: {risk_color};
    margin: 8px 0;
  }}
  .artifact-badge {{
    display: inline-block;
    border-radius: 6px;
    padding: 7px 10px;
    color: #fff;
    font-weight: 800;
    font-size: 13px;
    line-height: 1;
    white-space: nowrap;
  }}
  .meta-grid {{
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 8px 24px;
    text-align: left;
    margin-top: 16px;
    font-size: 14px;
    color: var(--text-muted);
  }}
  .meta-grid strong {{ color: var(--text); }}
  .section {{
    background: rgba(255,253,248,0.88);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 24px;
    margin-bottom: 22px;
    box-shadow: 0 10px 34px rgba(61, 45, 31, 0.06);
  }}
  .section h2 {{
    font-size: 20px;
    margin-bottom: 16px;
    padding-bottom: 8px;
    border-bottom: 1px solid var(--border);
  }}
  table {{ width: 100%; border-collapse: collapse; font-size: 14px; }}
  th, td {{ padding: 10px 12px; text-align: left; border-bottom: 1px solid var(--border); }}
  th {{ background: var(--surface2); color: var(--text-muted); font-weight: 650; text-transform: uppercase; font-size: 12px; }}
  tr:hover {{ background: rgba(193,95,60,0.05); }}
  .status-ok {{ color: var(--green); font-weight: 600; }}
  .status-warn {{ color: var(--yellow); font-weight: 600; }}
  .verdict-red {{ color: var(--red); font-weight: 700; }}
  .verdict-yellow {{ color: var(--yellow); font-weight: 700; }}
  .verdict-green {{ color: var(--green); font-weight: 700; }}
  .detail-card {{
    background: var(--surface2);
    border-radius: 8px;
    padding: 16px;
    margin-bottom: 12px;
    border-left: 4px solid var(--accent);
  }}
  .detail-card[open] {{ background: #f8efe2; }}
  .detail-header {{ display: flex; align-items: center; gap: 12px; flex-wrap: wrap; }}
  .detail-summary {{ cursor: pointer; list-style: none; }}
  .detail-summary::-webkit-details-marker {{ display: none; }}
  .detail-summary::after {{
    content: "展开详情";
    margin-left: auto;
    color: var(--accent);
    font-size: 12px;
    font-weight: 600;
  }}
  .detail-card[open] .detail-summary::after {{ content: "收起详情"; }}
  .section-hint {{ color: var(--text-muted); margin: -6px 0 14px; font-size: 14px; }}
  .evidence-summary {{ border-left: 4px solid var(--red); }}
  .suspicion-list {{
    display: flex;
    flex-direction: column;
    gap: 10px;
  }}
  .suspicion-card {{
    background: #fff8ed;
    border: 1px solid var(--border);
    border-left: 4px solid var(--red);
    border-radius: 8px;
    padding: 12px 14px;
  }}
  .suspicion-summary {{
    cursor: pointer;
    list-style: none;
    display: grid;
    grid-template-columns: 44px minmax(76px, 120px) minmax(180px, 1fr) auto;
    gap: 8px 12px;
    align-items: center;
  }}
  .suspicion-summary::-webkit-details-marker {{ display: none; }}
  .suspicion-summary::after {{
    content: "展开详情";
    grid-column: 4;
    grid-row: 2;
    justify-self: end;
    color: var(--accent);
    font-size: 12px;
    font-weight: 700;
  }}
  .suspicion-card[open] .suspicion-summary::after {{ content: "收起详情"; }}
  .suspicion-rank {{
    background: rgba(248,113,113,0.16);
    color: var(--red);
    border-radius: 999px;
    padding: 4px 8px;
    text-align: center;
    font-weight: 800;
    font-size: 13px;
  }}
  .suspicion-title {{
    font-weight: 700;
    color: var(--text);
    overflow-wrap: anywhere;
  }}
  .suspicion-score {{
    color: var(--text-muted);
    font-size: 12px;
    white-space: nowrap;
  }}
  .suspicion-brief {{
    grid-column: 2 / 4;
    color: var(--text-muted);
    font-size: 13px;
    line-height: 1.45;
    overflow-wrap: anywhere;
  }}
  .suspicion-body {{
    margin-top: 12px;
  }}
  .action-section {{
    border-left: 5px solid var(--accent);
  }}
  .action-list {{
    display: grid;
    gap: 10px;
  }}
  .action-card {{
    display: grid;
    grid-template-columns: 48px minmax(0, 1fr) 140px;
    gap: 12px;
    align-items: start;
    padding: 14px;
    background: #fff8ed;
    border: 1px solid var(--border);
    border-radius: 12px;
  }}
  .action-rank {{
    background: rgba(193,95,60,0.14);
    color: var(--accent);
    border-radius: 999px;
    padding: 4px 8px;
    font-weight: 800;
    text-align: center;
  }}
  .action-card p {{
    color: var(--text-muted);
    margin-top: 4px;
    font-size: 13px;
  }}
  .action-card a, .review-overview a {{
    color: var(--accent);
    text-decoration: none;
  }}
  .action-source {{
    color: var(--text-muted);
    font-size: 12px;
    text-align: right;
  }}
  .review-overview {{
    border-left: 5px solid var(--accent2);
  }}
  .overview-grid {{
    display: grid;
    grid-template-columns: repeat(3, minmax(0, 1fr));
    gap: 10px;
    margin-bottom: 16px;
  }}
  .overview-grid div {{
    border: 1px solid var(--border);
    background: #fff8ed;
    border-radius: 6px;
    padding: 10px;
  }}
  .overview-grid span {{
    display: block;
    color: var(--text-muted);
    font-size: 12px;
  }}
  .overview-grid strong {{
    display: block;
    margin-top: 3px;
    overflow-wrap: anywhere;
  }}
  .overview-columns, .limitation-grid {{
    display: grid;
    grid-template-columns: minmax(0, 1.2fr) minmax(0, 1fr);
    gap: 18px;
  }}
  .overview-actions {{
    padding-left: 22px;
  }}
  .overview-actions li {{
    margin-bottom: 10px;
  }}
  .overview-actions li span {{
    color: var(--text-muted);
    font-weight: 800;
    margin-right: 6px;
  }}
  .overview-actions p, .artifact-list {{
    color: var(--text-muted);
    font-size: 13px;
  }}
  .artifact-list {{
    padding-left: 18px;
    overflow-wrap: anywhere;
  }}
  .limitation-panel {{
    border-left: 5px solid var(--yellow);
  }}
  .checks-table {{ table-layout: fixed; }}
  .checks-table th:nth-child(1), .checks-table td:nth-child(1) {{ width: 44px; text-align: center; }}
  .checks-table th:nth-child(2), .checks-table td:nth-child(2) {{ width: 120px; }}
  .evidence-table th:nth-child(4), .evidence-table td:nth-child(4),
  .evidence-table th:nth-child(5), .evidence-table td:nth-child(5) {{ width: 28%; }}
  .evidence-cell, .reason-cell {{
    color: var(--text);
    line-height: 1.55;
    word-break: break-word;
    overflow-wrap: anywhere;
  }}
  .muted {{ color: var(--text-muted); text-align: center; padding: 18px; }}
  .detail-num {{
    background: var(--accent);
    color: var(--bg);
    border-radius: 50%;
    width: 28px; height: 28px;
    display: flex; align-items: center; justify-content: center;
    font-weight: 700; font-size: 13px;
  }}
  .detail-cat {{ color: var(--accent); font-weight: 600; }}
  .detail-item {{ flex: 1; }}
  .detail-verdict {{ font-size: 14px; }}
  .detail-brief {{
    flex-basis: 100%;
    color: var(--text-muted);
    font-size: 13px;
    padding-left: 40px;
  }}
  .detail-body {{ margin-top: 12px; }}
  .detail-evidence {{
    margin-top: 10px;
    padding: 10px 14px;
    background: rgba(183,121,31,0.10);
    border-radius: 6px;
    font-size: 14px;
    color: #6f4b12;
  }}
  .detail-evidence blockquote {{
    margin-top: 8px;
    white-space: pre-wrap;
    color: #6f4b12;
  }}
  .table-hint {{
    display: inline-block;
    color: var(--accent);
    font-weight: 600;
    margin-right: 8px;
  }}
  .summary-excerpt {{
    color: var(--text-muted);
    display: block;
    margin-top: 4px;
  }}
  .data-table-details {{
    margin-top: 8px;
  }}
  .data-table-details summary {{
    cursor: pointer;
    color: var(--accent);
    font-weight: 700;
    margin-bottom: 8px;
  }}
  .data-table-wrap {{
    margin-top: 8px;
    overflow-x: auto;
    border: 1px solid var(--border);
    border-radius: 6px;
    background: #fffdf8;
  }}
  .data-table {{
    width: max-content;
    min-width: 100%;
    border-collapse: collapse;
    font-size: 13px;
  }}
  .data-table th,
  .data-table td {{
    min-width: 96px;
    max-width: 320px;
    padding: 8px 10px;
    border: 1px solid var(--border);
    vertical-align: top;
    white-space: normal;
    color: var(--text);
  }}
  .data-table th {{
    position: sticky;
    top: 0;
    background: #efe3d3;
    color: var(--text);
    z-index: 1;
  }}
  .detail-text {{
    margin-top: 8px;
    font-size: 14px;
    color: var(--text-muted);
    white-space: pre-wrap;
  }}
  .conclusion-section {{ border-left: 4px solid var(--green); }}
  .conclusion-text {{ font-size: 16px; white-space: pre-wrap; color: var(--text); }}
  .coverage-warning {{
    border-left: 5px solid var(--yellow);
    background: rgba(183,121,31,0.10);
  }}
  .coverage-warning code {{ background: var(--surface2); padding: 2px 6px; border-radius: 4px; }}
  .coverage-ok {{
    border-left: 5px solid var(--green);
    background: rgba(47,125,80,0.08);
  }}
  .error-block {{
    background: rgba(180,35,24,0.08);
    border: 1px solid var(--red);
    border-radius: 8px;
    padding: 16px;
    white-space: pre-wrap;
    font-family: monospace;
    font-size: 13px;
    color: #fca5a5;
    overflow-x: auto;
  }}
  .score-bar {{
    height: 8px;
    background: var(--surface2);
    border-radius: 4px;
    overflow: hidden;
    margin-top: 8px;
  }}
  .score-fill {{
    height: 100%;
    border-radius: 4px;
    background: var(--red);
    transition: none;
  }}
  .score-breakdown {{
    margin-top: 6px;
    color: var(--text-muted);
    font-size: 13px;
  }}
  .reference-list {{
    display: flex;
    flex-direction: column;
    gap: 10px;
  }}
  .cross-file-list {{
    display: flex;
    flex-direction: column;
    gap: 10px;
  }}
  .cross-file-card {{
    border: 1px solid var(--border);
    border-left: 4px solid var(--yellow);
    border-radius: 10px;
    background: #fff8ed;
    padding: 12px 14px;
  }}
  .cross-file-card[open] {{
    background: #fbfbfb;
  }}
  .cross-file-summary {{
    cursor: pointer;
    list-style: none;
    display: grid;
    grid-template-columns: 46px minmax(110px, 160px) minmax(180px, 1fr);
    gap: 8px 12px;
    align-items: center;
  }}
  .cross-file-summary::-webkit-details-marker {{ display: none; }}
  .cross-file-rank {{
    color: var(--accent);
    font-weight: 800;
  }}
  .cross-file-severity {{
    font-weight: 800;
    border-radius: 999px;
    padding: 3px 8px;
    text-align: center;
    background: rgba(183,121,31,0.12);
  }}
  .cross-file-strong {{ color: var(--red); }}
  .cross-file-medium {{ color: var(--yellow); }}
  .cross-file-weak {{ color: var(--text-muted); }}
  .cross-file-title {{
    font-weight: 700;
    overflow-wrap: anywhere;
  }}
  .cross-file-reason {{
    grid-column: 2 / 4;
    color: var(--text-muted);
    font-size: 13px;
  }}
  .cross-file-body {{
    margin-top: 12px;
    color: var(--text-muted);
    font-size: 14px;
    display: grid;
    gap: 10px;
  }}
  .cross-file-body p {{
    margin-top: 4px;
    overflow-wrap: anywhere;
  }}
  .reference-card {{
    border: 1px solid var(--border);
    border-radius: 10px;
    background: #fff8ed;
    padding: 12px 14px;
  }}
  .reference-summary {{
    cursor: pointer;
    list-style: none;
    display: grid;
    grid-template-columns: 46px minmax(98px, 140px) minmax(220px, 1fr) auto;
    gap: 8px 12px;
    align-items: center;
  }}
  .reference-summary::-webkit-details-marker {{ display: none; }}
  .reference-index {{
    color: var(--accent);
    font-weight: 800;
  }}
  .reference-status {{
    font-weight: 800;
    border-radius: 999px;
    padding: 3px 8px;
    text-align: center;
    background: rgba(47,111,115,0.10);
  }}
  .reference-verified, .reference-likely {{ color: var(--green); }}
  .reference-weak, .reference-not_found, .reference-error {{ color: var(--red); }}
  .reference-title {{
    font-weight: 700;
    overflow-wrap: anywhere;
  }}
  .reference-confidence {{
    color: var(--text-muted);
    font-size: 12px;
    white-space: nowrap;
  }}
  .reference-issues {{
    grid-column: 2 / 5;
    color: var(--text-muted);
    font-size: 13px;
  }}
  .reference-body {{
    margin-top: 12px;
    color: var(--text-muted);
    font-size: 14px;
  }}
  .reference-matches {{
    margin-top: 8px;
    padding-left: 22px;
  }}
  .reference-matches li {{
    margin: 8px 0;
  }}
  .reference-matches a, .image-section a {{
    color: var(--accent);
    font-weight: 700;
  }}
  .image-table code {{
    color: var(--text-muted);
    white-space: normal;
    overflow-wrap: anywhere;
  }}
  .detector-hint {{
    display: inline-block;
    color: var(--accent);
    font-weight: 800;
    font-size: 12px;
    margin-bottom: 4px;
  }}
  .web-action-section {{
    border-left: 5px solid #2563eb;
  }}
  .web-action-toolbar {{
    display: flex;
    gap: 10px;
    flex-wrap: wrap;
    margin: 12px 0;
  }}
  .action-button, .secondary-button {{
    border: 1px solid #cbd5e1;
    background: #ffffff;
    color: var(--text-main);
    border-radius: 8px;
    padding: 9px 12px;
    font-weight: 800;
    cursor: pointer;
  }}
  .action-button {{
    background: #2563eb;
    border-color: #2563eb;
    color: #ffffff;
  }}
  .draft-language-select {{
    border: 1px solid #cbd5e1;
    background: #ffffff;
    color: var(--text-main);
    border-radius: 8px;
    padding: 9px 12px;
    font-weight: 800;
  }}
  .identity-grid {{
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(210px, 1fr));
    gap: 10px;
    margin: 12px 0;
  }}
  .identity-grid label, .custom-concern-label, .inline-control {{
    display: flex;
    flex-direction: column;
    gap: 5px;
    color: var(--text-muted);
    font-size: 12px;
    font-weight: 800;
  }}
  .identity-grid input, .custom-concern-input {{
    border: 1px solid #cbd5e1;
    border-radius: 8px;
    padding: 9px 10px;
    color: var(--text-main);
    background: #ffffff;
    font: inherit;
  }}
  .evidence-picker {{
    margin: 12px 0;
  }}
  .evidence-choice-list {{
    display: grid;
    gap: 8px;
    margin-top: 8px;
  }}
  .evidence-choice {{
    display: grid;
    grid-template-columns: 20px 1fr;
    gap: 8px;
    align-items: start;
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 9px;
    background: #ffffff;
  }}
  .evidence-choice small {{
    display: block;
    color: var(--text-muted);
    line-height: 1.45;
    margin-top: 3px;
  }}
  .custom-concern-input {{
    width: 100%;
    min-height: 70px;
    resize: vertical;
  }}
  .manual-confirmation {{
    display: grid;
    grid-template-columns: 20px 1fr;
    gap: 8px;
    align-items: start;
    margin: 10px 0;
    color: var(--text-main);
    font-size: 13px;
  }}
  .existing-followups {{
    color: var(--text-muted);
    font-size: 13px;
    margin: 8px 0;
  }}
  .action-button:hover, .secondary-button:hover {{
    filter: brightness(0.96);
  }}
  .web-action-status {{
    color: var(--text-muted);
    font-size: 13px;
    margin: 8px 0;
  }}
  .web-action-status.error {{
    color: #b42318;
    font-weight: 700;
  }}
  .generated-draft {{
    width: 100%;
    min-height: 260px;
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 12px;
    font: 14px/1.6 ui-monospace, SFMono-Regular, Menlo, Consolas, monospace;
    color: var(--text-main);
    background: #fff;
    resize: vertical;
  }}
  .footer {{
    text-align: center;
    color: var(--text-muted);
    font-size: 12px;
    margin-top: 32px;
    padding: 16px;
  }}
  /* Compact grayscale report skin: evidence first, decoration last. */
  :root {{
    --bg: #f6f6f6;
    --paper: #ffffff;
    --surface: #ffffff;
    --surface2: #eeeeee;
    --text: #171717;
    --text-muted: #666666;
    --accent: #111111;
    --accent2: #3f3f46;
    --border: #d4d4d4;
    --red: #b42318;
    --yellow: #8a5a00;
    --green: #166534;
    --shadow: none;
  }}
  body {{
    background: #f6f6f6;
    color: var(--text);
    line-height: 1.55;
    padding: 18px;
  }}
  .container {{ max-width: 1280px; }}
  .header, .section {{
    background: #ffffff;
    border: 1px solid var(--border);
    border-radius: 8px;
    box-shadow: none;
  }}
  .header {{
    padding: 18px;
    margin-bottom: 14px;
    text-align: left;
  }}
  .report-topline {{
    display: flex;
    justify-content: space-between;
    gap: 16px;
    align-items: flex-start;
    margin-bottom: 12px;
  }}
  .report-kicker {{
    color: var(--text-muted);
    font-size: 12px;
    font-weight: 800;
    letter-spacing: 0;
  }}
  .header h1 {{
    font-size: 24px;
    line-height: 1.2;
    margin: 3px 0 8px;
    letter-spacing: 0;
  }}
  .report-summary {{
    max-width: 820px;
    color: #333333;
    font-size: 14px;
  }}
  .risk-badge {{
    border-radius: 6px;
    padding: 7px 10px;
    margin: 0;
    font-size: 13px;
    line-height: 1;
    white-space: nowrap;
  }}
  .score-panel {{
    display: grid;
    grid-template-columns: 170px minmax(220px, 1fr);
    gap: 12px;
    align-items: center;
    padding: 12px 0;
    border-top: 1px solid var(--border);
    border-bottom: 1px solid var(--border);
  }}
  .score-value {{
    font-size: 34px;
    line-height: 1;
    font-weight: 900;
    color: {risk_color};
  }}
  .score-caption {{
    color: var(--text-muted);
    font-size: 12px;
    margin-top: 4px;
  }}
  .priority-label {{
    font-weight: 900;
    color: {risk_color};
    margin-bottom: 4px;
  }}
  .score-bar {{
    height: 7px;
    background: #e5e5e5;
    border-radius: 4px;
    margin-top: 0;
  }}
  .score-fill {{
    background: {risk_color};
    border-radius: 4px;
    transition: none;
  }}
  .score-breakdown {{
    margin-top: 8px;
    color: var(--text-muted);
    font-size: 12px;
  }}
  .meta-grid {{
    display: grid;
    grid-template-columns: repeat(4, minmax(0, 1fr));
    gap: 8px;
    margin-top: 12px;
  }}
  .meta-grid > div, .meta-item {{
    min-width: 0;
    border: 1px solid var(--border);
    border-radius: 6px;
    padding: 9px 10px;
    background: #fafafa;
    color: var(--text-muted);
    font-size: 12px;
  }}
  .meta-grid strong, .meta-item strong {{
    display: block;
    color: var(--text);
    font-size: 13px;
    margin-top: 3px;
    overflow-wrap: anywhere;
  }}
  .section {{
    padding: 16px;
    margin-bottom: 14px;
  }}
  .section h2 {{
    display: flex;
    align-items: center;
    gap: 8px;
    font-size: 17px;
    line-height: 1.25;
    margin-bottom: 12px;
    padding-bottom: 8px;
  }}
  .section-hint {{
    margin: -4px 0 12px;
    color: var(--text-muted);
    font-size: 13px;
  }}
  table {{ font-size: 13px; }}
  th, td {{
    padding: 8px 9px;
    border-bottom: 1px solid var(--border);
    vertical-align: top;
  }}
  th {{
    background: #f1f1f1;
    color: #444444;
    font-size: 11px;
    letter-spacing: 0;
  }}
  .action-card, .suspicion-card, .detail-card, .reference-card {{
    background: #ffffff;
    border: 1px solid var(--border);
    border-left: 3px solid #111111;
    border-radius: 6px;
    padding: 10px 12px;
  }}
  .cross-file-card {{
    background: #ffffff;
    border: 1px solid var(--border);
    border-left: 3px solid var(--yellow);
    border-radius: 6px;
    padding: 10px 12px;
  }}
  .evidence-summary {{ border-left: 3px solid var(--red); }}
  .suspicion-summary {{
    grid-template-columns: 44px minmax(72px, 112px) minmax(180px, 1fr) 96px;
    gap: 7px 10px;
  }}
  .suspicion-rank, .action-rank, .detail-num {{
    background: #f1f1f1;
    color: #111111;
    border: 1px solid var(--border);
    border-radius: 5px;
  }}
  .suspicion-brief {{
    grid-column: 2 / 4;
    color: var(--text-muted);
  }}
  .summary-action {{
    grid-column: 4;
    grid-row: 2;
    justify-self: end;
    color: #111111;
    font-size: 12px;
    font-weight: 800;
    white-space: nowrap;
  }}
  .detail-summary .summary-action {{
    grid-column: auto;
    grid-row: auto;
    margin-left: auto;
  }}
  .detail-card[open], .suspicion-card[open] {{
    background: #fbfbfb;
  }}
  .cross-file-card[open] {{
    background: #fbfbfb;
  }}
  .detail-summary::after, .suspicion-summary::after {{
    content: none;
  }}
  .detail-body, .suspicion-body {{
    border-top: 1px solid var(--border);
    margin-top: 10px;
    padding-top: 10px;
  }}
  .detail-evidence {{
    background: #f7f7f7;
    color: #222222;
    border: 1px solid var(--border);
    border-radius: 6px;
  }}
  .detail-evidence blockquote {{
    color: #222222;
  }}
  .detail-text {{
    color: #4b4b4b;
  }}
  .data-table-wrap {{
    background: #ffffff;
    border-radius: 6px;
  }}
  .data-table th {{
    background: #ededed;
  }}
  .merged-group {{
    border: 1px solid var(--border);
    border-radius: 6px;
    padding: 10px;
    margin-bottom: 10px;
    background: #fff;
  }}
  .merged-group summary {{
    cursor: pointer;
    font-weight: 800;
    color: #111111;
  }}
  .coverage-warning, .coverage-ok, .conclusion-section, .action-section, .web-action-section {{
    border-left-width: 3px;
    background: #ffffff;
  }}
  .web-action-section {{ border-left-color: #2563eb; }}
  .action-button, .secondary-button {{
    border-radius: 6px;
    padding: 8px 11px;
    color: #111111;
  }}
  .action-button {{ color: #ffffff; }}
  .generated-draft {{
    border-radius: 6px;
    min-height: 220px;
    color: #111111;
  }}
  .image-thumb {{
    max-width: 84px;
    max-height: 64px;
    object-fit: contain;
    border: 1px solid var(--border);
    border-radius: 4px;
    background: #ffffff;
  }}
  .muted-inline {{
    color: var(--text-muted);
    font-size: 12px;
  }}
  @media (max-width: 900px) {{
    body {{ padding: 10px; }}
    .report-topline {{ flex-direction: column; }}
    .score-panel {{ grid-template-columns: 1fr; }}
    .meta-grid {{ grid-template-columns: 1fr 1fr; }}
    .overview-grid {{ grid-template-columns: 1fr 1fr; }}
    .overview-columns, .limitation-grid {{ grid-template-columns: 1fr; }}
    .suspicion-summary, .reference-summary {{
      grid-template-columns: 40px minmax(70px, 100px) 1fr;
    }}
    .suspicion-score, .reference-confidence, .summary-action {{
      grid-column: 2 / 4;
      justify-self: start;
    }}
    .suspicion-brief, .reference-issues {{
      grid-column: 1 / 4;
    }}
  }}
  @media (max-width: 560px) {{
    .meta-grid {{ grid-template-columns: 1fr; }}
    .overview-grid {{ grid-template-columns: 1fr; }}
    .checks-table {{ table-layout: auto; }}
  }}
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <div class="report-topline">
      <div>
        <div class="report-kicker">Paper Audit / Veritas</div>
        <h1>学术论文审查报告</h1>
        <div class="report-summary">{summary_text}</div>
      </div>
      <div class="artifact-badge" style="background:{artifact_color};">{_html_escape(artifact_badge)}</div>
    </div>
    <div class="score-panel">
      <div>
        <div class="score-value">{report.get('detection_score', 0)}</div>
        <div class="score-caption">证据风险分 / 100，越高表示越需要优先复核</div>
      </div>
      <div>
        <div class="priority-label">复核优先级：{_html_escape(risk_icon)}</div>
        <div class="score-bar"><div class="score-fill" style="width:{min(report.get('detection_score', 0), 100)}%; background:{risk_color};"></div></div>
      {score_breakdown_html}
      </div>
    </div>
    <div class="meta-grid">
      <div><span>文件</span><strong>{_html_escape(pdf_path)}</strong></div>
      <div><span>产物类型</span><strong>{_html_escape(artifact_label)}</strong></div>
      <div><span>Prompt版本</span><strong>{_html_escape(meta.get('prompt_version', PROMPT_VERSION))}</strong></div>
      <div><span>Schema版本</span><strong>{_html_escape(meta.get('schema_version', SCHEMA_VERSION))}</strong></div>
      <div><span>Adapter版本</span><strong>{_html_escape(meta.get('adapter_version', ADAPTER_VERSION))}</strong></div>
      <div><span>规则版本</span><strong>{_html_escape(meta.get('risk_rule_version', report.get('rule_version', RISK_RULE_VERSION)))}</strong></div>
      <div><span>文件大小</span><strong>{meta.get('size_mb', 'N/A')} MB</strong></div>
      <div><span>提取字符数</span><strong>{extracted_chars}</strong></div>
      <div><span>提取方式</span><strong>{extraction_method}</strong></div>
      {chunk_info if chunk_info else ''}
      <div><span>审查时间</span><strong>{_html_escape(runtime.get('local_time') or time.strftime('%Y-%m-%d %H:%M:%S'))}</strong></div>
      <div><span>运行时UTC年份</span><strong>{_html_escape(runtime.get('utc_year', runtime_utc_year()))}</strong></div>
    </div>
  </div>

  {review_overview_html}
  {coverage_banner}
  {limited_notice}
  {action_summary_html}

  <div class="section" id="local-statistics">
    <h2>本地统计检测结果</h2>
    <table>
      <thead><tr><th>检测项</th><th>结果</th><th>状态</th></tr></thead>
      <tbody>
        <tr><td>Benford分布偏差</td><td>{benford_val}</td><td>{benford_status}</td></tr>
        <tr><td>p值数量/异常</td><td>{stat_result['p_value_count']} / {stat_result['p_value_abnormal']}个&gt;0.05</td><td><span class="{p_status_class}">{'⚠️异常' if p_abnormal else '✅正常'}</span></td></tr>
        <tr><td>标准差提及</td><td>{stat_result['sd_count']}处</td><td>N/A</td></tr>
        <tr><td>提取数字数</td><td>{stat_result['number_count']}</td><td>-</td></tr>
        {number_consistency}
      </tbody>
    </table>
  </div>

  {checks_html}
  {web_action_panel_html}
  {conclusion_html}
  {evidence_chain_audit_html}
  {image_audit_html}
  {cross_file_audit_html}
  {resource_audit_html}
  {reference_audit_html}

  <div class="footer">
    Generated by <strong>Veritas</strong> — 学术论文自动审查工具（耿同学标准） | {time.strftime('%Y-%m-%d %H:%M:%S')}
  </div>
</div>
</body>
</html>"""
    return html


def update_patterns(comments_file):
    """从PubPeer评论文本中用LLM提取新的欺诈模式，更新知识库
    
    comments_file: 包含PubPeer评论文本的文件路径
    """
    from datetime import datetime
    
    comments_path = Path(comments_file)
    if not comments_path.exists():
        print(f"❌ 评论文本文件不存在: {comments_path}")
        return 1
    
    with open(comments_path, "r", encoding="utf-8") as f:
        comments_text = f.read()
    
    if len(comments_text.strip()) < 20:
        print("❌ 评论文本内容过少，请提供更完整的PubPeer评论内容")
        return 1
    
    print(f"📖 已读取评论文本: {len(comments_text)}字符")
    print("🤖 正在用LLM分析评论，提取欺诈模式...")
    
    # 构建提取prompt
    extract_prompt = f"""分析以下来自PubPeer的学术评论，提取其中涉及的学术论文造假/可疑手法。

要求：
1. 每个造假手法提取为一个独立的模式条目
2. 按JSON数组格式输出，每个条目包含：id(英文大写下划线), category(分类), name(中文名), description(详细描述), detection_hint(检测提示), risk_level(高/中/低)
3. 只提取确实存在的造假手法，不要臆造
4. 合并相似的造假手法

PubPeer评论内容：
{comments_text}

输出格式：
[
  {{
    "id": "PATTERN_ID",
    "category": "图片与图表/数据与结果/方法论/结构与引用/作者与期刊",
    "name": "手法名称",
    "description": "手法描述",
    "detection_hint": "审查时如何检测此手法",
    "risk_level": "高/中/低"
  }}
]"""
    
    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": "你是一个学术论文打假专家，擅长从PubPeer评论中识别和归纳造假手法。"},
            {"role": "user", "content": extract_prompt}
        ],
        "temperature": 0.3,
    }
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {LLM_API_KEY}",
    }
    
    req = urllib.request.Request(
        LLM_API_URL,
        data=json.dumps(payload).encode(),
        headers=headers,
        method="POST",
    )
    
    try:
        resp = urllib.request.urlopen(req, timeout=60)
        result = json.loads(resp.read().decode("utf-8"))
        content = result["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"❌ LLM调用失败: {e}")
        return 1
    
    # 解析LLM输出的JSON
    json_match = re.search(r'\[[\s\S]*\]', content)
    if not json_match:
        print("❌ LLM未能输出有效的JSON格式，请重试")
        print(f"原始输出: {content[:500]}")
        return 1
    
    try:
        new_patterns = json.loads(json_match.group())
    except json.JSONDecodeError as e:
        print(f"❌ JSON解析失败: {e}")
        return 1
    
    if not new_patterns:
        print("⚠️ 未能从评论中提取到新的欺诈模式")
        return 0
    
    # 加载现有知识库
    if FRAUD_PATTERNS_PATH.exists():
        with open(FRAUD_PATTERNS_PATH, "r", encoding="utf-8") as f:
            kb_data = json.load(f)
        existing_ids = {p["id"] for p in kb_data.get("patterns", [])}
    else:
        kb_data = {"schema_version": "1.0", "last_updated": "", "contributors": ["community"], "patterns": []}
        existing_ids = set()
    
    # 去重合并
    added = 0
    for p in new_patterns:
        if p.get("id") and p["id"] not in existing_ids:
            kb_data["patterns"].append(p)
            existing_ids.add(p["id"])
            added += 1
            print(f"  ✅ 新增: [{p.get('risk_level','?')}] {p.get('name','?')}")
        else:
            print(f"  ⏭️ 跳过已存在: {p.get('name','?')} ({p.get('id','?')})")
    
    if added > 0:
        kb_data["last_updated"] = datetime.now().strftime("%Y-%m-%d")
        with open(FRAUD_PATTERNS_PATH, "w", encoding="utf-8") as f:
            json.dump(kb_data, f, ensure_ascii=False, indent=2)
        print(f"\n🎉 知识库已更新！新增{added}条模式，总计{len(kb_data['patterns'])}条")
    else:
        print("\n⚠️ 无新增模式，知识库未变更")
    
    return 0


# ══════════════════════════════════════════════════════════════
# 腾讯朱雀AI文本检测辅助功能
# ══════════════════════════════════════════════════════════════

ZHUQUE_URL = "https://matrix.tencent.com/ai-detect/"


def copy_to_clipboard(text: str) -> bool:
    """跨平台复制文本到系统剪贴板"""
    system = platform.system()
    try:
        if system == "Windows":
            # Windows: 使用clip命令
            process = subprocess.Popen(
                ["clip.exe"], stdin=subprocess.PIPE
            )
            process.communicate(text.encode("utf-16"))
            return process.returncode == 0
        elif system == "Darwin":  # macOS
            process = subprocess.Popen(
                ["pbcopy"], stdin=subprocess.PIPE
            )
            process.communicate(text.encode("utf-8"))
            return process.returncode == 0
        else:  # Linux
            # 优先尝试xclip，其次xsel
            for cmd in [["xclip", "-selection", "clipboard"], ["xsel", "--clipboard", "--input"]]:
                try:
                    process = subprocess.Popen(cmd, stdin=subprocess.PIPE)
                    process.communicate(text.encode("utf-8"))
                    if process.returncode == 0:
                        return True
                except FileNotFoundError:
                    continue
            return False
    except Exception as e:
        print(f"⚠️ 剪贴板写入失败: {e}")
        return False


def launch_zhuque_ai_detect(text: str):
    """启动腾讯朱雀AI文本检测：复制文本到剪贴板 → 打开检测页面 → 弹窗提醒"""
    print("\n" + "=" * 60)
    print("🤖 腾讯朱雀AI文本检测")
    print("=" * 60)

    # 1) 复制文本到剪贴板
    # 朱雀检测有字数限制，截取前8000字符
    detect_text = text[:8000]
    if len(text) > 8000:
        print(f"⚠️ 文本较长({len(text)}字符)，仅复制前8000字符到剪贴板（朱雀字数限制）")

    clip_ok = copy_to_clipboard(detect_text)
    if clip_ok:
        print("✅ 文本已复制到剪贴板")
    else:
        print("❌ 剪贴板写入失败，请手动复制论文文本")

    # 2) 打开浏览器
    print(f"🌐 正在打开朱雀AI检测页面...")
    webbrowser.open(ZHUQUE_URL)

    # 3) 弹窗提醒
    system = platform.system()
    try:
        if system == "Windows":
            import ctypes
            ctypes.windll.user32.MessageBoxW(
                0,
                "论文文本已复制到剪贴板！\n\n"
                "请在打开的朱雀AI检测页面中粘贴文本并点击检测。\n"
                "检测完成后，点击确定继续后续审查流程。",
                "🤖 朱雀AI文本检测",
                0x40  # MB_ICONINFORMATION
            )
        elif system == "Darwin":
            subprocess.run([
                "osascript", "-e",
                'display dialog "论文文本已复制到剪贴板！\n\n请在朱雀AI检测页面中粘贴文本并点击检测。\n检测完成后，点击确定继续后续审查流程。" '
                'buttons {"确定"} default button "确定" with title "🤖 朱雀AI文本检测" with icon note'
            ])
        else:  # Linux
            # 尝试zenity
            try:
                subprocess.run([
                    "zenity", "--info", "--title=🤖 朱雀AI文本检测", "--width=400",
                    "--text=论文文本已复制到剪贴板！\n\n请在朱雀AI检测页面中粘贴文本并点击检测。\n检测完成后，点击确定继续后续审查流程。"
                ])
            except FileNotFoundError:
                # 降级为终端提示
                input("\n⏸️ 论文文本已复制到剪贴板，请在浏览器中粘贴检测。\n检测完成后按回车继续...")
    except Exception:
        # 最终降级：终端等待
        input("\n⏸️ 论文文本已复制到剪贴板，请在浏览器中粘贴检测。\n检测完成后按回车继续...")

    print("✅ 朱雀AI检测流程结束，继续后续审查...")


# ──────────────────────────────────────────────────────────────
# AI图片检测（imagedetector.com）
# ──────────────────────────────────────────────────────────────

IMAGE_DETECT_URL = "https://imagedetector.com/"
IMAGE_DETECT_UPLOAD_BASE = "https://ai-image-detector-prod.nyc3.digitaloceanspaces.com"
GLM_IMAGE_MAX_BYTES = 5 * 1024 * 1024

# 支持的图片扩展名
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif", ".webp"}


def _dedupe_paths(paths):
    seen = set()
    result = []
    for path in paths:
        try:
            resolved = str(Path(path).resolve())
        except Exception:
            resolved = str(path)
        if resolved not in seen:
            seen.add(resolved)
            result.append(resolved)
    return result


def _image_output_dir(input_path: str):
    p = Path(input_path)
    if p.is_file():
        return p.parent / "_paper_audit_images"
    return p / "_paper_audit_images"


def _extract_images_from_mineru_zip(zip_path: Path, output_dir: Path) -> List[str]:
    images = []
    try:
        import zipfile
        output_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(zip_path) as zf:
            for name in zf.namelist():
                suffix = Path(name).suffix.lower()
                if suffix not in IMAGE_EXTENSIONS:
                    continue
                info = zf.getinfo(name)
                if info.file_size < MIN_IMAGE_BYTES:
                    continue
                out_name = f"{_safe_name(zip_path.stem)}_{_safe_name(Path(name).stem)}{suffix}"
                out_path = output_dir / out_name
                out_path.write_bytes(zf.read(name))
                images.append(str(out_path))
    except Exception as e:
        print(f"  ⚠️ MinerU图片提取失败 {zip_path}: {e}")
    return images


def collect_mineru_image_files(input_path: str, output_dir=None) -> List[str]:
    """Collect images saved inside MinerU zip artifacts generated by extraction."""
    p = Path(input_path)
    base = Path(output_dir) if output_dir else (p.parent if p.is_file() else p)
    zips = []
    if base.exists():
        zips.extend(base.glob("*.mineru.zip"))
    if p.is_dir():
        zips.extend(p.rglob("*.mineru.zip"))
    zips = _latest_mineru_zips(zips)
    images_dir = _image_output_dir(str(p))
    images = []
    for zip_path in _dedupe_paths(zips):
        images.extend(_extract_images_from_mineru_zip(Path(zip_path), images_dir))
    if zips:
        return _dedupe_paths(images)
    if images_dir.exists():
        for ext in IMAGE_EXTENSIONS:
            for image_path in images_dir.rglob(f"*{ext}"):
                try:
                    if image_path.stat().st_size >= MIN_IMAGE_BYTES:
                        images.append(str(image_path))
                except Exception:
                    continue
    return _dedupe_paths(images)


def _latest_mineru_zips(paths):
    latest = {}
    for path in paths:
        p = Path(path)
        name = p.name
        if not name.endswith(".mineru.zip"):
            continue
        key = name[:-len(".mineru.zip")]
        key = re.sub(r"\.[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$", "", key, flags=re.I)
        try:
            mtime = p.stat().st_mtime
        except Exception:
            mtime = 0
        current = latest.get(key)
        if current is None or mtime > current[0]:
            latest[key] = (mtime, str(p))
    return [item[1] for item in latest.values()]


def analyze_image_reasonability(image_path: str):
    """Lightweight local image sanity checks before external AI-image review."""
    result = {
        "path": str(image_path),
        "file": Path(image_path).name,
        "size_bytes": 0,
        "width": None,
        "height": None,
        "format": "",
        "risk": "needs_online_check",
        "issues": [],
    }
    try:
        path = Path(image_path)
        result["size_bytes"] = path.stat().st_size
        if result["size_bytes"] < MIN_IMAGE_BYTES:
            result["issues"].append("too_small")
        try:
            from PIL import Image, ImageStat
            with Image.open(path) as img:
                result["width"], result["height"] = img.size
                result["format"] = img.format or path.suffix.lstrip(".")
                if result["width"] < 120 or result["height"] < 120:
                    result["issues"].append("low_resolution")
                ratio = max(result["width"], result["height"]) / max(1, min(result["width"], result["height"]))
                if ratio > 8:
                    result["issues"].append("extreme_aspect_ratio")
                stat = ImageStat.Stat(img.convert("L").resize((128, 128)))
                if stat.stddev and stat.stddev[0] < 3:
                    result["issues"].append("near_blank_or_flat")
                if stat.stddev and stat.stddev[0] > 85:
                    result["issues"].append("very_high_noise_or_contrast")
        except ImportError:
            result["issues"].append("pillow_not_installed")
        except Exception as e:
            result["issues"].append(f"image_parse_error:{type(e).__name__}")
    except Exception as e:
        result["issues"].append(f"file_error:{type(e).__name__}")

    severe = {"low_resolution", "near_blank_or_flat", "image_parse_error:UnidentifiedImageError"}
    if any(issue in severe or issue.startswith("file_error") for issue in result["issues"]):
        result["risk"] = "local_warning"
    elif not result["issues"]:
        result["risk"] = "local_ok"
    return result


def _image_file_fingerprint(image_path: str):
    return _image_file_fingerprint_from_namespace(globals(), image_path)


def _image_semantic_cache_key(image_path: str, api_url=None, model=None, cache_version=None):
    return _image_semantic_cache_key_from_namespace(
        globals(),
        image_path,
        api_url=api_url,
        model=model,
        cache_version=cache_version,
    )


def _image_to_data_url(image_path: str):
    path = Path(image_path)
    try:
        from PIL import Image
        with Image.open(path) as img:
            img = img.convert("RGB")
            img.thumbnail((1400, 1400))
            side = max(512, img.width, img.height)
            side = min(side, 1600)
            canvas = Image.new("RGB", (side, side), "white")
            x = max(0, (side - img.width) // 2)
            y = max(0, (side - img.height) // 2)
            canvas.paste(img, (x, y))
            buf = io.BytesIO()
            canvas.save(buf, format="JPEG", quality=88, optimize=True)
            data = base64.b64encode(buf.getvalue()).decode("ascii")
            return f"data:image/jpeg;base64,{data}"
    except Exception:
        mime = mimetypes.guess_type(str(path))[0] or "image/png"
        data = base64.b64encode(path.read_bytes()).decode("ascii")
        return f"data:{mime};base64,{data}"


def _prepare_detector_upload_file(image_path: str):
    path = Path(image_path)
    mime = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
    if mime in {"image/jpeg", "image/png", "image/webp"}:
        return path.name, mime, path.read_bytes()
    try:
        from PIL import Image
        with Image.open(path) as img:
            img = img.convert("RGB")
            img.thumbnail((1600, 1600))
            buf = io.BytesIO()
            buf_name = f"{path.stem}.jpg"
            img.save(buf, format="JPEG", quality=90, optimize=True)
            return buf_name, "image/jpeg", buf.getvalue()
    except Exception:
        return path.name, mime, path.read_bytes()


class _ExternalCapabilityTimeout(BaseException):
    pass


def _run_with_alarm_timeout(func, timeout, timeout_result):
    """Bound third-party calls that may ignore socket timeouts."""
    try:
        seconds = max(1, int(timeout or 1))
    except Exception:
        seconds = 1
    if threading.current_thread() is not threading.main_thread() or not hasattr(signal, "SIGALRM"):
        return func()

    previous_handler = signal.getsignal(signal.SIGALRM)

    def _raise_timeout(signum, frame):
        raise _ExternalCapabilityTimeout(f"operation exceeded {seconds}s")

    signal.signal(signal.SIGALRM, _raise_timeout)
    signal.alarm(seconds)
    try:
        return func()
    except _ExternalCapabilityTimeout:
        return timeout_result()
    finally:
        signal.alarm(0)
        signal.signal(signal.SIGALRM, previous_handler)


def call_imagedetector(image_path: str, timeout=60):
    """Upload an image to imagedetector.com using the site's public web flow."""
    def _call():
        return _call_imagedetector_unbounded(image_path, timeout=timeout)

    return _run_with_alarm_timeout(_call, timeout, lambda: _detector_timeout_result(timeout))


def _call_imagedetector_unbounded(image_path: str, timeout=60):
    try:
        file_name, mime, content = _prepare_detector_upload_file(image_path)
        if len(content) < 1024:
            return {
                "status": "skipped",
                "provider": "imagedetector.com",
                "reason": "too_small",
                "summary": "图片小于imagedetector网页最小上传要求，跳过自动检测。",
            }
        if len(content) > 10 * 1024 * 1024:
            return {
                "status": "skipped",
                "provider": "imagedetector.com",
                "reason": "too_large",
                "summary": "图片超过imagedetector网页10MB限制，跳过自动检测。",
            }
        query = urllib.parse.urlencode({"fileName": file_name, "fileType": mime})
        headers = {
            "Accept": "application/json",
            "Referer": IMAGE_DETECT_URL,
            "User-Agent": "PaperAudit/1.0",
        }
        data, _ = _http_request(
            f"{IMAGE_DETECT_URL.rstrip('/')}/api/get-presigned-url?{query}",
            "GET",
            headers=headers,
            timeout=timeout,
        )
        upload_info = json.loads(data.decode("utf-8", errors="replace"))
        presigned_url = upload_info.get("presignedUrl")
        file_path = upload_info.get("filePath")
        expected_type = upload_info.get("expectedContentType") or mime
        if not presigned_url or not file_path:
            return {
                "status": "error",
                "provider": "imagedetector.com",
                "reason": "missing_upload_url",
                "summary": "imagedetector未返回可用上传地址。",
            }
        _http_request(
            presigned_url,
            "PUT",
            headers={"Content-Type": expected_type, "x-amz-acl": "private"},
            data=content,
            timeout=timeout,
        )
        image_url = f"{IMAGE_DETECT_UPLOAD_BASE.rstrip('/')}/{file_path.lstrip('/')}"
        detect_payload = json.dumps({"imageUrl": image_url}).encode("utf-8")
        data, _ = _http_request(
            f"{IMAGE_DETECT_URL.rstrip('/')}/api/detect",
            "POST",
            headers={
                "Content-Type": "application/json",
                "Accept": "application/json",
                "Referer": IMAGE_DETECT_URL,
                "User-Agent": "PaperAudit/1.0",
            },
            data=detect_payload,
            timeout=timeout,
        )
        return _normalize_detector_result(json.loads(data.decode("utf-8", errors="replace")))
    except Exception as e:
        return {
            "status": "error",
            "provider": "imagedetector.com",
            "reason": type(e).__name__,
            "summary": f"imagedetector自动检测失败：{type(e).__name__}",
        }


def call_glm_image_semantics(image_path: str, timeout=45, api_key=None, model=None):
    """Use the configured image semantic model to flag visual reasonability risks."""
    api_key = api_key or GLM_API_KEY
    model = model or GLM_VISION_MODEL

    def _call():
        return _call_glm_image_semantics_unbounded(image_path, timeout=timeout, api_key=api_key, model=model)

    return _run_with_alarm_timeout(_call, timeout, lambda: _glm_timeout_result(model, timeout))


def _call_glm_image_semantics_unbounded(image_path: str, timeout=45, api_key=None, model=None):
    """Unbounded implementation; call through call_glm_image_semantics in orchestration."""
    path = Path(image_path)
    if not api_key:
        return {
            "status": "skipped",
            "model": model,
            "summary": "图像语义分析API Key未配置，已跳过图像语义分析。",
            "risks": ["glm_key_missing"],
            "confidence": 0,
        }
    try:
        if path.exists() and path.stat().st_size > GLM_IMAGE_MAX_BYTES:
            return {
                "status": "skipped",
                "model": model,
                "summary": "图片超过图像语义分析的本地压缩前安全上限，已跳过。",
                "reasonability": "需人工核对",
                "risks": ["glm_image_too_large"],
                "manual_checks": ["人工核对该图原图、图注和正文结论是否一致。"],
                "confidence": 0,
            }
    except Exception:
        pass

    prompt = (
        "你是科研论文图像审查助手。请只基于这张图片本身做语义理解与合理性审查。"
        "不要输出推理过程、解释、Markdown或代码块；只返回一个合法JSON对象。"
        "不要把低分辨率、OCR错误、压缩噪声、表格截断或排版问题直接当作造假证据。"
        "如果图片是表格/局部截图，请重点说明可读内容和截断风险。"
        "reasonability字段必须严格取值为：合理、需人工核对、可疑。"
        "请返回严格JSON："
        "{\"summary\":\"一句话描述图片内容\","
        "\"image_type\":\"图/表/显微图/热图/流程图/照片/其他\","
        "\"scientific_context\":\"可能对应的科研用途\","
        "\"visible_text\":\"能读出的关键文字，读不出写空字符串\","
        "\"reasonability\":\"合理/需人工核对/可疑\","
        "\"risks\":[\"可疑点短语\"],"
        "\"manual_checks\":[\"建议人工核对事项\"],"
        "\"confidence\":0到1}"
    )
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": _image_to_data_url(image_path)}},
                ],
            }
        ],
        "temperature": 0.1,
        "max_tokens": 10000,
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    try:
        data, _ = _http_request(_chat_completions_endpoint(GLM_API_URL), "POST", headers=headers, data=json.dumps(payload).encode("utf-8"), timeout=timeout)
        result = json.loads(data.decode("utf-8", errors="replace"))
        message = ((result.get("choices") or [{}])[0].get("message") or {})
        content = (
            message.get("content")
            or message.get("reasoning_content")
            or message.get("reasoning")
            or ""
        ).strip()
        parsed = _extract_json_object(content)
        if not isinstance(parsed, dict):
            parsed = {"summary": _brief_text(content, 260), "risks": ["glm_json_parse_failed"], "confidence": 0}
        return _normalize_glm_image_result(parsed, model)
    except Exception as e:
        return _glm_error_result(e, model)


def build_image_audit(
    input_path: str,
    output_dir=None,
    limit=None,
    semantic=True,
    semantic_limit=None,
    semantic_timeout=45,
    semantic_cache=None,
    semantic_cache_save=None,
    detector=True,
    detector_limit=None,
    detector_timeout=60,
    detector_cache=None,
    detector_cache_save=None,
):
    images = collect_image_files(input_path, include_pdf=False, include_mineru=True, output_dir=output_dir)
    analyses = sorted((analyze_image_reasonability(path) for path in images), key=_image_audit_sort_key)
    analyses = analyses[:_effective_limit(limit, len(analyses))]
    semantic_cache = semantic_cache if isinstance(semantic_cache, dict) else {}
    detector_cache = detector_cache if isinstance(detector_cache, dict) else {}
    semantic_checked = 0
    if semantic:
        semantic_candidates = sorted(analyses, key=_image_semantic_priority_key)
        semantic_queue = semantic_candidates[:_effective_limit(semantic_limit, len(semantic_candidates))]
        for idx, item in enumerate(semantic_queue, 1):
            cache_key = _image_semantic_cache_key(item.get("path", ""))
            semantic_result = semantic_cache.get(cache_key)
            if isinstance(semantic_result, dict) and semantic_result.get("status") == "error":
                semantic_cache.pop(cache_key, None)
                _flush_image_cache(semantic_cache_save, "图像语义")
                semantic_result = None
            if not semantic_result:
                print(f"  🖼️ 图像语义分析 [{idx}/{len(semantic_queue)}] {item.get('file', '')}")
                semantic_result = call_glm_image_semantics(item.get("path", ""), timeout=semantic_timeout)
                if semantic_result.get("status") != "error":
                    semantic_cache[cache_key] = semantic_result
                    _flush_image_cache(semantic_cache_save, "图像语义")
            item["semantic"] = semantic_result
            semantic_checked += 1
    detector_checked = 0
    if detector:
        detector_candidates = sorted(analyses, key=_image_detector_priority_key)
        detector_queue = detector_candidates[:_effective_limit(detector_limit, len(detector_candidates))]
        for idx, item in enumerate(detector_queue, 1):
            cache_key = _image_file_fingerprint(item.get("path", "")) + ":imagedetector_v1"
            detector_result = detector_cache.get(cache_key)
            if isinstance(detector_result, dict) and detector_result.get("status") == "error":
                detector_cache.pop(cache_key, None)
                _flush_image_cache(detector_cache_save, "imagedetector")
                detector_result = None
            if not detector_result:
                print(f"  🖼️ imagedetector自动检测 [{idx}/{len(detector_queue)}] {item.get('file', '')}")
                detector_result = call_imagedetector(item.get("path", ""), timeout=detector_timeout)
                if detector_result.get("status") != "error":
                    detector_cache[cache_key] = detector_result
                    _flush_image_cache(detector_cache_save, "imagedetector")
            item["detector"] = detector_result
            detector_checked += 1
    return {
        "enabled": bool(analyses),
        "site": IMAGE_DETECT_URL,
        "semantic_enabled": bool(semantic),
        "semantic_model": GLM_VISION_MODEL,
        "semantic_checked": semantic_checked,
        "detector_enabled": bool(detector),
        "detector_checked": detector_checked,
        "image_count": len(images),
        "checked_count": len(analyses),
        "images": analyses,
        "note": "本地做尺寸、空白、噪声/对比度筛查；图像语义分析模型做图片语义理解；imagedetector.com子工具自动上传并记录AI概率。",
    }


def format_image_audit_html(image_audit):
    return _format_image_audit_html(image_audit, image_detect_url=IMAGE_DETECT_URL)


def format_image_audit_markdown(image_audit):
    return _format_image_audit_markdown(image_audit, image_detect_url=IMAGE_DETECT_URL)


def save_image_review_manifest(image_audit, output_dir):
    return _save_image_review_manifest(image_audit, output_dir, image_detect_url=IMAGE_DETECT_URL)


def extract_images_from_pdf(pdf_path: str) -> List[str]:
    """从PDF中提取内嵌图片到临时目录，返回图片路径列表

    优先使用PyMuPDF(fitz)，降级使用pdf2image整页渲染
    """
    images = []
    tmp_dir = os.path.join(os.path.dirname(pdf_path), "_veritas_images_tmp")
    os.makedirs(tmp_dir, exist_ok=True)

    # 方案1：PyMuPDF提取内嵌图片
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        img_count = 0
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            img_list = page.get_images(full=True)
            for img_idx, img_info in enumerate(img_list):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    if base_image and base_image.get("image"):
                        ext = base_image.get("ext", "png")
                        if ext not in ("png", "jpg", "jpeg", "bmp", "tiff", "webp"):
                            ext = "png"
                        fname = f"page{page_idx + 1}_img{img_idx + 1}.{ext}"
                        fpath = os.path.join(tmp_dir, fname)
                        with open(fpath, "wb") as f:
                            f.write(base_image["image"])
                        # 过滤掉太小的图片（图标、装饰等）
                        if os.path.getsize(fpath) > 5000:
                            images.append(fpath)
                            img_count += 1
                except Exception:
                    continue
        doc.close()
        if img_count > 0:
            print(f"  📎 PyMuPDF提取 {img_count} 张内嵌图片 → {tmp_dir}")
            return images
    except ImportError:
        pass
    except Exception as e:
        print(f"  ⚠️ PyMuPDF提取失败: {e}")

    # 方案2：pdf2image整页渲染
    try:
        from pdf2image import convert_from_path
        pages = convert_from_path(pdf_path, dpi=200)
        for i, page_img in enumerate(pages):
            fname = f"page{i + 1}_full.png"
            fpath = os.path.join(tmp_dir, fname)
            page_img.save(fpath, "PNG")
            if os.path.getsize(fpath) > 10000:
                images.append(fpath)
        if images:
            print(f"  📎 pdf2image渲染 {len(images)} 页 → {tmp_dir}")
            return images
    except ImportError:
        pass
    except Exception as e:
        print(f"  ⚠️ pdf2image渲染失败: {e}")

    return images


def collect_image_files(input_path: str, include_pdf=True, include_mineru=True, output_dir=None) -> List[str]:
    """收集论文相关图片文件：目录图片、PDF内嵌图，以及MinerU zip中的图片。"""
    images = []
    p = Path(input_path)

    if include_mineru:
        images.extend(collect_mineru_image_files(input_path, output_dir=output_dir))

    if include_pdf and p.is_file() and p.suffix.lower() == ".pdf":
        # PDF文件：提取内嵌图片
        print("  📸 从PDF中提取图片...")
        extracted = extract_images_from_pdf(str(p))
        images.extend(extracted)
    elif p.is_dir():
        # 目录：扫描所有图片文件
        for ext in IMAGE_EXTENSIONS:
            for f in p.rglob(f"*{ext}"):
                if ".paper_audit_resume" in str(f) or "_paper_audit_images" in str(f):
                    continue
                if f.stat().st_size > MIN_IMAGE_BYTES:  # 过滤小图标
                    images.append(str(f))

    return _dedupe_paths(images)


def launch_image_ai_detect(
    input_path: str,
    output_dir=None,
    limit=None,
    semantic=True,
    semantic_limit=None,
    semantic_timeout=45,
    semantic_cache=None,
    detector=True,
    detector_limit=None,
    detector_timeout=60,
    detector_cache=None,
):
    """Run the automatic image audit subtool and save the review manifest."""
    print("\n" + "=" * 60)
    print("🖼️ AI图片检测子工具 (图像语义分析 + imagedetector.com)")
    print("=" * 60)

    target_output_dir = output_dir or (Path(input_path).parent if Path(input_path).is_file() else Path(input_path))
    image_audit = build_image_audit(
        input_path,
        output_dir=target_output_dir,
        limit=limit,
        semantic=semantic,
        semantic_limit=semantic_limit,
        semantic_timeout=semantic_timeout,
        semantic_cache=semantic_cache,
        detector=detector,
        detector_limit=detector_limit,
        detector_timeout=detector_timeout,
        detector_cache=detector_cache,
    )
    if not image_audit.get("image_count"):
        print("⚠️ 未找到可检测的图片文件")
        return image_audit

    manifest_path = save_image_review_manifest(image_audit, target_output_dir)
    if manifest_path:
        print(f"  🧾 图片AI检测结果清单: {manifest_path}")
    print(
        f"✅ 图片子工具完成: 本地{image_audit.get('checked_count')}/{image_audit.get('image_count')}张；"
        f"图像语义分析 {image_audit.get('semantic_checked')}张；imagedetector {image_audit.get('detector_checked')}张"
    )

    # 清理临时提取目录
    tmp_dir = os.path.join(os.path.dirname(input_path) if os.path.isfile(input_path) else input_path, "_veritas_images_tmp")
    if os.path.isdir(tmp_dir):
        try:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)
            print("  🧹 已清理临时图片目录")
        except Exception:
            pass
    return image_audit



def _failed_artifact_options(input_path: Path, output_dir: Path, args) -> Dict[str, Any]:
    base = explicit_output_path_from_args(args)
    if base is None:
        return {}
    return {"output_dir": base.parent, "output_stem": base.name}


def run_audit(run_request: RunRequest, args=None) -> RunResult:
    global _RESUME_EVENTS_ENABLED
    args = args if args is not None else run_request.to_args()
    input_path = run_request.input_path
    if not input_path.exists():
        print(f"❌ 路径不存在: {input_path}")
        failure = AuditFailure(
            capability="input",
            error_class="input_path_not_found",
            message=f"路径不存在: {input_path}",
            retry_command=retry_command_from_args(args, input_path),
        )
        return RunResult.failed(failure, {}, meta={"input_path": str(input_path)})

    output_dir, output_stem = get_output_base(input_path)
    setup_run_logging(input_path)
    print(f"📁 所有输出将保存到: {output_dir}")
    runtime_config = apply_runtime_config(load_runtime_config())
    print(f"🤖 当前LLM: model={LLM_MODEL}, url={LLM_API_URL}")
    config_errors = runtime_config.validation_errors()
    if config_errors:
        detail = "; ".join(f"{e['capability']}.{e['field']}" for e in config_errors)
        print(f"⚠️ 关键配置缺失: {detail}。关键能力预检会在对应正式阶段前停止并生成失败诊断。")
    resume_dir = get_resume_dir(output_dir, output_stem)
    _RESUME_EVENTS_ENABLED = not bool(getattr(args, "no_resume", False))
    if getattr(args, "fresh", False):
        import shutil
        print(f"🧹 --fresh: 清空断点续作缓存目录 {resume_dir}")
        shutil.rmtree(resume_dir, ignore_errors=True)
        resume_dir = get_resume_dir(output_dir, output_stem)
    if args.no_resume:
        print("♻️ 已禁用断点续作缓存，本次将重新执行提取和LLM审查")
    else:
        print(f"🔁 断点续作缓存目录: {resume_dir}")
    retry_command = retry_command_from_args(args, input_path)
    failed_artifact_kwargs = _failed_artifact_options(input_path, output_dir, args)
    run_runtime = runtime_metadata()
    run_workspace = create_run_workspace(input_path, output_dir, output_stem)
    print(f"🗂️ 本次运行工作区: {run_workspace['run_dir']}")
    record_run_workspace_json(run_workspace, "input_manifest.json", {
        "input": str(input_path),
        "resolved_input": str(input_path.resolve()),
        "input_type": "directory" if input_path.is_dir() else "file",
        "exists": input_path.exists(),
        "size_bytes": input_path.stat().st_size if input_path.is_file() else None,
        "created_at": run_runtime["local_time"],
        "runtime": run_runtime,
    })
    allow_llm_cache_read = _allow_llm_cache_read(args.no_resume, getattr(args, "llm_cache_only", False))
    allow_llm_cache_write = not args.no_resume
    pdf_suffixes = {".pdf"}
    has_pdf_input = input_path.suffix.lower() in pdf_suffixes
    if input_path.is_dir():
        try:
            has_pdf_input = any(p.is_file() and p.suffix.lower() in pdf_suffixes for p in input_path.rglob("*"))
        except Exception:
            has_pdf_input = False
    use_mineru_default = has_pdf_input and not args.no_mineru
    if use_mineru_default and not args.mineru:
        print("📡 检测到PDF输入，默认启用MinerU提取；如需原始PDF文本提取请使用 --no-mineru")

    output_override_preview = explicit_output_path_from_args(args)
    preview_md, preview_html, preview_json = audit_artifact_paths(input_path, output_path=output_override_preview)
    if input_path.is_dir():
        extraction_route = "directory_multi_format"
    elif input_path.suffix.lower() == ".pdf":
        extraction_route = "mineru_pdf" if use_mineru_default else "raw_pdf_stream"
    elif input_path.suffix.lower() == ".docx":
        extraction_route = "direct_docx"
    elif input_path.suffix.lower() in {".xlsx", ".xlsm", ".csv"}:
        extraction_route = "spreadsheet_text"
    else:
        extraction_route = f"{input_path.suffix.lower().lstrip('.') or 'file'}_text"
    scope_flags = []
    for attr, label in (
        ("no_mineru", "--no-mineru"),
        ("no_reference_online", "--no-reference-online"),
        ("no_image_semantic", "--no-image-semantic"),
        ("no_image_detector", "--no-image-detector"),
        ("llm_cache_only", "--llm-cache-only"),
    ):
        if getattr(args, attr, False):
            scope_flags.append(label)
    for attr, label in (
        ("reference_online_limit", "--reference-online-limit"),
        ("image_audit_limit", "--image-audit-limit"),
        ("image_semantic_limit", "--image-semantic-limit"),
        ("image_detector_limit", "--image-detector-limit"),
    ):
        if getattr(args, attr, None) is not None:
            scope_flags.append(f"{label}={getattr(args, attr)}")
    print("🧭 运行摘要:")
    print(f"  - 输入: {input_path} ({'目录' if input_path.is_dir() else '单文件'})")
    print(f"  - 提取路线: {extraction_route}")
    print(f"  - 输出目录/产物: {preview_md.parent} / {preview_md.stem}")
    print(f"  - HTML/JSON预期: {preview_html.name} / {preview_json.name}")
    print(f"  - 断点续作缓存: {resume_dir}")
    print(f"  - 范围限制开关: {', '.join(scope_flags) if scope_flags else '无，默认尝试完整审查'}")

    resume_event(resume_dir, "init", "done", f"input={input_path}; llm={LLM_MODEL}; url={LLM_API_URL}; max_chars={args.max_chars}; use_mineru={use_mineru_default}")
    record_run_workspace_json(run_workspace, "cache_use.json", {
        "shared_resume_dir": str(resume_dir),
        "no_resume": bool(args.no_resume),
        "allow_llm_cache_read": bool(allow_llm_cache_read),
        "allow_llm_cache_write": bool(allow_llm_cache_write),
        "extract_cache_version": EXTRACT_CACHE_VERSION,
        "image_semantic_cache_version": IMAGE_SEMANTIC_CACHE_VERSION,
    })
    completed_stages = ["init", "runtime_config_loaded"]
    preflight_state = {}
    preflight_results = []

    def _record_preflight(result: PreflightResult):
        preflight_results.append(result.to_dict())
        record_run_workspace_json(run_workspace, "preflight.json", {
            "results": preflight_results,
            "updated_at": time.strftime("%F %T"),
        })
        resume_event(
            resume_dir,
            f"preflight_{result.capability}",
            "ok" if result.ok else "failed",
            result.message or "ok",
            error_class=result.error_class,
        )

    if use_mineru_default:
        print("🧪 关键能力预检: MinerU")
        mineru_preflight = run_preflight_once(preflight_state, "mineru", lambda: preflight_mineru(timeout=10))
        _record_preflight(mineru_preflight)
        if not mineru_preflight.ok:
            failure = preflight_failure_to_audit_failure(
                mineru_preflight,
                retry_command,
                completed_stages,
            )
            md_path, json_path = save_failed_audit_diagnostics(
                failure,
                input_path,
                **failed_artifact_kwargs,
                meta={"preflight_results": preflight_results},
            )
            record_run_workspace_artifacts(
                run_workspace,
                "failed",
                [md_path, json_path],
                meta={"preflight_results": preflight_results, "completed_stages": completed_stages},
            )
            print(f"❌ MinerU预检失败，未生成完整审查报告。失败诊断已保存: {md_path}, {json_path}")
            return RunResult.failed(
                failure,
                {"markdown": str(md_path), "json": str(json_path)},
                workspace=run_workspace,
                meta={"input_path": str(input_path), "preflight_results": preflight_results},
            )
        completed_stages.append("mineru_preflight")
    progress_bar(0, 5, "初始化完成")

    # ─── 阶段1：文本提取（支持单个文件/整个论文目录） ───
    extract_cache_path = resume_dir / "stage1_extract.json"
    cached_extract = None if args.no_resume else _json_load(extract_cache_path)
    extracted_file_texts = []
    if cached_extract and cached_extract.get("input") == str(input_path.resolve()) and cached_extract.get("use_mineru") == use_mineru_default and cached_extract.get("cache_version") == EXTRACT_CACHE_VERSION:
        full_text = cached_extract.get("full_text", "")
        meta = cached_extract.get("meta", {})
        extracted_file_texts = cached_extract.get("file_texts") or []
        raw_pdf = None
        use_mineru = cached_extract.get("use_mineru", use_mineru_default)
        print(f"🔁 断点续作：复用阶段1文本缓存 {extract_cache_path} ({len(full_text)}字符)")
        resume_event(resume_dir, "stage1_extract", "cache_hit", f"chars={len(full_text)}", cache=str(extract_cache_path))
        progress_bar(1, 5, "阶段1/5 文本提取缓存命中")
    else:
        full_text = None
        meta = {}
        raw_pdf = None
        use_mineru = use_mineru_default

    if full_text is None and input_path.is_dir():
        print(f"📂 检测到输入为目录，正在扫描所有论文相关文件...")
        file_classes, all_files = find_project_files(input_path)
        print(f"✅ 找到 {len(all_files)} 个相关文件:")
        for cat, files in file_classes.items():
            if not files:
                continue
            if isinstance(files, Path):
                print(f"  - {cat}: {files.name}")
            else:
                print(f"  - {cat}: {len(files)} 个文件")
        
        reference_file_set = set(file_classes.get("references") or [])
        audit_files = [p for p in all_files if p not in reference_file_set]

        def _file_audit_category(path):
            if path == file_classes.get("main_paper"):
                return "main_text"
            if path in set(file_classes.get("supplements") or []):
                return "supplement"
            if path in set(file_classes.get("data_files") or []):
                return "data_file"
            return "other"

        # 提取所有非参考文献文件文本合并；参考文献文件单独校检，避免污染主体审查
        full_text = ""
        reference_file_texts = []
        extracted_file_texts = []
        total_files = len(audit_files)
        for idx, file_path in enumerate(audit_files, 1):
            print(f"  📝 提取主体文件 [{idx}/{total_files}] {file_path.name}...")
            progress_bar(idx - 1, max(total_files, 1), f"阶段1/5 提取主体文件: {file_path.name}")
            dependency, install_command = optional_dependency_for_extension(file_path.suffix)
            if dependency:
                failure = AuditFailure(
                    capability="input_extraction",
                    error_class="missing_optional_dependency",
                    message=f"目录审查中的审查相关文件 {file_path.name} 需要安装可选依赖 {dependency}。",
                    fix_hints=[f"运行 `{install_command}` 后重试。", "或转换该文件为 PDF/文本格式后重新运行审查。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                    details={
                        "file": str(file_path),
                        "extension": file_path.suffix.lower(),
                        "dependency": dependency,
                        "install_command": install_command,
                        "resume_dir": str(resume_dir),
                    },
                )
                md_path, json_path = save_failed_audit_diagnostics(
                    failure,
                    input_path,
                    **failed_artifact_kwargs,
                    meta={"runtime": run_runtime, "preflight_results": preflight_results},
                )
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
            file_content = extract_text_from_file(file_path, max_chars_per_file=None,
                                                  use_mineru=use_mineru,
                                                  mineru_lang=args.mineru_lang,
                                                  output_dir=output_dir)
            body_text = extracted_body_text(file_content, file_path.name)
            if not body_text or body_text.startswith("[文件解析失败:"):
                failure = AuditFailure(
                    capability="input_extraction",
                    error_class="no_extractable_text",
                    message=f"未能从目录审查相关文件 {file_path.name} 提取到可审查文本。",
                    fix_hints=["检查文件是否为空、损坏、加密或需要额外依赖。", "转换该文件为 PDF/文本格式后重试，或在未来显式排除该文件。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                    details={
                        "file": str(file_path),
                        "extension": file_path.suffix.lower(),
                        "category": _file_audit_category(file_path),
                        "extract_preview": _brief_text(file_content, 240),
                        "resume_dir": str(resume_dir),
                    },
                )
                md_path, json_path = save_failed_audit_diagnostics(
                    failure,
                    input_path,
                    **failed_artifact_kwargs,
                    meta={"runtime": run_runtime, "preflight_results": preflight_results},
                )
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
            try:
                rel_path = str(file_path.relative_to(input_path))
            except Exception:
                rel_path = file_path.name
            extracted_file_texts.append({
                "file": file_path.name,
                "path": rel_path,
                "category": _file_audit_category(file_path),
                "text": file_content,
            })
            progress_bar(idx, max(total_files, 1), f"阶段1/5 已完成: {file_path.name}")
            full_text += f"\n\n=== 文件: {file_path.name} 路径: {file_path.relative_to(input_path)} ==="
            full_text += "\n" + file_content

        for idx, file_path in enumerate(file_classes.get("references") or [], 1):
            print(f"  📚 提取参考文献文件 [{idx}/{len(reference_file_set)}] {file_path.name}...")
            dependency, install_command = optional_dependency_for_extension(file_path.suffix)
            if dependency:
                failure = AuditFailure(
                    capability="input_extraction",
                    error_class="missing_optional_dependency",
                    message=f"目录审查中的参考文献文件 {file_path.name} 需要安装可选依赖 {dependency}。",
                    fix_hints=[f"运行 `{install_command}` 后重试。", "或转换该文件为 PDF/文本格式后重新运行审查。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                    details={
                        "file": str(file_path),
                        "extension": file_path.suffix.lower(),
                        "dependency": dependency,
                        "install_command": install_command,
                        "resume_dir": str(resume_dir),
                    },
                )
                md_path, json_path = save_failed_audit_diagnostics(
                    failure,
                    input_path,
                    **failed_artifact_kwargs,
                    meta={"runtime": run_runtime, "preflight_results": preflight_results},
                )
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
            reference_content = extract_text_from_file(file_path, max_chars_per_file=None,
                                                       use_mineru=use_mineru,
                                                       mineru_lang=args.mineru_lang,
                                                       output_dir=output_dir)
            if not extracted_body_text(reference_content, file_path.name):
                failure = AuditFailure(
                    capability="input_extraction",
                    error_class="no_extractable_text",
                    message=f"未能从目录审查相关参考文献文件 {file_path.name} 提取到可审查文本。",
                    fix_hints=["检查文件是否为空、损坏、加密或需要额外依赖。", "转换该文件为 PDF/文本格式后重试。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                    details={"file": str(file_path), "extension": file_path.suffix.lower(), "resume_dir": str(resume_dir)},
                )
                md_path, json_path = save_failed_audit_diagnostics(
                    failure,
                    input_path,
                    **failed_artifact_kwargs,
                    meta={"runtime": run_runtime, "preflight_results": preflight_results},
                )
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
            reference_file_texts.append(reference_content)
        
        def _class_count(v):
            if v is None:
                return 0
            if isinstance(v, Path):
                return str(v)
            try:
                return len(v)
            except TypeError:
                return str(v)

        meta = {
            "input_type": "directory",
            "total_files": len(all_files),
            "audit_files": len(audit_files),
            "reference_files": len(reference_file_set),
            "file_classes": {k: _class_count(v) for k, v in file_classes.items()},
            "total_chars": len(full_text),
            "extractor": "directory_multi_format",
            "extraction_method": "directory_multi_format",
            "size_mb": round(sum(p.stat().st_size for p in all_files if p.exists()) / 1024 / 1024, 2),
            "reference_file_text": "\n\n".join(reference_file_texts),
        }
        print(f"\n✅ 所有文件提取完成，总长度: {len(full_text)} 字符")
        progress_bar(1, 5, "阶段1/5 文本提取完成")
    elif full_text is None:
        # 单个文件走原有流程
        pdf_path = input_path
        print(f"📄 检测到输入为单个文件: {pdf_path.name}")
        single_suffix = pdf_path.suffix.lower()

        if single_suffix not in SUPPORTED_TEXT_FILE_EXTENSIONS:
            if single_suffix == ".doc":
                message = "暂不支持旧版二进制Word .doc 文件直接输入；请转换为 .docx 或 PDF 后重试。"
                hints = ["用 Word/WPS/LibreOffice 另存为 .docx。", "或导出为 PDF 后重新运行审查。"]
                error_class = "unsupported_legacy_doc"
            else:
                message = f"不支持的单文件输入类型: {single_suffix or '(无扩展名)'}。"
                hints = ["请使用 PDF、.docx、Excel、CSV、TXT 或 Markdown 文件。", "也可以把论文相关文件放入目录后进行目录审查。"]
                error_class = "unsupported_file_type"
            failure = AuditFailure(
                capability="input_extraction",
                error_class=error_class,
                message=message,
                fix_hints=hints,
                completed_stages=completed_stages,
                retry_command=retry_command,
            )
            md_path, json_path = save_failed_audit_diagnostics(failure, input_path, **failed_artifact_kwargs)
            record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
            return RunResult.failed(
                failure,
                {"markdown": str(md_path), "json": str(json_path)},
                workspace=run_workspace,
                meta={"input_path": str(input_path)},
            )

        if single_suffix != ".pdf":
            missing_dependency = (
                single_suffix == ".docx" and not DOCX_SUPPORTED
            ) or (
                single_suffix in {".xlsx", ".xlsm"} and not EXCEL_SUPPORTED
            )
            if missing_dependency:
                dependency, install_command = optional_dependency_for_extension(single_suffix)
                failure = AuditFailure(
                    capability="input_extraction",
                    error_class="missing_optional_dependency",
                    message=f"读取 {single_suffix} 文件需要安装可选依赖 {dependency}。",
                    fix_hints=[f"运行 `{install_command}` 后重试。", "或转换为 PDF 后重新运行审查。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                    details={"dependency": dependency, "install_command": install_command, "resume_dir": str(resume_dir)},
                )
                md_path, json_path = save_failed_audit_diagnostics(failure, input_path, **failed_artifact_kwargs)
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
            print(f"📖 正在提取{single_suffix}文件文本: {pdf_path}")
            full_text = extract_text_from_file(
                pdf_path,
                max_chars_per_file=None,
                use_mineru=False,
                mineru_lang=args.mineru_lang,
                output_dir=output_dir,
            )
            body_text = extracted_body_text(full_text, pdf_path.name)
            if not body_text:
                failure = AuditFailure(
                    capability="input_extraction",
                    error_class="no_extractable_text",
                    message=f"未能从 {single_suffix} 文件中提取到可审查文本。",
                    fix_hints=["检查文件是否为空、损坏或受保护。", "尝试另存为 .docx/PDF 后重试。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                )
                md_path, json_path = save_failed_audit_diagnostics(failure, input_path, **failed_artifact_kwargs)
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
            meta = {
                "input_type": "file",
                "source_file": pdf_path.name,
                "size_mb": round(pdf_path.stat().st_size / 1024 / 1024, 2),
                "total_chars": len(full_text),
                "chars_sent": len(full_text),
                "extractor": "single_file_multi_format",
                "extraction_method": f"{single_suffix.lstrip('.')}_text",
            }
            extracted_file_texts = [{
                "file": pdf_path.name,
                "path": pdf_path.name,
                "category": "main_text",
                "text": full_text,
            }]
            raw_pdf = None
            print(f"✅ 提取完成: {meta['total_chars']} 字符（全文保留）")
            progress_bar(1, 5, "阶段1/5 单文件文本提取完成")

        if single_suffix == ".pdf" and use_mineru:
            print(f"📡 [MinerU] 正在将PDF转为Markdown: {pdf_path.name}")
            md_text, md_meta = mineru_extract(pdf_path, language=args.mineru_lang, output_dir=output_dir)
            if md_text:
                full_text = md_text  # 保留全文
                meta = {
                    "size_mb": round(pdf_path.stat().st_size / 1024 / 1024, 2),
                    "total_chars": len(md_text),
                    "chars_sent": len(md_text),
                    "extraction_method": f"mineru_{md_meta.get('source', 'unknown')}",
                }
                if md_meta.get("batch_id"):
                    meta["mineru_batch_id"] = md_meta["batch_id"]
                if md_meta.get("task_id"):
                    meta["mineru_task_id"] = md_meta["task_id"]
                extracted_file_texts = [{
                    "file": pdf_path.name,
                    "path": pdf_path.name,
                    "category": "main_text",
                    "text": full_text,
                }]
                print(f"✅ MinerU提取完成: {len(md_text)} 字符（全文保留）")
                progress_bar(1, 5, "阶段1/5 MinerU文本提取完成")
            else:
                err = md_meta.get("error", "未知错误") if md_meta else "未知错误"
                print(f"❌ MinerU提取失败: {err}")
                print(f"⚠️ 降级使用原始PDF文本提取...")
                use_mineru = False

        if single_suffix == ".pdf" and (not use_mineru or full_text is None):
            print(f"📖 正在提取PDF文本: {pdf_path}")
            # extract_pdf_text的max_chars参数传大值以获取全文
            full_text, meta, raw_pdf = extract_pdf_text(str(pdf_path), max_chars=999999)
            if not full_text:
                print("❌ 未能从PDF中提取到文本（可能是扫描件或加密PDF）")
                print("💡 建议: 使用 --mineru 参数通过MinerU API提取（支持OCR）")
                failure = AuditFailure(
                    capability="input_extraction",
                    error_class="no_extractable_text",
                    message="未能从PDF中提取到文本（可能是扫描件或加密PDF）。",
                    fix_hints=["检查PDF是否加密或为扫描件。", "确认MinerU配置可用后重试。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                )
                md_path, json_path = save_failed_audit_diagnostics(failure, input_path, **failed_artifact_kwargs)
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
            print(f"✅ 提取完成: {meta['total_chars']} 字符（全文保留）")
            extracted_file_texts = [{
                "file": pdf_path.name,
                "path": pdf_path.name,
                "category": "main_text",
                "text": full_text,
            }]
            progress_bar(1, 5, "阶段1/5 PDF文本提取完成")

    meta = normalize_run_meta(meta, input_path, full_text)
    meta["runtime"] = run_runtime
    meta["paper_identity"] = extract_paper_identity(full_text, input_path)
    meta["preflight_results"] = preflight_results
    meta["cross_file_consistency_audit"] = build_cross_file_consistency_audit(extracted_file_texts, root_path=input_path)
    completed_stages.append("stage1_text_extraction")

    if not args.no_resume and full_text:
        _json_save(extract_cache_path, {
            "input": str(input_path.resolve()),
            "cache_version": EXTRACT_CACHE_VERSION,
            "use_mineru": use_mineru,
            "mineru_lang": args.mineru_lang,
            "full_text": full_text,
            "meta": meta,
            "file_texts": extracted_file_texts,
            "saved_at": time.strftime("%F %T"),
        })
        resume_event(resume_dir, "stage1_extract", "saved", f"chars={len(full_text)}; use_mineru={use_mineru}", cache=str(extract_cache_path))

    # ─── 朱雀AI文本检测（可选） ───
    if args.ai_detect:
        launch_zhuque_ai_detect(full_text)

    # ─── AI图片检测兼容参数 ───
    if args.image_detect:
        print("ℹ️ --image-detect 已改为兼容参数；图片检测将在阶段4自动调用图像语义分析与imagedetector子工具，不会打开网页或要求手动上传。")

    # ─── 参考文献剥离与单独校检 ───
    audit_text, references_text = split_references_from_text(full_text)
    reference_file_text = meta.pop("reference_file_text", "")
    if reference_file_text:
        references_text = (references_text + "\n\n" + reference_file_text).strip()
    reference_online_cache_path = resume_dir / "reference_online_cache.json"
    reference_online_enabled = bool(references_text) and not args.no_reference_online
    reference_online_cache = {} if args.no_resume else (_json_load(reference_online_cache_path, {}) or {})
    if reference_online_enabled:
        print(f"🔎 参考文献在线检索已启用: 上限{args.reference_online_limit}条, 超时{args.reference_timeout}s")
    reference_audit = audit_references(
        references_text,
        online=reference_online_enabled,
        online_limit=args.reference_online_limit,
        timeout=args.reference_timeout,
        cache=reference_online_cache,
    )
    if reference_online_enabled and not args.no_resume:
        _json_save(reference_online_cache_path, reference_online_cache)
        resume_event(
            resume_dir,
            "stage1_reference_online",
            "saved",
            f"checked={reference_audit.get('online_checked', 0)}; cache_entries={len(reference_online_cache)}",
            cache=str(reference_online_cache_path),
        )
    if references_text:
        meta["references_excluded_from_main_audit"] = True
        meta["reference_chars"] = len(references_text)
        meta["reference_count"] = reference_audit.get("reference_count", 0)
        meta["reference_audit"] = reference_audit
        print(f"📚 已从主体审查中剥离参考文献: {meta['reference_count']}条, {len(references_text)}字符；将单独校检")
        resume_event(resume_dir, "stage1_references", "done", f"refs={meta['reference_count']}; chars={len(references_text)}")
    else:
        meta["references_excluded_from_main_audit"] = False
        meta["reference_audit"] = reference_audit
        print("📚 未识别到独立参考文献章节，主体审查不做引用剥离")
    completed_stages.append("stage1_reference_audit")

    # ─── 代码仓库与在线部署资源可用性校检 ───
    resource_online_cache_path = resume_dir / "resource_online_cache.json"
    resource_online_cache = {} if args.no_resume else (_json_load(resource_online_cache_path, {}) or {})
    resource_online_enabled = not getattr(args, "no_resource_online", False)
    resource_audit = audit_resources(
        full_text,
        online=resource_online_enabled,
        timeout=getattr(args, "resource_timeout", 10),
        cache=resource_online_cache,
    )
    if resource_online_enabled and not args.no_resume:
        _json_save(resource_online_cache_path, resource_online_cache)
        resume_event(
            resume_dir,
            "stage1_resource_online",
            "saved",
            f"checked={resource_audit.get('online_checked', 0)}; cache_entries={len(resource_online_cache)}",
            cache=str(resource_online_cache_path),
        )
    meta["resource_count"] = resource_audit.get("resource_count", 0)
    meta["resource_audit"] = resource_audit
    if resource_audit.get("resource_count"):
        print(f"🔗 已识别代码仓库/在线资源: {resource_audit.get('resource_count')}项；在线检测{resource_audit.get('online_checked', 0)}项")
    else:
        print("🔗 未识别到代码仓库或论文部署的在线资源链接")
    completed_stages.append("stage1_resource_audit")

    # ─── 阶段2：本地统计检测（使用全文，统计不截断） ───
    progress_bar(1, 5, "阶段2/5 开始本地统计检测")
    print(f"🔢 正在执行本地统计检测...")
    stat_result = local_stat_check(audit_text)
    benford_str = f"{round(stat_result['benford_deviation'],3)}" if stat_result['benford_deviation'] else 'N/A'
    print(f"✅ 统计检测完成: Benford偏差={benford_str}, p值异常={stat_result['p_value_abnormal']}, 数字数={stat_result['number_count']}")
    resume_event(resume_dir, "stage2_stat", "done", f"numbers={stat_result['number_count']}; benford={benford_str}")
    progress_bar(2, 5, "阶段2/5 本地统计检测完成")
    completed_stages.append("stage2_stat_check")

    # ─── 阶段3：智能分块 + LLM语义审查（冗余机制） ───
    print("🧪 关键能力预检: 文本语义审查LLM")
    text_llm_preflight = run_preflight_once(preflight_state, "text_llm", lambda: preflight_text_llm(timeout=min(30, LLM_TIMEOUT)))
    _record_preflight(text_llm_preflight)
    meta["preflight_results"] = preflight_results
    if not text_llm_preflight.ok:
        failure = preflight_failure_to_audit_failure(
            text_llm_preflight,
            retry_command,
            completed_stages,
        )
        md_path, json_path = save_failed_audit_diagnostics(
            failure,
            input_path,
            **failed_artifact_kwargs,
            meta=meta,
        )
        record_run_workspace_artifacts(
            run_workspace,
            "failed",
            [md_path, json_path],
            meta={"preflight_results": preflight_results, "completed_stages": completed_stages},
        )
        print(f"❌ 文本LLM预检失败，未生成完整审查报告。失败诊断已保存: {md_path}, {json_path}")
        return RunResult.failed(
            failure,
            {"markdown": str(md_path), "json": str(json_path)},
            workspace=run_workspace,
            meta={"input_path": str(input_path), "preflight_results": preflight_results},
        )
    completed_stages.append("text_llm_preflight")

    chunk_size = min(int(args.max_chars), 4096)  # LLM单块硬上限4096字符
    overlap = min(512, chunk_size // 8)  # 重叠区约12.5%，最多512字符

    chunks = smart_chunk_text(audit_text, chunk_size=chunk_size, overlap=overlap)
    total_chunks = len(chunks)
    llm_cache_key = _text_fingerprint(audit_text, f"{LLM_API_URL}|{LLM_MODEL}|{chunk_size}|{overlap}|refs_excluded")
    llm_cache_dir = resume_dir / f"llm_{llm_cache_key}"
    llm_cache_dir.mkdir(parents=True, exist_ok=True)
    resume_event(resume_dir, "stage3_llm", "start", f"chunks={total_chunks}; chunk_size={chunk_size}; overlap={overlap}", cache_dir=str(llm_cache_dir))

    progress_bar(2, 5, f"阶段3/5 开始LLM审查：{total_chunks}块")

    if total_chunks == 1:
        # 短论文：直接全文审查
        print(f"🔍 论文长度({len(audit_text)}字符，已排除参考文献)在单块范围内，直接审查...")
        single_cache = llm_cache_dir / "chunk_0000.json"
        cached = _json_load(single_cache) if allow_llm_cache_read else None
        if cached:
            print(f"🔁 断点续作：复用LLM审查缓存 {single_cache}")
            resume_event(resume_dir, "stage3_llm_chunk", "cache_hit", "chunk=1/1", cache=str(single_cache))
            report = cached.get("report", {"parse_error": True, "raw_output": "缓存格式异常"})
        else:
            try:
                raw_content = ""
                report = {"parse_error": True, "raw_output": ""}
                schema_errors = []
                for schema_attempt in range(2):
                    raw_content = call_llm(audit_text)
                    report = parse_report(raw_content)
                    if not report.get("parse_error"):
                        break
                    schema_errors = report.get("schema_errors") or [report.get("raw_output", "")[:180]]
                    if schema_attempt == 0:
                        print(f"  ↻ LLM证据schema不合格，重试1次: {schema_errors}")
                if report.get("parse_error"):
                    raise RuntimeError(f"LLM返回结构不符合证据schema: {schema_errors}")
                if allow_llm_cache_write:
                    _json_save(single_cache, {"report": report, "raw_content": raw_content, "saved_at": time.strftime("%F %T")})
                    resume_event(resume_dir, "stage3_llm_chunk", "saved", "chunk=1/1", cache=str(single_cache))
            except Exception as e:
                print(f"❌ LLM调用失败: {e}")
                failure = AuditFailure(
                    capability="text_llm",
                    error_class="schema_error",
                    message=f"LLM语义审查失败或返回结构不符合证据schema: {e}",
                    fix_hints=["检查文本LLM服务稳定性和提示词输出格式。", "稍后重试或更换稳定的文本语义审查服务。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                    details={"raw_error": str(e), "chunk": "1/1"},
                )
                md_path, json_path = save_failed_audit_diagnostics(failure, input_path, **failed_artifact_kwargs, meta=meta)
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
    else:
        # 长论文：分块审查 + 合并
        print(f"🔍 论文较长({len(audit_text)}字符，已排除参考文献)，分为{total_chunks}块(每块≤{chunk_size}字符，重叠{overlap}字符)进行审查...")
        chunk_reports = [None] * total_chunks
        failed_chunks = []

        def _run_chunk_once(chunk_text, chunk_idx, retry=False):
            chunk_cache = llm_cache_dir / f"chunk_{chunk_idx:04d}.json"
            print(("  🔁 重试" if retry else "  📝 审查") + f"第{chunk_idx+1}/{total_chunks}块({len(chunk_text)}字符)...")
            raw_content = call_llm(chunk_text, chunk_info=(chunk_idx, total_chunks))
            chunk_report = parse_report(raw_content)
            if chunk_report.get("parse_error"):
                raise RuntimeError(f"LLM返回解析失败: {str(chunk_report.get('raw_output',''))[:180]}")
            if allow_llm_cache_write:
                _json_save(chunk_cache, {"report": chunk_report, "raw_content": raw_content, "saved_at": time.strftime("%F %T"), "chunk_index": chunk_idx, "total_chunks": total_chunks, "status": "ok", "retry": retry})
                resume_event(resume_dir, "stage3_llm_chunk", "retry_saved" if retry else "saved", f"chunk={chunk_idx+1}/{total_chunks}; chars={len(chunk_text)}", cache=str(chunk_cache))
            return chunk_report

        for chunk_text, chunk_idx, _ in chunks:
            progress_bar(chunk_idx, total_chunks, f"阶段3/5 LLM审查中：第{chunk_idx+1}/{total_chunks}块")
            chunk_cache = llm_cache_dir / f"chunk_{chunk_idx:04d}.json"
            cached = _json_load(chunk_cache) if allow_llm_cache_read else None
            if cached and cached.get("status") == "ok" and cached.get("report") and not cached.get("report", {}).get("parse_error"):
                print(f"     ↳ 断点续作：复用第{chunk_idx+1}块成功LLM缓存")
                resume_event(resume_dir, "stage3_llm_chunk", "cache_hit", f"chunk={chunk_idx+1}/{total_chunks}", cache=str(chunk_cache))
                chunk_reports[chunk_idx] = cached.get("report")
            elif getattr(args, "llm_cache_only", False):
                print(f"     ↳ cache-only：第{chunk_idx+1}块无成功缓存，跳过API调用")
                failed_chunks.append((chunk_text, chunk_idx, "cache_only_no_success_cache"))
                resume_event(resume_dir, "stage3_llm_chunk", "cache_only_miss", f"chunk={chunk_idx+1}/{total_chunks}", cache=str(chunk_cache))
            else:
                try:
                    chunk_reports[chunk_idx] = _run_chunk_once(chunk_text, chunk_idx, retry=False)
                except Exception as e:
                    print(f"  ⚠️ 第{chunk_idx+1}块LLM调用/解析失败，先记录并继续其他块: {e}")
                    failed_chunks.append((chunk_text, chunk_idx, str(e)))
                    if allow_llm_cache_write:
                        _json_save(chunk_cache, {"report": {"parse_error": True, "raw_output": str(e)}, "raw_content": str(e), "saved_at": time.strftime("%F %T"), "chunk_index": chunk_idx, "total_chunks": total_chunks, "status": "failed_pending_retry"})
                        resume_event(resume_dir, "stage3_llm_chunk", "failed_pending_retry", f"chunk={chunk_idx+1}/{total_chunks}; error={e}", cache=str(chunk_cache))
            if chunk_reports[chunk_idx] and not chunk_reports[chunk_idx].get("parse_error"):
                print(f"     → 第{chunk_idx+1}块风险: {chunk_reports[chunk_idx].get('risk_level', '未知')}")
            progress_bar(chunk_idx + 1, total_chunks, f"阶段3/5 LLM审查完成：第{chunk_idx+1}/{total_chunks}块")

        if failed_chunks:
            print(f"🔁 首轮完成，按顺序重试失败块: {[idx+1 for _, idx, _ in failed_chunks]}")
            resume_event(resume_dir, "stage3_llm_retry", "start", f"failed_chunks={[idx+1 for _, idx, _ in failed_chunks]}; cache_only={getattr(args, 'llm_cache_only', False)}")
            still_failed = []
            if getattr(args, "llm_cache_only", False):
                still_failed = [(idx, first_error) for _, idx, first_error in failed_chunks]
                print("⚠️ cache-only模式：不调用API重试，直接用已有成功缓存生成部分报告。")
            else:
                for chunk_text, chunk_idx, first_error in failed_chunks:
                    try:
                        chunk_reports[chunk_idx] = _run_chunk_once(chunk_text, chunk_idx, retry=True)
                        print(f"     ✅ 第{chunk_idx+1}块重试成功")
                    except Exception as e:
                        print(f"     ❌ 第{chunk_idx+1}块重试仍失败: {e}")
                        still_failed.append((chunk_idx, str(e)))
                        chunk_cache = llm_cache_dir / f"chunk_{chunk_idx:04d}.json"
                        if allow_llm_cache_write:
                            _json_save(chunk_cache, {"report": {"parse_error": True, "raw_output": str(e)}, "raw_content": str(e), "saved_at": time.strftime("%F %T"), "chunk_index": chunk_idx, "total_chunks": total_chunks, "status": "failed_final", "first_error": first_error})
                            resume_event(resume_dir, "stage3_llm_chunk", "failed_final", f"chunk={chunk_idx+1}/{total_chunks}; error={e}", cache=str(chunk_cache))
            if still_failed:
                failed_nums = [idx + 1 for idx, _ in still_failed]
                detail = "; ".join([f"第{idx+1}块: {err}" for idx, err in still_failed])
                resume_event(resume_dir, "stage3_llm_retry", "still_failed", f"still_failed={failed_nums}; strict={args.strict_failed_chunks}")
                failure = AuditFailure(
                    capability="text_llm",
                    error_class="schema_error",
                    message="LLM分块重试后仍失败，停止生成完整审查报告: " + detail,
                    fix_hints=["检查文本LLM服务稳定性和严格证据schema输出。", "更换稳定服务或稍后重试。"],
                    completed_stages=completed_stages,
                    retry_command=retry_command,
                    details={"failed_chunks": failed_nums, "detail": detail},
                )
                md_path, json_path = save_failed_audit_diagnostics(failure, input_path, **failed_artifact_kwargs, meta=meta)
                record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
                return RunResult.failed(
                    failure,
                    {"markdown": str(md_path), "json": str(json_path)},
                    workspace=run_workspace,
                    meta={"input_path": str(input_path)},
                )
            else:
                resume_event(resume_dir, "stage3_llm_retry", "done", "all failed chunks recovered")

        successful_count = sum(1 for r in chunk_reports if r is not None and not r.get("parse_error"))
        failed_final = []
        for idx in range(total_chunks):
            if chunk_reports[idx] is None or chunk_reports[idx].get("parse_error"):
                failed_final.append(idx + 1)
        meta["llm_success_chunks"] = successful_count
        meta["llm_failed_chunks"] = failed_final
        meta["llm_coverage"] = f"{successful_count}/{total_chunks}"
        meta["llm_partial_report"] = bool(failed_final)

        chunk_reports = [r for r in chunk_reports if r is not None and not r.get("parse_error")]
        if not chunk_reports:
            message = f"所有LLM分块均失败，无法生成语义审查报告。失败块: {failed_final}。"
            resume_event(resume_dir, "stage4_merge", "skipped_no_success", message)
            failure = AuditFailure(
                capability="text_llm",
                error_class="schema_error",
                message=message,
                fix_hints=["检查文本LLM服务和证据schema输出。", "更换稳定服务后重试。"],
                completed_stages=completed_stages,
                retry_command=retry_command,
                details={"failed_chunks": failed_final},
            )
            md_path, json_path = save_failed_audit_diagnostics(failure, input_path, **failed_artifact_kwargs, meta=meta)
            record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
            return RunResult.failed(
                failure,
                {"markdown": str(md_path), "json": str(json_path)},
                workspace=run_workspace,
                meta={"input_path": str(input_path)},
            )
        else:

            progress_bar(3, 5, "阶段3/5 LLM审查完成")
            # 合并所有块的审查结果
            progress_bar(3, 5, "阶段4/5 开始合并审查结果")
            print(f"🔗 正在合并{len(chunk_reports)}块审查结果...")
            report = merge_chunk_reports(chunk_reports, stat_result)
            if report.get("_merged_from"):
                print(f"✅ 合并完成: 来自{report['_merged_from']}块，共{len(report.get('checks', []))}个检查项")
            meta["chunk_count"] = total_chunks
            meta["chunk_size"] = chunk_size
            meta["overlap"] = overlap
            if meta.get("llm_partial_report"):
                warning = f"注意：本报告仅覆盖 {meta.get('llm_coverage')} 个LLM分块；失败块: {meta.get('llm_failed_chunks')}。结论不完整，建议换稳定API后断点续跑。"
                report["_partial_warning"] = warning
                report["summary"] = warning + " " + str(report.get("summary", ""))
            resume_event(resume_dir, "stage4_merge", "done", f"checks={len(report.get('checks', [])) if isinstance(report, dict) else 'N/A'}; coverage={meta.get('llm_coverage')}")
            progress_bar(4, 5, "阶段4/5 审查结果合并完成")

    # ─── 图像合理性检测：使用MinerU已保存zip中的图片/目录图片生成报告清单 ───
    image_semantic_cache_path = resume_dir / "image_semantic_cache.json"
    image_semantic_local_cache_path = output_dir / "image_semantic_cache.json"
    image_semantic_cache = {} if args.no_resume else _load_merged_json_dicts(
        image_semantic_local_cache_path,
        image_semantic_cache_path,
    )
    image_detector_cache_path = resume_dir / "image_detector_cache.json"
    image_detector_cache = {} if args.no_resume else (_json_load(image_detector_cache_path, {}) or {})
    image_semantic_enabled = not args.no_image_semantic and bool(GLM_API_KEY)
    image_detector_enabled = not args.no_image_detector
    if not args.no_image_semantic and not GLM_API_KEY:
        print("⚠️ 图像语义分析API Key未配置，图像语义分析将跳过；本地合理性检测和imagedetector清单仍会生成")
    image_semantic_cache_save = None
    if image_semantic_enabled and not args.no_resume:
        def image_semantic_cache_save():
            _json_save(image_semantic_cache_path, image_semantic_cache)
            _json_save(image_semantic_local_cache_path, image_semantic_cache)
    image_detector_cache_save = None
    if image_detector_enabled and not args.no_resume:
        image_detector_cache_save = lambda: _json_save(image_detector_cache_path, image_detector_cache)
    image_audit = build_image_audit(
        str(input_path),
        output_dir=output_dir,
        limit=args.image_audit_limit,
        semantic=image_semantic_enabled,
        semantic_limit=args.image_semantic_limit,
        semantic_timeout=args.image_semantic_timeout,
        semantic_cache=image_semantic_cache,
        semantic_cache_save=image_semantic_cache_save,
        detector=image_detector_enabled,
        detector_limit=args.image_detector_limit,
        detector_timeout=args.image_detector_timeout,
        detector_cache=image_detector_cache,
        detector_cache_save=image_detector_cache_save,
    )
    if image_semantic_enabled and not args.no_resume:
        _json_save(image_semantic_cache_path, image_semantic_cache)
        _json_save(image_semantic_local_cache_path, image_semantic_cache)
        resume_event(
            resume_dir,
            "stage4_image_semantic",
            "saved",
            f"semantic_checked={image_audit.get('semantic_checked', 0)}; cache_entries={len(image_semantic_cache)}; local_cache={image_semantic_local_cache_path}",
            cache=str(image_semantic_cache_path),
        )
    if image_detector_enabled and not args.no_resume:
        _json_save(image_detector_cache_path, image_detector_cache)
        resume_event(
            resume_dir,
            "stage4_image_detector",
            "saved",
            f"detector_checked={image_audit.get('detector_checked', 0)}; cache_entries={len(image_detector_cache)}",
            cache=str(image_detector_cache_path),
        )
    meta["image_audit"] = image_audit
    if image_audit.get("image_count"):
        print(f"🖼️ 图像检测完成: 本地{image_audit.get('checked_count')}/{image_audit.get('image_count')}张；图像语义分析 {image_audit.get('semantic_checked')}张；imagedetector {image_audit.get('detector_checked')}张")
        manifest_path = save_image_review_manifest(image_audit, output_dir)
        if manifest_path:
            meta["image_review_manifest"] = str(manifest_path)
            print(f"🖼️ 图像AI检测结果清单已保存: {manifest_path}")
    failed_capability, failed_message, failed_details = coverage_blocking_failure(meta)
    if failed_capability:
        failure = AuditFailure(
            capability=failed_capability,
            error_class="provider_unavailable",
            message=failed_message,
            fix_hints=["检查第三方服务配置、网络情况和服务商状态后重试。"],
            completed_stages=completed_stages,
            retry_command=retry_command,
            details=failed_details,
        )
        md_path, json_path = save_failed_audit_diagnostics(failure, input_path, **failed_artifact_kwargs, meta=meta)
        record_run_workspace_artifacts(run_workspace, "failed", [md_path, json_path], meta={"completed_stages": completed_stages})
        return RunResult.failed(
            failure,
            {"markdown": str(md_path), "json": str(json_path)},
            workspace=run_workspace,
            meta={"input_path": str(input_path)},
        )
    report = apply_risk_rules(report, stat_result=stat_result, image_audit=meta.get("image_audit"))
    meta["risk_rule_version"] = RISK_RULE_VERSION
    meta["evidence_chain_audit"] = build_evidence_chain_audit(
        full_text,
        extracted_file_texts,
        report,
        meta,
        stat_result,
    )

    # ─── 阶段5：生成报告 ───
    progress_bar(4, 5, "阶段5/5 开始生成报告")
    limited_reasons = audit_limited_reasons(args, meta, has_pdf_input=has_pdf_input)
    apply_audit_artifact_type(meta, limited_reasons)
    report_actions_port = int(getattr(args, "report_actions_port", 8765) or 8765)
    meta["report_actions"] = {
        "host": "127.0.0.1",
        "port": report_actions_port,
        "url": report_action_service_url("127.0.0.1", report_actions_port),
        "auto_start": not bool(getattr(args, "no_open", False)),
    }
    # 确定输出路径（优先HTML）
    output_override = explicit_output_path_from_args(args)
    output_path, html_output_path, json_path = audit_artifact_paths(
        input_path,
        artifact_type=meta.get("artifact_type", "complete"),
        output_path=output_override,
    )
    meta["artifact_paths"] = {
        "markdown": str(output_path),
        "html": str(html_output_path),
        "json": str(json_path),
    }
    meta["followups_dir"] = str(html_output_path.parent / "followups")
    report_input = str(input_path)
    md_report = format_report(report, report_input, meta, stat_result)

    # 生成HTML报告
    html_report = format_html_report(report, report_input, meta, stat_result)

    # 写入Markdown报告
    output_path.write_text(md_report, encoding="utf-8")
    print(f"✅ Markdown报告已保存: {output_path}")
    resume_event(resume_dir, "stage5_report", "markdown_saved", str(output_path))

    # 写入HTML报告
    html_output_path.write_text(html_report, encoding="utf-8")
    print(f"✅ HTML报告已保存: {html_output_path}")
    resume_event(resume_dir, "stage5_report", "html_saved", str(html_output_path))

    if args.json:
        json_path.write_text(
            json.dumps({"report_type": meta.get("artifact_type", "complete"), "llm_report": report, "stat_result": stat_result, "meta": meta, "reference_audit": reference_audit, "resource_audit": resource_audit, "cross_file_consistency_audit": meta.get("cross_file_consistency_audit"), "evidence_chain_audit": meta.get("evidence_chain_audit")},
                       ensure_ascii=False, indent=2),
            encoding="utf-8")
        print(f"✅ 原始JSON已保存: {json_path}")

    record_run_workspace_artifacts(
        run_workspace,
        meta.get("artifact_type", "complete"),
        [output_path, html_output_path, json_path],
        meta={"artifact_type": meta.get("artifact_type"), "limited_reasons": meta.get("limited_reasons", [])},
    )

    resume_event(resume_dir, "all", "done", "audit completed")
    progress_bar(5, 5, "阶段5/5 全部完成")
    print(f"🧾 完整日志: {_RUN_LOG_FILE}")

    # 自动打开HTML报告
    if args.no_open:
        print(f"🌐 已跳过自动打开HTML报告: {html_output_path}")
    else:
        action_status = ensure_report_action_service(
            port=report_actions_port,
            log_path=Path(run_workspace["run_dir"]) / "report_actions.log",
        )
        meta["report_actions"]["service_status"] = action_status.get("status")
        if action_status.get("ok"):
            if action_status.get("status") == "already_running":
                print(f"🌐 HTML动作服务已在运行: {action_status.get('url')}")
            else:
                print(f"🌐 HTML动作服务已自动启动: {action_status.get('url')}")
        else:
            print(f"⚠️ HTML动作服务自动启动失败: {action_status.get('error') or action_status.get('status')}，报告仍可打开")
        try:
            open_html_artifact(html_output_path)
            print(f"🌐 已在浏览器中打开HTML报告")
        except Exception as e:
            print(f"⚠️ 自动打开浏览器失败: {e}，请手动打开: {html_output_path}")

    # 打印摘要
    if not report.get("parse_error"):
        risk = report.get("risk_level", "未知")
        print(f"\n📊 复核优先级: {risk} | 总评: {report.get('summary', 'N/A')}")

    artifact_paths = {"markdown": str(output_path), "html": str(html_output_path)}
    if json_path.exists():
        artifact_paths["json"] = str(json_path)
    result_factory = RunResult.limited if meta.get("artifact_type") == "limited" else RunResult.complete
    return result_factory(artifact_paths, workspace=run_workspace, meta=meta)

def main():
    global LLM_TIMEOUT, LLM_RETRIES
    parser = argparse.ArgumentParser(
        description="学术论文自动审查工具（耿同学标准 + MinerU）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 默认：PDF优先使用MinerU提取 + 完整审查
  python paper_audit.py paper.pdf

  # 调试/范围受限：禁用MinerU，不作为完整正式审查
  python paper_audit.py paper.pdf --no-mineru

  # 指定输出路径
  python paper_audit.py paper.pdf --mineru -o report.md --json
  
  # 更新欺诈模式知识库（从PubPeer评论自动提取新pattern）
  python paper_audit.py --update-patterns pubpeer_comments.txt
""")
    parser.add_argument("pdf_path", nargs='?', help="待审查的文件路径或论文目录路径（支持PDF、Word .docx、Excel、Supplement等，更新/服务模式下无需提供）")
    parser.add_argument("--serve-report-actions", action="store_true",
                        help="启动本机HTML报告动作服务：一键生成PubPeer comment和期刊letter")
    parser.add_argument("--report-actions-port", type=int, default=8765,
                        help="HTML报告动作服务端口（默认8765，仅监听127.0.0.1）")
    parser.add_argument("--serve-web", action="store_true",
                        help="启动本机Web Runner工作台：通过浏览器运行审查、查看日志和打开产物")
    parser.add_argument("--gui", action="store_true",
                        help="启动本机桌面GUI软件：选择输入、运行审查并直接打开报告产物")
    parser.add_argument("--web-port", type=int, default=8765,
                        help="Web Runner端口（默认8765，仅监听127.0.0.1）")
    parser.add_argument("--update-patterns", metavar="COMMENTS_FILE", 
                        help="从PubPeer评论文本文件中自动提取新的欺诈模式，更新知识库")
    parser.add_argument("--mineru", action="store_true",
                        help="使用MinerU API将PDF转为Markdown再审查（PDF默认已启用，保留该参数用于兼容旧命令）")
    parser.add_argument("--mineru-model", default="vlm",
                        choices=["pipeline", "vlm", "MinerU-HTML"],
                        help="MinerU模型版本（默认vlm，仅Precision API生效）")
    parser.add_argument("--mineru-lang", default="ch",
                        help="MinerU OCR语言（默认ch=中英，en=英文，japan=日文）")
    parser.add_argument("--no-mineru", action="store_true",
                        help="调试/范围受限：禁用MinerU；不能作为完整正式审查")
    parser.add_argument("--max-chars", type=int, default=4096,
                        help="LLM分块单块最大字符数（默认4096；超过4096会自动压到4096）")
    parser.add_argument("--output", "-o", help="输出报告文件路径（默认输出到同目录）")
    parser.add_argument("--json", action="store_true", help="同时保存原始JSON结果")
    parser.add_argument("--ai-detect", action="store_true", help="开启腾讯朱雀AI文本检测：自动复制文本到剪贴板+打开检测页面")
    parser.add_argument("--image-detect", action="store_true", help="兼容旧参数：图片检测已默认自动执行，不再打开网页或要求手动上传")
    parser.add_argument("--image-audit-limit", type=int, default=None,
                        help="报告中纳入图像合理性检测的图片数量上限（默认全部；设置后为范围受限审查）")
    parser.add_argument("--no-image-semantic", action="store_true",
                        help="调试/范围受限：关闭图像语义分析；存在可检测图片时不能作为完整正式审查")
    parser.add_argument("--image-semantic-limit", type=int, default=None,
                        help="调用图像语义分析的图片数量上限（默认全部；设置后为范围受限审查）")
    parser.add_argument("--image-semantic-timeout", type=int, default=45,
                        help="单张图片图像语义分析请求超时时间秒数（默认45）")
    parser.add_argument("--no-image-detector", action="store_true",
                        help="调试/范围受限：关闭imagedetector.com自动图片AI概率检测；存在可检测图片时不能作为完整正式审查")
    parser.add_argument("--image-detector-limit", type=int, default=None,
                        help="自动调用imagedetector.com检测的图片数量上限（默认全部；设置后为范围受限审查）")
    parser.add_argument("--image-detector-timeout", type=int, default=60,
                        help="单张图片imagedetector自动检测超时时间秒数（默认60）")
    parser.add_argument("--no-reference-online", action="store_true",
                        help="调试/范围受限：关闭参考文献在线真实性检索；识别到参考文献时不能作为完整正式审查")
    parser.add_argument("--reference-online-limit", type=int, default=None,
                        help="参考文献在线检索条数上限（默认全部；设置后为范围受限审查）")
    parser.add_argument("--reference-timeout", type=int, default=10,
                        help="单个参考文献外部检索源超时时间秒数（默认10）")
    parser.add_argument("--no-resource-online", action="store_true",
                        help="调试/范围受限：关闭代码仓库与在线资源可用性校检；识别到资源时不能作为完整正式审查")
    parser.add_argument("--resource-timeout", type=int, default=10,
                        help="单个代码仓库/在线资源可用性请求超时时间秒数（默认10）")
    parser.add_argument("--no-resume", action="store_true",
                        help="禁用断点续作缓存，强制重新提取文本和重新LLM审查")
    parser.add_argument("--fresh", action="store_true",
                        help="运行前清空本输入的断点续作缓存，然后重新开始；默认不清空缓存并自动续跑")
    parser.add_argument("--llm-timeout", type=int, default=LLM_TIMEOUT,
                        help=f"单次LLM请求超时时间秒数（默认{LLM_TIMEOUT}；不稳定网关可调小以更快跳过）")
    parser.add_argument("--llm-retries", type=int, default=LLM_RETRIES,
                        help=f"每次LLM调用内部重试次数（默认{LLM_RETRIES}，即最多{LLM_RETRIES + 1}次尝试）")
    parser.add_argument("--strict-failed-chunks", action="store_true",
                        help="严格模式：失败块首轮+补跑仍失败时停止生成报告；当前默认仍生成覆盖率受限报告，后续将改为失败诊断")
    parser.add_argument("--llm-cache-only", action="store_true",
                        help="调试/范围受限：只复用已有成功LLM分块缓存，不再调用API；不能作为完整正式审查")
    parser.add_argument("--no-open", action="store_true",
                        help="生成报告后不自动打开HTML报告，适合CI、服务器和批处理环境")
    args = parser.parse_args()
    if getattr(args, "max_chars", 4096) > 4096:
        print(f"⚠️ --max-chars={args.max_chars} 超过4096，已自动调整为4096")
        args.max_chars = 4096
    if getattr(args, "max_chars", 4096) < 512:
        print(f"⚠️ --max-chars={args.max_chars} 过小，已自动调整为512")
        args.max_chars = 512
    LLM_TIMEOUT = max(10, int(getattr(args, "llm_timeout", LLM_TIMEOUT) or LLM_TIMEOUT))
    LLM_RETRIES = max(0, int(getattr(args, "llm_retries", LLM_RETRIES) or 0))

    # ─── 知识库更新模式 ───
    if args.update_patterns:
        apply_runtime_config(load_runtime_config())
        return update_patterns(args.update_patterns)
    if args.serve_report_actions:
        apply_runtime_config(load_runtime_config())
        return serve_report_actions(port=args.report_actions_port)
    if args.serve_web:
        apply_runtime_config(load_runtime_config())
        return serve_web_runner(port=args.web_port, open_browser=not args.no_open)
    if args.gui:
        return gui_main()

    # ─── 正常审查模式 ───
    if not args.pdf_path:
        parser.error("审查模式需要提供path参数（文件或目录，或使用 --update-patterns / --serve-web / --gui 更新知识库或启动服务）")

    run_request = RunRequest.from_args(args)
    return run_audit(run_request, args).exit_code


if __name__ == "__main__":
    exit(main())
