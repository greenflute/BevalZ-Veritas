#!/usr/bin/env python3
"""Paper Audit - 学术论文自动审查工具 [耿同学版]
基于3个开源项目思路开发：
- wooly99/geng-academic-fraud-detector 耿同学六式
- NeoSpecies/AcademicIntegrityHunter 本地统计算法
- jingshouyan/academic-integrity-geng 五维审查体系
输入PDF路径 → MinerU转Markdown → 本地统计检测 + LLM语义分析 → 输出md格式报告
用法: python paper_audit.py <pdf_path> [--mineru] [--max-chars 8000] [--output report.md]
"""
import re, json, time, argparse, urllib.request, zlib, math, collections, os, mimetypes, fnmatch, csv, platform, webbrowser, subprocess, sys
from pathlib import Path
from typing import Tuple, Dict, List

# 可选依赖：处理Word/Excel/Supplement文件
try:
    from docx import Document
    DOCX_SUPPORTED = True
except ImportError:
    DOCX_SUPPORTED = False
try:
    from openpyxl import load_workbook
    EXCEL_SUPPORTED = True
except ImportError:
    EXCEL_SUPPORTED = False

# ══════════════════════════════════════════════════════════════
# 配置区
# ══════════════════════════════════════════════════════════════
import importlib

# 尝试加载外部配置文件（优先config.py）
try:
    config = importlib.import_module("config")
    LLM_API_KEY = getattr(config, "LLM_API_KEY", "")
    LLM_API_URL = getattr(config, "LLM_API_URL", "https://api.openai.com/v1/chat/completions")
    LLM_MODEL = getattr(config, "LLM_MODEL", "gpt-3.5-turbo")
    MINERU_TOKEN = getattr(config, "MINERU_TOKEN", "")
except ImportError:
    # 未找到config.py使用默认OpenAI兼容配置，按实际使用修改
    LLM_API_KEY = ""
    LLM_API_URL = "https://api.openai.com/v1/chat/completions"
    LLM_MODEL = "gpt-3.5-turbo"
    MINERU_TOKEN = ""

MINERU_BASE = "https://mineru.net"

# ─── 欺诈模式知识库加载 ───
FRAUD_PATTERNS_PATH = Path(__file__).parent / "fraud_patterns.json"
FRAUD_PATTERNS = []
PATTERN_HINTS = ""
if FRAUD_PATTERNS_PATH.exists():
    try:
        with open(FRAUD_PATTERNS_PATH, "r", encoding="utf-8") as f:
            pattern_data = json.load(f)
            FRAUD_PATTERNS = pattern_data.get("patterns", [])
        print(f"✅ 加载欺诈模式知识库成功，共{len(FRAUD_PATTERNS)}条检测模式")
        # 构建提示词片段
        PATTERN_HINTS = "\n## 最新欺诈模式知识库（社区贡献+PubPeer案例汇总）\n"
        for idx, p in enumerate(FRAUD_PATTERNS, 1):
            PATTERN_HINTS += f"{idx}. [{p['risk_level']}风险] {p['name']}：{p['detection_hint']}\n"
    except Exception as e:
        print(f"⚠️ 知识库加载失败: {e}，使用默认检测规则")

# ══════════════════════════════════════════════════════════════
# 审查体系配置 - LLM System Prompt
# ══════════════════════════════════════════════════════════════
SYSTEM_PROMPT_TPL = """你是一个严厉的学术论文审查专家（耿同学标准）。
你需要结合以下维度对输入的论文文本进行审查，输出严格的JSON格式：

## 审查维度
1. 数据与结果自洽性 — 数字前后矛盾、统计量不一致、图表数据不匹配
2. 图片与图表异常 — 描述性分析图片可疑特征（旋转复用、背景一致、拼接痕迹）
3. 方法论严谨性 — 样本量不足、缺乏多重比较校正、实验设计缺陷
4. 结构与引用规范性 — 自引率异常、引用质量差、逻辑谬误
5. 作者与期刊可信度 — 产出异常、利益冲突未披露、同行评审缺失

## 检查项（耿同学六式 + 7类红旗）
- 耿同学六式：图片复用/数据造假/图片拼接/统计异常/产出异常/方法矛盾
- 7类红旗：引用质量差/逻辑谬误/方法论缺陷/可疑结论/同行评审缺失/利益冲突未披露/语言质量差

请按以下JSON格式输出（确保JSON合法，无多余内容）：
{{
  "summary": "一句话总评",
  "risk_level": "高/中/低/可疑黑产",
  "detection_score": 0,
  "checks": [
    {{
      "category": "数据与结果/图片与图表/方法论/结构与引用/作者与期刊",
      "item": "检查项名称",
      "verdict": "🚩红旗/⚠️疑点/✅通过",
      "evidence": "具体证据，引用原文片段",
      "detail": "详细分析说明"
    }}
  ],
  "conclusion": "综合结论与行动建议"
}}
{pattern_hints}
"""

# 动态构建系统提示词
SYSTEM_PROMPT = SYSTEM_PROMPT_TPL.format(pattern_hints=PATTERN_HINTS)

# ══════════════════════════════════════════════════════════════
# MinerU API 模块 — PDF转Markdown
# ══════════════════════════════════════════════════════════════

def _http_request(url, method="GET", headers=None, data=None, timeout=60):
    """通用HTTP请求封装（纯标准库）"""
    if headers is None:
        headers = {}
    req = urllib.request.Request(url, data=data, headers=headers, method=method)
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        return resp.read(), resp.status


def mineru_precision_extract_by_url(pdf_url, model_version="vlm", language="ch",
                                     poll_interval=10, poll_timeout=600):
    """🎯 Precision API — 通过URL解析PDF（需要Token，≤200MB/200页）

    流程：POST创建任务 → GET轮询结果 → 下载zip中的Markdown
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    print(f"  🎯 [MinerU Precision] 提交URL任务: {pdf_url[:80]}...")

    # 1. 创建提取任务
    create_url = f"{MINERU_BASE}/api/v4/extract/task"
    payload = json.dumps({"url": pdf_url, "model_version": model_version}).encode()
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MINERU_TOKEN}"
    }
    try:
        resp_data, status = _http_request(create_url, "POST", headers, payload, timeout=30)
        result = json.loads(resp_data.decode())
    except Exception as e:
        return None, {"error": f"创建任务失败: {e}"}

    if result.get("code") != 0 and not result.get("data", {}).get("batch_id"):
        return None, {"error": f"创建任务返回异常: {result}"}

    batch_id = result.get("data", {}).get("batch_id")
    if not batch_id:
        return None, {"error": f"未获取到batch_id: {result}"}
    print(f"  ✅ 任务已创建: batch_id={batch_id}")

    # 2. 轮询任务状态
    poll_url = f"{MINERU_BASE}/api/v4/extract/task/{batch_id}"
    start = time.time()
    state_labels = {"processing": "处理中", "queued": "排队中"}

    while time.time() - start < poll_timeout:
        try:
            resp_data, _ = _http_request(poll_url, "GET", headers, timeout=30)
            result = json.loads(resp_data.decode())
        except Exception as e:
            print(f"  ⚠️ 轮询异常: {e}")
            time.sleep(poll_interval)
            continue

        task_list = result.get("data", {}).get("task_list", [])
        if not task_list:
            # 单文件模式
            state = result.get("data", {}).get("state", "unknown")
        else:
            state = task_list[0].get("state", "unknown")

        elapsed = int(time.time() - start)

        if state == "done":
            # 获取zip下载链接
            zip_url = task_list[0].get("zip_url") if task_list else result.get("data", {}).get("zip_url")
            if not zip_url:
                return None, {"error": "任务完成但未获取到下载链接"}

            print(f"  ✅ [{elapsed}s] 解析完成，下载Markdown...")
            markdown = _download_zip_and_extract_md(zip_url)
            if markdown:
                meta = {"source": "mineru_precision", "batch_id": batch_id,
                        "model": model_version, "chars": len(markdown)}
                return markdown, meta
            else:
                return None, {"error": "下载或解压zip失败"}

        elif state == "failed":
            err = task_list[0].get("err_msg", "未知") if task_list else "未知"
            return None, {"error": f"任务失败: {err}"}

        label = state_labels.get(state, state)
        print(f"  ⏳ [{elapsed}s] {label}...")
        time.sleep(poll_interval)

    return None, {"error": f"轮询超时({poll_timeout}s), batch_id={batch_id}"}


def mineru_agent_extract_by_file(file_path, language="ch",
                                  poll_interval=5, poll_timeout=300):
    """⚡ Agent API — 上传本地文件解析（无需Token，IP限流≤10MB/20页）

    流程：POST上传文件 → 轮询 → 下载Markdown
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    file_path = Path(file_path)
    file_size_mb = file_path.stat().st_size / 1024 / 1024
    print(f"  ⚡ [MinerU Agent] 上传文件: {file_path.name} ({file_size_mb:.1f}MB)")

    if file_size_mb > 10:
        print(f"  ⚠️ 文件>{10}MB，Agent API不支持，自动切换Precision API...")
        return mineru_precision_extract_by_file(file_path, language=language)

    # 1. 上传文件到 Agent API
    upload_url = f"{MINERU_BASE}/api/v1/agent/parse/file"

    # 构建 multipart/form-data
    boundary = f"----PaperAudit{int(time.time()*1000)}"
    filename = file_path.name
    mime_type = mimetypes.guess_type(str(file_path))[0] or "application/pdf"

    with open(file_path, "rb") as f:
        file_bytes = f.read()

    body_parts = []
    # language 字段
    body_parts.append(f"--{boundary}\r\n"
                      f"Content-Disposition: form-data; name=\"language\"\r\n\r\n"
                      f"{language}".encode())
    # file 字段
    body_parts.append(f"--{boundary}\r\n"
                      f"Content-Disposition: form-data; name=\"file\"; filename=\"{filename}\"\r\n"
                      f"Content-Type: {mime_type}\r\n\r\n".encode())
    body_parts.append(file_bytes)
    body_parts.append(f"\r\n--{boundary}--\r\n".encode())

    body = b"".join(body_parts)
    headers = {
        "Content-Type": f"multipart/form-data; boundary={boundary}",
    }

    try:
        resp_data, status = _http_request(upload_url, "POST", headers, body, timeout=120)
        result = json.loads(resp_data.decode())
    except Exception as e:
        return None, {"error": f"Agent上传失败: {e}"}

    if result.get("code") != 0 and not result.get("data", {}).get("task_id"):
        return None, {"error": f"Agent上传返回异常: {result}"}

    task_id = result.get("data", {}).get("task_id")
    if not task_id:
        return None, {"error": f"未获取到task_id: {result}"}
    print(f"  ✅ 上传成功: task_id={task_id}")

    # 2. 轮询任务状态
    return _mineru_agent_poll(task_id, poll_interval, poll_timeout)


def mineru_precision_extract_by_file(file_path, model_version="vlm", language="ch",
                                     poll_interval=10, poll_timeout=600):
    """🎯 Precision API — 上传本地文件解析（需要Token，≤200MB/200页）

    先上传文件获取URL，再创建提取任务。适用于大文件。
    """
    file_path = Path(file_path)
    file_size_mb = file_path.stat().st_size / 1024 / 1024
    print(f"  🎯 [MinerU Precision] 上传文件: {file_path.name} ({file_size_mb:.1f}MB)")

    # 方案：通过 Agent API 的 /file 端点获取文件URL，再用 Precision API 处理
    # 更简洁：直接使用 Agent file upload → 拿到文件URL → Precision 创建任务
    # 但实际上 Precision API 需要一个可访问的 URL

    # 退而求其次：如果文件 ≤ 10MB，用 Agent API；否则尝试用 file-urls 端点
    # Precision API 还有一个 /api/v4/file-urls/batch 端点用于上传本地文件
    # 这里我们用最通用的方式：先获取上传凭证，上传文件，再创建任务

    # 1. 获取文件上传凭证
    batch_url = f"{MINERU_BASE}/api/v4/file-urls/batch"
    payload = json.dumps({"file_names": [file_path.name]}).encode()
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MINERU_TOKEN}"
    }

    try:
        resp_data, _ = _http_request(batch_url, "POST", headers, payload, timeout=30)
        result = json.loads(resp_data.decode())
    except Exception as e:
        return None, {"error": f"获取上传凭证失败: {e}"}

    batch_data = result.get("data", {})
    batch_id = batch_data.get("batch_id")
    file_urls = batch_data.get("file_urls", [])

    if not batch_id or not file_urls:
        return None, {"error": f"上传凭证异常: {result}"}

    upload_info = file_urls[0]
    put_url = upload_info.get("upload_url") or upload_info.get("url")
    file_url = upload_info.get("url") or upload_info.get("put_url")

    # 2. PUT上传文件
    print(f"  📤 上传文件到MinerU存储...")
    with open(file_path, "rb") as f:
        file_data = f.read()

    try:
        _http_request(put_url, "PUT",
                      {"Content-Type": "application/pdf"},
                      file_data, timeout=120)
    except Exception as e:
        return None, {"error": f"PUT上传文件失败: {e}"}

    print(f"  ✅ 文件已上传，创建提取任务...")

    # 3. 创建提取任务
    # 使用已上传文件的URL（若put_url是minio的，file_url就是下载链接）
    # 尝试用 batch_id 提取
    task_url = f"{MINERU_BASE}/api/v4/extract/task"
    task_payload = json.dumps({
        "url": file_url,
        "model_version": model_version
    }).encode()
    task_headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MINERU_TOKEN}"
    }

    try:
        resp_data, _ = _http_request(task_url, "POST", task_headers, task_payload, timeout=30)
        result = json.loads(resp_data.decode())
    except Exception as e:
        return None, {"error": f"创建提取任务失败: {e}"}

    extract_batch_id = result.get("data", {}).get("batch_id", batch_id)
    print(f"  ✅ 任务已创建: batch_id={extract_batch_id}")

    # 4. 轮询
    poll_url = f"{MINERU_BASE}/api/v4/extract/task/{extract_batch_id}"
    start = time.time()
    state_labels = {"processing": "处理中", "queued": "排队中"}

    while time.time() - start < poll_timeout:
        try:
            resp_data, _ = _http_request(poll_url, "GET", task_headers, timeout=30)
            result = json.loads(resp_data.decode())
        except Exception as e:
            print(f"  ⚠️ 轮询异常: {e}")
            time.sleep(poll_interval)
            continue

        task_list = result.get("data", {}).get("task_list", [])
        state = task_list[0].get("state", "unknown") if task_list else result.get("data", {}).get("state", "unknown")
        elapsed = int(time.time() - start)

        if state == "done":
            zip_url = task_list[0].get("zip_url") if task_list else result.get("data", {}).get("zip_url")
            if zip_url:
                print(f"  ✅ [{elapsed}s] 解析完成，下载Markdown...")
                markdown = _download_zip_and_extract_md(zip_url)
                if markdown:
                    return markdown, {"source": "mineru_precision_file", "batch_id": extract_batch_id,
                                      "model": model_version, "chars": len(markdown)}
            return None, {"error": "完成但未获取到下载链接"}

        elif state == "failed":
            err = task_list[0].get("err_msg", "未知") if task_list else "未知"
            return None, {"error": f"任务失败: {err}"}

        label = state_labels.get(state, state)
        print(f"  ⏳ [{elapsed}s] {label}...")
        time.sleep(poll_interval)

    return None, {"error": f"轮询超时({poll_timeout}s)"}


def _mineru_agent_poll(task_id, poll_interval=5, poll_timeout=300):
    """轮询Agent API任务状态"""
    poll_url = f"{MINERU_BASE}/api/v1/agent/parse/{task_id}"
    start = time.time()
    state_labels = {"processing": "处理中", "queued": "排队中"}

    while time.time() - start < poll_timeout:
        try:
            resp_data, _ = _http_request(poll_url, "GET", timeout=30)
            result = json.loads(resp_data.decode())
        except Exception as e:
            print(f"  ⚠️ 轮询异常: {e}")
            time.sleep(poll_interval)
            continue

        data = result.get("data", {})
        state = data.get("state", "unknown")
        elapsed = int(time.time() - start)

        if state == "done":
            markdown_url = data.get("markdown_url")
            if markdown_url:
                print(f"  ✅ [{elapsed}s] 解析完成，下载Markdown...")
                try:
                    md_data, _ = _http_request(markdown_url, "GET", timeout=30)
                    markdown = md_data.decode("utf-8", errors="ignore")
                    return markdown, {"source": "mineru_agent", "task_id": task_id,
                                      "chars": len(markdown)}
                except Exception as e:
                    return None, {"error": f"下载Markdown失败: {e}"}
            return None, {"error": "完成但未获取到markdown_url"}

        elif state == "failed":
            err = data.get("err_msg", "未知错误")
            return None, {"error": f"Agent任务失败: {err}"}

        label = state_labels.get(state, state)
        print(f"  ⏳ [{elapsed}s] {label}...")
        time.sleep(poll_interval)

    return None, {"error": f"轮询超时({poll_timeout}s), task_id={task_id}"}


def _download_zip_and_extract_md(zip_url):
    """下载zip并提取Markdown文件（纯标准库实现）"""
    try:
        zip_data, _ = _http_request(zip_url, "GET", timeout=60)
    except Exception as e:
        print(f"  ❌ 下载zip失败: {e}")
        return None

    # 用 zipfile 从内存解析
    import zipfile, io
    try:
        with zipfile.ZipFile(io.BytesIO(zip_data)) as zf:
            # 找到 .md 文件
            md_files = [n for n in zf.namelist() if n.endswith(".md")]
            if not md_files:
                # 降级：找 .txt 或其他文本
                text_files = [n for n in zf.namelist() if n.endswith((".txt", ".mdown", ".markdown"))]
                md_files = text_files
            if not md_files:
                print(f"  ⚠️ zip中未找到Markdown文件: {zf.namelist()[:10]}")
                # 尝试任何非图片文件
                for n in zf.namelist():
                    if not any(n.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".svg", ".html"]):
                        try:
                            content = zf.read(n).decode("utf-8", errors="ignore")
                            if len(content) > 100:
                                return content
                        except:
                            continue
                return None

            # 读取最大的 .md 文件
            best = None
            best_len = 0
            for md_file in md_files:
                content = zf.read(md_file).decode("utf-8", errors="ignore")
                if len(content) > best_len:
                    best = content
                    best_len = len(content)
            return best
    except zipfile.BadZipFile:
        # 不是zip？尝试直接作为文本
        try:
            return zip_data.decode("utf-8", errors="ignore")
        except:
            return None


def mineru_extract(file_path, language="ch"):
    """MinerU统一入口：自动选择API路径

    - 本地文件 ≤10MB → Agent API（免Token，快速）
    - 本地文件 >10MB → Precision API（需Token，功能强）
    - URL → Precision API
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    file_path = Path(file_path)
    file_size_mb = file_path.stat().st_size / 1024 / 1024

    if file_size_mb <= 10:
        return mineru_agent_extract_by_file(file_path, language=language)
    else:
        return mineru_precision_extract_by_file(file_path, language=language)


# ══════════════════════════════════════════════════════════════
# 本地统计检测模块
# ══════════════════════════════════════════════════════════════

def benford_analysis(numbers):
    """Benford定律分析：识别异常数字分布（伪造数据首位数字偏离Benford分布）"""
    if len(numbers) < 100:
        return None, "样本不足(需≥100)"
    digits = [str(abs(int(n)))[0] for n in numbers if abs(int(n)) >= 1]
    if not digits:
        return None, "无有效数字"
    counts = collections.Counter(digits)
    total = len(digits)
    expected = {str(d): math.log10(1 + 1/d) * total for d in range(1, 10)}
    deviations = {}
    for d in range(1, 10):
        d_str = str(d)
        actual = counts.get(d_str, 0)
        exp = expected[d_str]
        deviation = abs(actual - exp) / exp
        deviations[d_str] = deviation
    avg_deviation = sum(deviations.values()) / 9
    return avg_deviation, "高偏差⚠️" if avg_deviation > 0.3 else "正常✅"


def extract_all_numbers(text):
    """提取文本中所有数字（排除年份、页码等噪声）"""
    exclude = {2018, 2019, 2020, 2021, 2022, 2023, 2024, 2025, 2026, 2027,
               1, 2, 3, 4, 5, 6, 7, 8, 9, 10}
    nums = []
    for match in re.finditer(r'\b(\d+\.?\d*)\b', text):
        try:
            n = float(match.group(1))
            if n not in exclude and n > 0:
                nums.append(n)
        except:
            pass
    return nums


def local_stat_check(text):
    """本地统计检测，无需LLM

    包含：Benford定律分析、p值异常检测、标准差异常、数字自洽性
    """
    result = {
        "benford_deviation": None,
        "benford_status": None,
        "p_value_count": 0,
        "p_value_abnormal": 0,
        "p_value_details": [],
        "sd_count": 0,
        "sd_abnormal": 0,
        "number_count": 0,
        "number_consistency": None,
    }
    # 提取数字
    nums = extract_all_numbers(text)
    result["number_count"] = len(nums)
    # Benford分析
    if nums:
        dev, status = benford_analysis(nums)
        result["benford_deviation"] = dev
        result["benford_status"] = status
    # p值检测
    p_matches = re.findall(r'p\s*[=<]\s*(\d+\.?\d*)', text, re.IGNORECASE)
    result["p_value_count"] = len(p_matches)
    for p in p_matches:
        try:
            pv = float(p)
            if pv > 0.05:
                result["p_value_abnormal"] += 1
                result["p_value_details"].append(f"p={'<='+p if pv<=0.001 else '='+p}")
        except:
            pass
    # 标准差异常
    sd_matches = re.findall(r'(?:std|sd|标准差|SE|SEM)\s*[=:≈]\s*(\d+\.?\d*)', text, re.IGNORECASE)
    result["sd_count"] = len(sd_matches)
    # 数字自洽性检查：提取"n=XX"样本量，检查是否有矛盾
    n_matches = re.findall(r'(?:n|N|sample|样本)\s*[=:]\s*(\d+)', text, re.IGNORECASE)
    if len(set(n_matches)) > 1:
        result["number_consistency"] = f"检测到不同样本量: {set(n_matches)}"
    return result


# ══════════════════════════════════════════════════════════════
# 目录级综合分析模块
# ══════════════════════════════════════════════════════════════

def find_project_files(root_path: Path) -> Tuple[Dict, List[Path]]:
    """递归扫描目录，识别论文项目相关的所有文件
    返回：(文件分类字典, 所有有效文件列表)
    """
    SUPPORTED_EXTS = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt", ".md"}
    SUPPLEMENT_KEYWORDS = {"supplement", "supp", "补充材料", "原始数据", "data", "source", "appendix"}
    
    file_categories = {
        "main_paper": None,
        "supplements": [],
        "data_files": [],
        "other": []
    }
    all_files = []
    
    for root, _, files in os.walk(root_path):
        for file in files:
            fpath = Path(root) / file
            ext = fpath.suffix.lower()
            if ext not in SUPPORTED_EXTS:
                continue
            
            fname = fpath.name.lower()
            all_files.append(fpath)
            
            # 分类
            if ext == ".pdf" and (file_categories["main_paper"] is None or 
                                 ("main" in fname or "paper" in fname or "article" in fname)):
                file_categories["main_paper"] = fpath
            elif any(kw in fname for kw in SUPPLEMENT_KEYWORDS):
                file_categories["supplements"].append(fpath)
            elif ext in {".xlsx", ".xlsm", ".csv"}:
                file_categories["data_files"].append(fpath)
            else:
                file_categories["other"].append(fpath)
    
    # 未找到明确主论文则取最大的PDF
    if file_categories["main_paper"] is None:
        pdf_files = [f for f in all_files if f.suffix.lower() == ".pdf"]
        if pdf_files:
            file_categories["main_paper"] = sorted(pdf_files, key=lambda x: x.stat().st_size, reverse=True)[0]
    
    return file_categories, all_files


def extract_text_from_file(file_path: Path, max_chars_per_file=5000) -> str:
    """从任意支持的文件类型中提取文本"""
    ext = file_path.suffix.lower()
    text = f"\n\n=== 文件: {file_path.name} ==="
    
    try:
        if ext == ".pdf":
            text += "\n" + extract_pdf_text(file_path, max_chars=max_chars_per_file)
        elif ext == ".docx" and DOCX_SUPPORTED:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += "\n" + para.text
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join([cell.text for cell in row.cells])
        elif ext in {".xlsx", ".xlsm"} and EXCEL_SUPPORTED:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                text += f"\n[工作表: {sheet_name}]"
                sheet = wb[sheet_name]
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if i > 1000:
                        text += "\n[数据过多，已截断]"
                        break
                    row_str = " | ".join([str(v) for v in row if v is not None])
                    if row_str.strip():
                        text += "\n" + row_str
        elif ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if i > 1000:
                        text += "\n[数据过多，已截断]"
                        break
                    text += "\n" + " | ".join(row)
        elif ext in {".txt", ".md"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text += "\n" + f.read(max_chars_per_file)
    except Exception as e:
        text += f"\n[文件解析失败: {str(e)}]"
    
    return text[:max_chars_per_file + len(f"=== 文件: {file_path.name} ===")]


# ══════════════════════════════════════════════════════════════
# PDF原始提取模块（MinerU不可用时的降级方案）
# ══════════════════════════════════════════════════════════════

def extract_pdf_text(filepath, max_chars=8000):
    """从PDF文件中提取文本（纯标准库实现，MinerU的降级方案）"""
    with open(filepath, "rb") as f:
        raw = f.read()
    parts = []
    for s in re.findall(rb"stream\r?\n(.*?)\r?\nendstream", raw, re.DOTALL):
        try:
            dec = zlib.decompress(s)
            for x in re.findall(rb"\((.*?)\)\s*Tj", dec):
                d = x.decode("latin-1", errors="ignore")
                if len(d.strip()) > 1:
                    parts.append(d)
            for bt in re.findall(rb"BT(.*?)ET", dec, re.DOTALL):
                for x in re.findall(rb"\((.*?)\)", bt):
                    d = x.decode("latin-1", errors="ignore")
                    if len(d.strip()) > 1:
                        parts.append(d)
        except:
            pass
    text = re.sub(r"\s+", " ", " ".join(parts)).strip()
    meta = {"size_mb": round(len(raw) / 1024 / 1024, 2), "total_chars": len(text),
            "extraction_method": "raw_pdf_stream"}
    return text[:max_chars], meta, raw


# ══════════════════════════════════════════════════════════════
# LLM调用模块
# ══════════════════════════════════════════════════════════════

def smart_chunk_text(text, chunk_size=8000, overlap=1000):
    """智能分块：按段落边界切割，保留重叠区确保上下文连贯

    返回: [(chunk_text, chunk_index, total_chunks), ...]
    """
    if len(text) <= chunk_size:
        return [(text, 0, 1)]

    # 按双换行分段
    paragraphs = re.split(r'\n{2,}', text)
    # 合并过短的段落
    merged_paras = []
    buf = ""
    for p in paragraphs:
        if buf and len(buf) + len(p) + 2 > chunk_size * 0.3:
            merged_paras.append(buf)
            buf = p
        else:
            buf = (buf + "\n\n" + p).strip() if buf else p
    if buf:
        merged_paras.append(buf)

    # 按chunk_size组装块
    chunks = []
    current = ""
    for para in merged_paras:
        if current and len(current) + len(para) + 2 > chunk_size:
            chunks.append(current)
            # 保留overlap：从当前块尾部取overlap长度作为下一块开头
            if overlap > 0 and len(current) > overlap:
                current = current[-overlap:] + "\n\n" + para
            else:
                current = para
        else:
            current = (current + "\n\n" + para).strip() if current else para
    if current:
        chunks.append(current)

    total = len(chunks)
    return [(c, i, total) for i, c in enumerate(chunks)]


def call_llm(text, max_retries=3, chunk_info=None):
    """调用Mimo API进行语义审查

    chunk_info: (chunk_index, total_chunks) 分块信息，None表示全文
    """
    # 构建用户消息
    if chunk_info and chunk_info[1] > 1:
        idx, total = chunk_info
        user_msg = (
            f"审查以下论文文本（第{idx+1}/{total}段，请重点关注本段内容，"
            f"同时注意与其他段落的逻辑连贯性）：\n\n{text}"
        )
    else:
        user_msg = f"审查以下论文文本：\n\n{text}"

    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_msg}
        ],
        "temperature": 0.2
    }
    data = json.dumps(payload).encode()
    req = urllib.request.Request(
        LLM_API_URL, data=data,
        headers={"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"},
        method="POST"
    )
    for attempt in range(max_retries + 1):
        try:
            with urllib.request.urlopen(req, timeout=180) as resp:
                raw = resp.read().decode("utf-8")
            result = json.loads(raw)
            return result["choices"][0]["message"]["content"]
        except Exception as e:
            if attempt < max_retries:
                time.sleep(3 * (attempt + 1))
            else:
                raise RuntimeError(f"API调用失败({attempt+1}次): {str(e)[:100]}...")


def merge_chunk_reports(reports, stat_result=None):
    """合并多块审查结果：去重、合并检查项、重新评估风险等级

    reports: [parse_report返回的dict, ...]
    """
    if len(reports) == 1:
        return reports[0]

    # 1. 收集所有检查项，按 (category, item) 去重
    seen_keys = set()
    all_checks = []
    for i, r in enumerate(reports):
        if r.get("parse_error"):
            continue
        for c in r.get("checks", []):
            key = (c.get("category", ""), c.get("item", ""))
            if key not in seen_keys:
                seen_keys.add(key)
                c["_source_chunk"] = i + 1
                all_checks.append(c)
            else:
                # 合并：如果已有同key项，补充来源信息
                for existing in all_checks:
                    ekey = (existing.get("category", ""), existing.get("item", ""))
                    if ekey == key:
                        # 保留更严重的判定
                        severity = {"🚩红旗": 3, "⚠️疑点": 2, "✅通过": 1}
                        old_s = severity.get(existing.get("verdict", ""), 0)
                        new_s = severity.get(c.get("verdict", ""), 0)
                        if new_s > old_s:
                            existing["verdict"] = c["verdict"]
                            existing["evidence"] = c.get("evidence", existing.get("evidence", ""))
                            existing["detail"] = c.get("detail", existing.get("detail", ""))
                        # 补充证据
                        if c.get("evidence") and c["evidence"] not in existing.get("evidence", ""):
                            existing["evidence"] = (existing.get("evidence", "") + 
                                                     f" [第{c.get('_source_chunk', i+1)}段补充: {c['evidence']}]")
                        break

    # 2. 统计红旗/疑点数量
    red_flags = sum(1 for c in all_checks if "红旗" in c.get("verdict", ""))
    warnings = sum(1 for c in all_checks if "疑点" in c.get("verdict", ""))

    # 3. 重新评估风险等级
    if red_flags >= 3:
        risk_level = "高"
    elif red_flags >= 1 or warnings >= 3:
        risk_level = "中"
    elif warnings >= 1:
        risk_level = "低"
    else:
        risk_level = "低"

    # 结合统计结果调整
    if stat_result:
        if stat_result.get("benford_deviation") and stat_result["benford_deviation"] > 0.3:
            risk_level = "高" if risk_level == "中" else risk_level
        if stat_result.get("p_value_abnormal", 0) > 2:
            risk_level = max(risk_level, "中", key=["低", "中", "高", "可疑黑产"].index)

    # 4. 计算打假得分
    detection_score = red_flags * 30 + warnings * 10
    if stat_result and stat_result.get("benford_deviation"):
        detection_score += int(stat_result["benford_deviation"] * 50)

    # 5. 综合所有summary
    summaries = [r.get("summary", "") for r in reports if not r.get("parse_error")]
    merged_summary = " | ".join(s for s in summaries if s)[:200]
    if len(summaries) > 1:
        merged_summary = f"[合并{len(reports)}段审查] {merged_summary}"

    # 6. 综合conclusion
    conclusions = [r.get("conclusion", "") for r in reports if not r.get("parse_error") and r.get("conclusion")]
    merged_conclusion = "\n\n".join(conclusions) if conclusions else ""

    # 清理临时字段
    for c in all_checks:
        c.pop("_source_chunk", None)

    return {
        "summary": merged_summary,
        "risk_level": risk_level,
        "detection_score": detection_score,
        "checks": all_checks,
        "conclusion": merged_conclusion,
        "_merged_from": len(reports),
    }


# ══════════════════════════════════════════════════════════════
# 报告解析与格式化
# ══════════════════════════════════════════════════════════════

def parse_report(content):
    """解析LLM返回的JSON报告，容错处理"""
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass
    m = re.search(r'\{[\s\S]*\}', content)
    if m:
        try:
            return json.loads(m.group())
        except json.JSONDecodeError:
            pass
    return {"raw_output": content, "parse_error": True}


def format_report(report, pdf_path, meta, stat_result):
    """将审查结果格式化为Markdown报告"""
    risk_icons = {"高": "🔴", "中": "🟡", "低": "🟢", "可疑黑产": "⚫️"}
    lines = [
        f"# 📄 学术论文审查报告 [耿同学标准]",
        f"",
        f"**文件**: `{pdf_path}`",
        f"**文件大小**: {meta.get('size_mb', 'N/A')} MB",
        f"**提取字符数**: {meta.get('total_chars', meta.get('chars', 'N/A'))}",
        f"**提取方式**: {meta.get('extraction_method', meta.get('source', 'N/A'))}",
    ]
    # 显示分块信息（如果是分块审查）
    if meta.get("chunk_count") and meta["chunk_count"] > 1:
        lines.append(f"**审查方式**: 分块审查 | {meta['chunk_count']}块 | 单块上限{meta['chunk_size']}字符 | 重叠{meta['overlap']}字符")
    lines.append(f"**审查时间**: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    lines.extend([
        f"",
        f"## 📊 本地统计检测结果",
        f"| 检测项 | 结果 | 状态 |",
        f"|--------|------|------|",
        f"| Benford分布偏差 | {round(stat_result['benford_deviation'],3) if stat_result['benford_deviation'] else '样本不足'} | {stat_result['benford_status'] or 'N/A'} |",
        f"| p值数量/异常 | {stat_result['p_value_count']} / {stat_result['p_value_abnormal']}个>0.05 | {'⚠️异常' if stat_result['p_value_abnormal'] else '✅正常'} |",
        f"| 标准差提及 | {stat_result['sd_count']}处 | N/A |",
    ])
    lines.append(f"| 提取数字数 | {stat_result['number_count']} | - |")

    if stat_result.get("number_consistency"):
        lines.append(f"| 数字自洽性 | {stat_result['number_consistency']} | ⚠️矛盾 |")

    lines.append("")

    if report.get("parse_error"):
        lines.append("## ⚠️ LLM报告解析失败（原始输出）")
        lines.append(f"```\n{report['raw_output']}\n```")
        return "\n".join(lines)

    lines.append(f"## 总评: {report.get('summary', 'N/A')}")
    risk = report.get('risk_level', '未知')
    lines.append(f"**风险等级**: {risk_icons.get(risk, '⚪')} {risk}")
    lines.append(f"**打假得分**: {report.get('detection_score', 0)} (越高越可疑)")
    lines.append("")

    checks = report.get("checks", [])
    if checks:
        lines.append("## 🔍 逐项检查")
        lines.append("")
        lines.append("| # | 分类 | 检查项 | 判定 |")
        lines.append("|---|------|--------|------|")
        for i, c in enumerate(checks, 1):
            lines.append(f"| {i} | {c.get('category', 'N/A')} | {c.get('item', 'N/A')} | {c.get('verdict', 'N/A')} |")
        lines.append("")

        lines.append("## 📋 详细分析")
        lines.append("")
        for i, c in enumerate(checks, 1):
            lines.append(f"### {i}. {c.get('category', 'N/A')} - {c.get('item', 'N/A')} — {c.get('verdict', 'N/A')}")
            if c.get("evidence"):
                lines.append(f"> **证据**: {c['evidence']}")
            if c.get("detail"):
                lines.append(f"\n{c['detail']}")
            lines.append("")

    if report.get("conclusion"):
        lines.append("## 📝 综合结论")
        lines.append(f"\n{report['conclusion']}")

    return "\n".join(lines)


def format_html_report(report, pdf_path, meta, stat_result):
    """将审查结果格式化为美观的HTML报告"""
    risk_colors = {"高": "#dc2626", "中": "#f59e0b", "低": "#16a34a", "可疑黑产": "#7c3aed"}
    risk_icons = {"高": "🔴", "中": "🟡", "低": "🟢", "可疑黑产": "⚫️"}
    risk = report.get('risk_level', '未知')
    risk_color = risk_colors.get(risk, "#6b7280")
    risk_icon = risk_icons.get(risk, "⚪")

    # 统计检测状态
    benford_val = round(stat_result['benford_deviation'], 3) if stat_result['benford_deviation'] else '样本不足'
    benford_status = stat_result.get('benford_status', 'N/A') or 'N/A'
    p_abnormal = stat_result['p_value_abnormal']
    p_status_class = "status-warn" if p_abnormal else "status-ok"

    # 分块信息
    chunk_info = ""
    if meta.get("chunk_count") and meta["chunk_count"] > 1:
        chunk_info = f"""
        <div class="info-row">
            <span class="info-label">审查方式</span>
            <span class="info-value">分块审查 | {meta['chunk_count']}块 | 单块上限{meta['chunk_size']}字符 | 重叠{meta['overlap']}字符</span>
        </div>"""

    # 数字自洽性
    number_consistency = ""
    if stat_result.get("number_consistency"):
        number_consistency = f"""
        <tr>
            <td>数字自洽性</td>
            <td>{stat_result['number_consistency']}</td>
            <td><span class="status-warn">⚠️ 矛盾</span></td>
        </tr>"""

    # 解析失败
    if report.get("parse_error"):
        checks_html = f"""
        <div class="section">
            <h2>⚠️ LLM报告解析失败（原始输出）</h2>
            <pre class="error-block">{_html_escape(report.get('raw_output', ''))}</pre>
        </div>"""
        conclusion_html = ""
    else:
        # 逐项检查表
        checks = report.get("checks", [])
        checks_table_rows = ""
        for i, c in enumerate(checks, 1):
            verdict = c.get('verdict', 'N/A')
            verdict_class = "verdict-red" if "红旗" in verdict else ("verdict-yellow" if "疑点" in verdict else "verdict-green")
            checks_table_rows += f"""
            <tr>
                <td>{i}</td>
                <td>{_html_escape(c.get('category', 'N/A'))}</td>
                <td>{_html_escape(c.get('item', 'N/A'))}</td>
                <td><span class="{verdict_class}">{_html_escape(verdict)}</span></td>
            </tr>"""

        # 详细分析
        detail_cards = ""
        for i, c in enumerate(checks, 1):
            verdict = c.get('verdict', 'N/A')
            verdict_class = "verdict-red" if "红旗" in verdict else ("verdict-yellow" if "疑点" in verdict else "verdict-green")
            detail_cards += f"""
            <div class="detail-card">
                <div class="detail-header">
                    <span class="detail-num">#{i}</span>
                    <span class="detail-cat">{_html_escape(c.get('category', 'N/A'))}</span>
                    <span class="detail-item">{_html_escape(c.get('item', 'N/A'))}</span>
                    <span class="{verdict_class} detail-verdict">{_html_escape(verdict)}</span>
                </div>
                {f'<div class="detail-evidence">📋 <strong>证据</strong>: {_html_escape(c["evidence"])}</div>' if c.get("evidence") else ''}
                {f'<div class="detail-text">{_html_escape(c["detail"])}</div>' if c.get("detail") else ''}
            </div>"""

        checks_html = f"""
        <div class="section">
            <h2>🔍 逐项检查</h2>
            <table class="checks-table">
                <thead><tr><th>#</th><th>分类</th><th>检查项</th><th>判定</th></tr></thead>
                <tbody>{checks_table_rows}</tbody>
            </table>
        </div>
        <div class="section">
            <h2>📋 详细分析</h2>
            {detail_cards}
        </div>"""

        conclusion_html = ""
        if report.get("conclusion"):
            conclusion_html = f"""
            <div class="section conclusion-section">
                <h2>📝 综合结论</h2>
                <p class="conclusion-text">{_html_escape(report['conclusion'])}</p>
            </div>"""

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>📄 学术论文审查报告</title>
<style>
  :root {{
    --bg: #0f172a;
    --surface: #1e293b;
    --surface2: #334155;
    --text: #e2e8f0;
    --text-muted: #94a3b8;
    --accent: #38bdf8;
    --border: #475569;
    --red: #ef4444;
    --yellow: #f59e0b;
    --green: #22c55e;
  }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Noto Sans SC', sans-serif;
    background: var(--bg);
    color: var(--text);
    line-height: 1.6;
    padding: 20px;
  }}
  .container {{ max-width: 960px; margin: 0 auto; }}
  .header {{
    background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
    border: 1px solid var(--border);
    border-radius: 16px;
    padding: 32px;
    margin-bottom: 24px;
    text-align: center;
  }}
  .header h1 {{ font-size: 28px; margin-bottom: 16px; }}
  .risk-badge {{
    display: inline-block;
    font-size: 24px;
    font-weight: 700;
    padding: 8px 24px;
    border-radius: 999px;
    color: #fff;
    background: {risk_color};
    margin: 8px 0;
  }}
  .meta-grid {{
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 8px 24px;
    text-align: left;
    margin-top: 16px;
    font-size: 14px;
    color: var(--text-muted);
  }}
  .meta-grid strong {{ color: var(--text); }}
  .section {{
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 12px;
    padding: 24px;
    margin-bottom: 20px;
  }}
  .section h2 {{
    font-size: 20px;
    margin-bottom: 16px;
    padding-bottom: 8px;
    border-bottom: 1px solid var(--border);
  }}
  table {{ width: 100%; border-collapse: collapse; font-size: 14px; }}
  th, td {{ padding: 10px 12px; text-align: left; border-bottom: 1px solid var(--border); }}
  th {{ background: var(--surface2); color: var(--text-muted); font-weight: 600; text-transform: uppercase; font-size: 12px; }}
  tr:hover {{ background: rgba(56,189,248,0.05); }}
  .status-ok {{ color: var(--green); font-weight: 600; }}
  .status-warn {{ color: var(--yellow); font-weight: 600; }}
  .verdict-red {{ color: var(--red); font-weight: 700; }}
  .verdict-yellow {{ color: var(--yellow); font-weight: 700; }}
  .verdict-green {{ color: var(--green); font-weight: 700; }}
  .detail-card {{
    background: var(--surface2);
    border-radius: 8px;
    padding: 16px;
    margin-bottom: 12px;
    border-left: 4px solid var(--accent);
  }}
  .detail-header {{ display: flex; align-items: center; gap: 12px; flex-wrap: wrap; }}
  .detail-num {{
    background: var(--accent);
    color: var(--bg);
    border-radius: 50%;
    width: 28px; height: 28px;
    display: flex; align-items: center; justify-content: center;
    font-weight: 700; font-size: 13px;
  }}
  .detail-cat {{ color: var(--accent); font-weight: 600; }}
  .detail-item {{ flex: 1; }}
  .detail-verdict {{ font-size: 14px; }}
  .detail-evidence {{
    margin-top: 10px;
    padding: 10px 14px;
    background: rgba(245,158,11,0.1);
    border-radius: 6px;
    font-size: 14px;
    color: #fbbf24;
  }}
  .detail-text {{
    margin-top: 8px;
    font-size: 14px;
    color: var(--text-muted);
    white-space: pre-wrap;
  }}
  .conclusion-section {{ border-left: 4px solid var(--green); }}
  .conclusion-text {{ font-size: 16px; white-space: pre-wrap; color: var(--text); }}
  .error-block {{
    background: rgba(239,68,68,0.1);
    border: 1px solid var(--red);
    border-radius: 8px;
    padding: 16px;
    white-space: pre-wrap;
    font-family: monospace;
    font-size: 13px;
    color: #fca5a5;
    overflow-x: auto;
  }}
  .score-bar {{
    height: 8px;
    background: var(--surface2);
    border-radius: 4px;
    overflow: hidden;
    margin-top: 8px;
  }}
  .score-fill {{
    height: 100%;
    border-radius: 4px;
    background: linear-gradient(90deg, var(--green), var(--yellow), var(--red));
    transition: width 0.5s;
  }}
  .footer {{
    text-align: center;
    color: var(--text-muted);
    font-size: 12px;
    margin-top: 32px;
    padding: 16px;
  }}
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <h1>📄 学术论文审查报告</h1>
    <div class="risk-badge">{risk_icon} {risk}</div>
    <div style="margin-top:8px; color:var(--text-muted);">
      打假得分: <strong style="color:{risk_color}">{report.get('detection_score', 0)}</strong>（越高越可疑）
      <div class="score-bar"><div class="score-fill" style="width:{min(report.get('detection_score', 0), 100)}%"></div></div>
    </div>
    <div class="meta-grid">
      <div><strong>文件</strong>: {_html_escape(pdf_path)}</div>
      <div><strong>文件大小</strong>: {meta.get('size_mb', 'N/A')} MB</div>
      <div><strong>提取字符数</strong>: {meta.get('total_chars', meta.get('chars', 'N/A'))}</div>
      <div><strong>提取方式</strong>: {meta.get('extraction_method', meta.get('source', 'N/A'))}</div>
      {chunk_info.replace('class="info-row"', '').replace('class="info-label"', '').replace('class="info-value"', '').strip().replace('<span>', '<div>').replace('</span>', '</div>') if chunk_info else '<div></div><div></div>'}
      <div><strong>审查时间</strong>: {time.strftime('%Y-%m-%d %H:%M:%S')}</div>
    </div>
  </div>

  <div class="section">
    <h2>📊 本地统计检测结果</h2>
    <table>
      <thead><tr><th>检测项</th><th>结果</th><th>状态</th></tr></thead>
      <tbody>
        <tr><td>Benford分布偏差</td><td>{benford_val}</td><td>{benford_status}</td></tr>
        <tr><td>p值数量/异常</td><td>{stat_result['p_value_count']} / {stat_result['p_value_abnormal']}个&gt;0.05</td><td><span class="{p_status_class}">{'⚠️异常' if p_abnormal else '✅正常'}</span></td></tr>
        <tr><td>标准差提及</td><td>{stat_result['sd_count']}处</td><td>N/A</td></tr>
        <tr><td>提取数字数</td><td>{stat_result['number_count']}</td><td>-</td></tr>
        {number_consistency}
      </tbody>
    </table>
  </div>

  {checks_html}
  {conclusion_html}

  <div class="footer">
    Generated by <strong>Veritas</strong> — 学术论文自动审查工具（耿同学标准） | {time.strftime('%Y-%m-%d %H:%M:%S')}
  </div>
</div>
</body>
</html>"""
    return html


def _html_escape(text):
    """HTML特殊字符转义"""
    if not text:
        return ""
    return (str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("\n", "<br>"))

def update_patterns(comments_file):
    """从PubPeer评论文本中用LLM提取新的欺诈模式，更新知识库
    
    comments_file: 包含PubPeer评论文本的文件路径
    """
    from datetime import datetime
    
    comments_path = Path(comments_file)
    if not comments_path.exists():
        print(f"❌ 评论文本文件不存在: {comments_path}")
        return 1
    
    with open(comments_path, "r", encoding="utf-8") as f:
        comments_text = f.read()
    
    if len(comments_text.strip()) < 20:
        print("❌ 评论文本内容过少，请提供更完整的PubPeer评论内容")
        return 1
    
    print(f"📖 已读取评论文本: {len(comments_text)}字符")
    print("🤖 正在用LLM分析评论，提取欺诈模式...")
    
    # 构建提取prompt
    extract_prompt = f"""分析以下来自PubPeer的学术评论，提取其中涉及的学术论文造假/可疑手法。

要求：
1. 每个造假手法提取为一个独立的模式条目
2. 按JSON数组格式输出，每个条目包含：id(英文大写下划线), category(分类), name(中文名), description(详细描述), detection_hint(检测提示), risk_level(高/中/低)
3. 只提取确实存在的造假手法，不要臆造
4. 合并相似的造假手法

PubPeer评论内容：
{comments_text}

输出格式：
[
  {{
    "id": "PATTERN_ID",
    "category": "图片与图表/数据与结果/方法论/结构与引用/作者与期刊",
    "name": "手法名称",
    "description": "手法描述",
    "detection_hint": "审查时如何检测此手法",
    "risk_level": "高/中/低"
  }}
]"""
    
    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": "你是一个学术论文打假专家，擅长从PubPeer评论中识别和归纳造假手法。"},
            {"role": "user", "content": extract_prompt}
        ],
        "temperature": 0.3,
    }
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {LLM_API_KEY}",
    }
    
    req = urllib.request.Request(
        LLM_API_URL,
        data=json.dumps(payload).encode(),
        headers=headers,
        method="POST",
    )
    
    try:
        resp = urllib.request.urlopen(req, timeout=60)
        result = json.loads(resp.read().decode("utf-8"))
        content = result["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"❌ LLM调用失败: {e}")
        return 1
    
    # 解析LLM输出的JSON
    json_match = re.search(r'\[[\s\S]*\]', content)
    if not json_match:
        print("❌ LLM未能输出有效的JSON格式，请重试")
        print(f"原始输出: {content[:500]}")
        return 1
    
    try:
        new_patterns = json.loads(json_match.group())
    except json.JSONDecodeError as e:
        print(f"❌ JSON解析失败: {e}")
        return 1
    
    if not new_patterns:
        print("⚠️ 未能从评论中提取到新的欺诈模式")
        return 0
    
    # 加载现有知识库
    if FRAUD_PATTERNS_PATH.exists():
        with open(FRAUD_PATTERNS_PATH, "r", encoding="utf-8") as f:
            kb_data = json.load(f)
        existing_ids = {p["id"] for p in kb_data.get("patterns", [])}
    else:
        kb_data = {"schema_version": "1.0", "last_updated": "", "contributors": ["community"], "patterns": []}
        existing_ids = set()
    
    # 去重合并
    added = 0
    for p in new_patterns:
        if p.get("id") and p["id"] not in existing_ids:
            kb_data["patterns"].append(p)
            existing_ids.add(p["id"])
            added += 1
            print(f"  ✅ 新增: [{p.get('risk_level','?')}] {p.get('name','?')}")
        else:
            print(f"  ⏭️ 跳过已存在: {p.get('name','?')} ({p.get('id','?')})")
    
    if added > 0:
        kb_data["last_updated"] = datetime.now().strftime("%Y-%m-%d")
        with open(FRAUD_PATTERNS_PATH, "w", encoding="utf-8") as f:
            json.dump(kb_data, f, ensure_ascii=False, indent=2)
        print(f"\n🎉 知识库已更新！新增{added}条模式，总计{len(kb_data['patterns'])}条")
    else:
        print("\n⚠️ 无新增模式，知识库未变更")
    
    return 0


# ══════════════════════════════════════════════════════════════
# 腾讯朱雀AI文本检测辅助功能
# ══════════════════════════════════════════════════════════════

ZHUQUE_URL = "https://matrix.tencent.com/ai-detect/"


def copy_to_clipboard(text: str) -> bool:
    """跨平台复制文本到系统剪贴板"""
    system = platform.system()
    try:
        if system == "Windows":
            # Windows: 使用clip命令
            process = subprocess.Popen(
                ["clip"], stdin=subprocess.PIPE, shell=True
            )
            process.communicate(text.encode("utf-16"))
            return process.returncode == 0
        elif system == "Darwin":  # macOS
            process = subprocess.Popen(
                ["pbcopy"], stdin=subprocess.PIPE
            )
            process.communicate(text.encode("utf-8"))
            return process.returncode == 0
        else:  # Linux
            # 优先尝试xclip，其次xsel
            for cmd in [["xclip", "-selection", "clipboard"], ["xsel", "--clipboard", "--input"]]:
                try:
                    process = subprocess.Popen(cmd, stdin=subprocess.PIPE)
                    process.communicate(text.encode("utf-8"))
                    if process.returncode == 0:
                        return True
                except FileNotFoundError:
                    continue
            return False
    except Exception as e:
        print(f"⚠️ 剪贴板写入失败: {e}")
        return False


def launch_zhuque_ai_detect(text: str):
    """启动腾讯朱雀AI文本检测：复制文本到剪贴板 → 打开检测页面 → 弹窗提醒"""
    print("\n" + "=" * 60)
    print("🤖 腾讯朱雀AI文本检测")
    print("=" * 60)

    # 1) 复制文本到剪贴板
    # 朱雀检测有字数限制，截取前8000字符
    detect_text = text[:8000]
    if len(text) > 8000:
        print(f"⚠️ 文本较长({len(text)}字符)，仅复制前8000字符到剪贴板（朱雀字数限制）")

    clip_ok = copy_to_clipboard(detect_text)
    if clip_ok:
        print("✅ 文本已复制到剪贴板")
    else:
        print("❌ 剪贴板写入失败，请手动复制论文文本")

    # 2) 打开浏览器
    print(f"🌐 正在打开朱雀AI检测页面...")
    webbrowser.open(ZHUQUE_URL)

    # 3) 弹窗提醒
    system = platform.system()
    try:
        if system == "Windows":
            import ctypes
            ctypes.windll.user32.MessageBoxW(
                0,
                "论文文本已复制到剪贴板！\n\n"
                "请在打开的朱雀AI检测页面中粘贴文本并点击检测。\n"
                "检测完成后，点击确定继续后续审查流程。",
                "🤖 朱雀AI文本检测",
                0x40  # MB_ICONINFORMATION
            )
        elif system == "Darwin":
            subprocess.run([
                "osascript", "-e",
                'display dialog "论文文本已复制到剪贴板！\n\n请在朱雀AI检测页面中粘贴文本并点击检测。\n检测完成后，点击确定继续后续审查流程。" '
                'buttons {"确定"} default button "确定" with title "🤖 朱雀AI文本检测" with icon note'
            ])
        else:  # Linux
            # 尝试zenity
            try:
                subprocess.run([
                    "zenity", "--info", "--title=🤖 朱雀AI文本检测", "--width=400",
                    "--text=论文文本已复制到剪贴板！\n\n请在朱雀AI检测页面中粘贴文本并点击检测。\n检测完成后，点击确定继续后续审查流程。"
                ])
            except FileNotFoundError:
                # 降级为终端提示
                input("\n⏸️ 论文文本已复制到剪贴板，请在浏览器中粘贴检测。\n检测完成后按回车继续...")
    except Exception:
        # 最终降级：终端等待
        input("\n⏸️ 论文文本已复制到剪贴板，请在浏览器中粘贴检测。\n检测完成后按回车继续...")

    print("✅ 朱雀AI检测流程结束，继续后续审查...")


# ──────────────────────────────────────────────────────────────
# AI图片检测（imagedetector.com）
# ──────────────────────────────────────────────────────────────

IMAGE_DETECT_URL = "https://imagedetector.com/"

# 支持的图片扩展名
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif", ".webp"}


def extract_images_from_pdf(pdf_path: str) -> List[str]:
    """从PDF中提取内嵌图片到临时目录，返回图片路径列表

    优先使用PyMuPDF(fitz)，降级使用pdf2image整页渲染
    """
    images = []
    tmp_dir = os.path.join(os.path.dirname(pdf_path), "_veritas_images_tmp")
    os.makedirs(tmp_dir, exist_ok=True)

    # 方案1：PyMuPDF提取内嵌图片
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        img_count = 0
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            img_list = page.get_images(full=True)
            for img_idx, img_info in enumerate(img_list):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    if base_image and base_image.get("image"):
                        ext = base_image.get("ext", "png")
                        if ext not in ("png", "jpg", "jpeg", "bmp", "tiff", "webp"):
                            ext = "png"
                        fname = f"page{page_idx + 1}_img{img_idx + 1}.{ext}"
                        fpath = os.path.join(tmp_dir, fname)
                        with open(fpath, "wb") as f:
                            f.write(base_image["image"])
                        # 过滤掉太小的图片（图标、装饰等）
                        if os.path.getsize(fpath) > 5000:
                            images.append(fpath)
                            img_count += 1
                except Exception:
                    continue
        doc.close()
        if img_count > 0:
            print(f"  📎 PyMuPDF提取 {img_count} 张内嵌图片 → {tmp_dir}")
            return images
    except ImportError:
        pass
    except Exception as e:
        print(f"  ⚠️ PyMuPDF提取失败: {e}")

    # 方案2：pdf2image整页渲染
    try:
        from pdf2image import convert_from_path
        pages = convert_from_path(pdf_path, dpi=200)
        for i, page_img in enumerate(pages):
            fname = f"page{i + 1}_full.png"
            fpath = os.path.join(tmp_dir, fname)
            page_img.save(fpath, "PNG")
            if os.path.getsize(fpath) > 10000:
                images.append(fpath)
        if images:
            print(f"  📎 pdf2image渲染 {len(images)} 页 → {tmp_dir}")
            return images
    except ImportError:
        pass
    except Exception as e:
        print(f"  ⚠️ pdf2image渲染失败: {e}")

    return images


def collect_image_files(input_path: str) -> List[str]:
    """收集论文相关图片文件（从目录扫描 + PDF提取）

    返回所有图片的绝对路径列表
    """
    images = []
    p = Path(input_path)

    if p.is_file() and p.suffix.lower() == ".pdf":
        # PDF文件：提取内嵌图片
        print("  📸 从PDF中提取图片...")
        extracted = extract_images_from_pdf(str(p))
        images.extend(extracted)
    elif p.is_dir():
        # 目录：扫描所有图片文件
        for ext in IMAGE_EXTENSIONS:
            for f in p.rglob(f"*{ext}"):
                if f.stat().st_size > 5000:  # 过滤小图标
                    images.append(str(f))

    return images


def launch_image_ai_detect(input_path: str):
    """启动imagedetector.com图片AI检测：收集图片 → 打开浏览器 → 逐张提醒上传"""
    print("\n" + "=" * 60)
    print("🖼️ AI图片检测 (imagedetector.com)")
    print("=" * 60)

    # 1) 收集图片
    images = collect_image_files(input_path)

    if not images:
        print("⚠️ 未找到可检测的图片文件")
        print("💡 提示: 可手动访问 https://imagedetector.com/ 上传图片检测")
        return

    # 限制数量（太多图片会非常耗时）
    MAX_IMAGES = 30
    if len(images) > MAX_IMAGES:
        print(f"⚠️ 图片数量过多({len(images)}张)，仅处理前{MAX_IMAGES}张（小图标已过滤）")
        images = images[:MAX_IMAGES]

    print(f"  📋 共 {len(images)} 张图片待检测")

    # 2) 打开浏览器
    print(f"🌐 正在打开 imagedetector.com 检测页面...")
    webbrowser.open(IMAGE_DETECT_URL)
    time.sleep(2)  # 等待浏览器加载

    # 3) 逐张提醒上传
    system = platform.system()
    for idx, img_path in enumerate(images, 1):
        img_name = os.path.basename(img_path)
        msg = (
            f"图片 [{idx}/{len(images)}]: {img_name}\n\n"
            f"路径: {img_path}\n\n"
            f"请在 imagedetector.com 页面中上传此图片并点击检测。\n"
            f"检测完成后，点击确定继续下一张图片。"
        )

        print(f"\n  📸 [{idx}/{len(images)}] {img_name}")
        print(f"     路径: {img_path}")

        try:
            if system == "Windows":
                import ctypes
                result = ctypes.windll.user32.MessageBoxW(
                    0,
                    msg,
                    f"🖼️ AI图片检测 ({idx}/{len(images)})",
                    0x41  # MB_OKCANCEL + MB_ICONWARNING
                )
                if result == 2:  # IDCANCEL
                    print("  ⏹️ 用户取消图片检测流程")
                    break
            elif system == "Darwin":
                subprocess.run([
                    "osascript", "-e",
                    f'display dialog "{msg}" '
                    f'buttons {{"取消", "确定"}} default button "确定" with title "🖼️ AI图片检测 ({idx}/{len(images)})" with icon note'
                ])
            else:  # Linux
                try:
                    subprocess.run([
                        "zenity", "--info", "--title=🖼️ AI图片检测", "--width=450",
                        f"--text={msg}"
                    ])
                except FileNotFoundError:
                    input(f"  ⏸️ 按回车继续下一张...")
        except Exception:
            input(f"  ⏸️ 检测完成后按回车继续下一张...")

    print("✅ AI图片检测流程结束，继续后续审查...")

    # 清理临时提取目录
    tmp_dir = os.path.join(os.path.dirname(input_path) if os.path.isfile(input_path) else input_path, "_veritas_images_tmp")
    if os.path.isdir(tmp_dir):
        try:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)
            print("  🧹 已清理临时图片目录")
        except Exception:
            pass


def main():
    parser = argparse.ArgumentParser(
        description="学术论文自动审查工具（耿同学标准 + MinerU）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # MinerU提取 + 完整审查（推荐）
  python paper_audit.py paper.pdf --mineru

  # 仅原始PDF文本提取 + 审查
  python paper_audit.py paper.pdf

  # 指定输出路径
  python paper_audit.py paper.pdf --mineru -o report.md --json
  
  # 更新欺诈模式知识库（从PubPeer评论自动提取新pattern）
  python paper_audit.py --update-patterns pubpeer_comments.txt
""")
    parser.add_argument("pdf_path", nargs='?', help="待审查的文件路径或论文目录路径（支持PDF/Word/Excel/Supplement等，更新模式下无需提供）")
    parser.add_argument("--update-patterns", metavar="COMMENTS_FILE", 
                        help="从PubPeer评论文本文件中自动提取新的欺诈模式，更新知识库")
    parser.add_argument("--mineru", action="store_true",
                        help="使用MinerU API将PDF转为Markdown再审查（推荐，质量更高）")
    parser.add_argument("--mineru-model", default="vlm",
                        choices=["pipeline", "vlm", "MinerU-HTML"],
                        help="MinerU模型版本（默认vlm，仅Precision API生效）")
    parser.add_argument("--mineru-lang", default="ch",
                        help="MinerU OCR语言（默认ch=中英，en=英文，japan=日文）")
    parser.add_argument("--no-mineru", action="store_true",
                        help="强制使用原始PDF文本提取，禁用MinerU")
    parser.add_argument("--max-chars", type=int, default=12000,
                        help="提取文本最大字符数（默认12000）")
    parser.add_argument("--output", "-o", help="输出报告文件路径（默认输出到同目录）")
    parser.add_argument("--json", action="store_true", help="同时保存原始JSON结果")
    parser.add_argument("--ai-detect", action="store_true", help="开启腾讯朱雀AI文本检测：自动复制文本到剪贴板+打开检测页面")
    parser.add_argument("--image-detect", action="store_true", help="开启AI图片检测（imagedetector.com）：自动提取图片+打开检测页面+逐张提醒上传")
    args = parser.parse_args()

    # ─── 知识库更新模式 ───
    if args.update_patterns:
        return update_patterns(args.update_patterns)

    # ─── 正常审查模式 ───
    if not args.pdf_path:
        parser.error("审查模式需要提供path参数（文件或目录，或使用 --update-patterns 更新知识库）")

    input_path = Path(args.pdf_path)
    if not input_path.exists():
        print(f"❌ 路径不存在: {input_path}")
        return 1

    # ─── 阶段1：文本提取（支持单个文件/整个论文目录） ───
    full_text = None
    meta = {}
    raw_pdf = None
    use_mineru = args.mineru and not args.no_mineru

    if input_path.is_dir():
        print(f"📂 检测到输入为目录，正在扫描所有论文相关文件...")
        file_classes, all_files = find_project_files(input_path)
        print(f"✅ 找到 {len(all_files)} 个相关文件:")
        for cat, files in file_classes.items():
            if files: print(f"  - {cat}: {len(files)} 个文件")
        
        # 提取所有文件文本合并
        full_text = ""
        total_files = len(all_files)
        for idx, file_path in enumerate(all_files, 1):
            print(f"  📝 提取 [{idx}/{total_files}] {file_path.name}...")
            file_content = extract_text_from_file(file_path, use_mineru=use_mineru)
            full_text += f"\n\n=== 文件: {file_path.name} 路径: {file_path.relative_to(input_path)} ==="
            full_text += "\n" + file_content
        
        meta = {
            "input_type": "directory",
            "total_files": total_files,
            "file_classes": {k: len(v) for k, v in file_classes.items()},
            "total_chars": len(full_text),
            "extractor": "directory_multi_format"
        }
        print(f"\n✅ 所有文件提取完成，总长度: {len(full_text)} 字符")
    else:
        # 单个文件走原有流程
        pdf_path = input_path
        print(f"📄 检测到输入为单个文件: {pdf_path.name}")

        if use_mineru:
            print(f"📡 [MinerU] 正在将PDF转为Markdown: {pdf_path.name}")
            md_text, md_meta = mineru_extract(pdf_path, language=args.mineru_lang)
            if md_text:
                full_text = md_text  # 保留全文
                meta = {
                    "size_mb": round(pdf_path.stat().st_size / 1024 / 1024, 2),
                    "total_chars": len(md_text),
                    "chars_sent": len(md_text),
                    "extraction_method": f"mineru_{md_meta.get('source', 'unknown')}",
                }
                if md_meta.get("batch_id"):
                    meta["mineru_batch_id"] = md_meta["batch_id"]
                if md_meta.get("task_id"):
                    meta["mineru_task_id"] = md_meta["task_id"]
                print(f"✅ MinerU提取完成: {len(md_text)} 字符（全文保留）")
            else:
                err = md_meta.get("error", "未知错误") if md_meta else "未知错误"
                print(f"❌ MinerU提取失败: {err}")
                print(f"⚠️ 降级使用原始PDF文本提取...")
                use_mineru = False

        if not use_mineru or full_text is None:
            print(f"📖 正在提取PDF文本: {pdf_path}")
            # extract_pdf_text的max_chars参数传大值以获取全文
            full_text, meta, raw_pdf = extract_pdf_text(str(pdf_path), max_chars=999999)
            if not full_text:
                print("❌ 未能从PDF中提取到文本（可能是扫描件或加密PDF）")
                print("💡 建议: 使用 --mineru 参数通过MinerU API提取（支持OCR）")
                return 1
            print(f"✅ 提取完成: {meta['total_chars']} 字符（全文保留）")

    # ─── 朱雀AI文本检测（可选） ───
    if args.ai_detect:
        launch_zhuque_ai_detect(full_text)

    # ─── AI图片检测（可选） ───
    if args.image_detect:
        launch_image_ai_detect(args.path)

    # ─── 阶段2：本地统计检测（使用全文，统计不截断） ───
    print(f"🔢 正在执行本地统计检测...")
    stat_result = local_stat_check(full_text)
    benford_str = f"{round(stat_result['benford_deviation'],3)}" if stat_result['benford_deviation'] else 'N/A'
    print(f"✅ 统计检测完成: Benford偏差={benford_str}, p值异常={stat_result['p_value_abnormal']}, 数字数={stat_result['number_count']}")

    # ─── 阶段3：智能分块 + LLM语义审查（冗余机制） ───
    chunk_size = args.max_chars  # 复用max_chars作为单块上限
    overlap = min(1000, chunk_size // 8)  # 重叠区约12.5%

    chunks = smart_chunk_text(full_text, chunk_size=chunk_size, overlap=overlap)
    total_chunks = len(chunks)

    if total_chunks == 1:
        # 短论文：直接全文审查
        print(f"🔍 论文长度({len(full_text)}字符)在单块范围内，直接审查...")
        try:
            raw_content = call_llm(full_text)
            report = parse_report(raw_content)
        except Exception as e:
            print(f"❌ LLM调用失败: {e}")
            report = {"parse_error": True, "raw_output": str(e)}
    else:
        # 长论文：分块审查 + 合并
        print(f"🔍 论文较长({len(full_text)}字符)，分为{total_chunks}块(每块≤{chunk_size}字符，重叠{overlap}字符)进行审查...")
        chunk_reports = []
        for chunk_text, chunk_idx, _ in chunks:
            print(f"  📝 审查第{chunk_idx+1}/{total_chunks}块({len(chunk_text)}字符)...")
            try:
                raw_content = call_llm(chunk_text, chunk_info=(chunk_idx, total_chunks))
                chunk_report = parse_report(raw_content)
                chunk_reports.append(chunk_report)
                # 简要反馈
                if not chunk_report.get("parse_error"):
                    risk = chunk_report.get("risk_level", "未知")
                    print(f"     → 第{chunk_idx+1}块风险: {risk}")
                else:
                    print(f"     → 第{chunk_idx+1}块解析异常")
            except Exception as e:
                print(f"  ⚠️ 第{chunk_idx+1}块LLM调用失败: {e}")
                chunk_reports.append({"parse_error": True, "raw_output": str(e)})

        # 合并所有块的审查结果
        print(f"🔗 正在合并{len(chunk_reports)}块审查结果...")
        report = merge_chunk_reports(chunk_reports, stat_result)
        if report.get("_merged_from"):
            print(f"✅ 合并完成: 来自{report['_merged_from']}块，共{len(report.get('checks', []))}个检查项")
            # 更新meta
            meta["chunk_count"] = total_chunks
            meta["chunk_size"] = chunk_size
            meta["overlap"] = overlap

    # ─── 阶段4：生成报告 ───
    md_report = format_report(report, str(pdf_path), meta, stat_result)

    # 生成HTML报告
    html_report = format_html_report(report, str(pdf_path), meta, stat_result)

    # 确定输出路径（优先HTML）
    if args.output:
        output_path = Path(args.output)
        # 如果用户指定了.md后缀，同时生成HTML
        html_output_path = output_path.with_suffix(".html")
    else:
        output_path = pdf_path.with_suffix(".audit.md")
        html_output_path = pdf_path.with_suffix(".audit.html")

    # 写入Markdown报告
    output_path.write_text(md_report, encoding="utf-8")
    print(f"✅ Markdown报告已保存: {output_path}")

    # 写入HTML报告
    html_output_path.write_text(html_report, encoding="utf-8")
    print(f"✅ HTML报告已保存: {html_output_path}")

    if args.json:
        json_path = pdf_path.with_suffix(".audit.json")
        json_path.write_text(
            json.dumps({"llm_report": report, "stat_result": stat_result, "meta": meta},
                       ensure_ascii=False, indent=2),
            encoding="utf-8")
        print(f"✅ 原始JSON已保存: {json_path}")

    # 自动打开HTML报告
    try:
        html_abs = str(html_output_path.resolve())
        webbrowser.open(f"file:///{html_abs}" if platform.system() == "Windows" else f"file://{html_abs}")
        print(f"🌐 已在浏览器中打开HTML报告")
    except Exception as e:
        print(f"⚠️ 自动打开浏览器失败: {e}，请手动打开: {html_output_path}")

    # 打印摘要
    if not report.get("parse_error"):
        risk = report.get("risk_level", "未知")
        print(f"\n📊 风险等级: {risk} | 总评: {report.get('summary', 'N/A')}")

    return 0


if __name__ == "__main__":
    exit(main())
